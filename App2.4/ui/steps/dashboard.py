"""Dashboard-style result explorer with chart, table, and map blocks.

The dashboard is for decision makers. It focuses on crash patterns,
spatial-unit rankings, crash-density/HIN outputs, and selected read-only map
views. It does not overwrite workflow results.
"""

import io
import re
import html
import base64
import zipfile
from datetime import datetime
try:
    from zoneinfo import ZoneInfo
except Exception:  # pragma: no cover
    ZoneInfo = None

import pandas as pd
import streamlit as st
import plotly.express as px
import plotly.io as pio

try:
    import geopandas as gpd
except Exception:  # pragma: no cover
    gpd = None

try:
    import folium
    from folium.features import GeoJsonTooltip
    from streamlit_folium import st_folium
except Exception:  # pragma: no cover
    folium = None
    GeoJsonTooltip = None
    st_folium = None

try:
    from docx import Document
    from docx.shared import Inches
except Exception:  # pragma: no cover
    Document = None
    Inches = None

APP_PALETTE = ["#2563eb", "#16a34a", "#f97316", "#dc2626", "#7c3aed", "#0891b2", "#ca8a04", "#475569", "#db2777", "#0f766e"]
KABCO_COLOR_MAP = {
    "K": "#dc2626",
    "A": "#f97316",
    "B": "#eab308",
    "C": "#16a34a",
    "O": "#2563eb",
    "UNKNOWN": "#64748b",
    "Unknown": "#64748b",
}
MONTH_ORDER = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]


def _drop_geometry(df):
    return df.drop(columns="geometry", errors="ignore") if df is not None else None


def _safe_scalar(value):
    """Make values safe for Streamlit Arrow, JSON, HTML, and Word export."""
    if value is None:
        return None
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass
    if isinstance(value, (pd.Timestamp, datetime)):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if hasattr(value, "isoformat") and value.__class__.__name__ in ["date", "datetime", "time"]:
        try:
            return value.isoformat()
        except Exception:
            return str(value)
    if isinstance(value, (list, tuple, set, dict)):
        return str(value)
    return value


def _safe_dataframe_for_display(df):
    """Return a copy that Streamlit can display without pyarrow mixed-type errors."""
    if df is None:
        return pd.DataFrame()
    out = _drop_geometry(df).copy()
    for col in out.columns:
        if pd.api.types.is_datetime64_any_dtype(out[col]):
            out[col] = out[col].dt.strftime("%Y-%m-%d %H:%M:%S")
        elif out[col].dtype == "object":
            out[col] = out[col].map(_safe_scalar)
            # Mixed object columns, such as Location containing text and ints,
            # trigger ArrowTypeError. String conversion is safer for dashboard UI.
            out[col] = out[col].astype(str).replace({"None": "", "nan": "", "NaT": ""})
    return out


def _safe_geojson_gdf(gdf):
    """Return GeoDataFrame with JSON-safe attributes and Leaflet-safe geometry.

    TIGER and uploaded layers sometimes contain empty geometry, GeometryCollection,
    or mixed object/date attributes. Folium/streamlit-folium can fail when a
    GeoJSON feature does not have a simple ``coordinates`` member, so this
    helper explodes multipart features and removes empty/unsupported geometry.
    """
    if gdf is None:
        return gdf
    work = gdf.copy()
    if gpd is not None and not isinstance(work, gpd.GeoDataFrame):
        try:
            work = gpd.GeoDataFrame(work, geometry="geometry")
        except Exception:
            return gdf
    if "geometry" not in work.columns:
        return work
    try:
        work = work[work.geometry.notna() & (~work.geometry.is_empty)].copy()
    except Exception:
        pass
    try:
        # Split multipart roads/corridors; ignore index so Folium ids are simple.
        work = work.explode(index_parts=False, ignore_index=True)
    except Exception:
        pass
    try:
        # Folium's get_bounds handles Point/LineString/Polygon/Multi* well, but
        # GeometryCollection can trigger KeyError: 'coordinates'.
        ok_types = {
            "Point", "MultiPoint", "LineString", "MultiLineString",
            "Polygon", "MultiPolygon"
        }
        work = work[work.geometry.geom_type.isin(ok_types)].copy()
    except Exception:
        pass
    for col in list(work.columns):
        if col == "geometry":
            continue
        try:
            if pd.api.types.is_datetime64_any_dtype(work[col]):
                work[col] = work[col].dt.strftime("%Y-%m-%d %H:%M:%S")
            elif work[col].dtype == "object":
                work[col] = work[col].map(_safe_scalar)
                work[col] = work[col].astype(str).replace({"None": "", "nan": "", "NaT": ""})
        except Exception:
            work[col] = work[col].astype(str)
    return work


def _safe_geojson_string(gdf):
    """Convert a GeoDataFrame to a GeoJSON string, returning None on failure."""
    try:
        clean = _safe_geojson_gdf(gdf)
        if clean is None or getattr(clean, "empty", True):
            return None
        return clean.to_json()
    except Exception:
        return None


def _polish_figure(fig):
    """Apply a consistent, readable dashboard palette.

    Single-series bar charts stay one calm blue. Multi-trace charts, stacked
    charts, pie/donut charts, treemaps, heatmaps, and line charts keep multiple
    colors so KABCO/type/mode splits are clear.
    """
    try:
        fig.update_layout(
            template="plotly_white",
            colorway=APP_PALETTE,
            font=dict(color="#1f2937"),
            paper_bgcolor="white",
            plot_bgcolor="white",
        )
        bar_traces = [tr for tr in fig.data if getattr(tr, "type", "") == "bar"]
        line_traces = [tr for tr in fig.data if getattr(tr, "type", "") == "scatter"]
        for i, trace in enumerate(fig.data):
            t = getattr(trace, "type", "")
            marker = getattr(trace, "marker", None)
            name = str(getattr(trace, "name", ""))
            if t == "bar" and marker is not None:
                if len(bar_traces) <= 1 and (name in ["", "None"]):
                    marker.color = "#2563eb"
                else:
                    marker.color = KABCO_COLOR_MAP.get(name, APP_PALETTE[i % len(APP_PALETTE)])
            elif t == "pie" and marker is not None:
                marker.colors = APP_PALETTE
            elif t == "treemap" and marker is not None:
                try:
                    marker.colors = APP_PALETTE
                except Exception:
                    pass
            elif t == "scatter":
                try:
                    trace.line.color = KABCO_COLOR_MAP.get(name, APP_PALETTE[i % len(APP_PALETTE)])
                    if marker is not None:
                        marker.color = KABCO_COLOR_MAP.get(name, APP_PALETTE[i % len(APP_PALETTE)])
                except Exception:
                    pass
    except Exception:
        pass
    return fig

def _safe_name(value):
    return re.sub(r"[^a-z0-9]+", "_", str(value).lower()).strip("_") or "data"


def _report_time_text(timezone_name=None):
    """Return report timestamp in the user's selected/local timezone.

    Streamlit Cloud often runs in UTC, which made reports appear several
    hours later than the user's local machine. The dashboard export uses a
    selected timezone, defaulting to America/Denver for Colorado safety
    projects, and falls back to the server local clock if zoneinfo is not
    available.
    """
    tz = timezone_name or "America/Denver"
    try:
        if ZoneInfo is not None:
            return datetime.now(ZoneInfo(tz)).strftime("%Y-%m-%d %H:%M %Z")
    except Exception:
        pass
    try:
        return datetime.now().astimezone().strftime("%Y-%m-%d %H:%M %Z")
    except Exception:
        return datetime.now().strftime("%Y-%m-%d %H:%M")


def _enrich_assigned_crashes(st):
    """Return assigned crashes with original crash attributes when IDs match.

    Some workflow tables keep only assignment fields. Dashboard summaries need the
    original crash type/year/severity/reason fields, so we merge them back when a
    common crash ID is available. If no match is possible, the assigned table is
    returned unchanged.
    """
    assigned = st.session_state.get("assigned_crashes")
    crashes = st.session_state.get("crashes")
    if assigned is None:
        return None
    assigned_df = _drop_geometry(assigned).copy()
    if crashes is None:
        return assigned_df
    crash_df = _drop_geometry(crashes).copy()
    candidates = ["CrashID", "SourceCrashID", "CRASH_ID", "OBJECTID", "CaseID", "CrashNumber"]
    for key in candidates:
        if key in assigned_df.columns and key in crash_df.columns:
            add_cols = [c for c in crash_df.columns if c not in assigned_df.columns or c == key]
            try:
                return assigned_df.merge(crash_df[add_cols], on=key, how="left", suffixes=("", "_src"))
            except Exception:
                return assigned_df
    return assigned_df

def _available_tables(st):
    tables = {}

    if st.session_state.get("crashes") is not None:
        tables["Uploaded crashes"] = _drop_geometry(st.session_state["crashes"])

    assigned_enriched = _enrich_assigned_crashes(st)
    if assigned_enriched is not None:
        tables["Assigned crashes"] = assigned_enriched

    if st.session_state.get("spatial_units_density_map") is not None:
        tables["Crash density results"] = _drop_geometry(st.session_state["spatial_units_density_map"])

    if st.session_state.get("kabco_result") is not None:
        tables["Crash count / severity summary"] = _drop_geometry(st.session_state["kabco_result"])

    if st.session_state.get("section7_results") is not None:
        results = st.session_state["section7_results"]
        if results.get("risk_segments") is not None:
            tables["HIN risk segments"] = _drop_geometry(results["risk_segments"])
        if results.get("risk_windows") is not None:
            tables["Sliding windows"] = _drop_geometry(results["risk_windows"])
        if results.get("risk_corridors") is not None:
            tables["HIN corridors"] = _drop_geometry(results["risk_corridors"])
        if results.get("route_summary") is not None:
            tables["Sliding-window route summary"] = _drop_geometry(results["route_summary"])

    if st.session_state.get("corridors") is not None:
        tables["Generated corridors"] = _drop_geometry(st.session_state["corridors"])

    if st.session_state.get("final_corridors") is not None:
        tables["Filtered corridors"] = _drop_geometry(st.session_state["final_corridors"])

    return {
        name: table
        for name, table in tables.items()
        if table is not None and not getattr(table, "empty", True)
    }


def _fallback_crs_from_session(st):
    """Find a trusted CRS from workflow layers when a result layer lost CRS."""
    for key in [
        "selected_roads",
        "corridor_roads",
        "selected_boundary",
        "spatial_units",
        "corridors",
        "final_corridors",
    ]:
        layer = st_obj.session_state.get(key) if st_obj is not None else None
        if layer is not None and hasattr(layer, "crs") and layer.crs is not None:
            return layer.crs
    return "EPSG:4326"


def _repair_gdf_crs(gdf, st=None):
    """Return a GeoDataFrame with a usable CRS for dashboard maps.

    Some workflow result copies can lose CRS. If their coordinates are outside
    longitude/latitude range, setting EPSG:4326 directly makes the dashboard map
    zoom to the world. In that case, use the road/boundary CRS from session when
    available before converting to web map coordinates.
    """
    if gdf is None or getattr(gdf, "empty", True) or not hasattr(gdf, "geometry"):
        return gdf
    work = gdf.copy()
    try:
        bounds = work.total_bounds
        looks_lonlat = (
            -180 <= float(bounds[0]) <= 180
            and -90 <= float(bounds[1]) <= 90
            and -180 <= float(bounds[2]) <= 180
            and -90 <= float(bounds[3]) <= 90
        )
    except Exception:
        looks_lonlat = True
    if work.crs is None:
        if looks_lonlat:
            work = work.set_crs(epsg=4326)
        else:
            fallback = _fallback_crs_from_session(st) if st is not None else "EPSG:4326"
            work = work.set_crs(fallback)
    return work


def _available_maps(st):
    maps = {}
    density = st.session_state.get("spatial_units_density_map")
    if density is not None and not getattr(density, "empty", True):
        maps["Crash density map"] = _repair_gdf_crs(density, st)

    results = st.session_state.get("section7_results")
    if results is not None:
        risk_segments = results.get("risk_segments")
        if risk_segments is not None and not getattr(risk_segments, "empty", True):
            maps["HIN priority map"] = _repair_gdf_crs(risk_segments, st)

    corridors = st.session_state.get("final_corridors", st.session_state.get("corridors"))
    if corridors is not None and not getattr(corridors, "empty", True):
        maps["Corridor map"] = _repair_gdf_crs(corridors, st)

    return maps


def _numeric_cols(df):
    return [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]


def _groupable_cols(df):
    cols = []
    n = max(len(df), 1)
    max_unique = max(30, min(250, int(n * 0.45)))
    important_words = [
        "type", "year", "month", "date", "time", "severity", "kabco",
        "unit", "route", "corridor", "segment", "intersection", "road",
        "manner", "weather", "light", "surface", "reason", "cause",
        "class", "injury", "crash", "collision", "factor", "direction",
    ]
    for col in df.columns:
        if col == "geometry":
            continue
        lower = str(col).lower()
        nunique = df[col].nunique(dropna=True)
        if any(word in lower for word in important_words):
            cols.append(col)
        elif not pd.api.types.is_numeric_dtype(df[col]):
            try:
                if df[col].astype(str).str.len().mean() <= 100:
                    cols.append(col)
            except Exception:
                cols.append(col)
        elif nunique <= max_unique:
            cols.append(col)
    seen = set()
    out = []
    for c in cols:
        if c not in seen:
            out.append(c)
            seen.add(c)
    return out


def _find_col(df, keywords):
    for col in df.columns:
        lower = str(col).lower()
        if all(k in lower for k in keywords):
            return col
    for col in df.columns:
        lower = str(col).lower()
        if any(k in lower for k in keywords):
            return col
    return None


def _default_metric(num_cols):
    preferred = [
        "HIN_Priority_Index", "RiskScore", "CrashDensity", "CrashDensity_per_mile",
        "Crash_Count", "CrashCount", "TotalCrashes", "EPDO", "KSI_Count",
        "Fatal_Injury_Count", "Total",
    ]
    for col in preferred:
        if col in num_cols:
            return col
    return num_cols[0] if num_cols else None


def _rank_table(df, metric=None):
    out = df.copy()
    if metric and metric in out.columns:
        out = out.sort_values(metric, ascending=False).copy()
    if "Rank" not in out.columns:
        out.insert(0, "Rank", range(1, len(out) + 1))
    return out


def _aggregate(df, group_col, metric, aggregation, top_n):
    if not group_col:
        return pd.DataFrame()
    work = df.copy()
    work[group_col] = work[group_col].fillna("Unknown").astype(str)

    if aggregation == "Count" or metric is None or metric not in work.columns:
        out = work.groupby(group_col, dropna=False).size().reset_index(name="Count")
        value_col = "Count"
    else:
        work["__metric__"] = pd.to_numeric(work[metric], errors="coerce")
        agg_map = {
            "Sum": "sum",
            "Average": "mean",
            "Mean": "mean",
            "Median": "median",
            "Minimum": "min",
            "Maximum": "max",
        }
        agg_func = agg_map.get(aggregation, "sum")
        out = work.groupby(group_col, dropna=False)["__metric__"].agg(agg_func).reset_index()
        value_col = f"{aggregation} {metric}"
        out = out.rename(columns={"__metric__": value_col})

    out = out.sort_values(value_col, ascending=False).head(top_n).reset_index(drop=True)
    out.insert(0, "Rank", range(1, len(out) + 1))
    return out




def _rank_units_for_chart(df, metric, top_n=15):
    """Rank existing spatial-unit rows without summing the metric.

    Crash-density/HIN result tables already have one row per spatial unit. Using
    groupby-sum can create labels like "Sum CrashDensity" and can accidentally
    group by a numeric length field. This helper keeps the original unit ID and
    ranks by the selected metric directly.
    """
    if df is None or getattr(df, "empty", True) or metric not in df.columns:
        return pd.DataFrame(), None, None
    unit_col = _unit_col(df)
    work = df.copy()
    if unit_col is None:
        unit_col = "DashboardUnitID"
        work[unit_col] = [f"UNIT_{i + 1}" for i in range(len(work))]
    else:
        lower_unit = str(unit_col).lower()
        if any(bad in lower_unit for bad in ["length", "mile", "density", "count", "score", "index"]):
            unit_col = "DashboardUnitID"
            work[unit_col] = [f"UNIT_{i + 1}" for i in range(len(work))]
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work[unit_col] = work[unit_col].astype(str)
    keep_cols = [unit_col, metric]
    for extra in ["CrashCount", "Crash_Count", "UnitType", "City", "Length_Miles", "Length_mi", "Route", "FULLNAME", "RoadName", "RoadName1", "RoadName2", "CorridorID"]:
        if extra in work.columns and extra not in keep_cols:
            keep_cols.append(extra)
    out = work[keep_cols].sort_values(metric, ascending=False).head(top_n).reset_index(drop=True)
    out.insert(0, "Rank", range(1, len(out) + 1))
    return out, unit_col, metric


def _nice_metric_label(metric):
    labels = {
        "CrashDensity": "Crash density",
        "CrashDensity_per_mile": "Crash density",
        "CrashCount": "Crash count",
        "Crash_Count": "Crash count",
        "HIN_Priority_Index": "HIN priority index",
        "RiskScore": "Risk score",
    }
    return labels.get(str(metric), str(metric).replace("_", " "))

def _style(st):
    st.markdown(
        """
        <style>
        /* Dashboard and Crash insights need normal page scrolling.
           The workflow map shell uses fixed-height panes, but dashboard tabs
           contain long charts/tables. Keep the dashboard on the browser/page
           scrollbar instead of a fragile internal tab-panel scrollbar. */
        html, body, .stApp, [data-testid="stAppViewContainer"],
        [data-testid="stMain"], section[data-testid="stMain"],
        [data-testid="stMainBlockContainer"],
        [data-testid="stAppViewBlockContainer"],
        .main, .block-container {
            height: auto !important;
            max-height: none !important;
            min-height: 100vh !important;
            overflow-y: auto !important;
            overflow-x: hidden !important;
        }
        .block-container { padding-top: .65rem; padding-bottom: 6rem !important; max-width: 1900px; }
        [data-testid="stVerticalBlock"] { overflow: visible !important; }
        .dashboard-scroll-note { color:#64748b; font-size:.86rem; margin-top:-.2rem; margin-bottom:.4rem; }
        .dashboard-hero {
            border-radius: 18px;
            padding: 1.0rem 1.2rem;
            background: linear-gradient(135deg, #f8fafc 0%, #edf5f1 100%);
            border: 1px solid #dbe7e1;
            margin-bottom: .9rem;
        }
        .dashboard-hero h1 { margin: 0 0 .25rem 0; letter-spacing: -0.03em; font-size: 1.55rem; }
        .dashboard-hero p { margin: 0; max-width: 980px; color: #52615a; }
        div[data-testid="stMetric"] {
            background: #ffffff;
            border: 1px solid #e3e8e5;
            padding: .75rem .9rem;
            border-radius: 15px;
            box-shadow: 0 8px 22px rgba(15, 43, 33, .055);
        }
        .dashboard-card {
            background: #ffffff;
            border: 1px solid #e3e8e5;
            border-radius: 16px;
            padding: 1rem;
            margin-bottom: .75rem;
        }
        .dashboard-section-title {
            font-size: 1.38rem;
            font-weight: 750;
            letter-spacing: -0.02em;
            padding: .3rem 0 .45rem 0;
            margin: .55rem 0 .4rem 0;
            border-bottom: 1px solid #e6ece8;
        }
        .dashboard-section-title span {
            color: #4a5b53;
            font-size: .88rem;
            font-weight: 450;
            margin-left: .45rem;
        }
        .small-muted { color: #62716a; font-size: .9rem; }
        .dashboard-chart-card {
            border: 1px solid #e5e7eb;
            border-radius: 14px;
            padding: .55rem .6rem .25rem .6rem;
            margin-bottom: .75rem;
            background: white;
        }
        .stTabs [data-baseweb="tab-panel"],
        .stTabs [role="tabpanel"],
        div[data-baseweb="tab-panel"] {
            max-height: none !important;
            overflow-y: visible !important;
            overflow-x: visible !important;
            padding-right: 0 !important;
            padding-bottom: 2rem !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def _metric_cards(st, tables):
    analysis_type = st.session_state.get("analysis_type", st.session_state.get("spatial_unit", "Spatial Unit"))
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.metric("Crash records", f"{len(crashes):,}" if crashes is not None else "0")
    with c2:
        if crashes is not None:
            year_col = _find_col(crashes, ["year"])
            if year_col:
                years = pd.to_numeric(crashes[year_col], errors="coerce").dropna()
                label = f"{int(years.min())}–{int(years.max())}" if not years.empty else "Not found"
            else:
                label = "Not found"
            st.metric("Crash years", label)
        else:
            st.metric("Crash years", "Not available")
    with c3:
        if density is not None and "CrashDensity" in density.columns:
            st.metric("Avg crash density", f"{pd.to_numeric(density['CrashDensity'], errors='coerce').mean():,.2f}")
        elif density is not None and "CrashCount" in density.columns:
            st.metric("Avg crash count", f"{pd.to_numeric(density['CrashCount'], errors='coerce').mean():,.2f}")
        else:
            st.metric("Crash density", "Run crash density")
    with c4:
        if str(analysis_type).lower().startswith("intersection") and hin is None:
            st.metric("Risk method", "Crash count/density")
            st.caption("Intersection workflow does not use Sliding Window HIN.")
        elif hin is not None:
            metric = _default_metric(_numeric_cols(hin))
            if metric:
                top = pd.to_numeric(hin[metric], errors="coerce").max()
                st.metric("Highest risk value", f"{top:,.2f}")
            else:
                st.metric("Highest risk value", "Not available")
        else:
            st.metric("Risk method", "Run HIN")



def _unit_col(df):
    """Find a readable spatial-unit identifier column.

    Important: choose ID columns before loose words like "segment". In the
    segment workflow, columns such as SegmentLength_Mile contain the word
    "segment" but are numeric measures, not IDs. This function prevents those
    fields from appearing as the ranking y-axis.
    """
    if df is None:
        return None
    exact_priority = [
        "UnitID", "SpatialUnitID", "RiskSegmentID", "WindowID", "SegmentID", "SourceSegmentID",
        "IntersectionID", "CorridorID", "CorridorID_Final", "FeatureID", "ObjectID",
        "Route", "FULLNAME", "RoadName", "Name",
    ]
    norm = {str(c).lower().replace("_", "").replace(" ", ""): c for c in df.columns}
    for wanted in exact_priority:
        key = wanted.lower().replace("_", "").replace(" ", "")
        if key in norm:
            return norm[key]

    # Prefer text/object ID-like fields. Avoid length/mile/density/count fields.
    bad_words = ["length", "mile", "density", "count", "area", "score", "index", "risk", "rank"]
    for c in df.columns:
        lower = str(c).lower()
        if any(bad in lower for bad in bad_words):
            continue
        if lower.endswith("id") or any(w in lower for w in ["unit", "intersection", "corridor", "route", "roadname"]):
            return c
    return None


def _crash_type_col(df):
    if df is None:
        return None
    return (
        _find_col(df, ["crash", "type"])
        or _find_col(df, ["collision", "type"])
        or _find_col(df, ["manner"])
        or _find_col(df, ["type"])
        or _find_col(df, ["kabco"])
        or _find_col(df, ["severity"])
    )




def _kabco_col(df):
    """Find the KABCO / injury severity column when present."""
    if df is None:
        return None
    for col in df.columns:
        lower = str(col).lower().replace("_", "")
        if lower == "kabco" or "kabco" in lower:
            return col
    for col in df.columns:
        lower = str(col).lower()
        if "severity" in lower or "injury" in lower:
            return col
    return None


def _order_kabco(df, kabco_col):
    """Sort KABCO values in engineering severity order when possible."""
    if df is None or kabco_col not in df.columns:
        return df
    order = {"K": 1, "A": 2, "B": 3, "C": 4, "O": 5, "PDO": 5, "UNKNOWN": 9}
    out = df.copy()
    out["__kabco_order__"] = out[kabco_col].astype(str).str.upper().map(order).fillna(8)
    out = out.sort_values(["__kabco_order__", kabco_col]).drop(columns="__kabco_order__")
    return out

def _crash_count_col(df):
    if df is None:
        return None
    preferred = ["CrashCount", "Crash_Count", "TotalCrashes", "CrashCnt", "Crashes", "Count"]
    for c in preferred:
        if c in df.columns:
            return c
    for c in df.columns:
        lower = str(c).lower().replace("_", "")
        if "crash" in lower and ("count" in lower or "cnt" in lower or "total" in lower):
            return c
    return None


def _road_class_col(df):
    """Road-class helper used only as a fallback after the user-selected filter column."""
    if df is None:
        return None
    preferred = [
        "RoadClassFilterCol", "RoadClassFilterValues", "RoadStyleClass",
        "FunctionalClass", "RoadClass", "RoadType", "Functional_Class",
        "Road_Type", "RoadType", "F_SYSTEM", "FUNC_CLASS", "MTFCC", "highway"
    ]
    norm = {str(c).lower().replace("_", "").replace(" ", ""): c for c in df.columns}
    for p in preferred:
        key = p.lower().replace("_", "").replace(" ", "")
        if key in norm:
            return norm[key]
    for c in df.columns:
        lower = str(c).lower()
        if ("road" in lower and ("class" in lower or "type" in lower)) or "functional" in lower:
            return c
    return None


def _active_road_class_column(st_obj=None):
    """Return the exact road-class column selected in Step 1, if the filter is active."""
    try:
        if st_obj is None:
            return None
        if not bool(st_obj.session_state.get("road_class_layer_enabled", False)):
            return None
        col = st_obj.session_state.get("analysis_road_class_col") or st_obj.session_state.get("road_class_viz_col")
        if col is None or str(col).strip() == "":
            return None
        return str(col)
    except Exception:
        return None


def _apply_selected_road_classes(work, class_col, st_obj=None):
    selected = (st_obj.session_state.get("analysis_road_class_values") or st_obj.session_state.get("road_class_viz_values")) if st_obj is not None else None
    if selected and class_col in work.columns:
        selected_text = {str(v) for v in selected}
        work = work[work[class_col].astype(str).isin(selected_text)].copy()
    return work


def _road_class_summary_table(crashes=None, density=None, roads=None, top_n=15, st_obj=None):
    """Create road-class chart only when the Step 1 road-class filter is enabled.

    This avoids misleading charts from unrelated fields such as OSM signal
    ``highway=traffic_signals``. The chart uses the exact column selected by the
    user in Step 1; if crash counts are not joined to that field, it falls back to
    total selected roadway length by that same field.
    """
    class_col = _active_road_class_column(st_obj)
    if not class_col:
        return None, pd.DataFrame(), None, None

    candidates = []
    if density is not None and not getattr(density, "empty", True):
        candidates.append(("density", density))
    if crashes is not None and not getattr(crashes, "empty", True):
        candidates.append(("crashes", crashes))
    # Prefer the Step 1 selected/display road network for roadway length context.
    for key in ["roads_class_display", "selected_roads", "analysis_roads"]:
        layer = st_obj.session_state.get(key) if st_obj is not None else None
        if layer is not None and not getattr(layer, "empty", True):
            candidates.append(("roads", layer))

    for kind, df in candidates:
        work = _drop_geometry(df).copy()
        use_col = class_col if class_col in work.columns else None
        if use_col is None and class_col == "RoadStyleClass" and "RoadStyleClass" in work.columns:
            use_col = "RoadStyleClass"
        if use_col is None:
            continue
        work = _apply_selected_road_classes(work, use_col, st_obj)
        if work.empty:
            continue
        count_col = _crash_count_col(work)
        length_col = _normal_col(work, ["Length_Miles", "Length_Mi", "length_mi", "Miles", "SegmentLength_Mile"])
        if count_col and count_col in work.columns:
            work[count_col] = pd.to_numeric(work[count_col], errors="coerce").fillna(0)
            out = work.groupby(use_col, dropna=False)[count_col].sum().reset_index(name="Crash count")
            value_col = "Crash count"
            title = f"Crashes by road class ({use_col})"
        elif kind == "roads" and length_col and length_col in work.columns:
            work[length_col] = pd.to_numeric(work[length_col], errors="coerce").fillna(0)
            out = work.groupby(use_col, dropna=False)[length_col].sum().reset_index(name="Total length (mi)")
            value_col = "Total length (mi)"
            title = f"Roadway length by road class ({use_col})"
        elif kind == "roads":
            out = work.groupby(use_col, dropna=False).size().reset_index(name="Road feature count")
            value_col = "Road feature count"
            title = f"Road features by road class ({use_col})"
        else:
            continue
        out[use_col] = out[use_col].astype(str).replace({"nan": "Unknown", "None": "Unknown", "": "Unknown"})
        out = out.sort_values(value_col, ascending=False).head(top_n).reset_index(drop=True)
        out.insert(0, "Rank", range(1, len(out) + 1))
        return title, out, use_col, value_col

    return None, pd.DataFrame(), None, None

def _chart_height():
    # Compact height so the dashboard page can show more charts without feeling cut off.
    return 280




def _time_col(df):
    """Find a date/time column that can be used for monthly crash trends."""
    if df is None:
        return None
    preferred = [
        "CrashDate", "Crash_Date", "CrashDateTime", "Crash_Date_Time", "Date", "Time",
        "ReportDate", "ReportedDate", "datetime", "timestamp",
    ]
    for c in preferred:
        if c in df.columns:
            return c
    for c in df.columns:
        lower = str(c).lower()
        if any(w in lower for w in ["date", "time", "month"]):
            return c
    return None


def _month_trend_table(crashes):
    """Build Jan-Dec crash counts with one colored line per year."""
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), None, None, None
    df = crashes.copy()
    year_col = _find_col(df, ["year"])
    month_col = _find_col(df, ["month"])
    date_col = _time_col(df)
    tmp = pd.DataFrame()
    if date_col:
        dt = pd.to_datetime(df[date_col], errors="coerce")
        if dt.notna().any():
            tmp = pd.DataFrame({"Year": dt.dt.year, "Month": dt.dt.month}).dropna()
    if tmp.empty and year_col and month_col:
        tmp = df[[year_col, month_col]].copy()
        tmp["Year"] = pd.to_numeric(tmp[year_col], errors="coerce")
        tmp["Month"] = pd.to_numeric(tmp[month_col], errors="coerce")
        tmp = tmp[["Year", "Month"]].dropna()
    if tmp.empty:
        return pd.DataFrame(), None, None, None
    tmp["Year"] = tmp["Year"].astype(int).astype(str)
    tmp["Month"] = tmp["Month"].astype(int)
    tmp = tmp[(tmp["Month"] >= 1) & (tmp["Month"] <= 12)].copy()
    if tmp.empty:
        return pd.DataFrame(), None, None, None
    out = tmp.groupby(["Year", "Month"], dropna=False).size().reset_index(name="Count")
    years = sorted(out["Year"].unique(), key=lambda x: int(x) if str(x).isdigit() else str(x))
    full = pd.MultiIndex.from_product([years, range(1, 13)], names=["Year", "Month"]).to_frame(index=False)
    out = full.merge(out, on=["Year", "Month"], how="left").fillna({"Count": 0})
    out["Count"] = out["Count"].astype(int)
    out["Month label"] = pd.Categorical(
        pd.to_datetime(out["Month"].astype(str), format="%m").dt.strftime("%b"),
        categories=MONTH_ORDER,
        ordered=True,
    )
    out = out.sort_values(["Year", "Month"])
    return out, "Month label", "Count", "Year"

def _year_kabco_table(crashes):
    """Return crash counts by year and KABCO for stacked annual bars."""
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), None, None
    year_col = _find_col(crashes, ["year"])
    kabco_col = _kabco_col(crashes)
    if not year_col or not kabco_col:
        return pd.DataFrame(), None, None
    work = crashes[[year_col, kabco_col]].copy()
    work[year_col] = work[year_col].fillna("Unknown").astype(str)
    work[kabco_col] = work[kabco_col].fillna("Unknown").astype(str)
    out = work.groupby([year_col, kabco_col], dropna=False).size().reset_index(name="Count")
    out = _order_kabco(out, kabco_col)
    try:
        out["__year_sort__"] = pd.to_numeric(out[year_col], errors="coerce")
        out = out.sort_values(["__year_sort__", kabco_col]).drop(columns="__year_sort__")
    except Exception:
        pass
    return out, year_col, kabco_col


def _road_class_kabco_table(crashes, st_obj=None):
    """Return road-class by KABCO crash counts when the road-class filter is active."""
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), None, None
    road_col = _active_road_class_column(st_obj)
    kabco_col = _kabco_col(crashes)
    if not road_col or not kabco_col or road_col not in crashes.columns:
        return pd.DataFrame(), None, None
    work = crashes[[road_col, kabco_col]].copy()
    work = _apply_selected_road_classes(work, road_col, st_obj)
    if work.empty:
        return pd.DataFrame(), None, None
    work[road_col] = work[road_col].fillna("Unknown").astype(str)
    work[kabco_col] = work[kabco_col].fillna("Unknown").astype(str)
    out = work.groupby([road_col, kabco_col], dropna=False).size().reset_index(name="Count")
    return out, road_col, kabco_col


def _crash_type_kabco_table(crashes):
    """Return crash-type by KABCO counts for severity pattern screening."""
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), None, None
    type_col = _crash_type_col(crashes)
    kabco_col = _kabco_col(crashes)
    if not type_col or not kabco_col or type_col == kabco_col:
        return pd.DataFrame(), None, None
    work = crashes[[type_col, kabco_col]].copy()
    work[type_col] = work[type_col].fillna("Unknown").astype(str)
    work[kabco_col] = work[kabco_col].fillna("Unknown").astype(str)
    # Keep top crash types to avoid unreadable charts.
    top_types = work[type_col].value_counts().head(10).index.tolist()
    work = work[work[type_col].isin(top_types)].copy()
    out = work.groupby([type_col, kabco_col], dropna=False).size().reset_index(name="Count")
    return out, type_col, kabco_col


def _mobility_mode_table(crashes):
    """Summarize travel-mode severity patterns as KABCO x mode counts."""
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), None, None
    type_col = _crash_type_col(crashes)
    kabco_col = _kabco_col(crashes)
    if not type_col or not kabco_col:
        return pd.DataFrame(), None, None
    work = crashes[[kabco_col, type_col]].copy()
    t = work[type_col].fillna("").astype(str).str.lower()
    work["Mode"] = "Motor vehicle / other"
    work.loc[t.str.contains("ped", na=False), "Mode"] = "Pedestrian"
    work.loc[t.str.contains("bicycle|bike", na=False), "Mode"] = "Bicycle"
    work.loc[t.str.contains("motorcycle|motor bike", na=False), "Mode"] = "Motorcycle"
    work[kabco_col] = work[kabco_col].fillna("Unknown").astype(str)
    if (work["Mode"] != "Motor vehicle / other").sum() == 0:
        return pd.DataFrame(), None, None
    out = work.groupby([kabco_col, "Mode"], dropna=False).size().reset_index(name="Count")
    out = _order_kabco(out, kabco_col)
    return out, kabco_col, "Mode"

def _hin_route_rank_for_chart(hin, metric, top_n=15):
    """Rank HIN by route using the highest segment/window score per route.

    This avoids stacked bars caused by repeated route names. Each route appears
    once, with hover/table context showing the selected segment/window ID and
    length when available.
    """
    if hin is None or getattr(hin, "empty", True) or metric not in hin.columns:
        return pd.DataFrame(), None, None
    route_col = None
    for c in ["Route", "FULLNAME", "RoadName", "CorridorName", "Name"]:
        if c in hin.columns:
            route_col = c
            break
    if route_col is None:
        return _rank_units_for_chart(hin, metric, top_n)
    work = hin.copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work[route_col] = work[route_col].fillna("Unknown route").astype(str)
    idx = work.groupby(route_col, dropna=False)[metric].idxmax()
    best = work.loc[idx].copy().sort_values(metric, ascending=False).head(top_n).reset_index(drop=True)
    best.insert(0, "Rank", range(1, len(best) + 1))
    id_col = _unit_col(hin)
    if id_col and id_col not in best.columns:
        id_col = None
    length_col = _normal_col(best, ["Length_Miles", "Length_Mi", "SegmentLength_Mile", "WindowLength_Miles", "length_mi"])
    keep = ["Rank", route_col, metric]
    for c in [id_col, length_col, _crash_count_col(best)]:
        if c and c in best.columns and c not in keep:
            keep.append(c)
    return best[keep], route_col, metric

def _render_pattern_charts(st, tables):
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))

    st.markdown("<div class='dashboard-section-title'>Crash patterns <span>years, severity, crash type, mode, and roadway context</span></div>", unsafe_allow_html=True)
    left, right = st.columns(2)

    with left:
        if crashes is not None:
            year_kabco, year_col, kabco_col = _year_kabco_table(crashes)
            if not year_kabco.empty:
                fig = px.bar(
                    year_kabco,
                    x=year_col,
                    y="Count",
                    color=kabco_col,
                    color_discrete_map=KABCO_COLOR_MAP,
                    title=f"Crashes by year and {kabco_col}",
                )
                fig.update_layout(
                    height=_chart_height(),
                    margin=dict(l=20, r=20, t=45, b=35),
                    barmode="stack",
                    xaxis_title="Year",
                    yaxis_title="Crash count",
                )
                st.plotly_chart(_polish_figure(fig), width="stretch")
            else:
                year_col = _find_col(crashes, ["year"])
                if year_col:
                    year_df = _aggregate(crashes, year_col, None, "Count", 20)
                    fig = px.bar(year_df.sort_values(year_col), x=year_col, y="Count", title="Crashes by year")
                    fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), xaxis_title="Year", yaxis_title="Crash count")
                    st.plotly_chart(_polish_figure(fig), width="stretch")
                else:
                    st.info("No crash year field was detected.")
        else:
            st.info("No crash table is available yet.")

    with right:
        if crashes is not None:
            type_col = _crash_type_col(crashes)
            if type_col:
                type_df = _aggregate(crashes, type_col, None, "Count", 10)
                fig = px.pie(type_df, names=type_col, values="Count", hole=0.38, title=f"Crash type share by {type_col}")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), legend=dict(orientation="v"))
                st.plotly_chart(_polish_figure(fig), width="stretch")
            else:
                st.info("No crash type field was detected.")
        else:
            st.info("No crash table is available yet.")

    left2, right2 = st.columns(2)
    with left2:
        monthly_df, period_col, value_col, color_col = _month_trend_table(crashes)
        if not monthly_df.empty:
            fig = px.line(monthly_df, x=period_col, y=value_col, color=color_col, markers=True, title="Monthly crash trend by year")
            fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), xaxis_title="Month", yaxis_title="Crash count")
            st.plotly_chart(_polish_figure(fig), width="stretch")
    with right2:
        mode_df, mode_kabco_col, mode_col = _mobility_mode_table(crashes)
        if not mode_df.empty:
            fig = px.line(mode_df, x=mode_kabco_col, y="Count", color=mode_col, markers=True, title="Travel mode distribution by KABCO")
            fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), xaxis_title="KABCO", yaxis_title="Crash count")
            st.plotly_chart(_polish_figure(fig), width="stretch")

    # Road class by KABCO heatmap appears only when the Step 1 road-class filter is active
    # and the assigned crash table carries the selected road-class field.
    road_kabco, road_col, road_kabco_col = _road_class_kabco_table(crashes, st_obj=st)
    if not road_kabco.empty:
        pivot = road_kabco.pivot_table(index=road_col, columns=road_kabco_col, values="Count", aggfunc="sum", fill_value=0)
        # Preserve KABCO order when possible.
        order = [c for c in ["K", "A", "B", "C", "O"] if c in pivot.columns]
        other = [c for c in pivot.columns if c not in order]
        pivot = pivot[order + other]
        fig = px.imshow(
            pivot,
            text_auto=True,
            aspect="auto",
            title=f"Road class by {road_kabco_col}",
            labels=dict(x=road_kabco_col, y="Road class", color="Crash count"),
        )
        fig.update_layout(height=330, margin=dict(l=20, r=20, t=45, b=35))
        st.plotly_chart(_polish_figure(fig), width="stretch")

    crash_kabco, crash_type_col, crash_kabco_col = _crash_type_kabco_table(crashes)
    if not crash_kabco.empty:
        pivot = crash_kabco.pivot_table(index=crash_type_col, columns=crash_kabco_col, values="Count", aggfunc="sum", fill_value=0)
        order = [c for c in ["K", "A", "B", "C", "O"] if c in pivot.columns]
        other = [c for c in pivot.columns if c not in order]
        pivot = pivot[order + other]
        fig = px.imshow(
            pivot,
            text_auto=True,
            aspect="auto",
            title=f"Crash type by {crash_kabco_col}",
            labels=dict(x=crash_kabco_col, y="Crash type", color="Crash count"),
            color_continuous_scale="YlOrRd",
        )
        fig.update_layout(height=420, margin=dict(l=20, r=20, t=45, b=35))
        st.plotly_chart(_polish_figure(fig), width="stretch")

    st.markdown("<div class='dashboard-section-title'>Risk and spatial-unit ranking <span>crash density, crash count, and HIN priority ranking</span></div>", unsafe_allow_html=True)
    left, right = st.columns(2)

    with left:
        if density is not None:
            metric = "CrashDensity" if "CrashDensity" in density.columns else _default_metric(_numeric_cols(density))
            rank_df, unit_col, value_col = _rank_units_for_chart(density, metric, 15) if metric else (pd.DataFrame(), None, None)
            if not rank_df.empty:
                rank_df = rank_df.sort_values(value_col, ascending=True)
                fig = px.bar(rank_df, y=unit_col, x=value_col, orientation="h", title="Top spatial units by crash density")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), yaxis_title="Spatial unit ID", xaxis_title="Crash density")
                st.plotly_chart(_polish_figure(fig), width="stretch")
            else:
                st.info("Run crash-density analysis to rank spatial units.")
        else:
            st.info("Crash-density results are not available yet. Run crash-density analysis first.")

    with right:
        if density is not None:
            count_col = _crash_count_col(density)
            rank_df, unit_col, value_col = _rank_units_for_chart(density, count_col, 15) if count_col else (pd.DataFrame(), None, None)
            if not rank_df.empty:
                rank_df = rank_df.sort_values(value_col, ascending=True)
                fig = px.bar(rank_df, y=unit_col, x=value_col, orientation="h", title="Top spatial units by crash count")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), yaxis_title="Spatial unit ID", xaxis_title="Crash count")
                st.plotly_chart(_polish_figure(fig), width="stretch")
            else:
                st.info("Crash count field was not found in the crash-density results.")
        else:
            st.info("Run crash-density analysis to show crash-count ranking.")

    if hin is not None:
        st.markdown("<div class='dashboard-section-title'>HIN priority ranking <span>highest HIN segment/window per route</span></div>", unsafe_allow_html=True)
        metric = _default_metric(_numeric_cols(hin))
        rank_df, route_col, value_col = _hin_route_rank_for_chart(hin, metric, 15) if metric else (pd.DataFrame(), None, None)
        if not rank_df.empty:
            rank_df = rank_df.sort_values(value_col, ascending=True)
            hover_cols = [c for c in rank_df.columns if c not in [route_col, value_col, "Rank"]]
            fig = px.bar(rank_df, y=route_col, x=value_col, orientation="h", hover_data=hover_cols, title=f"Top routes by highest {_nice_metric_label(metric)}")
            fig.update_layout(height=330, margin=dict(l=20, r=20, t=45, b=35), yaxis_title="Route", xaxis_title=_nice_metric_label(metric))
            st.plotly_chart(_polish_figure(fig), width="stretch")
            st.caption("Each route is shown once using the highest-scoring HIN segment/window on that route. Hover to see the selected segment/window and length when available.")
    st.caption("Scroll down to view all charts, maps, and tables. Use the Dashboard Builder tab to select charts/maps for export.")

def _infer_dataset_from_request(request, tables):
    q = str(request).lower()
    preferences = []
    if any(w in q for w in ["hin", "risk", "priority", "sliding"]):
        preferences += ["HIN risk segments", "HIN corridors", "Sliding windows"]
    if any(w in q for w in ["density", "top", "rank", "intersection", "segment", "corridor", "spatial unit", "spatial"]):
        preferences += ["Crash density results", "Crash count / severity summary"]
    if any(w in q for w in ["type", "year", "severity", "reason", "cause", "factor", "crash"]):
        preferences += ["Assigned crashes", "Uploaded crashes"]
    preferences += list(tables.keys())
    for name in preferences:
        if name in tables:
            return name
    return next(iter(tables.keys()))


def _infer_chart_type(request, group_col=None):
    q = str(request).lower()
    if any(w in q for w in ["table", "list", "records"]):
        return "Table"
    if any(w in q for w in ["line", "trend over time", "time series"]):
        return "Line chart"
    if any(w in q for w in ["heatmap", "heat map"]):
        return "Heatmap"
    if any(w in q for w in ["tree", "treemap", "tree map"]):
        return "Treemap"
    if any(w in q for w in ["radar", "spider"]):
        return "Radar chart"
    if any(w in q for w in ["pie", "donut", "share", "percent", "percentage", "composition"]):
        return "Pie chart"
    # KABCO/crash-type distribution is categorical, so a bar chart is clearer
    # than a histogram. Keep histogram only when the user explicitly asks for it.
    if "histogram" in q:
        return "Histogram"
    if any(w in q for w in ["box", "spread"]):
        return "Box plot"
    if any(w in q for w in ["stack", "stacked", "color", "colored", "split by"]):
        return "Stacked bar chart"
    if any(w in q for w in ["top", "rank", "highest", "worst", "most", "risky"]):
        return "Horizontal rank bar"
    return "Bar chart"


def _infer_aggregation(request):
    q = str(request).lower()
    if any(w in q for w in ["average", "avg", "mean"]):
        return "Average"
    if "median" in q:
        return "Median"
    if any(w in q for w in ["sum", "total"]):
        return "Sum"
    if any(w in q for w in ["minimum", "min", "lowest"]):
        return "Minimum"
    if any(w in q for w in ["maximum", "max", "highest"]):
        return "Maximum"
    return "Count"


def _infer_top_n(request, default=15):
    m = re.search(r"top\s+(\d+)", str(request).lower())
    if m:
        return max(1, min(100, int(m.group(1))))
    return default



def _infer_color_col(request, df, group_col=None):
    """Find a second grouping/color column, especially for commands like
    'count crashes in each intersection colored by crash type'.
    """
    q = str(request).lower()
    if not any(w in q for w in ["color", "colored", "stack", "stacked", "by crash type", "each intersection"]):
        return None
    if any(w in q for w in ["kabco", "severity", "injury"]):
        col = _kabco_col(df) or _find_col(df, ["severity", "injury"])
        if col and col != group_col:
            return col
    if any(w in q for w in ["crash type", "type", "collision", "manner"]):
        col = _crash_type_col(df)
        if col and col != group_col:
            return col
    # Look after words such as colored by / stacked by.
    m = re.search(r"(?:colored by|color by|stacked by|split by)\s+([a-zA-Z0-9_ ]+)", q)
    if m:
        col, _ = _best_column(m.group(1), _groupable_cols(df), minimum=4)
        if col and col != group_col:
            return col
    return None

def _infer_map_layers(request, maps):
    q = str(request).lower()
    selected = []
    for name in maps:
        lname = name.lower()
        if ("density" in q and "density" in lname) or ("hin" in q and "hin" in lname) or ("risk" in q and "hin" in lname) or ("corridor" in q and "corridor" in lname) or ("map" in q and not selected):
            selected.append(name)
    if "map" in q and not selected:
        selected = list(maps.keys())[:1]
    return selected[:2]


def _figure_to_png_bytes(fig):
    try:
        return pio.to_image(fig, format="png", width=1200, height=720, scale=2)
    except Exception:
        return None


def _render_chart_download_buttons(st, fig, data, base_name, key_prefix):
    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button(
            "Download data CSV",
            data=data.to_csv(index=False).encode("utf-8"),
            file_name=f"{_safe_name(base_name)}.csv",
            mime="text/csv",
            key=f"{key_prefix}_csv",
        )
    if fig is not None:
        with c2:
            st.download_button(
                "Download chart HTML",
                data=pio.to_html(fig, include_plotlyjs="cdn", full_html=True).encode("utf-8"),
                file_name=f"{_safe_name(base_name)}.html",
                mime="text/html",
                key=f"{key_prefix}_html",
            )
        with c3:
            png = _figure_to_png_bytes(fig)
            if png:
                st.download_button(
                    "Download chart PNG",
                    data=png,
                    file_name=f"{_safe_name(base_name)}.png",
                    mime="image/png",
                    key=f"{key_prefix}_png",
                )
            else:
                st.caption("PNG export needs kaleido.")



def _generate_assistant_chart(request, tables):
    dataset = _infer_dataset_from_request(request, tables)
    df = tables[dataset].copy()
    num_cols = _numeric_cols(df)
    group_cols = _groupable_cols(df)
    aggregation = _infer_aggregation(request)
    top_n = _infer_top_n(request)
    q = str(request).lower()

    wants_year = any(w in q for w in ["year", "annual", "annually", "over time", "trend"])
    wants_month = any(w in q for w in ["month", "monthly"])
    wants_kabco = any(w in q for w in ["kabco", "severity", "injury"])
    wants_type = any(w in q for w in ["crash type", "collision type", "manner", "type"])
    wants_stack = any(w in q for w in ["stack", "stacked", "colored", "color", "split by"])
    wants_line = any(w in q for w in ["line", "trend", "over time", "monthly"])

    group_col, group_candidates = _best_column(request, group_cols, minimum=5)

    if wants_month:
        # Use exact month column if present; otherwise date parsing happens in the default dashboard,
        # while assistant uses the closest time field available.
        group_col = _find_col(df, ["month"]) or _time_col(df) or _find_col(df, ["year"])
    elif wants_year:
        group_col = _find_col(df, ["year"]) or group_col
    elif any(w in q for w in ["intersection", "segment", "corridor", "spatial unit", "spatial"]):
        group_col = _unit_col(df) or group_col
    elif wants_kabco:
        group_col = _kabco_col(df) or _find_col(df, ["severity", "injury"]) or group_col
    elif wants_type:
        group_col = _crash_type_col(df) or group_col

    value_col = None
    value_candidates = []
    if aggregation != "Count" or any(w in q for w in ["density", "hin", "risk", "score", "index", "epdo", "ksi"]):
        value_col, value_candidates = _best_column(request, num_cols, minimum=5)
        if value_col is None:
            value_col = _default_metric(num_cols)

    color_col = _infer_color_col(request, df, group_col=group_col)
    if (wants_year or wants_month) and wants_kabco:
        color_col = _kabco_col(df) or _find_col(df, ["severity", "injury"])
    elif (wants_year or wants_month) and wants_type:
        color_col = _crash_type_col(df)
    elif wants_stack and wants_kabco:
        color_col = _kabco_col(df) or color_col
    elif wants_stack and wants_type:
        color_col = _crash_type_col(df) or color_col

    if color_col == group_col:
        color_col = None

    chart_type = _infer_chart_type(request, group_col)
    if wants_type and wants_kabco and any(w in q for w in ["heat", "map", "matrix", "by"]):
        chart_type = "Heatmap"
    elif color_col and group_col and wants_stack:
        chart_type = "Stacked bar chart"
    elif wants_line:
        chart_type = "Line chart"
    elif wants_type and any(w in q for w in ["share", "pie", "donut", "percent", "percentage"]):
        chart_type = "Pie chart"
    elif any(w in q for w in ["top", "rank", "highest", "worst", "most", "risky"]):
        chart_type = "Horizontal rank bar"

    value_col_for_render = None if aggregation == "Count" else value_col
    needs_clarification = group_col is None and chart_type not in ["Histogram", "Box plot", "Table"]
    if chart_type in ["Histogram", "Box plot"] and value_col is None:
        needs_clarification = True

    return {
        "dataset": dataset,
        "df": df,
        "group_col": group_col,
        "color_col": color_col,
        "value_col": value_col_for_render,
        "aggregation": aggregation,
        "chart_type": chart_type,
        "top_n": top_n,
        "needs_clarification": needs_clarification,
        "group_candidates": group_candidates,
        "value_candidates": value_candidates,
    }

def _render_smart_dashboard_assistant(st, tables):
    maps = _available_maps(st)
    st.markdown("<div class='dashboard-section-title'>Dashboard assistant <span>ask for charts, tables, maps, dashboards, or reports in plain language</span></div>", unsafe_allow_html=True)
    st.caption("Type what you want. The assistant searches all uploaded and generated datasets. If a requested column is unclear, it asks for the exact dataset column name.")

    request = st.chat_input("Example: Show crash type share as a pie chart and add crash density map to the dashboard")
    if request:
        history = st.session_state.get("dashboard_assistant_messages", [])
        history.append({"role": "user", "content": request})
        st.session_state["dashboard_assistant_messages"] = history[-12:]
        st.session_state["dashboard_assistant_request"] = request

    for msg in st.session_state.get("dashboard_assistant_messages", []):
        with st.chat_message(msg["role"]):
            st.write(msg["content"])

    current_request = st.session_state.get("dashboard_assistant_request", "")
    if not current_request:
        st.info("Try: 'create a pie chart of crash type', 'top 10 intersections by crash density', or 'make a dashboard with crash years, crash type, top risky units, and crash density map'.")
        return

    plan = _generate_assistant_chart(current_request, tables)
    selected_maps = _infer_map_layers(current_request, maps)
    q = current_request.lower()
    dashboard_request = any(w in q for w in ["dashboard", "report", "include", "add", "put"])

    if plan["needs_clarification"]:
        st.warning("I found the dataset, but I need one exact column name before creating the chart.")
        st.write(f"Dataset selected: **{plan['dataset']}**")
        df = plan["df"]
        group_options = _groupable_cols(df)
        num_options = _numeric_cols(df)
        if plan["group_col"] is None and plan["chart_type"] not in ["Histogram", "Box plot", "Table"]:
            candidates = [c for c, score in plan["group_candidates"] if score > 0]
            default_idx = 0
            options = candidates + [c for c in group_options if c not in candidates]
            if not options:
                options = group_options
            chosen_group = st.selectbox("Which exact column should be the category/group field?", options, key="smart_assistant_group_confirm")
            plan["group_col"] = chosen_group
        if plan["chart_type"] in ["Histogram", "Box plot"] and plan["value_col"] is None:
            chosen_value = st.selectbox("Which exact numeric column should be used as the value field?", num_options, key="smart_assistant_value_confirm")
            plan["value_col"] = chosen_value
        if not st.button("Create chart with this column", key="smart_assistant_create_after_confirm"):
            return

    st.markdown("**Assistant plan**")
    st.write(
        f"Dataset: **{plan['dataset']}** · Chart: **{plan['chart_type']}** · "
        f"Group: **{plan['group_col'] or 'None'}** · Calculation: **{plan['aggregation']}**"
        + (f" · Color: **{plan.get('color_col')}**" if plan.get('color_col') else "")
        + (f" · Value: **{plan['value_col']}**" if plan['value_col'] else "")
    )

    chart_data, fig = _render_chart(
        st=st,
        df=plan["df"],
        chart_type=plan["chart_type"],
        value_field=plan["value_col"],
        aggregation=plan["aggregation"],
        group_col=plan["group_col"],
        top_n=plan["top_n"],
        color_col=plan.get("color_col"),
    )
    title = "Assistant table" if fig is None else (fig.layout.title.text or "Assistant chart")
    _render_chart_download_buttons(st, fig, chart_data, title, "smart_assistant_current")

    c1, c2, c3 = st.columns(3)
    with c1:
        if fig is not None and st.button("Add chart to dashboard", key="smart_assistant_add_chart"):
            saved = st.session_state.get("dashboard_custom_figures", [])
            saved.append((str(title), fig, chart_data.copy()))
            st.session_state["dashboard_custom_figures"] = saved[-12:]
            st.success("Chart added to Dashboard Builder.")
    with c2:
        if selected_maps:
            st.write("Map layers found from your request:")
            st.write(", ".join(selected_maps))
            if st.button("Add map layers to dashboard", key="smart_assistant_add_maps"):
                st.session_state["dash_builder_map_layers"] = selected_maps
                st.success("Map layers added to Dashboard Builder.")
    with c3:
        if dashboard_request:
            if st.button("Build dashboard from this request", key="smart_assistant_build_dashboard"):
                if fig is not None:
                    saved = st.session_state.get("dashboard_custom_figures", [])
                    saved.append((str(title), fig, chart_data.copy()))
                    st.session_state["dashboard_custom_figures"] = saved[-12:]
                if selected_maps:
                    st.session_state["dash_builder_map_layers"] = selected_maps
                st.success("Dashboard Builder updated. Open the Dashboard builder tab to review/export.")

    if selected_maps:
        st.markdown("**Preview requested map layer**")
        map_name = selected_maps[0]
        _render_dashboard_map(st, map_name, maps[map_name], key=f"smart_assistant_map_{_safe_name(map_name)}", height=430, overlay_layers={name: gdf for name, gdf in _workflow_overlay_sources(st).items() if name in ["Roads", "Signals"]})

def _render_chart(st, df, chart_type, value_field, aggregation, group_col, top_n, color_col=None):
    if chart_type == "Table":
        ranked = _rank_table(df, value_field)
        default_cols = list(ranked.columns)[: min(14, len(ranked.columns))]
        show_cols = st.multiselect("Columns to show", list(ranked.columns), default=default_cols, key="dash_columns_to_show")
        st.dataframe(_safe_dataframe_for_display(ranked[show_cols] if show_cols else ranked), width="stretch", hide_index=True)
        return ranked, None

    if chart_type == "Summary cards":
        numeric = _numeric_cols(df)
        selected = st.multiselect("Value fields", numeric, default=[c for c in [value_field] if c in numeric] or numeric[:4], key="dash_summary_metric_fields")
        if not selected:
            st.info("Choose at least one numeric field.")
            return df, None
        card_cols = st.columns(min(4, len(selected)))
        for i, col in enumerate(selected):
            values = pd.to_numeric(df[col], errors="coerce")
            with card_cols[i % len(card_cols)]:
                st.metric(f"Average {col}", f"{values.mean():,.2f}")
                st.caption(f"sum {values.sum():,.2f} | median {values.median():,.2f}")
        return df, None

    if chart_type == "Stacked bar chart":
        if not group_col or not color_col:
            st.warning("Choose both a category/x-axis field and a color/stack field.")
            return df, None
        work = df.copy()
        work[group_col] = work[group_col].fillna("Unknown").astype(str)
        work[color_col] = work[color_col].fillna("Unknown").astype(str)
        chart_df = work.groupby([group_col, color_col], dropna=False).size().reset_index(name="Count")
        is_year = "year" in str(group_col).lower()
        if is_year:
            # Time/category trend: x = year, y = count, color/stack = KABCO/type.
            chart_df = chart_df.sort_values(group_col)
            fig = px.bar(
                chart_df,
                x=group_col,
                y="Count",
                color=color_col,
                color_discrete_map=KABCO_COLOR_MAP,
                title=f"Crash count by {group_col} and {color_col}",
            )
            fig.update_layout(xaxis_title="Year", yaxis_title="Crash count")
        else:
            # Ranking view: y = spatial unit/category, x = count, color/stack = type/severity.
            top_units = chart_df.groupby(group_col)["Count"].sum().sort_values(ascending=False).head(top_n).index.tolist()
            chart_df = chart_df[chart_df[group_col].isin(top_units)]
            order = chart_df.groupby(group_col)["Count"].sum().sort_values(ascending=True).index.tolist()
            chart_df[group_col] = pd.Categorical(chart_df[group_col], categories=order, ordered=True)
            fig = px.bar(
                chart_df.sort_values(group_col),
                y=group_col,
                x="Count",
                color=color_col,
                color_discrete_map=KABCO_COLOR_MAP,
                orientation="h",
                title=f"Crashes in each {group_col} by {color_col}",
            )
            fig.update_layout(xaxis_title="Crash count", yaxis_title=group_col)
        fig.update_layout(height=360, margin=dict(l=20, r=20, t=45, b=35), barmode="stack")
        st.plotly_chart(_polish_figure(fig), width="stretch")
        return chart_df, fig

    if chart_type == "Line chart":
        if not group_col:
            st.warning("Choose a time/category field for the line chart.")
            return df, None
        chart_df = _aggregate(df, group_col, value_field, aggregation, top_n=100)
        if chart_df.empty:
            st.warning("No data to display.")
            return df, None
        value_col = [c for c in chart_df.columns if c not in ["Rank", group_col]][0]
        try:
            chart_df = chart_df.sort_values(group_col)
        except Exception:
            pass
        fig = px.line(chart_df, x=group_col, y=value_col, markers=True, title=f"{aggregation} by {group_col}")
        fig.update_layout(height=360, margin=dict(l=20, r=20, t=45, b=40), xaxis_title=group_col, yaxis_title=_nice_metric_label(value_col))
        st.plotly_chart(_polish_figure(fig), width="stretch")
        return chart_df, fig

    if chart_type in ["Bar chart", "Horizontal rank bar", "Pie chart"]:
        if not group_col:
            st.warning("Choose a group/category field.")
            return df, None
        chart_df = _aggregate(df, group_col, value_field, aggregation, top_n)
        if chart_df.empty:
            st.warning("No data to display.")
            return df, None
        value_col = [c for c in chart_df.columns if c not in ["Rank", group_col]][0]
        if chart_type == "Pie chart":
            fig = px.pie(chart_df, names=group_col, values=value_col, hole=0.35, title=f"{group_col} by {aggregation.lower()}")
        elif chart_type == "Horizontal rank bar":
            fig = px.bar(chart_df, y=group_col, x=value_col, orientation="h", title=f"Top {top_n} {group_col}")
            fig.update_layout(yaxis={"categoryorder": "total ascending"})
        else:
            fig = px.bar(chart_df, x=group_col, y=value_col, title=f"Top {top_n} {group_col}")
        fig.update_layout(height=360, margin=dict(l=20, r=20, t=45, b=40))
        st.plotly_chart(_polish_figure(fig), width="stretch")
        return chart_df, fig

    if chart_type == "Heatmap":
        if not group_col:
            st.warning("Choose a row/category field for the heatmap.")
            return df, None
        if not color_col:
            # Use KABCO or crash type as the second dimension when possible.
            color_col = _kabco_col(df) or _crash_type_col(df)
        if not color_col or color_col == group_col:
            st.warning("Choose a second category field for the heatmap.")
            return df, None
        work = df.copy()
        work[group_col] = work[group_col].fillna("Unknown").astype(str)
        work[color_col] = work[color_col].fillna("Unknown").astype(str)
        chart_df = work.groupby([group_col, color_col], dropna=False).size().reset_index(name="Count")
        # Limit row categories by total count so the heatmap stays readable.
        top_groups = chart_df.groupby(group_col)["Count"].sum().sort_values(ascending=False).head(top_n).index.tolist()
        chart_df = chart_df[chart_df[group_col].isin(top_groups)]
        pivot = chart_df.pivot_table(index=group_col, columns=color_col, values="Count", aggfunc="sum", fill_value=0)
        order = [c for c in ["K", "A", "B", "C", "O"] if c in pivot.columns]
        other = [c for c in pivot.columns if c not in order]
        pivot = pivot[order + other]
        fig = px.imshow(
            pivot,
            text_auto=True,
            aspect="auto",
            title=f"{group_col} by {color_col}",
            labels=dict(x=color_col, y=group_col, color="Crash count"),
            color_continuous_scale="YlOrRd",
        )
        fig.update_layout(height=420, margin=dict(l=20, r=20, t=45, b=40))
        st.plotly_chart(_polish_figure(fig), width="stretch")
        return pivot.reset_index(), fig

    if chart_type == "Treemap":
        if not group_col:
            st.warning("Choose a category field for the treemap.")
            return df, None
        if not color_col:
            color_col = _kabco_col(df) or _crash_type_col(df)
        work = df.copy()
        work[group_col] = work[group_col].fillna("Unknown").astype(str)
        if color_col and color_col != group_col:
            work[color_col] = work[color_col].fillna("Unknown").astype(str)
            chart_df = work.groupby([group_col, color_col], dropna=False).size().reset_index(name="Count")
            top_groups = chart_df.groupby(group_col)["Count"].sum().sort_values(ascending=False).head(top_n).index.tolist()
            chart_df = chart_df[chart_df[group_col].isin(top_groups)]
            fig = px.treemap(chart_df, path=[group_col, color_col], values="Count", color=color_col, title=f"{group_col} by {color_col}")
        else:
            chart_df = _aggregate(work, group_col, value_field, aggregation, top_n)
            value_col = [c for c in chart_df.columns if c not in ["Rank", group_col]][0]
            fig = px.treemap(chart_df, path=[group_col], values=value_col, color=group_col, title=f"{group_col} treemap")
        fig.update_layout(height=430, margin=dict(l=20, r=20, t=45, b=20))
        st.plotly_chart(_polish_figure(fig), width="stretch")
        return chart_df, fig

    if chart_type == "Radar chart":
        st.warning("Radar is not used by default because many crash categories become hard to read. Use the heatmap or treemap option for crash type by KABCO.")
        return df, None

    if chart_type == "Histogram":
        if not value_field:
            st.warning("Choose a numeric value field for the histogram.")
            return df, None
        fig = px.histogram(df, x=value_field, nbins=25, title=f"Distribution of {value_field}")
        fig.update_layout(height=360, margin=dict(l=20, r=20, t=45, b=40))
        st.plotly_chart(_polish_figure(fig), width="stretch")
        return df[[value_field]].copy(), fig

    if chart_type == "Box plot":
        if not value_field:
            st.warning("Choose a numeric value field for the box plot.")
            return df, None
        if group_col:
            fig = px.box(df, x=group_col, y=value_field, title=f"{value_field} by {group_col}")
        else:
            fig = px.box(df, y=value_field, title=f"{value_field} distribution")
        fig.update_layout(height=360, margin=dict(l=20, r=20, t=45, b=40))
        st.plotly_chart(_polish_figure(fig), width="stretch")
        return df, fig

    return df, None


def _metric_for_map(gdf, map_name):
    cols = list(gdf.columns)
    if "HIN" in map_name:
        for c in ["HIN_Priority_Index", "RiskScore", "CrashDensity", "CrashCount"]:
            if c in cols:
                return c
    if "density" in map_name.lower():
        for c in ["CrashDensity", "CrashDensity_per_mile", "CrashCount", "Crash_Count"]:
            if c in cols:
                return c
    for c in ["CrashDensity", "HIN_Priority_Index", "CrashCount", "Crash_Count"]:
        if c in cols:
            return c
    nums = _numeric_cols(_drop_geometry(gdf))
    return nums[0] if nums else None


def _style_feature(value, min_value, max_value, highlight=False):
    """Dashboard/result map colors: same green-yellow-orange-red ramp as workflow maps."""
    if highlight:
        return {"color": "#111827", "weight": 5, "fillColor": "#f97316", "fillOpacity": 0.80}
    if value is None or pd.isna(value) or max_value == min_value:
        color = "#16a34a"
    else:
        ratio = max(0.0, min(1.0, (float(value) - float(min_value)) / (float(max_value) - float(min_value))))
        if ratio < 0.25:
            color = "#16a34a"      # low = green
        elif ratio < 0.50:
            color = "#facc15"      # moderate = yellow
        elif ratio < 0.75:
            color = "#f97316"      # high = orange
        else:
            color = "#dc2626"      # highest = red
    return {"color": color, "weight": 3, "fillColor": color, "fillOpacity": 0.62}




def _workflow_overlay_sources(st):
    """Return optional workflow layers that can be shown on dashboard maps."""
    sources = {}
    for label, state_key in [
        ("Roads", "selected_roads"),
        ("Road class layer", "roads_class_display"),
        ("Signals", "signals_clean"),
        ("Crash points", "crashes"),
        ("Corridors", "final_corridors"),
        ("Generated corridors", "corridors"),
        ("Study boundary", "selected_boundary"),
    ]:
        data = st.session_state.get(state_key)
        if data is not None and not getattr(data, "empty", True):
            sources[label] = data
    # Avoid duplicate corridor choices when final_corridors and corridors are the same purpose.
    if "Corridors" in sources and "Generated corridors" in sources:
        pass
    return sources


def _overlay_style(label):
    lname = label.lower()
    if "road" in lname:
        return {"color": "#6b7280", "weight": 2, "opacity": 0.65, "fillOpacity": 0.0}
    if "signal" in lname:
        return {"color": "#16a34a", "weight": 2, "opacity": 0.9, "fillColor": "#22c55e", "fillOpacity": 0.9, "radius": 5}
    if "crash" in lname:
        return {"color": "#111827", "weight": 1, "opacity": 0.85, "fillColor": "#ef4444", "fillOpacity": 0.7, "radius": 3}
    if "boundary" in lname:
        return {"color": "#111827", "weight": 2, "opacity": 0.85, "fillOpacity": 0.02}
    if "corridor" in lname:
        return {"color": "#7c3aed", "weight": 3, "opacity": 0.85, "fillOpacity": 0.05}
    return {"color": "#374151", "weight": 2, "opacity": 0.7, "fillOpacity": 0.1}


def _add_overlay_layer(fmap, gdf, label, show=False):
    if gdf is None or getattr(gdf, "empty", True) or folium is None:
        return
    try:
        layer = gdf.copy()
        if layer.crs is None:
            layer = layer.set_crs(epsg=4326)
        layer = layer.to_crs(epsg=4326)
        style = _overlay_style(label)
        group = folium.FeatureGroup(name=label, show=show)
        # Points are rendered as circle markers for readability.
        if all(getattr(geom, "geom_type", "") == "Point" for geom in layer.geometry.dropna().head(50)):
            for _, row in layer.iterrows():
                geom = row.geometry
                if geom is None or geom.is_empty:
                    continue
                popup_cols = [c for c in ["CrashID", "SignalID", "CorridorID", "Route", "FULLNAME"] if c in layer.columns]
                popup = "<br>".join([f"{c}: {row.get(c)}" for c in popup_cols]) if popup_cols else label
                folium.CircleMarker(
                    location=[geom.y, geom.x],
                    radius=style.get("radius", 4),
                    color=style.get("color", "#111827"),
                    weight=style.get("weight", 1),
                    fill=True,
                    fill_color=style.get("fillColor", style.get("color", "#111827")),
                    fill_opacity=style.get("fillOpacity", 0.8),
                    popup=popup,
                ).add_to(group)
        else:
            gj = _safe_geojson_string(layer)
            if gj:
                folium.GeoJson(
                    gj,
                    name=label,
                    style_function=lambda feature, stl=style: stl,
                ).add_to(group)
        group.add_to(fmap)
    except Exception:
        return

def _fit_bounds_for_layer(work):
    """Return finite bounds and padded bounds for Leaflet/Matplotlib."""
    try:
        clean = work[work.geometry.notna() & (~work.geometry.is_empty)].copy()
        if clean.empty:
            clean = work
        minx, miny, maxx, maxy = [float(v) for v in clean.total_bounds]
        if not all(pd.notna(v) for v in [minx, miny, maxx, maxy]):
            return None
        if minx == maxx:
            minx -= 0.002
            maxx += 0.002
        if miny == maxy:
            miny -= 0.002
            maxy += 0.002
        return minx, miny, maxx, maxy
    except Exception:
        return None


def _render_dashboard_map(st, map_name, gdf, key, highlight_value=None, height=460, overlay_layers=None):
    if folium is None or st_folium is None:
        st.info("Install folium and streamlit-folium to show dashboard maps.")
        return
    if gdf is None or getattr(gdf, "empty", True):
        st.info("No map data is available for this layer.")
        return
    work = _repair_gdf_crs(gdf, st)
    work = work.to_crs(epsg=4326)
    bounds = _fit_bounds_for_layer(work)
    if bounds is None:
        st.info("Map bounds could not be calculated for this layer.")
        return
    minx, miny, maxx, maxy = bounds
    center = [(miny + maxy) / 2, (minx + maxx) / 2]
    fmap = folium.Map(location=center, zoom_start=13, tiles="cartodbpositron")
    metric = _metric_for_map(work, map_name)
    values = pd.to_numeric(work[metric], errors="coerce") if metric else pd.Series([], dtype=float)
    min_value = values.min() if not values.empty else 0
    max_value = values.max() if not values.empty else 1
    unit_col = _unit_col(_drop_geometry(work))
    tooltip_cols = [c for c in [unit_col, metric, "CrashCount", "Crash_Count", "Rank"] if c and c in work.columns]

    def style_fn(feature):
        props = feature.get("properties", {})
        val = props.get(metric) if metric else None
        highlight = False
        if highlight_value is not None and unit_col and props.get(unit_col) is not None:
            highlight = str(props.get(unit_col)) == str(highlight_value)
        return _style_feature(val, min_value, max_value, highlight=highlight)

    tooltip = GeoJsonTooltip(fields=tooltip_cols, aliases=tooltip_cols) if GeoJsonTooltip and tooltip_cols else None
    gj = _safe_geojson_string(work)
    if not gj:
        return "<p>Map geometry could not be converted to displayable GeoJSON.</p>"
    folium.GeoJson(gj, name=map_name, style_function=style_fn, tooltip=tooltip, show=True).add_to(fmap)
    if metric:
        _add_map_legend(fmap, _nice_metric_label(metric), min_value, max_value)
    for overlay_name, overlay_gdf in (overlay_layers or {}).items():
        _add_overlay_layer(fmap, overlay_gdf, overlay_name, show=False)
    folium.LayerControl(collapsed=False).add_to(fmap)
    try:
        fmap.fit_bounds([[miny, minx], [maxy, maxx]], padding=(24, 24))
    except Exception:
        pass
    map_key = f"{key}_{round(minx,5)}_{round(miny,5)}_{round(maxx,5)}_{round(maxy,5)}"
    st_folium(fmap, height=height, width="100%", key=map_key, returned_objects=[])



def _build_default_figures(tables):
    figures = []
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))
    if crashes is not None:
        year_kabco, year_col, kabco_col = _year_kabco_table(crashes)
        if not year_kabco.empty:
            fig = px.bar(year_kabco, x=year_col, y="Count", color=kabco_col, color_discrete_map=KABCO_COLOR_MAP, title=f"Crashes by year and {kabco_col}")
            fig.update_layout(barmode="stack", xaxis_title="Year", yaxis_title="Crash count")
            figures.append((f"Crashes by year and {kabco_col}", fig, year_kabco))
        else:
            year_col = _find_col(crashes, ["year"])
            if year_col:
                year_df = _aggregate(crashes, year_col, None, "Count", 20)
                fig = px.bar(year_df.sort_values(year_col), x=year_col, y="Count", title="Crashes by year")
                figures.append(("Crashes by year", fig, year_df))
        type_col = _crash_type_col(crashes)
        if type_col:
            type_df = _aggregate(crashes, type_col, None, "Count", 12)
            pie = px.pie(type_df, names=type_col, values="Count", hole=0.38, title=f"Crash type share by {type_col}")
            figures.append((f"Crash type share by {type_col}", pie, type_df))
        monthly_df, period_col, value_col, color_col = _month_trend_table(crashes)
        if not monthly_df.empty:
            fig = px.line(monthly_df, x=period_col, y=value_col, color=color_col, markers=True, title="Monthly crash trend by year")
            fig.update_layout(xaxis_title="Month", yaxis_title="Crash count")
            figures.append(("Monthly crash trend by year", fig, monthly_df))
        road_kabco, road_col, road_kabco_col = _road_class_kabco_table(crashes, None)
        if not road_kabco.empty:
            pivot = road_kabco.pivot_table(index=road_col, columns=road_kabco_col, values="Count", aggfunc="sum", fill_value=0)
            fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Road class by {road_kabco_col}", labels=dict(x=road_kabco_col, y="Road class", color="Crash count"))
            figures.append((f"Road class by {road_kabco_col}", fig, pivot.reset_index()))
        crash_kabco, crash_type_col, crash_kabco_col = _crash_type_kabco_table(crashes)
        if not crash_kabco.empty:
            pivot = crash_kabco.pivot_table(index=crash_type_col, columns=crash_kabco_col, values="Count", aggfunc="sum", fill_value=0)
            order = [c for c in ["K", "A", "B", "C", "O"] if c in pivot.columns]
            other = [c for c in pivot.columns if c not in order]
            pivot = pivot[order + other]
            fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Crash type by {crash_kabco_col}", labels=dict(x=crash_kabco_col, y="Crash type", color="Crash count"), color_continuous_scale="YlOrRd")
            figures.append((f"Crash type by {crash_kabco_col}", fig, pivot.reset_index()))
    if density is not None:
        metric = "CrashDensity" if "CrashDensity" in density.columns else _default_metric(_numeric_cols(density))
        rank_df, unit_col, value_col = _rank_units_for_chart(density, metric, 15) if metric else (pd.DataFrame(), None, None)
        if not rank_df.empty:
            rank_df = rank_df.sort_values(value_col, ascending=True)
            fig = px.bar(rank_df, y=unit_col, x=value_col, orientation="h", title="Top spatial units by crash density")
            fig.update_layout(yaxis_title="Spatial unit ID", xaxis_title="Crash density")
            figures.append(("Top spatial units by crash density", fig, rank_df))
    if hin is not None:
        metric = _default_metric(_numeric_cols(hin))
        rank_df, unit_col, value_col = _hin_route_rank_for_chart(hin, metric, 15) if metric else (pd.DataFrame(), None, None)
        if not rank_df.empty:
            rank_df = rank_df.sort_values(value_col, ascending=True)
            fig = px.bar(rank_df, y=unit_col, x=value_col, orientation="h", title=f"Top routes by highest {_nice_metric_label(metric)}")
            fig.update_layout(yaxis_title="Route", xaxis_title=_nice_metric_label(metric))
            figures.append((f"Top routes by highest {_nice_metric_label(metric)}", fig, rank_df))
    return figures

def _normal_col(df, candidates):
    if df is None:
        return None
    lower_map = {str(c).lower().replace("_", "").replace(" ", ""): c for c in df.columns}
    for cand in candidates:
        key = cand.lower().replace("_", "").replace(" ", "")
        if key in lower_map:
            return lower_map[key]
    for c in df.columns:
        cl = str(c).lower()
        if any(cand.lower() in cl for cand in candidates):
            return c
    return None


def _is_bad_route_candidate(name):
    """Return True for fields that should not be used as route names.

    Dashboard charts should group by road/route name.  Very short severity
    fields such as B, C, O or injury-count fields can accidentally match
    substring searches and create y-axis labels such as 1B / 2B.
    """
    low = str(name).strip().lower()
    compact = low.replace("_", "").replace(" ", "")
    bad_exact = {
        "k", "a", "b", "c", "o", "kabco", "severity", "dashboardkabco",
        "fatalities", "fatals", "seriousinjuries", "minorinjuries",
        "possibleinjuries", "noinjury", "dashboardfatalities",
        "dashboardseriousinjuries", "dashboardminorinjuries",
        "dashboardpossibleinjuries", "dashboardnoinjury",
    }
    if compact in bad_exact:
        return True
    bad_words = [
        "injur", "fatal", "severity", "kabco", "crashcount", "crash_count",
        "count", "score", "index", "density", "mile", "length", "year",
    ]
    return any(w in low for w in bad_words)


def _dashboard_route_col(df):
    """Find a true route/road-name column for dashboard charts.

    This is stricter than _normal_col because _normal_col allows substring
    matching.  A candidate such as B can otherwise be selected as a route
    column because route-like fields are missing from a result table.
    """
    if df is None or not hasattr(df, "columns"):
        return None

    exact_candidates = [
        "Route", "FULLNAME", "RoadName", "Road_Name", "RouteName",
        "RouteName_Calc", "CorridorRoute", "Corridor_Route", "StreetName",
        "Street_Name", "Name", "NAME", "DisplayName", "Road", "FacilityName",
    ]
    norm_map = {str(c).lower().replace("_", "").replace(" ", ""): c for c in df.columns}
    for cand in exact_candidates:
        key = cand.lower().replace("_", "").replace(" ", "")
        col = norm_map.get(key)
        if col is not None and not _is_bad_route_candidate(col):
            try:
                vals = df[col].dropna()
                if not vals.empty and not pd.to_numeric(vals, errors="coerce").notna().all():
                    return col
            except Exception:
                return col

    # Controlled fallback: require route/road/street/name words and exclude
    # numeric/severity fields.
    for col in df.columns:
        low = str(col).lower()
        if _is_bad_route_candidate(col):
            continue
        if not any(w in low for w in ["route", "road", "street", "corridor", "fullname", "facility", "name"]):
            continue
        try:
            vals = df[col].dropna()
            if vals.empty:
                continue
            if pd.to_numeric(vals, errors="coerce").notna().mean() > 0.8:
                continue
        except Exception:
            pass
        return col

    return None


def _top_density_export_table(density, top_n=20):
    """Decision-ready top crash-density table with only useful columns."""
    if density is None or getattr(density, "empty", True):
        return pd.DataFrame()
    df = _drop_geometry(density).copy()
    unit_col = _unit_col(df) or _normal_col(df, ["UnitID", "IntersectionID", "CorridorID", "SegmentID", "Route"])
    unit_type_col = _normal_col(df, ["UnitType", "IntersectionType", "CorridorType", "SegmentType"])
    city_col = _normal_col(df, ["City", "city_name"])
    length_col = _normal_col(df, ["Length_Miles", "Length_Mi", "length_mi", "Miles"])
    count_col = _crash_count_col(df)
    density_col = _normal_col(df, ["CrashDensity", "Crash_Density", "crash_density"])
    if density_col is None:
        density_col = _default_metric(_numeric_cols(df))
    if density_col:
        df[density_col] = pd.to_numeric(df[density_col], errors="coerce").fillna(0)
        df = df.sort_values(density_col, ascending=False)
    out = pd.DataFrame()
    out["Rank"] = range(1, min(top_n, len(df)) + 1)
    use = df.head(top_n).reset_index(drop=True)
    out["Spatial unit id"] = use[unit_col].astype(str) if unit_col else use.index.astype(str)
    out["Unit type"] = use[unit_type_col].astype(str) if unit_type_col else ""
    out["City"] = use[city_col].astype(str) if city_col else ""
    out["Length_mi"] = pd.to_numeric(use[length_col], errors="coerce").round(3) if length_col else ""

    # Context columns requested for decision tables. Intersections show the two
    # crossing roads; corridors/segments show route name when available.
    road1_col = _normal_col(use, ["RoadName1", "Road1", "Route1", "Street1", "FromRoad"])
    road2_col = _normal_col(use, ["RoadName2", "Road2", "Route2", "Street2", "ToRoad"])
    route_col = _normal_col(use, ["Route", "FULLNAME", "RoadName", "RouteName", "CorridorRoute", "RouteName_Calc"])
    unit_type_text = str(use[unit_type_col].iloc[0]).lower() if unit_type_col and len(use) else ""
    if road1_col and road2_col:
        out["Road 1"] = use[road1_col].astype(str)
        out["Road 2"] = use[road2_col].astype(str)
    elif route_col:
        out["Route name"] = use[route_col].astype(str)

    out["Crash count"] = pd.to_numeric(use[count_col], errors="coerce").fillna(0).astype(int) if count_col else ""
    out["Crash density"] = pd.to_numeric(use[density_col], errors="coerce").round(3) if density_col else ""
    return out


def _top_hin_export_table(hin, top_n=20):
    if hin is None or getattr(hin, "empty", True):
        return pd.DataFrame()
    df = _drop_geometry(hin).copy()
    metric = _default_metric(_numeric_cols(df))
    unit_col = _unit_col(df) or _normal_col(df, ["UnitID", "RiskSegmentID", "WindowID", "SegmentID", "CorridorID", "Route"])
    if unit_col is not None and any(bad in str(unit_col).lower() for bad in ["length", "mile", "density", "count", "score", "index"]):
        unit_col = None
    if metric:
        df[metric] = pd.to_numeric(df[metric], errors="coerce").fillna(0)
        df = df.sort_values(metric, ascending=False)
    use = df.head(top_n).reset_index(drop=True)
    out = pd.DataFrame({"Rank": range(1, len(use) + 1)})
    out["Spatial unit id"] = use[unit_col].astype(str) if unit_col else [f"UNIT_{i + 1}" for i in range(len(use))]
    if "UnitType" in use.columns:
        out["Unit type"] = use["UnitType"].astype(str)
    route_col = _normal_col(use, ["Route", "FULLNAME", "RoadName", "RouteName", "RouteName_Calc"])
    if route_col:
        out["Route name"] = use[route_col].astype(str)
    if "Length_Miles" in use.columns:
        out["Length_mi"] = pd.to_numeric(use["Length_Miles"], errors="coerce").round(3)
    if metric:
        out[metric] = pd.to_numeric(use[metric], errors="coerce").round(3)
    return out


def _severity_summary_export(summary, top_n=20):
    if summary is None or getattr(summary, "empty", True):
        return pd.DataFrame()
    df = _drop_geometry(summary).copy()
    metric = _crash_count_col(df) or _normal_col(df, ["Total", "CrashCount", "Crash_Count", "CrashDensity"])
    if metric and metric in df.columns:
        df[metric] = pd.to_numeric(df[metric], errors="coerce").fillna(0)
        df = df.sort_values(metric, ascending=False)
    df = df.head(top_n).copy().reset_index(drop=True)
    if "Rank" not in df.columns:
        df.insert(0, "Rank", range(1, len(df) + 1))
    keep = [c for c in ["Rank", "UnitID", "UnitType", "City", "CrashCount", "CrashDensity", "K", "A", "B", "C", "O", "Total"] if c in df.columns]
    return df[keep] if keep else df.iloc[:, :min(8, len(df.columns))]


def _export_tables_only(tables, top_n=20):
    """Return only report-ready summary tables.

    Excludes raw uploaded/assigned rows and the KABCO distribution table because
    KABCO is already shown as a chart. Ranking is only used for spatial-unit
    priority tables.
    """
    out = {}
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))
    summary = tables.get("Crash count / severity summary")

    if crashes is not None:
        type_col = _crash_type_col(crashes)
        year_col = _find_col(crashes, ["year"])
        if type_col:
            out["Crash type summary"] = _aggregate(crashes, type_col, None, "Count", top_n)
        if year_col:
            out["Crash year summary"] = _aggregate(crashes, year_col, None, "Count", top_n)

    if density is not None:
        top_density = _top_density_export_table(density, top_n=top_n)
        if not top_density.empty:
            out["Top crash-density spatial units"] = top_density

    if summary is not None:
        sev = _severity_summary_export(summary, top_n=top_n)
        if not sev.empty:
            out["Severity summary by spatial unit"] = sev

    if hin is not None:
        top_hin = _top_hin_export_table(hin, top_n=top_n)
        if not top_hin.empty:
            out["Top HIN/risk spatial units"] = top_hin

    return {name: _safe_dataframe_for_display(df) for name, df in out.items()}

def _download_excel_bytes(tables):
    buffer = io.BytesIO()
    export_tables = _export_tables_only(tables)
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for name, df in export_tables.items():
            sheet = _safe_name(name)[:31] or "data"
            _drop_geometry(df).to_excel(writer, sheet_name=sheet, index=False)
    buffer.seek(0)
    return buffer.getvalue()





def _legend_html(title, min_value, max_value):
    try:
        min_label = f"{float(min_value):,.2f}"
        max_label = f"{float(max_value):,.2f}"
    except Exception:
        min_label, max_label = "Low", "High"
    return f"""
    <div style='position: fixed; bottom: 28px; right: 28px; z-index:9999;
         background: white; border: 1px solid #d1d5db; border-radius: 10px;
         padding: 10px 12px; font-size: 12px; box-shadow: 0 6px 18px rgba(0,0,0,.12);'>
      <div style='font-weight:700; margin-bottom:6px;'>{html.escape(str(title))}</div>
      <div><span style='background:#16a34a;display:inline-block;width:14px;height:10px;margin-right:6px;'></span>Low ({min_label})</div>
      <div><span style='background:#facc15;display:inline-block;width:14px;height:10px;margin-right:6px;'></span>Moderate</div>
      <div><span style='background:#f97316;display:inline-block;width:14px;height:10px;margin-right:6px;'></span>High</div>
      <div><span style='background:#dc2626;display:inline-block;width:14px;height:10px;margin-right:6px;'></span>Highest ({max_label})</div>
    </div>
    """


def _add_map_legend(fmap, title, min_value, max_value):
    try:
        from branca.element import Element
        fmap.get_root().html.add_child(Element(_legend_html(title, min_value, max_value)))
    except Exception:
        pass

def _dashboard_map_html(map_name, gdf, height=520, overlay_layers=None):
    if folium is None or gdf is None or getattr(gdf, "empty", True):
        return "<p>No map data available.</p>"
    work = _repair_gdf_crs(gdf, None)
    work = work.to_crs(epsg=4326)
    bounds = _fit_bounds_for_layer(work)
    if bounds is None:
        return "<p>Map bounds could not be calculated.</p>"
    minx, miny, maxx, maxy = bounds
    fmap = folium.Map(location=[(miny + maxy) / 2, (minx + maxx) / 2], zoom_start=13, tiles="cartodbpositron", height=height)
    metric = _metric_for_map(work, map_name)
    values = pd.to_numeric(work[metric], errors="coerce") if metric else pd.Series([], dtype=float)
    min_value = values.min() if not values.empty else 0
    max_value = values.max() if not values.empty else 1
    unit_col = _unit_col(_drop_geometry(work))
    tooltip_cols = [c for c in [unit_col, metric, "CrashCount", "Crash_Count", "Rank"] if c and c in work.columns]
    def style_fn(feature):
        props = feature.get("properties", {})
        return _style_feature(props.get(metric) if metric else None, min_value, max_value, highlight=False)
    tooltip = GeoJsonTooltip(fields=tooltip_cols, aliases=tooltip_cols) if GeoJsonTooltip and tooltip_cols else None
    gj = _safe_geojson_string(work)
    if not gj:
        return "<p>Map geometry could not be converted to displayable GeoJSON.</p>"
    folium.GeoJson(gj, name=map_name, style_function=style_fn, tooltip=tooltip, show=True).add_to(fmap)
    if metric:
        _add_map_legend(fmap, _nice_metric_label(metric), min_value, max_value)
    for overlay_name, overlay_gdf in (overlay_layers or {}).items():
        _add_overlay_layer(fmap, overlay_gdf, overlay_name, show=False)
    folium.LayerControl(collapsed=False).add_to(fmap)
    try:
        fmap.fit_bounds([[miny, minx], [maxy, maxx]], padding=(24, 24))
    except Exception:
        pass
    return fmap.get_root().render()

def _export_dashboard_html(tables, selected_blocks, selected_maps, maps=None, extra_figures=None, overlay_layers=None, report_timezone=None):
    figures = _build_default_figures(tables) + (extra_figures or [])
    parts = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        "<title>HIN dashboard report</title>",
        "<style>body{font-family:Arial,sans-serif;margin:32px;color:#1f2937} .card{border:1px solid #e5e7eb;border-radius:14px;padding:16px;margin:14px 0} h1{margin-bottom:4px} h2{margin-top:28px}</style>",
        "</head><body>",
        "<h1>HIN dashboard report</h1>",
        f"<p>Generated {_report_time_text(report_timezone)}</p>",
    ]
    for title, fig, data in figures:
        if title in selected_blocks or not selected_blocks:
            parts.append(f"<div class='card'><h2>{html.escape(title)}</h2>")
            parts.append(pio.to_html(fig, include_plotlyjs="cdn", full_html=False))
            parts.append("</div>")
    if selected_maps:
        parts.append("<div class='card'><h2>Selected map layers</h2>")
        for m in selected_maps:
            parts.append(f"<h3>{html.escape(m)}</h3>")
            if maps and m in maps:
                parts.append(_dashboard_map_html(m, maps[m], overlay_layers=overlay_layers))
            else:
                parts.append(f"<p>{html.escape(m)}</p>")
        parts.append("</div>")
    for name, df in _export_tables_only(tables).items():
        parts.append(f"<div class='card'><h2>{html.escape(name)}</h2>")
        parts.append(_drop_geometry(df).head(25).to_html(index=False, escape=True))
        parts.append("</div>")
    parts.append("</body></html>")
    return "\n".join(parts).encode("utf-8")



def _static_map_png(gdf, title="Map layer", overlay_layers=None):
    """Create a static map image for Word/PNG exports.

    This uses the same result geometry and optional workflow context layers
    (roads/signals/corridors/boundary) so the report contains an actual map
    image instead of only a text placeholder or an interactive HTML map.
    """
    if gdf is None or getattr(gdf, "empty", True):
        return None
    try:
        import matplotlib.pyplot as plt
        from matplotlib.lines import Line2D
        work = _repair_gdf_crs(gdf, None)
        work = work.to_crs(epsg=4326)
        metric = _metric_for_map(work, title)
        fig, ax = plt.subplots(figsize=(9.5, 6.2))

        # Light context layers first.
        for label, layer in (overlay_layers or {}).items():
            if layer is None or getattr(layer, "empty", True):
                continue
            try:
                ov = layer.copy()
                if ov.crs is None:
                    ov = ov.set_crs(epsg=4326)
                ov = ov.to_crs(epsg=4326)
                lname = str(label).lower()
                if "road" in lname:
                    ov.plot(ax=ax, color="#94a3b8", linewidth=0.55, alpha=0.55)
                elif "signal" in lname and all(getattr(g, "geom_type", "") == "Point" for g in ov.geometry.dropna().head(25)):
                    ov.plot(ax=ax, color="#16a34a", markersize=12, alpha=0.85)
                elif "crash" in lname and all(getattr(g, "geom_type", "") == "Point" for g in ov.geometry.dropna().head(25)):
                    ov.plot(ax=ax, color="#ef4444", markersize=5, alpha=0.45)
                elif "boundary" in lname:
                    ov.boundary.plot(ax=ax, color="#111827", linewidth=1.0, alpha=0.7)
                elif "corridor" in lname:
                    ov.plot(ax=ax, color="#7c3aed", linewidth=1.0, alpha=0.30)
            except Exception:
                pass

        work = _safe_geojson_gdf(work)
        geom_types = set(str(g.geom_type) for g in work.geometry.dropna().head(50))
        if metric and metric in work.columns:
            # Use same low-to-high color meaning as the workflow maps: green -> yellow -> orange -> red.
            # Plot the result layer last and with high z-order so it is visible
            # above context roads/signals in the Word report. For tiny
            # intersection/corridor polygons, add representative-point markers
            # colored by the same metric so the crash-density/HIN layer is clear.
            if geom_types and all("Line" in gt for gt in geom_types):
                work.plot(ax=ax, column=metric, legend=True, cmap="RdYlGn_r", linewidth=1.8, alpha=0.96, zorder=10, legend_kwds={"label": _nice_metric_label(metric), "shrink": 0.72})
            elif geom_types and all(gt == "Point" or gt == "MultiPoint" for gt in geom_types):
                work.plot(ax=ax, column=metric, legend=True, cmap="RdYlGn_r", markersize=46, alpha=0.96, zorder=10, legend_kwds={"label": _nice_metric_label(metric), "shrink": 0.72})
            else:
                work.plot(ax=ax, column=metric, legend=True, cmap="RdYlGn_r", linewidth=1.2, edgecolor="#374151", alpha=0.55, zorder=9, legend_kwds={"label": _nice_metric_label(metric), "shrink": 0.72})
                try:
                    pts = work.copy()
                    pts["geometry"] = pts.geometry.representative_point()
                    pts.plot(ax=ax, column=metric, cmap="RdYlGn_r", markersize=36, alpha=0.98, zorder=12)
                except Exception:
                    pass
        else:
            work.plot(ax=ax, color="#60a5fa", linewidth=1.2, edgecolor="#374151", alpha=0.75, zorder=10)
        bounds = _fit_bounds_for_layer(work)
        if bounds is None:
            return None
        minx, miny, maxx, maxy = bounds
        dx = max((maxx - minx) * 0.08, 0.002)
        dy = max((maxy - miny) * 0.08, 0.002)
        ax.set_xlim(minx - dx, maxx + dx)
        ax.set_ylim(miny - dy, maxy + dy)
        ax.set_title(title, fontsize=13, fontweight="bold")
        ax.set_axis_off()
        legend_items = []
        labels = set((overlay_layers or {}).keys())
        if "Roads" in labels or "Road class layer" in labels:
            legend_items.append(Line2D([0], [0], color="#94a3b8", lw=2, label="Roads"))
        if "Signals" in labels:
            legend_items.append(Line2D([0], [0], marker="o", color="w", markerfacecolor="#16a34a", markersize=7, label="Signals"))
        if "Crash points" in labels:
            legend_items.append(Line2D([0], [0], marker="o", color="w", markerfacecolor="#ef4444", markersize=6, label="Crash points"))
        if legend_items:
            ax.legend(handles=legend_items, loc="lower left", frameon=True, fontsize=8)
        buffer = io.BytesIO()
        fig.savefig(buffer, format="png", dpi=180, bbox_inches="tight")
        plt.close(fig)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception:
        return None

def _export_dashboard_docx(tables, selected_blocks, selected_maps, extra_figures=None, maps=None, overlay_layers=None, report_timezone=None):
    if Document is None:
        return None
    doc = Document()
    doc.add_heading("HIN dashboard report", level=0)
    doc.add_paragraph(f"Generated {_report_time_text(report_timezone)}.")
    doc.add_paragraph("This report summarizes crash patterns, crash-density rankings, selected maps, and HIN/risk results from the app dashboard.")

    figures = _build_default_figures(tables) + (extra_figures or [])
    for title, fig, data in figures:
        if selected_blocks and title not in selected_blocks:
            continue
        doc.add_heading(title, level=1)
        img = _figure_to_png_bytes(_polish_figure(fig))
        if img:
            img_buf = io.BytesIO(img)
            doc.add_picture(img_buf, width=Inches(6.5))
        else:
            doc.add_paragraph("Chart image could not be generated in this environment. The summary table is included below.")
        doc.add_paragraph("Summary table")
        table_df = _safe_dataframe_for_display(data.head(15).copy())
        if table_df.empty:
            continue
        table = doc.add_table(rows=1, cols=len(table_df.columns))
        table.style = "Table Grid"
        for i, col in enumerate(table_df.columns):
            table.rows[0].cells[i].text = str(col)
        for _, row in table_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(table_df.columns):
                cells[i].text = str(row[col])

    if selected_maps:
        doc.add_heading("Selected map layers", level=1)
        for m in selected_maps:
            doc.add_heading(str(m), level=2)
            if maps and m in maps:
                map_png = _static_map_png(maps[m], str(m), overlay_layers=overlay_layers)
                if map_png:
                    doc.add_picture(io.BytesIO(map_png), width=Inches(6.5))
                else:
                    doc.add_paragraph("Static map image could not be generated.")
            else:
                doc.add_paragraph("Map layer selected in dashboard builder.")

    doc.add_heading("Decision-ready result tables", level=1)
    for name, df in _export_tables_only(tables).items():
        doc.add_heading(name, level=2)
        table_df = _safe_dataframe_for_display(df).head(20).copy()
        if table_df.empty:
            doc.add_paragraph("No rows.")
            continue
        max_cols = min(10, len(table_df.columns))
        table_df = table_df.iloc[:, :max_cols]
        table = doc.add_table(rows=1, cols=len(table_df.columns))
        table.style = "Table Grid"
        for i, col in enumerate(table_df.columns):
            table.rows[0].cells[i].text = str(col)
        for _, row in table_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(table_df.columns):
                cells[i].text = str(row[col])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def _export_summary_image(tables, image_format="png", extra_figures=None):
    figures = _build_default_figures(tables) + (extra_figures or [])
    if not figures:
        return None
    try:
        from PIL import Image
        images = []
        for _, fig, _ in figures[:4]:
            img_bytes = pio.to_image(_polish_figure(fig), format="png", width=900, height=520, scale=2)
            images.append(Image.open(io.BytesIO(img_bytes)).convert("RGB"))
        if not images:
            return None
        cols = 2 if len(images) > 1 else 1
        rows = (len(images) + cols - 1) // cols
        w = max(im.width for im in images)
        h = max(im.height for im in images)
        canvas = Image.new("RGB", (cols * w, rows * h), "white")
        for i, im in enumerate(images):
            canvas.paste(im, ((i % cols) * w, (i // cols) * h))
        buffer = io.BytesIO()
        fmt = "JPEG" if image_format.lower() in ["jpg", "jpeg"] else "PNG"
        canvas.save(buffer, format=fmt, quality=92)
        return buffer.getvalue()
    except Exception:
        try:
            return pio.to_image(_polish_figure(figures[0][1]), format=image_format, width=1400, height=800, scale=2)
        except Exception:
            return None


def _render_dashboard_builder(st, tables):
    maps = _available_maps(st)
    custom_figures = st.session_state.get("dashboard_custom_figures", [])
    default_figures = _build_default_figures(tables) + custom_figures
    figure_titles = [title for title, _, _ in default_figures]

    st.markdown("<div class='dashboard-section-title'>Dashboard builder <span>choose charts, tables, and map layers for one review page</span></div>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([0.27, 0.27, 0.46], gap="large")
    with c1:
        st.markdown("**Charts and figures**")
        selected_blocks = st.multiselect(
            "Select dashboard charts",
            figure_titles,
            default=figure_titles[: min(4, len(figure_titles))],
            key="dash_builder_chart_blocks",
        )
        include_tables = st.checkbox("Include ranking/data table previews", value=True, key="dash_builder_include_tables")
        if custom_figures and st.button("Clear added custom charts", key="clear_custom_dashboard_charts"):
            st.session_state["dashboard_custom_figures"] = []
            st.rerun()
    with c2:
        st.markdown("**Map layers**")
        selected_maps = st.multiselect(
            "Select dashboard maps",
            list(maps.keys()),
            default=list(maps.keys())[: min(2, len(maps))],
            key="dash_builder_map_layers",
            help="Dashboard maps are read-only. Map editing and filtering stay in the Visualization section.",
        )
        overlay_sources = _workflow_overlay_sources(st)
        selected_overlays = st.multiselect(
            "Optional workflow layers on dashboard maps",
            list(overlay_sources.keys()),
            default=[name for name in ["Roads", "Signals"] if name in overlay_sources],
            key="dash_builder_overlay_layers",
            help="Turn on context layers such as roads, signals, crash points, corridors, or study boundary.",
        )
        # Highlight selector removed: dashboard maps are read-only context layers.
        highlight_layer = None
        highlight_value = None
    with c3:
        st.markdown("**Exports**")
        st.caption("Export the dashboard as a static PNG summary or a Word report with charts, map summaries, and decision-ready tables.")
        report_timezone = st.selectbox(
            "Report time zone",
            ["America/Denver", "Local/server time", "UTC", "America/Chicago", "America/Los_Angeles", "America/New_York"],
            index=0,
            key="dashboard_report_timezone",
            help="Streamlit Cloud often runs in UTC. Choose the local project timezone so the report timestamp matches your expected local time.",
        )
        report_tz_value = None if report_timezone == "Local/server time" else report_timezone
        d1, d2, d3 = st.columns(3)
        with d1:
            png_bytes = _export_summary_image(tables, "png", extra_figures=custom_figures)
            if png_bytes:
                st.download_button("Download PNG", data=png_bytes, file_name="hin_dashboard_summary.png", mime="image/png", key="dash_export_png")
            else:
                st.caption("PNG needs kaleido")
        with d2:
            docx_bytes = _export_dashboard_docx(tables, selected_blocks, selected_maps, extra_figures=custom_figures, maps=maps, overlay_layers={name: overlay_sources[name] for name in selected_overlays if name in overlay_sources}, report_timezone=report_tz_value)
            if docx_bytes is not None:
                st.download_button(
                    "Download Word report",
                    data=docx_bytes,
                    file_name=_report_docx_filename(tables),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dash_export_docx",
                )
            else:
                st.info("Install python-docx for Word export.")

    st.markdown("<div class='dashboard-section-title'>Generated dashboard</div>", unsafe_allow_html=True)
    chart_titles = set(selected_blocks)
    for i in range(0, len(default_figures), 2):
        cols = st.columns(2)
        for j, col in enumerate(cols):
            idx = i + j
            if idx >= len(default_figures):
                continue
            title, fig, data = default_figures[idx]
            if title not in chart_titles:
                continue
            with col:
                fig.update_layout(height=330, margin=dict(l=20, r=20, t=45, b=35))
                st.plotly_chart(_polish_figure(fig), width="stretch", key=f"dash_generated_fig_{idx}")

    if selected_maps:
        map_cols = st.columns(min(2, len(selected_maps)))
        for i, map_name in enumerate(selected_maps[:2]):
            with map_cols[i % len(map_cols)]:
                st.markdown(f"**{map_name}**")
                _render_dashboard_map(
                    st,
                    map_name,
                    maps[map_name],
                    key=f"dash_map_{_safe_name(map_name)}_{i}",
                    height=380,
                    overlay_layers={name: overlay_sources[name] for name in selected_overlays if name in overlay_sources},
                )

    if include_tables:
        table_name = st.selectbox("Dashboard table preview", list(tables.keys()), key="dash_builder_table_preview")
        df = tables[table_name].copy()
        metric = _default_metric(_numeric_cols(df))
        st.dataframe(_safe_dataframe_for_display(_rank_table(df, metric).head(50)), width="stretch", hide_index=True)


def _render_sliding_window_route_summary_table(st, tables):
    """Show route-level sliding-window summary in the dashboard."""

    route_summary = tables.get("Sliding-window route summary")
    if route_summary is None or getattr(route_summary, "empty", True):
        return

    st.markdown(
        "<div class='dashboard-section-title'>Sliding-window route summary "
        "<span>route length, window count, and route-level max scores</span></div>",
        unsafe_allow_html=True,
    )

    display = route_summary.copy()
    preferred_cols = [
        "Dashboard_Route_Name",
        "Route",
        "Route_Length_Miles",
        "Window_Count",
        "Assigned_Crash_Count",
        "Assigned_EPDO",
        "Max_Window_Crash_Count",
        "Max_Window_EPDO",
        "Max_High_Risk_Score",
        "Max_HIN_Priority_Index",
    ]
    cols = [c for c in preferred_cols if c in display.columns]
    if cols:
        display = display[cols].copy()

    for col in [
        "Route_Length_Miles",
        "Assigned_EPDO",
        "Max_Window_EPDO",
        "Max_High_Risk_Score",
        "Max_HIN_Priority_Index",
    ]:
        if col in display.columns:
            display[col] = pd.to_numeric(display[col], errors="coerce").round(3)

    st.dataframe(
        _safe_dataframe_for_display(display),
        width="stretch",
        hide_index=True,
    )


def render_dashboard_page(st):
    _style(st)

    st.markdown(
        """
        <div class='dashboard-hero'>
          <h1>Results dashboard</h1>
          <p>Build a decision-ready dashboard with crash trends, crash type summaries, spatial-unit rankings, crash-density maps, and HIN priority maps. Dashboard maps are read-only; analysis results remain unchanged.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    top_left, top_right = st.columns([0.78, 0.22])
    with top_right:
        if st.button("Back to workflow", key="dashboard_back_to_workflow", width="stretch"):
            st.session_state["dashboard_mode"] = False
            st.rerun()

    tables = _available_tables(st)
    if not tables:
        st.info("Run the workflow first. Dashboard options will appear after data or results are available.")
        return

    # Compact production dashboard: remove row/column status cards so charts have more room.

    tab_insights, tab_builder, tab_tables = st.tabs([
        "Crash insights",
        "Dashboard builder",
        "Data tables",
    ])

    st.markdown("<div class='dashboard-scroll-note'>Scroll down to view all charts, maps, and tables. Use the Dashboard builder tab to select charts/maps for export.</div>", unsafe_allow_html=True)

    with tab_insights:
        _render_pattern_charts(st, tables)
        _render_sliding_window_route_summary_table(st, tables)

    with tab_builder:
        _render_dashboard_builder(st, tables)

    # Chart Builder was removed from the dashboard navigation.
    # Use Crash insights for default figures, Dashboard assistant for plain-language
    # chart creation, and Dashboard builder to select final charts/maps.

    with tab_tables:
        table_name = st.selectbox("Table", list(tables.keys()), key="dash_full_table_dataset")
        df = tables[table_name].copy()
        search_text = st.text_input("Search table text", key="dash_table_search")
        if search_text:
            mask = df.astype(str).apply(lambda col: col.str.contains(search_text, case=False, na=False)).any(axis=1)
            df = df[mask]
        metric = _default_metric(_numeric_cols(df))
        ranked = _rank_table(df, metric)
        st.dataframe(_safe_dataframe_for_display(ranked), width="stretch", hide_index=True)

        d1, d2 = st.columns(2)
        with d1:
            st.download_button(
                "Download this table CSV",
                data=_safe_dataframe_for_display(ranked).to_csv(index=False).encode("utf-8"),
                file_name=f"dashboard_{_safe_name(table_name)}.csv",
                mime="text/csv",
                key="dashboard_download_table_csv",
            )
        with d2:
            st.download_button(
                "Download all dashboard tables Excel",
                data=_download_excel_bytes(tables),
                file_name="hin_dashboard_tables.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="dashboard_download_all_excel",
            )


# --- V28 dashboard chart and HIN summary overrides ---

def _kabco_order_values(series):
    order = ["K", "A", "B", "C", "O"]
    vals = [str(v) for v in series.dropna().astype(str).unique().tolist()]
    return [v for v in order if v in vals] + [v for v in vals if v not in order]


def _apply_kabco_trace_colors(fig):
    """Force K/A/B/C/O traces to use distinct colors after Plotly creation."""
    try:
        for i, trace in enumerate(fig.data):
            name = str(getattr(trace, "name", ""))
            color = KABCO_COLOR_MAP.get(name, APP_PALETTE[i % len(APP_PALETTE)])
            if getattr(trace, "type", "") == "bar":
                trace.marker.color = color
            elif getattr(trace, "type", "") == "scatter":
                trace.line.color = color
                trace.marker.color = color
        fig.update_layout(colorway=APP_PALETTE)
    except Exception:
        pass
    return fig


def _context_cols_for_hover(df):
    cols = []
    for c in [
        "Route", "FULLNAME", "RoadName", "RouteName", "CorridorRoute",
        "FromMile", "ToMile", "From_Mile", "To_Mile", "from_mile", "to_mile",
        "Length_Miles", "Length_Mi", "SegmentLength_Mile", "WindowLength_Miles",
        "CrashCount", "Crash_Count", "CrashDensity", "HIN_Priority_Index",
    ]:
        if df is not None and c in df.columns and c not in cols:
            cols.append(c)
    return cols


def _rank_units_for_chart(df, metric, top_n=15):
    """Rank rows and keep route/milepost hover context."""
    if df is None or getattr(df, "empty", True) or metric not in df.columns:
        return pd.DataFrame(), None, None
    unit_col = _unit_col(df)
    work = df.copy()
    if unit_col is None:
        unit_col = "DashboardUnitID"
        work[unit_col] = [f"UNIT_{i + 1}" for i in range(len(work))]
    else:
        lower_unit = str(unit_col).lower()
        if any(bad in lower_unit for bad in ["length", "mile", "density", "count", "score", "index"]):
            unit_col = "DashboardUnitID"
            work[unit_col] = [f"UNIT_{i + 1}" for i in range(len(work))]
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work[unit_col] = work[unit_col].astype(str)
    keep_cols = [unit_col, metric]
    for extra in _context_cols_for_hover(work) + ["UnitType", "City", "RoadName1", "RoadName2", "CorridorID"]:
        if extra in work.columns and extra not in keep_cols:
            keep_cols.append(extra)
    out = work[keep_cols].sort_values(metric, ascending=False).head(top_n).reset_index(drop=True)
    out.insert(0, "Rank", range(1, len(out) + 1))
    return out, unit_col, metric


def _hin_segment_rank_for_chart(hin, metric, top_n=20):
    """Show each HIN segment/window as one bar, not stacked or grouped by route."""
    if hin is None or getattr(hin, "empty", True) or metric not in hin.columns:
        return pd.DataFrame(), None, None
    work = hin.copy()
    id_col = _unit_col(work)
    if id_col is None or any(bad in str(id_col).lower() for bad in ["length", "mile", "density", "count", "score", "index"]):
        id_col = _normal_col(work, ["RiskSegmentID", "WindowID", "SegmentID", "UnitID"])
    if id_col is None:
        id_col = "HIN segment ID"
        work[id_col] = [f"HIN_{i + 1}" for i in range(len(work))]
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work[id_col] = work[id_col].astype(str)
    keep_cols = [id_col, metric]
    for c in _context_cols_for_hover(work):
        if c not in keep_cols:
            keep_cols.append(c)
    out = work[keep_cols].sort_values(metric, ascending=False).head(top_n).reset_index(drop=True)
    out.insert(0, "Rank", range(1, len(out) + 1))
    return out, id_col, metric


def _mode_severity_bubble_table(crashes):
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), None, None
    type_col = _crash_type_col(crashes)
    kabco_col = _kabco_col(crashes)
    if not type_col or not kabco_col:
        return pd.DataFrame(), None, None
    work = crashes[[kabco_col, type_col]].copy()
    text = work[type_col].fillna("").astype(str).str.lower()
    work["Mode"] = "Motor vehicle / other"
    work.loc[text.str.contains("ped", na=False), "Mode"] = "Pedestrian"
    work.loc[text.str.contains("bicycle|bike", na=False), "Mode"] = "Bicycle"
    work.loc[text.str.contains("motorcycle|motor bike", na=False), "Mode"] = "Motorcycle"
    work[kabco_col] = work[kabco_col].fillna("Unknown").astype(str)
    if (work["Mode"] != "Motor vehicle / other").sum() == 0:
        return pd.DataFrame(), None, None
    out = work.groupby([kabco_col, "Mode"], dropna=False).size().reset_index(name="Count")
    out = _order_kabco(out, kabco_col)
    return out, kabco_col, "Mode"


def _summary_kpi_values(crashes):
    vals = {
        "Total crashes": 0,
        "Fatal crashes": 0,
        "Fatalities": 0,
        "Serious injury crashes": 0,
        "Serious injuries": 0,
    }
    if crashes is None or getattr(crashes, "empty", True):
        return vals
    vals["Total crashes"] = int(len(crashes))
    kabco = _kabco_col(crashes)
    if kabco and kabco in crashes.columns:
        k = crashes[kabco].fillna("").astype(str).str.upper().str.strip()
        vals["Fatal crashes"] = int((k == "K").sum())
        vals["Serious injury crashes"] = int((k == "A").sum())
    fatal_cols = [c for c in crashes.columns if any(w in str(c).lower() for w in ["fatalit", "fatalities", "fatal_count", "killed"])]
    serious_cols = [c for c in crashes.columns if any(w in str(c).lower() for w in ["serious", "suspected_serious", "a_inj", "incapac"])]
    if fatal_cols:
        vals["Fatalities"] = int(pd.to_numeric(crashes[fatal_cols[0]], errors="coerce").fillna(0).sum())
    else:
        vals["Fatalities"] = vals["Fatal crashes"]
    if serious_cols:
        vals["Serious injuries"] = int(pd.to_numeric(crashes[serious_cols[0]], errors="coerce").fillna(0).sum())
    else:
        vals["Serious injuries"] = vals["Serious injury crashes"]
    return vals


def _render_kpi_strip(st, crashes):
    vals = _summary_kpi_values(crashes)
    st.markdown("<div class='dashboard-section-title'>Crash summary <span>selected study period and filters</span></div>", unsafe_allow_html=True)
    c1, c2, c3, c4, c5 = st.columns(5)
    icons = ["🚗", "🛑", "💔", "⚠️", "🏥"]
    for col, (label, value), icon in zip([c1, c2, c3, c4, c5], vals.items(), icons):
        with col:
            st.markdown(
                f"""
                <div style='border:1px solid #e5e7eb;border-radius:14px;padding:13px 14px;background:#ffffff;box-shadow:0 1px 4px rgba(15,23,42,.05)'>
                  <div style='font-size:1.45rem'>{icon}</div>
                  <div style='font-size:.80rem;color:#64748b'>{html.escape(label)}</div>
                  <div style='font-size:1.65rem;font-weight:700;color:#0f172a'>{int(value):,}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )


def _render_hin_network_summary(st, hin, crashes):
    if hin is None or getattr(hin, "empty", True):
        return
    metric = _default_metric(_numeric_cols(hin))
    if not metric:
        return
    st.markdown("<div class='dashboard-section-title'>High Injury Network summary <span>share of miles and crashes captured by selected high-risk network</span></div>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1.2, 1, 1])
    method = c1.selectbox("High-risk network threshold", ["Top 10% of miles", "Top 5% of miles", "HIN index >= 75", "HIN index >= 50", "Top 20 segments/windows"], key="hin_summary_threshold")
    work = hin.copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    length_col = _normal_col(work, ["Length_Miles", "Length_Mi", "SegmentLength_Mile", "WindowLength_Miles", "length_mi"])
    if length_col:
        work[length_col] = pd.to_numeric(work[length_col], errors="coerce").fillna(0)
    else:
        length_col = "__unit_length__"
        work[length_col] = 1.0
    total_mi = float(work[length_col].sum()) if work[length_col].sum() else float(len(work))
    sorted_work = work.sort_values(metric, ascending=False).copy()
    if method == "Top 10% of miles":
        limit = total_mi * 0.10
        selected = sorted_work[sorted_work[length_col].cumsum() <= limit].copy()
        if selected.empty and not sorted_work.empty:
            selected = sorted_work.head(1).copy()
    elif method == "Top 5% of miles":
        limit = total_mi * 0.05
        selected = sorted_work[sorted_work[length_col].cumsum() <= limit].copy()
        if selected.empty and not sorted_work.empty:
            selected = sorted_work.head(1).copy()
    elif method == "HIN index >= 75":
        selected = sorted_work[sorted_work[metric] >= 75].copy()
    elif method == "HIN index >= 50":
        selected = sorted_work[sorted_work[metric] >= 50].copy()
    else:
        selected = sorted_work.head(20).copy()
    high_mi = float(selected[length_col].sum()) if not selected.empty else 0.0
    pct_mi = high_mi / total_mi * 100 if total_mi else 0.0
    count_col = _crash_count_col(selected)
    high_crashes = int(pd.to_numeric(selected[count_col], errors="coerce").fillna(0).sum()) if count_col else 0
    total_crashes = int(pd.to_numeric(work[_crash_count_col(work)], errors="coerce").fillna(0).sum()) if _crash_count_col(work) else (len(crashes) if crashes is not None else 0)
    pct_crash = high_crashes / total_crashes * 100 if total_crashes else 0.0
    c2.metric("High-risk miles", f"{high_mi:,.2f} mi", f"{pct_mi:,.1f}% of analyzed miles")
    c3.metric("Crashes on selected HIN", f"{high_crashes:,}", f"{pct_crash:,.1f}% of assigned crashes")
    st.caption("Select a high-risk threshold such as top miles or HIN index. The summary uses available HIN length and crash-count fields; severity/mode percentages are shown when those fields exist in the HIN results.")


def _render_pattern_charts(st, tables):
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))

    _render_kpi_strip(st, crashes)

    st.markdown("<div class='dashboard-section-title'>Crash patterns <span>years, crash type, month, KABCO, mode, and roadway context</span></div>", unsafe_allow_html=True)
    left, right = st.columns(2)
    with left:
        if crashes is not None:
            year_kabco, year_col, kabco_col = _year_kabco_table(crashes)
            if not year_kabco.empty:
                year_kabco[kabco_col] = pd.Categorical(year_kabco[kabco_col].astype(str), categories=_kabco_order_values(year_kabco[kabco_col]), ordered=True)
                fig = px.bar(year_kabco.sort_values([year_col, kabco_col]), x=year_col, y="Count", color=kabco_col, color_discrete_map=KABCO_COLOR_MAP, title=f"Crashes by year and {kabco_col}", hover_data={"Count": True, kabco_col: True})
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), barmode="stack", xaxis_title="Year", yaxis_title="Crash count")
                st.plotly_chart(_apply_kabco_trace_colors(_polish_figure(fig)), width="stretch")
            else:
                year_col = _find_col(crashes, ["year"])
                if year_col:
                    year_df = _aggregate(crashes, year_col, None, "Count", 20)
                    fig = px.bar(year_df.sort_values(year_col), x=year_col, y="Count", title="Crashes by year")
                    fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), xaxis_title="Year", yaxis_title="Crash count")
                    st.plotly_chart(_polish_figure(fig), width="stretch")
        else:
            st.info("No crash table is available yet.")
    with right:
        if crashes is not None:
            type_col = _crash_type_col(crashes)
            if type_col:
                type_df = _aggregate(crashes, type_col, None, "Count", 10)
                fig = px.pie(type_df, names=type_col, values="Count", hole=0.38, title=f"Crash type share by {type_col}")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), legend=dict(orientation="v"))
                st.plotly_chart(_polish_figure(fig), width="stretch")

    left2, right2 = st.columns(2)
    with left2:
        monthly_df, period_col, value_col, color_col = _month_trend_table(crashes)
        if not monthly_df.empty:
            fig = px.line(monthly_df, x=period_col, y=value_col, color=color_col, markers=True, category_orders={period_col: MONTH_ORDER}, title="Monthly crash trend by year")
            fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), xaxis_title="Month", yaxis_title="Crash count")
            st.plotly_chart(_polish_figure(fig), width="stretch")
    with right2:
        mode_df, mode_kabco_col, mode_col = _mode_severity_bubble_table(crashes)
        if not mode_df.empty:
            fig = px.scatter(mode_df, x=mode_kabco_col, y=mode_col, size="Count", color=mode_col, size_max=48, title="Travel mode severity bubble chart", hover_data=["Count"])
            fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), xaxis_title="KABCO", yaxis_title="Mode")
            st.plotly_chart(_polish_figure(fig), width="stretch")

    road_kabco, road_col, road_kabco_col = _road_class_kabco_table(crashes, st_obj=st)
    if not road_kabco.empty:
        pivot = road_kabco.pivot_table(index=road_col, columns=road_kabco_col, values="Count", aggfunc="sum", fill_value=0)
        order = [c for c in ["K", "A", "B", "C", "O"] if c in pivot.columns]
        pivot = pivot[order + [c for c in pivot.columns if c not in order]]
        fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Road class by {road_kabco_col}", labels=dict(x=road_kabco_col, y="Road class", color="Crash count"), color_continuous_scale="YlOrRd")
        fig.update_layout(height=330, margin=dict(l=20, r=20, t=45, b=35))
        st.plotly_chart(_polish_figure(fig), width="stretch")

    crash_kabco, crash_type_col, crash_kabco_col = _crash_type_kabco_table(crashes)
    if not crash_kabco.empty:
        left3, right3 = st.columns(2)
        with left3:
            pivot = crash_kabco.pivot_table(index=crash_type_col, columns=crash_kabco_col, values="Count", aggfunc="sum", fill_value=0)
            order = [c for c in ["K", "A", "B", "C", "O"] if c in pivot.columns]
            pivot = pivot[order + [c for c in pivot.columns if c not in order]]
            fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Crash type by {crash_kabco_col}", labels=dict(x=crash_kabco_col, y="Crash type", color="Crash count"), color_continuous_scale="YlOrRd")
            fig.update_layout(height=420, margin=dict(l=20, r=20, t=45, b=35))
            st.plotly_chart(_polish_figure(fig), width="stretch")
        with right3:
            tree = crash_kabco.copy()
            tree["All crashes"] = "All crashes"
            fig = px.treemap(tree, path=["All crashes", crash_type_col, crash_kabco_col], values="Count", color=crash_kabco_col, color_discrete_map=KABCO_COLOR_MAP, title=f"Crash type and {crash_kabco_col} treemap")
            fig.update_layout(height=420, margin=dict(l=20, r=20, t=45, b=35))
            st.plotly_chart(_polish_figure(fig), width="stretch")

    st.markdown("<div class='dashboard-section-title'>Risk and spatial-unit ranking <span>crash density, crash count, and HIN priority ranking</span></div>", unsafe_allow_html=True)
    left, right = st.columns(2)
    with left:
        if density is not None:
            metric = "CrashDensity" if "CrashDensity" in density.columns else _default_metric(_numeric_cols(density))
            rank_df, unit_col, value_col = _rank_units_for_chart(density, metric, 15) if metric else (pd.DataFrame(), None, None)
            if not rank_df.empty:
                plot_df = rank_df.sort_values(value_col, ascending=True)
                fig = px.bar(plot_df, y=unit_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title="Top spatial units by crash density")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), yaxis_title="Spatial unit ID", xaxis_title="Crash density")
                st.plotly_chart(_polish_figure(fig), width="stretch")
            else:
                st.info("Run crash-density analysis to rank spatial units.")
        else:
            st.info("Crash-density results are not available yet. Run crash-density analysis first.")
    with right:
        if density is not None:
            count_col = _crash_count_col(density)
            rank_df, unit_col, value_col = _rank_units_for_chart(density, count_col, 15) if count_col else (pd.DataFrame(), None, None)
            if not rank_df.empty:
                plot_df = rank_df.sort_values(value_col, ascending=True)
                fig = px.bar(plot_df, y=unit_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title="Top spatial units by crash count")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), yaxis_title="Spatial unit ID", xaxis_title="Crash count")
                st.plotly_chart(_polish_figure(fig), width="stretch")

    if hin is not None:
        st.markdown("<div class='dashboard-section-title'>HIN priority ranking <span>each segment/window is one bar; hover for route and milepost context</span></div>", unsafe_allow_html=True)
        metric = _default_metric(_numeric_cols(hin))
        rank_df, seg_col, value_col = _hin_segment_rank_for_chart(hin, metric, 20) if metric else (pd.DataFrame(), None, None)
        if not rank_df.empty:
            plot_df = rank_df.sort_values(value_col, ascending=True)
            fig = px.bar(plot_df, y=seg_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title=f"Top HIN segments/windows by {_nice_metric_label(metric)}")
            fig.update_layout(height=420, margin=dict(l=20, r=20, t=45, b=35), yaxis_title="Segment / window ID", xaxis_title=_nice_metric_label(metric))
            st.plotly_chart(_polish_figure(fig), width="stretch")
        _render_hin_network_summary(st, hin, crashes)
    st.caption("Scroll down to view all charts, maps, and tables. Use the Dashboard Builder tab to select charts/maps for export.")


def _add_overlay_layer(fmap, gdf, label, show=False):
    """Add compact overlay layers. Signals use an icon marker instead of plain dots."""
    if gdf is None or getattr(gdf, "empty", True) or folium is None:
        return
    try:
        layer = gdf.copy()
        if layer.crs is None:
            layer = layer.set_crs(epsg=4326)
        layer = layer.to_crs(epsg=4326)
        style = _overlay_style(label)
        group = folium.FeatureGroup(name=label, show=show)
        is_point = all(getattr(geom, "geom_type", "") == "Point" for geom in layer.geometry.dropna().head(50))
        if is_point:
            for _, row in layer.iterrows():
                geom = row.geometry
                if geom is None or geom.is_empty:
                    continue
                popup_cols = [c for c in ["CrashID", "SignalID", "CorridorID", "Route", "FULLNAME"] if c in layer.columns]
                popup = "<br>".join([f"{c}: {row.get(c)}" for c in popup_cols]) if popup_cols else label
                if "signal" in label.lower():
                    folium.Marker(
                        location=[geom.y, geom.x],
                        popup=popup,
                        icon=folium.Icon(color="green", icon="traffic-light", prefix="fa"),
                    ).add_to(group)
                else:
                    folium.CircleMarker(
                        location=[geom.y, geom.x], radius=style.get("radius", 4),
                        color=style.get("color", "#111827"), weight=style.get("weight", 1), fill=True,
                        fill_color=style.get("fillColor", style.get("color", "#111827")), fill_opacity=style.get("fillOpacity", 0.8), popup=popup,
                    ).add_to(group)
        else:
            gj = _safe_geojson_string(layer)
            if gj:
                folium.GeoJson(gj, name=label, style_function=lambda feature, stl=style: stl).add_to(group)
        group.add_to(fmap)
    except Exception:
        return


def _render_dashboard_map(st, map_name, gdf, key, highlight_value=None, height=460, overlay_layers=None):
    if folium is None or st_folium is None:
        st.info("Install folium and streamlit-folium to show dashboard maps.")
        return
    if gdf is None or getattr(gdf, "empty", True):
        st.info("No map data is available for this layer.")
        return
    work = _repair_gdf_crs(gdf, st).to_crs(epsg=4326)
    bounds = _fit_bounds_for_layer(work)
    if bounds is None:
        st.info("Map bounds could not be calculated for this layer.")
        return
    minx, miny, maxx, maxy = bounds
    center = [(miny + maxy) / 2, (minx + maxx) / 2]
    fmap = folium.Map(location=center, zoom_start=13, tiles="cartodbpositron")
    metric = _metric_for_map(work, map_name)
    values = pd.to_numeric(work[metric], errors="coerce") if metric else pd.Series([], dtype=float)
    min_value = values.min() if not values.empty else 0
    max_value = values.max() if not values.empty else 1
    unit_col = _unit_col(_drop_geometry(work))
    tooltip_cols = [c for c in [unit_col, metric, "CrashCount", "Crash_Count", "Rank", "Route", "FULLNAME", "FromMile", "ToMile"] if c and c in work.columns]
    def style_fn(feature):
        props = feature.get("properties", {})
        val = props.get(metric) if metric else None
        highlight = False
        if highlight_value is not None and unit_col and props.get(unit_col) is not None:
            highlight = str(props.get(unit_col)) == str(highlight_value)
        return _style_feature(val, min_value, max_value, highlight=highlight)
    tooltip = GeoJsonTooltip(fields=tooltip_cols, aliases=tooltip_cols) if GeoJsonTooltip and tooltip_cols else None
    gj = _safe_geojson_string(work)
    if not gj:
        st.info("Map geometry could not be converted to displayable GeoJSON.")
        return
    folium.GeoJson(gj, name=map_name, style_function=style_fn, tooltip=tooltip, show=True).add_to(fmap)
    if metric:
        _add_map_legend(fmap, _nice_metric_label(metric), min_value, max_value)
    for overlay_name, overlay_gdf in (overlay_layers or {}).items():
        _add_overlay_layer(fmap, overlay_gdf, overlay_name, show=False)
    folium.LayerControl(collapsed=True).add_to(fmap)
    try:
        fmap.get_root().html.add_child(folium.Element("""
        <style>
        .leaflet-control-layers {font-size:11px; max-height:170px; overflow:auto;}
        .leaflet-control-layers-expanded {padding:6px 8px;}
        .leaflet-control-layers label {margin-bottom:2px;}
        </style>
        """))
        fmap.fit_bounds([[miny, minx], [maxy, maxx]], padding=(24, 24))
    except Exception:
        pass
    map_key = f"{key}_{round(minx,5)}_{round(miny,5)}_{round(maxx,5)}_{round(maxy,5)}"
    st_folium(fmap, height=height, width="100%", key=map_key, returned_objects=[])


def _build_default_figures(tables):
    """Default export figures aligned to the V28 dashboard charts."""
    figures = []
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))
    if crashes is not None:
        year_kabco, year_col, kabco_col = _year_kabco_table(crashes)
        if not year_kabco.empty:
            fig = px.bar(year_kabco, x=year_col, y="Count", color=kabco_col, color_discrete_map=KABCO_COLOR_MAP, title=f"Crashes by year and {kabco_col}")
            fig.update_layout(barmode="stack", xaxis_title="Year", yaxis_title="Crash count")
            figures.append((f"Crashes by year and {kabco_col}", _apply_kabco_trace_colors(_polish_figure(fig)), year_kabco))
        type_col = _crash_type_col(crashes)
        if type_col:
            type_df = _aggregate(crashes, type_col, None, "Count", 12)
            pie = px.pie(type_df, names=type_col, values="Count", hole=0.38, title=f"Crash type share by {type_col}")
            figures.append((f"Crash type share by {type_col}", _polish_figure(pie), type_df))
        monthly_df, period_col, value_col, color_col = _month_trend_table(crashes)
        if not monthly_df.empty:
            fig = px.line(monthly_df, x=period_col, y=value_col, color=color_col, markers=True, category_orders={period_col: MONTH_ORDER}, title="Monthly crash trend by year")
            fig.update_layout(xaxis_title="Month", yaxis_title="Crash count")
            figures.append(("Monthly crash trend by year", _polish_figure(fig), monthly_df))
        crash_kabco, crash_type_col, crash_kabco_col = _crash_type_kabco_table(crashes)
        if not crash_kabco.empty:
            pivot = crash_kabco.pivot_table(index=crash_type_col, columns=crash_kabco_col, values="Count", aggfunc="sum", fill_value=0)
            fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Crash type by {crash_kabco_col}", labels=dict(x=crash_kabco_col, y="Crash type", color="Crash count"), color_continuous_scale="YlOrRd")
            figures.append((f"Crash type by {crash_kabco_col}", _polish_figure(fig), pivot.reset_index()))
    if density is not None:
        metric = "CrashDensity" if "CrashDensity" in density.columns else _default_metric(_numeric_cols(density))
        rank_df, unit_col, value_col = _rank_units_for_chart(density, metric, 15) if metric else (pd.DataFrame(), None, None)
        if not rank_df.empty:
            plot_df = rank_df.sort_values(value_col, ascending=True)
            fig = px.bar(plot_df, y=unit_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title="Top spatial units by crash density")
            fig.update_layout(yaxis_title="Spatial unit ID", xaxis_title="Crash density")
            figures.append(("Top spatial units by crash density", _polish_figure(fig), rank_df))
    if hin is not None:
        metric = _default_metric(_numeric_cols(hin))
        rank_df, seg_col, value_col = _hin_segment_rank_for_chart(hin, metric, 20) if metric else (pd.DataFrame(), None, None)
        if not rank_df.empty:
            plot_df = rank_df.sort_values(value_col, ascending=True)
            fig = px.bar(plot_df, y=seg_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title=f"Top HIN segments/windows by {_nice_metric_label(metric)}")
            fig.update_layout(yaxis_title="Segment / window ID", xaxis_title=_nice_metric_label(metric))
            figures.append((f"Top HIN segments/windows by {_nice_metric_label(metric)}", _polish_figure(fig), rank_df))
    return figures

# --- V29 dashboard severity, HIN-table, KPI-builder, and report cleanup overrides ---

SEVERITY_COLUMN_PATTERNS = {
    "K": ["fatals", "fatalit", "fatalities", "fatal", "fatal injury", "fatal_crash", "killed", "death"],
    "A": ["level a", "serious", "incapac", "suspected serious", "a inj", "a_inj"],
    "B": ["level b", "non-incap", "non incapac", "evident", "minor", "b inj", "b_inj"],
    "C": ["level c", "possible", "complaint", "c inj", "c_inj"],
    "O": ["uninjured", "uninj", "no injury", "pdo", "property damage", "o inj", "o_inj"],
}


def _severity_count_columns(df):
    """Find person-count severity columns such as Fatalities, Level A Injuries, etc."""
    found = {}
    if df is None:
        return found
    for code, patterns in SEVERITY_COLUMN_PATTERNS.items():
        for col in df.columns:
            lower = str(col).lower().replace("_", " ")
            if any(p in lower for p in patterns):
                vals = pd.to_numeric(df[col], errors="coerce")
                if vals.notna().any():
                    found[code] = col
                    break
    return found


def _normalize_kabco_value(value):
    """Normalize common severity text into K/A/B/C/O.

    Important order: B must be checked before A because phrases such as
    "Evident Non-Incapacitating (B)" contain the substring "incapacitating".
    If A is checked first, B crashes are incorrectly counted as A and vanish
    from the stacked year/KABCO chart.
    """
    if pd.isna(value):
        return "Unknown"
    text = str(value).strip()
    upper = text.upper().strip()
    if upper in ["K", "A", "B", "C", "O"]:
        return upper
    lower = text.lower().replace("_", " ").replace("-", " ")

    # Direct code-in-label checks first.
    if "(k)" in lower or lower.startswith("k "):
        return "K"
    if "(a)" in lower or "level a" in lower or lower.startswith("a "):
        return "A"
    if "(b)" in lower or "level b" in lower or lower.startswith("b "):
        return "B"
    if "(c)" in lower or "level c" in lower or lower.startswith("c "):
        return "C"
    if "(o)" in lower or lower.startswith("o "):
        return "O"

    # B before A because non-incapacitating contains incapacitating.
    if "non incapac" in lower or "nonincap" in lower or "evident non" in lower or "minor" in lower or "level b" in lower:
        return "B"
    if "fatal" in lower or "killed" in lower or "death" in lower:
        return "K"
    if "serious" in lower or "incapac" in lower or "level a" in lower:
        return "A"
    if "possible" in lower or "complaint" in lower or "level c" in lower:
        return "C"
    if "pdo" in lower or "no injury" in lower or "uninj" in lower or "property" in lower:
        return "O"
    return text if text else "Unknown"


def _year_series_from_crashes(df):
    if df is None or getattr(df, "empty", True):
        return None
    year_col = _find_col(df, ["year"])
    if year_col:
        vals = pd.to_numeric(df[year_col], errors="coerce")
        if vals.notna().any():
            return vals.astype("Int64").astype(str).replace("<NA>", "Unknown")
        return df[year_col].fillna("Unknown").astype(str)
    date_col = _time_col(df)
    if date_col:
        dt = pd.to_datetime(df[date_col], errors="coerce")
        if dt.notna().any():
            return dt.dt.year.astype("Int64").astype(str).replace("<NA>", "Unknown")
    return None


def _year_kabco_table(crashes):
    """Return year x normalized-KABCO counts.

    If a KABCO field exists, rows are crash counts by crash severity. If the dataset
    instead has person-count columns such as Fatalities / Level A Injuries /
    Uninjured, counts are summed from those columns so other datasets still work.
    """
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), None, None
    years = _year_series_from_crashes(crashes)
    if years is None:
        return pd.DataFrame(), None, None
    kabco_col = _kabco_col(crashes)
    sev_cols = _severity_count_columns(crashes)
    if kabco_col and kabco_col in crashes.columns:
        work = pd.DataFrame({"Year": years, "KABCO": crashes[kabco_col].map(_normalize_kabco_value)})
        work = work[work["Year"].ne("Unknown")]
        if work.empty:
            return pd.DataFrame(), None, None
        out = work.groupby(["Year", "KABCO"], dropna=False).size().reset_index(name="Count")
    elif sev_cols:
        parts = []
        for code, col in sev_cols.items():
            vals = pd.to_numeric(crashes[col], errors="coerce").fillna(0)
            tmp = pd.DataFrame({"Year": years, "KABCO": code, "Count": vals})
            tmp = tmp[tmp["Year"].ne("Unknown") & (tmp["Count"] > 0)]
            if not tmp.empty:
                parts.append(tmp)
        if not parts:
            return pd.DataFrame(), None, None
        out = pd.concat(parts, ignore_index=True).groupby(["Year", "KABCO"], dropna=False)["Count"].sum().reset_index()
    else:
        return pd.DataFrame(), None, None
    out = _order_kabco(out, "KABCO")
    try:
        out["__year_sort__"] = pd.to_numeric(out["Year"], errors="coerce")
        out["__kabco_sort__"] = out["KABCO"].map({"K": 1, "A": 2, "B": 3, "C": 4, "O": 5}).fillna(9)
        out = out.sort_values(["__year_sort__", "__kabco_sort__"]).drop(columns=["__year_sort__", "__kabco_sort__"])
    except Exception:
        pass
    return out, "Year", "KABCO"


def _kabco_col(df):
    """Find a true KABCO-like crash severity field, avoiding person-count columns."""
    if df is None:
        return None
    for col in df.columns:
        lower = str(col).lower().replace("_", "")
        if lower == "kabco" or "kabco" in lower:
            return col
    candidates = []
    for col in df.columns:
        lower = str(col).lower()
        if any(bad in lower for bad in ["fatalities", "injuries", "injury count", "persons", "uninjured"]):
            continue
        if "severity" in lower or "injury" in lower:
            candidates.append(col)
    for col in candidates:
        vals = df[col].dropna().astype(str).head(200).map(_normalize_kabco_value).str.upper().unique().tolist()
        if any(v in ["K", "A", "B", "C", "O"] for v in vals):
            return col
    return None


def _summary_kpi_values(crashes):
    vals = {
        "Total crashes": 0,
        "Fatal crashes": 0,
        "Fatalities": 0,
        "Serious injury crashes": 0,
        "Serious injuries": 0,
    }
    if crashes is None or getattr(crashes, "empty", True):
        return vals
    vals["Total crashes"] = int(len(crashes))
    kabco = _kabco_col(crashes)
    sev_cols = _severity_count_columns(crashes)
    if kabco and kabco in crashes.columns:
        k = crashes[kabco].map(_normalize_kabco_value).astype(str).str.upper().str.strip()
        vals["Fatal crashes"] = int((k == "K").sum())
        vals["Serious injury crashes"] = int((k == "A").sum())
    if "K" in sev_cols:
        fatal_vals = pd.to_numeric(crashes[sev_cols["K"]], errors="coerce").fillna(0)
        vals["Fatalities"] = int(fatal_vals.sum())
        vals["Fatal crashes"] = int((fatal_vals > 0).sum()) if vals["Fatal crashes"] == 0 else vals["Fatal crashes"]
    else:
        vals["Fatalities"] = vals["Fatal crashes"]
    if "A" in sev_cols:
        serious_vals = pd.to_numeric(crashes[sev_cols["A"]], errors="coerce").fillna(0)
        vals["Serious injuries"] = int(serious_vals.sum())
        vals["Serious injury crashes"] = int((serious_vals > 0).sum()) if vals["Serious injury crashes"] == 0 else vals["Serious injury crashes"]
    else:
        vals["Serious injuries"] = vals["Serious injury crashes"]
    return vals


def _context_cols_for_hover(df):
    cols = []
    if df is None:
        return cols
    has_length = any(c in df.columns for c in ["Length_Miles", "Length_Mi", "WindowLength_Miles"])
    for c in [
        "Route", "FULLNAME", "RoadName", "RouteName", "CorridorRoute",
        "FromMile", "ToMile", "From_Mile", "To_Mile", "from_mile", "to_mile",
        "Length_Miles", "Length_Mi", "WindowLength_Miles",
        "CrashCount", "Crash_Count", "CrashDensity", "HIN_Priority_Index",
    ]:
        if c in df.columns and c not in cols:
            cols.append(c)
    if not has_length and "SegmentLength_Mile" in df.columns:
        cols.append("SegmentLength_Mile")
    return cols


def _hin_table_for_display(hin, metric, top_n=20):
    if hin is None or getattr(hin, "empty", True) or metric not in hin.columns:
        return pd.DataFrame()
    work = hin.copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work = work.sort_values(metric, ascending=False).head(top_n).reset_index(drop=True)
    id_col = _normal_col(work, ["RiskSegmentID", "WindowID", "SegmentID", "UnitID", "SourceSegmentID"])
    route_col = _normal_col(work, ["Route", "FULLNAME", "RoadName", "RouteName", "CorridorRoute", "Name"])
    length_col = _normal_col(work, ["Length_Miles", "Length_Mi", "WindowLength_Miles", "SegmentLength_Mile", "length_mi"])
    from_col = _normal_col(work, ["FromMile", "From_Mile", "from_mile", "BeginMile", "StartMile"])
    to_col = _normal_col(work, ["ToMile", "To_Mile", "to_mile", "EndMile"])
    out = pd.DataFrame()
    out["Rank"] = range(1, len(work) + 1)
    out["SegID"] = work[id_col].astype(str).values if id_col else [f"HIN_{i+1}" for i in range(len(work))]
    out["Seg length"] = pd.to_numeric(work[length_col], errors="coerce").round(3).values if length_col else ""
    out["Seg from mile"] = pd.to_numeric(work[from_col], errors="coerce").round(3).values if from_col else ""
    out["Seg to mile"] = pd.to_numeric(work[to_col], errors="coerce").round(3).values if to_col else ""
    out["Route"] = work[route_col].astype(str).values if route_col else ""
    # Approximate route context from the available HIN rows. If route data are not present, leave blank.
    if route_col and route_col in work.columns:
        full = hin.copy()
        if length_col and length_col in full.columns:
            full[length_col] = pd.to_numeric(full[length_col], errors="coerce").fillna(0)
            route_len = full.groupby(route_col, dropna=False)[length_col].sum().to_dict()
        else:
            route_len = {}
        if from_col and to_col and from_col in full.columns and to_col in full.columns:
            full[from_col] = pd.to_numeric(full[from_col], errors="coerce")
            full[to_col] = pd.to_numeric(full[to_col], errors="coerce")
            route_from = full.groupby(route_col, dropna=False)[from_col].min().to_dict()
            route_to = full.groupby(route_col, dropna=False)[to_col].max().to_dict()
        else:
            route_from, route_to = {}, {}
        out["Route length"] = [round(float(route_len.get(v, 0)), 3) if v in route_len else "" for v in work[route_col]]
        out["Route From mile"] = [round(float(route_from.get(v)), 3) if v in route_from and pd.notna(route_from.get(v)) else "" for v in work[route_col]]
        out["Route To mile"] = [round(float(route_to.get(v)), 3) if v in route_to and pd.notna(route_to.get(v)) else "" for v in work[route_col]]
    else:
        out["Route length"] = ""
        out["Route From mile"] = ""
        out["Route To mile"] = ""
    out["HIN index"] = pd.to_numeric(work[metric], errors="coerce").round(3).values
    return out


def _render_pattern_charts(st, tables):
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))

    _render_kpi_strip(st, crashes)

    st.markdown("<div class='dashboard-section-title'>Crash patterns <span>years, crash type, month, KABCO, mode, and roadway context</span></div>", unsafe_allow_html=True)
    left, right = st.columns(2)
    with left:
        year_kabco, year_col, kabco_col = _year_kabco_table(crashes)
        if not year_kabco.empty:
            fig = px.bar(year_kabco, x=year_col, y="Count", color=kabco_col, color_discrete_map=KABCO_COLOR_MAP, category_orders={kabco_col: ["K", "A", "B", "C", "O"]}, title=f"Crashes by year and {kabco_col}", hover_data={"Count": True, kabco_col: True})
            fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), barmode="stack", xaxis_title="Year", yaxis_title="Crash count")
            st.plotly_chart(_apply_kabco_trace_colors(_polish_figure(fig)), width="stretch")
        elif crashes is not None:
            years = _year_series_from_crashes(crashes)
            if years is not None:
                year_df = pd.DataFrame({"Year": years}).query("Year != 'Unknown'").groupby("Year").size().reset_index(name="Count")
                fig = px.bar(year_df.sort_values("Year"), x="Year", y="Count", title="Crashes by year")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), xaxis_title="Year", yaxis_title="Crash count")
                st.plotly_chart(_polish_figure(fig), width="stretch")
    with right:
        if crashes is not None:
            type_col = _crash_type_col(crashes)
            if type_col:
                type_df = _aggregate(crashes, type_col, None, "Count", 10)
                fig = px.pie(type_df, names=type_col, values="Count", hole=0.38, title=f"Crash type share by {type_col}")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), legend=dict(orientation="v"))
                st.plotly_chart(_polish_figure(fig), width="stretch")

    left2, right2 = st.columns(2)
    with left2:
        monthly_df, period_col, value_col, color_col = _month_trend_table(crashes)
        if not monthly_df.empty:
            fig = px.line(monthly_df, x=period_col, y=value_col, color=color_col, markers=True, category_orders={period_col: MONTH_ORDER}, title="Monthly crash trend by year")
            fig.update_layout(height=340, margin=dict(l=20, r=20, t=45, b=35), xaxis_title="Month", yaxis_title="Crash count")
            st.plotly_chart(_polish_figure(fig), width="stretch")
    with right2:
        mode_df, mode_kabco_col, mode_col = _mode_severity_bubble_table(crashes)
        if not mode_df.empty:
            fig = px.scatter(mode_df, x=mode_kabco_col, y=mode_col, size="Count", color=mode_col, size_max=56, title="Travel mode severity bubble chart", hover_data=["Count"])
            fig.update_traces(marker=dict(sizemin=5, opacity=0.78, line=dict(width=1, color="white")))
            fig.update_layout(height=360, margin=dict(l=20, r=20, t=45, b=35), xaxis_title="KABCO", yaxis_title="Mode")
            st.plotly_chart(_polish_figure(fig), width="stretch")

    road_kabco, road_col, road_kabco_col = _road_class_kabco_table(crashes, st_obj=st)
    if not road_kabco.empty:
        pivot = road_kabco.pivot_table(index=road_col, columns=road_kabco_col, values="Count", aggfunc="sum", fill_value=0)
        order = [c for c in ["K", "A", "B", "C", "O"] if c in pivot.columns]
        pivot = pivot[order + [c for c in pivot.columns if c not in order]]
        fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Road class by {road_kabco_col}", labels=dict(x=road_kabco_col, y="Road class", color="Crash count"), color_continuous_scale="YlOrRd")
        fig.update_layout(height=330, margin=dict(l=20, r=20, t=45, b=35))
        st.plotly_chart(_polish_figure(fig), width="stretch")

    crash_kabco, crash_type_col, crash_kabco_col = _crash_type_kabco_table(crashes)
    if not crash_kabco.empty:
        left3, right3 = st.columns(2)
        with left3:
            pivot = crash_kabco.pivot_table(index=crash_type_col, columns=crash_kabco_col, values="Count", aggfunc="sum", fill_value=0)
            order = [c for c in ["K", "A", "B", "C", "O"] if c in pivot.columns]
            pivot = pivot[order + [c for c in pivot.columns if c not in order]]
            fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Crash type by {crash_kabco_col}", labels=dict(x=crash_kabco_col, y="Crash type", color="Crash count"), color_continuous_scale="YlOrRd")
            fig.update_layout(height=420, margin=dict(l=20, r=20, t=45, b=35))
            st.plotly_chart(_polish_figure(fig), width="stretch")
        with right3:
            tree = crash_kabco.copy()
            tree["All crashes"] = "All crashes"
            fig = px.treemap(tree, path=["All crashes", crash_type_col, crash_kabco_col], values="Count", color=crash_kabco_col, color_discrete_map=KABCO_COLOR_MAP, title=f"Crash type and {crash_kabco_col} treemap")
            fig.update_layout(height=420, margin=dict(l=20, r=20, t=45, b=35))
            st.plotly_chart(_polish_figure(fig), width="stretch")

    st.markdown("<div class='dashboard-section-title'>Risk and spatial-unit ranking <span>crash density, crash count, and HIN priority ranking</span></div>", unsafe_allow_html=True)
    left, right = st.columns(2)
    with left:
        if density is not None:
            metric = "CrashDensity" if "CrashDensity" in density.columns else _default_metric(_numeric_cols(density))
            rank_df, unit_col, value_col = _rank_units_for_chart(density, metric, 15) if metric else (pd.DataFrame(), None, None)
            if not rank_df.empty:
                plot_df = rank_df.sort_values(value_col, ascending=True)
                fig = px.bar(plot_df, y=unit_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title="Top spatial units by crash density")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), yaxis_title="Spatial unit ID", xaxis_title="Crash density")
                st.plotly_chart(_polish_figure(fig), width="stretch")
    with right:
        if density is not None:
            count_col = _crash_count_col(density)
            rank_df, unit_col, value_col = _rank_units_for_chart(density, count_col, 15) if count_col else (pd.DataFrame(), None, None)
            if not rank_df.empty:
                plot_df = rank_df.sort_values(value_col, ascending=True)
                fig = px.bar(plot_df, y=unit_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title="Top spatial units by crash count")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), yaxis_title="Spatial unit ID", xaxis_title="Crash count")
                st.plotly_chart(_polish_figure(fig), width="stretch")

    if hin is not None:
        st.markdown("<div class='dashboard-section-title'>HIN priority ranking <span>table view with route and milepost context</span></div>", unsafe_allow_html=True)
        metric = "HIN_Priority_Index" if "HIN_Priority_Index" in hin.columns else _default_metric(_numeric_cols(hin))
        hin_table = _hin_table_for_display(hin, metric, 20) if metric else pd.DataFrame()
        if not hin_table.empty:
            st.dataframe(_safe_dataframe_for_display(hin_table), width="stretch", hide_index=True)
        _render_hin_network_summary(st, hin, crashes)
    st.caption("The green arrow in the HIN summary cards is a positive delta indicator. It shows the selected high-risk network's share of analyzed miles and assigned crashes, not a change from a previous year.")


def _build_kpi_figure(tables):
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    vals = _summary_kpi_values(crashes)
    df = pd.DataFrame({"Metric": list(vals.keys()), "Value": list(vals.values())})
    fig = px.bar(df, x="Metric", y="Value", text="Value", title="Crash summary KPI cards")
    fig.update_traces(marker_color="#2563eb", texttemplate="%{text:,}", textposition="outside")
    fig.update_layout(height=320, margin=dict(l=20, r=20, t=45, b=65), xaxis_title="", yaxis_title="Count")
    return "Crash summary KPI cards", _polish_figure(fig), df


def _build_default_figures(tables):
    figures = []
    figures.append(_build_kpi_figure(tables))
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))
    if crashes is not None:
        year_kabco, year_col, kabco_col = _year_kabco_table(crashes)
        if not year_kabco.empty:
            fig = px.bar(year_kabco, x=year_col, y="Count", color=kabco_col, color_discrete_map=KABCO_COLOR_MAP, category_orders={kabco_col: ["K", "A", "B", "C", "O"]}, title=f"Crashes by year and {kabco_col}")
            fig.update_layout(barmode="stack", xaxis_title="Year", yaxis_title="Crash count")
            figures.append((f"Crashes by year and {kabco_col}", _apply_kabco_trace_colors(_polish_figure(fig)), year_kabco))
        type_col = _crash_type_col(crashes)
        if type_col:
            type_df = _aggregate(crashes, type_col, None, "Count", 12)
            pie = px.pie(type_df, names=type_col, values="Count", hole=0.38, title=f"Crash type share by {type_col}")
            figures.append((f"Crash type share by {type_col}", _polish_figure(pie), type_df))
        monthly_df, period_col, value_col, color_col = _month_trend_table(crashes)
        if not monthly_df.empty:
            fig = px.line(monthly_df, x=period_col, y=value_col, color=color_col, markers=True, category_orders={period_col: MONTH_ORDER}, title="Monthly crash trend by year")
            fig.update_layout(xaxis_title="Month", yaxis_title="Crash count")
            figures.append(("Monthly crash trend by year", _polish_figure(fig), monthly_df))
        mode_df, mode_kabco_col, mode_col = _mode_severity_bubble_table(crashes)
        if not mode_df.empty:
            fig = px.scatter(mode_df, x=mode_kabco_col, y=mode_col, size="Count", color=mode_col, size_max=56, title="Travel mode severity bubble chart", hover_data=["Count"])
            fig.update_traces(marker=dict(sizemin=5, opacity=0.78, line=dict(width=1, color="white")))
            figures.append(("Travel mode severity bubble chart", _polish_figure(fig), mode_df))
        crash_kabco, crash_type_col, crash_kabco_col = _crash_type_kabco_table(crashes)
        if not crash_kabco.empty:
            pivot = crash_kabco.pivot_table(index=crash_type_col, columns=crash_kabco_col, values="Count", aggfunc="sum", fill_value=0)
            fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Crash type by {crash_kabco_col}", labels=dict(x=crash_kabco_col, y="Crash type", color="Crash count"), color_continuous_scale="YlOrRd")
            figures.append((f"Crash type by {crash_kabco_col}", _polish_figure(fig), pivot.reset_index()))
            tree = crash_kabco.copy(); tree["All crashes"] = "All crashes"
            fig = px.treemap(tree, path=["All crashes", crash_type_col, crash_kabco_col], values="Count", color=crash_kabco_col, color_discrete_map=KABCO_COLOR_MAP, title=f"Crash type and {crash_kabco_col} treemap")
            figures.append((f"Crash type and {crash_kabco_col} treemap", _polish_figure(fig), tree))
    if density is not None:
        metric = "CrashDensity" if "CrashDensity" in density.columns else _default_metric(_numeric_cols(density))
        rank_df, unit_col, value_col = _rank_units_for_chart(density, metric, 15) if metric else (pd.DataFrame(), None, None)
        if not rank_df.empty:
            plot_df = rank_df.sort_values(value_col, ascending=True)
            fig = px.bar(plot_df, y=unit_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title="Top spatial units by crash density")
            fig.update_layout(yaxis_title="Spatial unit ID", xaxis_title="Crash density")
            figures.append(("Top spatial units by crash density", _polish_figure(fig), rank_df))
    if hin is not None:
        metric = "HIN_Priority_Index" if "HIN_Priority_Index" in hin.columns else _default_metric(_numeric_cols(hin))
        hin_table = _hin_table_for_display(hin, metric, 20) if metric else pd.DataFrame()
        if not hin_table.empty:
            # Export as a table-like bar placeholder so the block can be selected; Word tables carry the detail.
            fig = px.bar(hin_table.sort_values("HIN index", ascending=True), y="SegID", x="HIN index", orientation="h", title="Top HIN segments/windows by HIN priority index")
            fig.update_layout(yaxis_title="SegID", xaxis_title="HIN priority index")
            figures.append(("Top HIN segments/windows table", _polish_figure(fig), hin_table))
    return figures


def _render_dashboard_builder(st, tables):
    maps = _available_maps(st)
    custom_figures = st.session_state.get("dashboard_custom_figures", [])
    default_figures = _build_default_figures(tables) + custom_figures
    figure_titles = [title for title, _, _ in default_figures]

    st.markdown("<div class='dashboard-section-title'>Dashboard builder <span>choose charts, tables, and map layers for one review page</span></div>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([0.27, 0.27, 0.46], gap="large")
    with c1:
        st.markdown("**Charts and figures**")
        selected_blocks = st.multiselect("Select dashboard charts", figure_titles, default=figure_titles[: min(5, len(figure_titles))], key="dash_builder_chart_blocks")
        include_tables = st.checkbox("Include ranking/data table previews", value=True, key="dash_builder_include_tables")
        if custom_figures and st.button("Clear added custom charts", key="clear_custom_dashboard_charts"):
            st.session_state["dashboard_custom_figures"] = []
            st.rerun()
    with c2:
        st.markdown("**Map layers**")
        selected_maps = st.multiselect("Select dashboard maps", list(maps.keys()), default=list(maps.keys())[: min(2, len(maps))], key="dash_builder_map_layers", help="Dashboard maps are read-only. Map editing and filtering stay in the Visualization section.")
        overlay_sources = _workflow_overlay_sources(st)
        selected_overlays = st.multiselect("Optional workflow layers on dashboard maps", list(overlay_sources.keys()), default=[name for name in ["Roads"] if name in overlay_sources], key="dash_builder_overlay_layers", help="Only selected context layers are included in the dashboard and report maps. Signals are not included unless you select Signals.")
    with c3:
        st.markdown("**Exports**")
        st.caption("Export the dashboard as a static PNG summary or a Word report with charts, map summaries, and decision-ready tables.")
        report_timezone = st.selectbox("Report time zone", ["America/Denver", "Local/server time", "UTC", "America/Chicago", "America/Los_Angeles", "America/New_York"], index=0, key="dashboard_report_timezone", help="Streamlit Cloud often runs in UTC. Choose the local project timezone so the report timestamp matches your expected local time.")
        report_tz_value = None if report_timezone == "Local/server time" else report_timezone
        d1, d2, d3 = st.columns(3)
        with d1:
            png_bytes = _export_summary_image(tables, "png", extra_figures=custom_figures)
            if png_bytes:
                st.download_button("Download PNG", data=png_bytes, file_name="hin_dashboard_summary.png", mime="image/png", key="dash_export_png")
            else:
                st.caption("PNG needs kaleido")
        with d2:
            docx_bytes = _export_dashboard_docx(tables, selected_blocks, selected_maps, extra_figures=custom_figures, maps=maps, overlay_layers={name: overlay_sources[name] for name in selected_overlays if name in overlay_sources}, report_timezone=report_tz_value)
            if docx_bytes is not None:
                st.download_button("Download Word report", data=docx_bytes, file_name=_report_docx_filename(tables), mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dash_export_docx")
            else:
                st.info("Install python-docx for Word export.")
        with d3:
            data_zip_bytes = _download_generated_data_zip(st, tables)
            if data_zip_bytes:
                st.download_button(
                    "Download all generated data ZIP",
                    data=data_zip_bytes,
                    file_name="hin_generated_data_export.zip",
                    mime="application/zip",
                    key="dash_export_generated_data_zip",
                    help="Exports workflow-generated tables and GIS layers as CSV and GeoJSON files. This does not change analysis results.",
                )
            else:
                st.caption("Run workflow steps before exporting data.")

    st.markdown("<div class='dashboard-section-title'>Generated dashboard</div>", unsafe_allow_html=True)
    chart_titles = set(selected_blocks)
    for i in range(0, len(default_figures), 2):
        cols = st.columns(2)
        for j, col in enumerate(cols):
            idx = i + j
            if idx >= len(default_figures):
                continue
            title, fig, data = default_figures[idx]
            if title not in chart_titles:
                continue
            with col:
                fig.update_layout(height=330, margin=dict(l=20, r=20, t=45, b=35))
                st.plotly_chart(_polish_figure(fig), width="stretch", key=f"dash_generated_fig_{idx}")
    if selected_maps:
        map_cols = st.columns(min(2, len(selected_maps)))
        for i, map_name in enumerate(selected_maps[:2]):
            with map_cols[i % len(map_cols)]:
                st.markdown(f"**{map_name}**")
                _render_dashboard_map(st, map_name, maps[map_name], key=f"dash_map_{_safe_name(map_name)}_{i}", height=420, overlay_layers={name: overlay_sources[name] for name in selected_overlays if name in overlay_sources})
    if include_tables:
        st.markdown("**Dashboard table preview**")
        compact_tables = _report_tables(tables)
        if compact_tables:
            table_name = st.selectbox("Preview table", list(compact_tables.keys()), key="dash_builder_preview_table")
            st.dataframe(_safe_dataframe_for_display(compact_tables[table_name]).head(25), width="stretch", hide_index=True)
        else:
            st.info("No report-ready result tables are available yet.")

# --- V30 overrides: FARS/date detection, report tables, KPI report, HIN table cleanup ---

CRASH_TYPE_VALUE_HINTS = [
    "rear end", "front-to-rear", "front to rear", "angle", "sideswipe", "head on",
    "front-to-front", "front to front", "rear-to-side", "rear to side", "opposite direction",
    "same direction", "broadside", "approach turn", "turning", "overturn", "rollover",
    "pedestrian", "bicycle", "parked vehicle", "fixed object",
]


def _exact_or_contains_col(df, exact_names=None, contains=None, avoid=None):
    if df is None:
        return None
    exact_names = exact_names or []
    contains = contains or []
    avoid = avoid or []
    norm = {str(c).lower().replace("_", "").replace(" ", ""): c for c in df.columns}
    for name in exact_names:
        key = str(name).lower().replace("_", "").replace(" ", "")
        if key in norm:
            return norm[key]
    for col in df.columns:
        lower = str(col).lower()
        if any(a in lower for a in avoid):
            continue
        if any(k in lower for k in contains):
            return col
    return None


def _crash_type_col(df):
    """Find crash type/manner field, including FARS man_collname.

    This intentionally avoids ID/code columns unless their values look like real
    crash-type descriptions. For FARS, man_collname is preferred over st_case or
    SourceCrashID.
    """
    if df is None or getattr(df, "empty", True):
        return None
    preferred = [
        "Crash_Type", "CrashType", "CollisionType", "Manner_of_Collision",
        "MannerOfCollision", "man_collname", "MAN_COLLNAME", "mancollname",
        "harm_evname", "FirstHarmfulEvent", "First_Harmful_Event",
    ]
    col = _exact_or_contains_col(df, preferred, contains=["man_collname", "collision", "crash_type", "crashtype"])
    if col:
        return col
    # Score text columns by whether their values contain collision/crash-type words.
    best_col, best_score = None, 0
    for c in df.columns:
        lower = str(c).lower()
        if any(bad in lower for bad in ["id", "case", "source", "year", "month", "day", "lat", "lon", "mile"]):
            continue
        try:
            sample = " ".join(df[c].dropna().astype(str).head(300).str.lower().tolist())
        except Exception:
            continue
        score = sum(1 for hint in CRASH_TYPE_VALUE_HINTS if hint in sample)
        if score > best_score:
            best_col, best_score = c, score
    if best_score >= 1:
        return best_col
    return (
        _find_col(df, ["crash", "type"])
        or _find_col(df, ["collision", "type"])
        or _find_col(df, ["manner"])
    )


def _year_series_from_crashes(df):
    if df is None or getattr(df, "empty", True):
        return None
    # Prefer true year columns, then FARS caseyear, then parsed dates.
    for name in ["year", "crash_year", "CrashYear", "u_Year", "caseyear"]:
        col = _exact_or_contains_col(df, [name])
        if col:
            vals = pd.to_numeric(df[col], errors="coerce")
            if vals.notna().any():
                return vals.astype("Int64").astype(str).replace("<NA>", "Unknown")
    date_col = _time_col(df)
    if date_col:
        dt = pd.to_datetime(df[date_col], errors="coerce")
        if dt.notna().any():
            return dt.dt.year.astype("Int64").astype(str).replace("<NA>", "Unknown")
    return None


def _month_series_from_crashes(df):
    if df is None or getattr(df, "empty", True):
        return None
    for name in ["month", "crash_month", "CrashMonth", "u_Month"]:
        col = _exact_or_contains_col(df, [name])
        if col:
            vals = pd.to_numeric(df[col], errors="coerce")
            if vals.notna().any():
                return vals.astype("Int64")
    # Month names such as January.
    for name in ["monthname", "MonthName"]:
        col = _exact_or_contains_col(df, [name])
        if col:
            vals = pd.to_datetime(df[col].astype(str), format="%B", errors="coerce").dt.month
            if vals.notna().any():
                return vals.astype("Int64")
    date_col = _time_col(df)
    if date_col:
        dt = pd.to_datetime(df[date_col], errors="coerce")
        if dt.notna().any():
            return dt.dt.month.astype("Int64")
    return None


def _time_col(df):
    """Find an actual date/datetime field, not monthname/hourname text."""
    if df is None:
        return None
    preferred = [
        "CrashDate", "Crash_Date", "CrashDateTime", "Crash_Date_Time", "Date",
        "ReportDate", "ReportedDate", "datetime", "timestamp",
    ]
    for c in preferred:
        if c in df.columns:
            return c
    for c in df.columns:
        lower = str(c).lower()
        if any(skip in lower for skip in ["monthname", "hourname", "minutename", "dayname"]):
            continue
        if "date" in lower or "datetime" in lower or "timestamp" in lower:
            return c
    return None


def _month_trend_table(crashes):
    """Build Jan-Dec crash counts with one colored line per year.

    Uses explicit year/month columns first. This prevents FARS monthname values
    like January from being parsed as year 1970.
    """
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), None, None, None
    years = _year_series_from_crashes(crashes)
    months = _month_series_from_crashes(crashes)
    if years is None or months is None:
        return pd.DataFrame(), None, None, None
    tmp = pd.DataFrame({"Year": years, "Month": months})
    tmp["YearNum"] = pd.to_numeric(tmp["Year"], errors="coerce")
    tmp["Month"] = pd.to_numeric(tmp["Month"], errors="coerce")
    tmp = tmp[tmp["YearNum"].notna() & tmp["Month"].between(1, 12)].copy()
    if tmp.empty:
        return pd.DataFrame(), None, None, None
    tmp["Year"] = tmp["YearNum"].astype(int).astype(str)
    tmp["Month"] = tmp["Month"].astype(int)
    out = tmp.groupby(["Year", "Month"], dropna=False).size().reset_index(name="Count")
    years_sorted = sorted(out["Year"].unique(), key=lambda x: int(x) if str(x).isdigit() else str(x))
    full = pd.MultiIndex.from_product([years_sorted, range(1, 13)], names=["Year", "Month"]).to_frame(index=False)
    out = full.merge(out, on=["Year", "Month"], how="left").fillna({"Count": 0})
    out["Count"] = out["Count"].astype(int)
    out["Month label"] = pd.Categorical(pd.to_datetime(out["Month"].astype(str), format="%m").dt.strftime("%b"), categories=MONTH_ORDER, ordered=True)
    out = out.sort_values(["Year", "Month"])
    return out, "Month label", "Count", "Year"


def _crash_id_col(df):
    if df is None:
        return None
    mapped_col = st.session_state.get("mapped_crash_id_col", "")
    if mapped_col and mapped_col in df.columns:
        return mapped_col
    preferred = [
        "DashboardCrashID", "SourceCrashID", "CrashID", "CrashId",
        "crash_id", "CRASH_ID", "Crash_ID", "st_case", "ST_CASE",
        "case_id", "CASE_ID", "CaseID", "CrashNumber", "OBJECTID"
    ]
    return _exact_or_contains_col(df, preferred, contains=["dashboardcrashid", "sourcecrashid", "crashid", "caseid", "stcase"])


def _unique_crash_count(df, mask=None):
    if df is None:
        return 0
    work = df.loc[mask].copy() if mask is not None else df.copy()
    key = _crash_id_col(work)
    if key and key in work.columns:
        return int(work[key].nunique(dropna=True))
    return int(len(work))


def _summary_kpi_values(crashes):
    vals = {
        "Total crashes": 0,
        "Fatal crashes": 0,
        "Fatalities": 0,
        "Serious injury crashes": 0,
        "Serious injuries": 0,
    }
    if crashes is None or getattr(crashes, "empty", True):
        return vals
    vals["Total crashes"] = _unique_crash_count(crashes)
    kabco = _kabco_col(crashes)
    sev_cols = _severity_count_columns(crashes)
    # Person-count columns, including FARS fatals/fatalities and local Level A injuries.
    if "K" in sev_cols:
        fatal_vals = pd.to_numeric(crashes[sev_cols["K"]], errors="coerce").fillna(0)
        vals["Fatalities"] = int(fatal_vals.sum())
        vals["Fatal crashes"] = _unique_crash_count(crashes, fatal_vals > 0)
    if "A" in sev_cols:
        serious_vals = pd.to_numeric(crashes[sev_cols["A"]], errors="coerce").fillna(0)
        vals["Serious injuries"] = int(serious_vals.sum())
        vals["Serious injury crashes"] = _unique_crash_count(crashes, serious_vals > 0)
    # Crash-level KABCO fields fallback.
    if kabco and kabco in crashes.columns:
        k = crashes[kabco].map(_normalize_kabco_value).astype(str).str.upper().str.strip()
        if vals["Fatal crashes"] == 0:
            vals["Fatal crashes"] = _unique_crash_count(crashes, k == "K")
        if vals["Serious injury crashes"] == 0:
            vals["Serious injury crashes"] = _unique_crash_count(crashes, k == "A")
        if vals["Fatalities"] == 0:
            vals["Fatalities"] = vals["Fatal crashes"]
        if vals["Serious injuries"] == 0:
            vals["Serious injuries"] = vals["Serious injury crashes"]
    return vals


def _year_kabco_table(crashes):
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), None, None
    years = _year_series_from_crashes(crashes)
    if years is None:
        return pd.DataFrame(), None, None
    kabco_col = _kabco_col(crashes)
    sev_cols = _severity_count_columns(crashes)
    if kabco_col and kabco_col in crashes.columns:
        work = pd.DataFrame({"Year": years, "KABCO": crashes[kabco_col].map(_normalize_kabco_value)})
        work = work[work["Year"].ne("Unknown")]
        if work.empty:
            return pd.DataFrame(), None, None
        out = work.groupby(["Year", "KABCO"], dropna=False).size().reset_index(name="Count")
    elif sev_cols:
        parts = []
        for code, col in sev_cols.items():
            vals = pd.to_numeric(crashes[col], errors="coerce").fillna(0)
            tmp = pd.DataFrame({"Year": years, "KABCO": code, "Count": vals})
            tmp = tmp[tmp["Year"].ne("Unknown") & (tmp["Count"] > 0)]
            if not tmp.empty:
                parts.append(tmp)
        if not parts:
            return pd.DataFrame(), None, None
        out = pd.concat(parts, ignore_index=True).groupby(["Year", "KABCO"], dropna=False)["Count"].sum().reset_index()
    else:
        return pd.DataFrame(), None, None
    out = _order_kabco(out, "KABCO")
    try:
        out["__year_sort__"] = pd.to_numeric(out["Year"], errors="coerce")
        out["__kabco_sort__"] = out["KABCO"].map({"K": 1, "A": 2, "B": 3, "C": 4, "O": 5}).fillna(9)
        out = out.sort_values(["__year_sort__", "__kabco_sort__"]).drop(columns=["__year_sort__", "__kabco_sort__"])
    except Exception:
        pass
    return out, "Year", "KABCO"


def _hin_table_for_display(hin, metric, top_n=20):
    """Decision table for top HIN windows/segments.

    For sliding windows, FromMile/ToMile are shown as the window mile range.
    Route context is limited to route name and total route length to avoid
    confusing route from/to fields.
    """
    if hin is None or getattr(hin, "empty", True) or metric not in hin.columns:
        return pd.DataFrame()
    work = hin.copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work = work.sort_values(metric, ascending=False).head(top_n).reset_index(drop=True)
    id_col = _normal_col(work, ["RiskSegmentID", "WindowID", "SlidingWindowID", "SegmentID", "UnitID", "SourceSegmentID"])
    route_col = _normal_col(work, ["Route", "FULLNAME", "RoadName", "RouteName", "CorridorRoute", "Name"])
    length_col = _normal_col(work, ["WindowLength_Miles", "Length_Miles", "Length_Mi", "SegmentLength_Mile", "length_mi"])
    from_col = _normal_col(work, ["FromMile", "From_Mile", "from_mile", "BeginMile", "StartMile", "WindowFromMile"])
    to_col = _normal_col(work, ["ToMile", "To_Mile", "to_mile", "EndMile", "WindowToMile"])
    out = pd.DataFrame()
    out["Rank"] = range(1, len(work) + 1)
    out["SegID"] = work[id_col].astype(str).values if id_col else [f"HIN_{i+1}" for i in range(len(work))]
    out["Seg/window length"] = pd.to_numeric(work[length_col], errors="coerce").round(3).values if length_col else ""
    out["From mile"] = pd.to_numeric(work[from_col], errors="coerce").round(3).values if from_col else ""
    out["To mile"] = pd.to_numeric(work[to_col], errors="coerce").round(3).values if to_col else ""
    out["Route"] = work[route_col].astype(str).values if route_col else ""
    if route_col and route_col in hin.columns:
        full = hin.copy()
        if length_col and length_col in full.columns:
            full[length_col] = pd.to_numeric(full[length_col], errors="coerce").fillna(0)
            route_len = full.groupby(route_col, dropna=False)[length_col].sum().to_dict()
        else:
            route_len = {}
        out["Route total length"] = [round(float(route_len.get(v, 0)), 3) if v in route_len else "" for v in work[route_col]]
    else:
        out["Route total length"] = ""
    out["HIN index"] = pd.to_numeric(work[metric], errors="coerce").round(3).values
    return out


def _top_hin_export_table(hin, top_n=20):
    if hin is None or getattr(hin, "empty", True):
        return pd.DataFrame()
    metric = "HIN_Priority_Index" if "HIN_Priority_Index" in hin.columns else _default_metric(_numeric_cols(hin))
    return _hin_table_for_display(hin, metric, top_n) if metric else pd.DataFrame()


def _report_tables(tables, top_n=20):
    """Report-ready tables for dashboard preview/export."""
    return _export_tables_only(tables, top_n=top_n)


def _build_default_figures(tables):
    """Default figures. KPI cards stay as dashboard/report cards, not a chart."""
    figures = []
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))
    if crashes is not None:
        year_kabco, year_col, kabco_col = _year_kabco_table(crashes)
        if not year_kabco.empty:
            fig = px.bar(year_kabco, x=year_col, y="Count", color=kabco_col, color_discrete_map=KABCO_COLOR_MAP, category_orders={kabco_col: ["K", "A", "B", "C", "O"]}, title=f"Crashes by year and {kabco_col}")
            fig.update_layout(barmode="stack", xaxis_title="Year", yaxis_title="Crash count")
            figures.append((f"Crashes by year and {kabco_col}", _apply_kabco_trace_colors(_polish_figure(fig)), year_kabco))
        type_col = _crash_type_col(crashes)
        if type_col:
            type_df = _aggregate(crashes, type_col, None, "Count", 12)
            pie = px.pie(type_df, names=type_col, values="Count", hole=0.38, title=f"Crash type share by {type_col}")
            figures.append((f"Crash type share by {type_col}", _polish_figure(pie), type_df))
        monthly_df, period_col, value_col, color_col = _month_trend_table(crashes)
        if not monthly_df.empty:
            fig = px.line(monthly_df, x=period_col, y=value_col, color=color_col, markers=True, category_orders={period_col: MONTH_ORDER}, title="Monthly crash trend by year")
            fig.update_layout(xaxis_title="Month", yaxis_title="Crash count")
            figures.append(("Monthly crash trend by year", _polish_figure(fig), monthly_df))
        mode_df, mode_kabco_col, mode_col = _mode_severity_bubble_table(crashes)
        if not mode_df.empty:
            fig = px.scatter(mode_df, x=mode_kabco_col, y=mode_col, size="Count", color=mode_col, size_max=56, title="Travel mode severity bubble chart", hover_data=["Count"])
            fig.update_traces(marker=dict(sizemin=5, opacity=0.78, line=dict(width=1, color="white")))
            figures.append(("Travel mode severity bubble chart", _polish_figure(fig), mode_df))
        crash_kabco, crash_type_col, crash_kabco_col = _crash_type_kabco_table(crashes)
        if not crash_kabco.empty:
            pivot = crash_kabco.pivot_table(index=crash_type_col, columns=crash_kabco_col, values="Count", aggfunc="sum", fill_value=0)
            fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Crash type by {crash_kabco_col}", labels=dict(x=crash_kabco_col, y="Crash type", color="Crash count"), color_continuous_scale="YlOrRd")
            figures.append((f"Crash type by {crash_kabco_col}", _polish_figure(fig), pivot.reset_index()))
            tree = crash_kabco.copy(); tree["All crashes"] = "All crashes"
            fig = px.treemap(tree, path=["All crashes", crash_type_col, crash_kabco_col], values="Count", color=crash_kabco_col, color_discrete_map=KABCO_COLOR_MAP, title=f"Crash type and {crash_kabco_col} treemap")
            figures.append((f"Crash type and {crash_kabco_col} treemap", _polish_figure(fig), tree))
    if density is not None:
        metric = "CrashDensity" if "CrashDensity" in density.columns else _default_metric(_numeric_cols(density))
        rank_df, unit_col, value_col = _rank_units_for_chart(density, metric, 15) if metric else (pd.DataFrame(), None, None)
        if not rank_df.empty:
            plot_df = rank_df.sort_values(value_col, ascending=True)
            fig = px.bar(plot_df, y=unit_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title="Top spatial units by crash density")
            fig.update_layout(yaxis_title="Spatial unit ID", xaxis_title="Crash density")
            figures.append(("Top spatial units by crash density", _polish_figure(fig), rank_df))
    if hin is not None:
        metric = "HIN_Priority_Index" if "HIN_Priority_Index" in hin.columns else _default_metric(_numeric_cols(hin))
        hin_table = _hin_table_for_display(hin, metric, 20) if metric else pd.DataFrame()
        if not hin_table.empty:
            # Keep a simple placeholder chart for visual dashboard selection; table is the main HIN presentation.
            fig = px.bar(hin_table.sort_values("HIN index", ascending=True), y="SegID", x="HIN index", orientation="h", title="Top HIN segments/windows by HIN priority index")
            fig.update_layout(yaxis_title="SegID", xaxis_title="HIN priority index")
            figures.append(("Top HIN segments/windows table", _polish_figure(fig), hin_table))
    return figures


def _export_dashboard_docx(tables, selected_blocks, selected_maps, extra_figures=None, maps=None, overlay_layers=None, report_timezone=None):
    if Document is None:
        return None
    doc = Document()
    doc.add_heading("HIN dashboard report", level=0)
    doc.add_paragraph(f"Generated {_report_time_text(report_timezone)}.")
    doc.add_paragraph("This report summarizes crash patterns, crash-density rankings, selected maps, and HIN/risk results from the app dashboard.")

    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    kpis = _summary_kpi_values(crashes)
    doc.add_heading("Crash summary", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Total crashes: {kpis.get('Total crashes', 0):,}    ").bold = True
    p.add_run(f"Fatal crashes: {kpis.get('Fatal crashes', 0):,}    ").bold = True
    p.add_run(f"Fatalities: {kpis.get('Fatalities', 0):,}    ").bold = True
    p.add_run(f"Serious injury crashes: {kpis.get('Serious injury crashes', 0):,}    ").bold = True
    p.add_run(f"Serious injuries: {kpis.get('Serious injuries', 0):,}").bold = True

    figures = _build_default_figures(tables) + (extra_figures or [])
    for title, fig, data in figures:
        if selected_blocks and title not in selected_blocks:
            continue
        doc.add_heading(title, level=1)
        img = _figure_to_png_bytes(_polish_figure(fig))
        if img:
            doc.add_picture(io.BytesIO(img), width=Inches(6.5))
        else:
            doc.add_paragraph("Chart image could not be generated in this environment. The summary table is included below.")
        table_df = _safe_dataframe_for_display(data.head(15).copy())
        if not table_df.empty:
            doc.add_paragraph("Summary table")
            table = doc.add_table(rows=1, cols=len(table_df.columns))
            table.style = "Table Grid"
            for i, col in enumerate(table_df.columns):
                table.rows[0].cells[i].text = str(col)
            for _, row in table_df.iterrows():
                cells = table.add_row().cells
                for i, col in enumerate(table_df.columns):
                    cells[i].text = str(row[col])

    if selected_maps:
        doc.add_heading("Selected map layers", level=1)
        for m in selected_maps:
            doc.add_heading(str(m), level=2)
            if maps and m in maps:
                map_png = _static_map_png(maps[m], str(m), overlay_layers=overlay_layers)
                if map_png:
                    doc.add_picture(io.BytesIO(map_png), width=Inches(6.5))
                else:
                    doc.add_paragraph("Static map image could not be generated.")
            else:
                doc.add_paragraph("Map layer selected in dashboard builder.")

    doc.add_heading("Decision-ready result tables", level=1)
    for name, df in _report_tables(tables).items():
        doc.add_heading(name, level=2)
        table_df = _safe_dataframe_for_display(df.head(20).copy())
        if table_df.empty:
            continue
        table = doc.add_table(rows=1, cols=len(table_df.columns))
        table.style = "Table Grid"
        for i, col in enumerate(table_df.columns):
            table.rows[0].cells[i].text = str(col)
        for _, row in table_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(table_df.columns):
                cells[i].text = str(row[col])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- V31 dashboard data-detection, severity table, chart, and OSM-safe overrides ---

def _crash_type_col(df):
    """Find crash type / manner-of-collision field from uploaded, local, or FARS data."""
    if df is None or getattr(df, "empty", True):
        return None
    exact_names = [
        "Crash_Type", "CrashType", "CollisionType", "Manner_of_Collision",
        "MannerOfCollision", "man_collname", "MAN_COLLNAME", "Man_CollName",
        "harm_evname", "HARM_EVNAME", "FirstHarmfulEvent", "First_Harmful_Event",
    ]
    for wanted in exact_names:
        for col in df.columns:
            if str(col).lower().replace("_", "") == wanted.lower().replace("_", ""):
                return col
    for col in df.columns:
        lower = str(col).lower()
        if any(key in lower for key in ["man_coll", "manner", "collision", "crash_type", "crashtype"]):
            if not any(bad in lower for bad in ["id", "case", "source"]):
                return col
    hints = [
        "rear", "front", "head on", "angle", "sideswipe", "broadside", "approach",
        "turn", "pedestrian", "bicycle", "fixed object", "parked", "overturn", "animal",
    ]
    best_col, best_score = None, 0
    for col in df.columns:
        lower = str(col).lower()
        if any(bad in lower for bad in ["id", "case", "source", "year", "month", "day", "lat", "lon", "mile", "time", "date"]):
            continue
        try:
            sample = " ".join(df[col].dropna().astype(str).head(500).str.lower().tolist())
        except Exception:
            continue
        score = sum(1 for h in hints if h in sample)
        if score > best_score:
            best_col, best_score = col, score
    return best_col if best_score >= 1 else None


def _summary_kpi_values(crashes):
    """Crash/person summary with proper crash-count vs person-count logic.

    Fatal crashes and serious-injury crashes count unique crash/case IDs when an
    ID field exists. Fatalities and serious injuries sum the person-count fields.
    """
    vals = {
        "Total crashes": 0,
        "Fatal crashes": 0,
        "Fatalities": 0,
        "Serious injury crashes": 0,
        "Serious injuries": 0,
    }
    if crashes is None or getattr(crashes, "empty", True):
        return vals
    vals["Total crashes"] = _unique_crash_count(crashes)
    sev_cols = _severity_count_columns(crashes)
    kabco = _kabco_col(crashes)
    if "K" in sev_cols:
        fatal_vals = pd.to_numeric(crashes[sev_cols["K"]], errors="coerce").fillna(0)
        vals["Fatalities"] = int(fatal_vals.sum())
        vals["Fatal crashes"] = _unique_crash_count(crashes, fatal_vals > 0)
    elif kabco and kabco in crashes.columns:
        k = crashes[kabco].map(_normalize_kabco_value).astype(str).str.upper().str.strip()
        vals["Fatal crashes"] = _unique_crash_count(crashes, k == "K")
        vals["Fatalities"] = vals["Fatal crashes"]
    if "A" in sev_cols:
        serious_vals = pd.to_numeric(crashes[sev_cols["A"]], errors="coerce").fillna(0)
        vals["Serious injuries"] = int(serious_vals.sum())
        vals["Serious injury crashes"] = _unique_crash_count(crashes, serious_vals > 0)
    elif kabco and kabco in crashes.columns:
        k = crashes[kabco].map(_normalize_kabco_value).astype(str).str.upper().str.strip()
        vals["Serious injury crashes"] = _unique_crash_count(crashes, k == "A")
        vals["Serious injuries"] = vals["Serious injury crashes"]
    return vals


def _mode_severity_bubble_table(crashes):
    """Travel-mode x severity table with a display-size field for readable bubbles."""
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), None, None
    kabco_col = _kabco_col(crashes)
    if kabco_col and kabco_col in crashes.columns:
        kabco_vals = crashes[kabco_col].map(_normalize_kabco_value)
        base = crashes.copy()
        base["KABCO_Normalized"] = kabco_vals
        sev_col = "KABCO_Normalized"
    else:
        sev_cols = _severity_count_columns(crashes)
        parts = []
        for code, col in sev_cols.items():
            tmp = crashes.copy()
            tmp["KABCO_Normalized"] = code
            tmp["__sev_count__"] = pd.to_numeric(tmp[col], errors="coerce").fillna(0)
            tmp = tmp[tmp["__sev_count__"] > 0]
            if not tmp.empty:
                parts.append(tmp)
        if not parts:
            return pd.DataFrame(), None, None
        base = pd.concat(parts, ignore_index=True)
        sev_col = "KABCO_Normalized"

    # Detect mode from crash type text and common vehicle/mode fields.
    text_cols = []
    for c in base.columns:
        lower = str(c).lower()
        if any(k in lower for k in ["crash", "type", "manner", "collision", "vehicle", "veh", "body", "mode", "person"]):
            text_cols.append(c)
    def mode_for_row(row):
        text = " ".join(str(row.get(c, "")) for c in text_cols).lower()
        if "ped" in text or "pedestrian" in text:
            return "Pedestrian"
        if "bike" in text or "bicycle" in text or "cycl" in text:
            return "Bicycle"
        if "motorcycle" in text or "motor cycle" in text:
            return "Motorcycle"
        if "bus" in text or "transit" in text or "school bus" in text:
            return "Bus / transit"
        return "Motor vehicle / other"
    base["Mode"] = base.apply(mode_for_row, axis=1)
    if "__sev_count__" in base.columns:
        out = base.groupby([sev_col, "Mode"], dropna=False)["__sev_count__"].sum().reset_index(name="Count")
    else:
        out = base.groupby([sev_col, "Mode"], dropna=False).size().reset_index(name="Count")
    out = out.rename(columns={sev_col: "KABCO"})
    out = out[out["KABCO"].astype(str).isin(["K", "A", "B", "C", "O"])]
    if out.empty:
        return pd.DataFrame(), None, None
    full = pd.MultiIndex.from_product([ ["K","A","B","C","O"], sorted(out["Mode"].unique()) ], names=["KABCO", "Mode"]).to_frame(index=False)
    out = full.merge(out, on=["KABCO", "Mode"], how="left").fillna({"Count": 0})
    out["Count"] = pd.to_numeric(out["Count"], errors="coerce").fillna(0)
    # sqrt scale makes small bike/pedestrian bubbles visible without letting vehicle PDO dominate.
    out["BubbleSize"] = (out["Count"].pow(0.5) * 8 + 8).round(2)
    out = _order_kabco(out, "KABCO")
    return out, "KABCO", "Mode"


def _spatial_unit_severity_table(tables, top_n=20):
    """Each row is one spatial unit; K/A/B/C/O columns contain crash/person counts."""
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame()
    unit_col = _unit_col(crashes) or _normal_col(crashes, ["UnitID", "IntersectionID", "CorridorID", "SegmentID", "SegID"])
    if not unit_col:
        return pd.DataFrame()
    work = _drop_geometry(crashes).copy()
    sev_cols = _severity_count_columns(work)
    rows = None
    if sev_cols:
        agg = work[[unit_col] + list(sev_cols.values())].copy()
        for code, col in sev_cols.items():
            agg[code] = pd.to_numeric(agg[col], errors="coerce").fillna(0)
        rows = agg.groupby(unit_col, dropna=False)[[c for c in ["K", "A", "B", "C", "O"] if c in agg.columns]].sum().reset_index()
    else:
        kabco = _kabco_col(work)
        if not kabco or kabco not in work.columns:
            return pd.DataFrame()
        work["KABCO_Normalized"] = work[kabco].map(_normalize_kabco_value)
        rows = work.groupby([unit_col, "KABCO_Normalized"], dropna=False).size().reset_index(name="Count")
        rows = rows[rows["KABCO_Normalized"].isin(["K", "A", "B", "C", "O"])]
        if rows.empty:
            return pd.DataFrame()
        rows = rows.pivot_table(index=unit_col, columns="KABCO_Normalized", values="Count", aggfunc="sum", fill_value=0).reset_index()
    for code in ["K", "A", "B", "C", "O"]:
        if code not in rows.columns:
            rows[code] = 0
    rows["Total"] = rows[["K", "A", "B", "C", "O"]].sum(axis=1)
    rows = rows.sort_values("Total", ascending=False).head(top_n).reset_index(drop=True)
    rows.insert(0, "Rank", range(1, len(rows) + 1))
    rows = rows.rename(columns={unit_col: "Spatial unit id"})
    return rows[["Rank", "Spatial unit id", "K", "A", "B", "C", "O", "Total"]]


def _export_tables_only(tables, top_n=20):
    out = {}
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))
    if crashes is not None:
        type_col = _crash_type_col(crashes)
        if type_col:
            out["Crash type summary"] = _aggregate(crashes, type_col, None, "Count", top_n)
        years = _year_series_from_crashes(crashes)
        if years is not None:
            ydf = pd.DataFrame({"Year": years})
            ydf = ydf[ydf["Year"].ne("Unknown")]
            out["Crash year summary"] = ydf.groupby("Year", dropna=False).size().reset_index(name="Count").sort_values("Year")
        sev = _spatial_unit_severity_table(tables, top_n=top_n)
        if not sev.empty:
            out["Severity summary by spatial unit"] = sev
    if density is not None:
        top_density = _top_density_export_table(density, top_n=top_n)
        if not top_density.empty:
            out["Top crash-density spatial units"] = top_density
    if hin is not None:
        top_hin = _top_hin_export_table(hin, top_n=top_n)
        if not top_hin.empty:
            out["Top HIN/risk spatial units"] = top_hin
    return {name: _safe_dataframe_for_display(df) for name, df in out.items()}


def _render_pattern_charts(st, tables):
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))

    _render_kpi_strip(st, crashes)
    st.markdown("<div class='dashboard-section-title'>Crash patterns <span>years, crash type, month, KABCO, mode, and roadway context</span></div>", unsafe_allow_html=True)
    left, right = st.columns(2)
    with left:
        if crashes is not None:
            year_kabco, year_col, kabco_col = _year_kabco_table(crashes)
            if not year_kabco.empty:
                fig = px.bar(year_kabco, x=year_col, y="Count", color=kabco_col, color_discrete_map=KABCO_COLOR_MAP, category_orders={kabco_col: ["K","A","B","C","O"]}, title=f"Crashes by year and {kabco_col}", hover_data={"Count": True, kabco_col: True})
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), barmode="stack", xaxis_title="Year", yaxis_title="Crash count")
                st.plotly_chart(_apply_kabco_trace_colors(_polish_figure(fig)), width="stretch")
            else:
                years = _year_series_from_crashes(crashes)
                if years is not None:
                    ydf = pd.DataFrame({"Year": years})
                    ydf = ydf[ydf["Year"].ne("Unknown")].groupby("Year").size().reset_index(name="Count")
                    fig = px.bar(ydf.sort_values("Year"), x="Year", y="Count", title="Crashes by year")
                    fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), xaxis_title="Year", yaxis_title="Crash count")
                    st.plotly_chart(_polish_figure(fig), width="stretch")
    with right:
        if crashes is not None:
            type_col = _crash_type_col(crashes)
            if type_col:
                type_df = _aggregate(crashes, type_col, None, "Count", 10)
                fig = px.pie(type_df, names=type_col, values="Count", hole=0.38, title=f"Crash type share by {type_col}")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), legend=dict(orientation="v"))
                st.plotly_chart(_polish_figure(fig), width="stretch")

    left2, right2 = st.columns(2)
    with left2:
        monthly_df, period_col, value_col, color_col = _month_trend_table(crashes)
        if not monthly_df.empty:
            fig = px.line(monthly_df, x=period_col, y=value_col, color=color_col, markers=True, category_orders={period_col: MONTH_ORDER}, title="Monthly crash trend by year")
            fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), xaxis_title="Month", yaxis_title="Crash count")
            st.plotly_chart(_polish_figure(fig), width="stretch")
    with right2:
        mode_df, mode_kabco_col, mode_col = _mode_severity_bubble_table(crashes)
        if not mode_df.empty:
            fig = px.scatter(mode_df, x=mode_kabco_col, y=mode_col, size="BubbleSize", color=mode_col, size_max=42, title="Travel mode severity bubble chart", hover_data={"Count": True, "BubbleSize": False})
            fig.update_traces(marker=dict(sizemin=7, opacity=0.80, line=dict(width=1, color="white")))
            fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), xaxis_title="Severity", yaxis_title="Mode")
            st.plotly_chart(_polish_figure(fig), width="stretch")

    road_kabco, road_col, road_kabco_col = _road_class_kabco_table(crashes, st_obj=st)
    if not road_kabco.empty:
        left3, right3 = st.columns(2)
        with left3:
            pivot = road_kabco.pivot_table(index=road_col, columns=road_kabco_col, values="Count", aggfunc="sum", fill_value=0)
            order = [c for c in ["K","A","B","C","O"] if c in pivot.columns]
            pivot = pivot[order + [c for c in pivot.columns if c not in order]]
            fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Road class by {road_kabco_col}", labels=dict(x=road_kabco_col, y="Road class", color="Crash count"), color_continuous_scale="YlOrRd")
            fig.update_layout(height=380, margin=dict(l=20, r=20, t=45, b=35))
            st.plotly_chart(_polish_figure(fig), width="stretch")
        with right3:
            tree = road_kabco.copy(); tree["All crashes"] = "All crashes"
            fig = px.treemap(tree, path=["All crashes", road_col, road_kabco_col], values="Count", color=road_kabco_col, color_discrete_map=KABCO_COLOR_MAP, title=f"Road class and {road_kabco_col} treemap")
            fig.update_layout(height=380, margin=dict(l=20, r=20, t=45, b=35))
            st.plotly_chart(_polish_figure(fig), width="stretch")

    crash_kabco, crash_type_col, crash_kabco_col = _crash_type_kabco_table(crashes)
    if not crash_kabco.empty:
        pivot = crash_kabco.pivot_table(index=crash_type_col, columns=crash_kabco_col, values="Count", aggfunc="sum", fill_value=0)
        order = [c for c in ["K","A","B","C","O"] if c in pivot.columns]
        pivot = pivot[order + [c for c in pivot.columns if c not in order]]
        fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Crash type by {crash_kabco_col}", labels=dict(x=crash_kabco_col, y="Crash type", color="Crash count"), color_continuous_scale="YlOrRd")
        fig.update_layout(height=430, margin=dict(l=20, r=20, t=45, b=35))
        st.plotly_chart(_polish_figure(fig), width="stretch")

    st.markdown("<div class='dashboard-section-title'>Risk and spatial-unit ranking <span>crash density, crash count, and HIN priority ranking</span></div>", unsafe_allow_html=True)
    left, right = st.columns(2)
    with left:
        if density is not None:
            metric = "CrashDensity" if "CrashDensity" in density.columns else _default_metric(_numeric_cols(density))
            rank_df, unit_col, value_col = _rank_units_for_chart(density, metric, 15) if metric else (pd.DataFrame(), None, None)
            if not rank_df.empty:
                plot_df = rank_df.sort_values(value_col, ascending=True)
                fig = px.bar(plot_df, y=unit_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title="Top spatial units by crash density")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), yaxis_title="Spatial unit ID", xaxis_title="Crash density")
                st.plotly_chart(_polish_figure(fig), width="stretch")
    with right:
        if density is not None:
            count_col = _crash_count_col(density)
            rank_df, unit_col, value_col = _rank_units_for_chart(density, count_col, 15) if count_col else (pd.DataFrame(), None, None)
            if not rank_df.empty:
                plot_df = rank_df.sort_values(value_col, ascending=True)
                fig = px.bar(plot_df, y=unit_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title="Top spatial units by crash count")
                fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35), yaxis_title="Spatial unit ID", xaxis_title="Crash count")
                st.plotly_chart(_polish_figure(fig), width="stretch")
    if hin is not None:
        metric = "HIN_Priority_Index" if "HIN_Priority_Index" in hin.columns else _default_metric(_numeric_cols(hin))
        hin_tbl = _hin_table_for_display(hin, metric, 20) if metric else pd.DataFrame()
        if not hin_tbl.empty:
            st.markdown("<div class='dashboard-section-title'>HIN priority ranking <span>table view with route and window/segment context</span></div>", unsafe_allow_html=True)
            st.dataframe(_safe_dataframe_for_display(hin_tbl), width="stretch", height=360)
    _render_hin_network_summary(st, hin, crashes)


def _build_default_figures(tables):
    figures = []
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))
    if crashes is not None:
        year_kabco, year_col, kabco_col = _year_kabco_table(crashes)
        if not year_kabco.empty:
            fig = px.bar(year_kabco, x=year_col, y="Count", color=kabco_col, color_discrete_map=KABCO_COLOR_MAP, category_orders={kabco_col: ["K","A","B","C","O"]}, title=f"Crashes by year and {kabco_col}")
            fig.update_layout(barmode="stack", xaxis_title="Year", yaxis_title="Crash count")
            figures.append((f"Crashes by year and {kabco_col}", _apply_kabco_trace_colors(_polish_figure(fig)), year_kabco))
        type_col = _crash_type_col(crashes)
        if type_col:
            type_df = _aggregate(crashes, type_col, None, "Count", 12)
            pie = px.pie(type_df, names=type_col, values="Count", hole=0.38, title=f"Crash type share by {type_col}")
            figures.append((f"Crash type share by {type_col}", _polish_figure(pie), type_df))
        monthly_df, period_col, value_col, color_col = _month_trend_table(crashes)
        if not monthly_df.empty:
            fig = px.line(monthly_df, x=period_col, y=value_col, color=color_col, markers=True, category_orders={period_col: MONTH_ORDER}, title="Monthly crash trend by year")
            figures.append(("Monthly crash trend by year", _polish_figure(fig), monthly_df))
        mode_df, mode_kabco_col, mode_col = _mode_severity_bubble_table(crashes)
        if not mode_df.empty:
            fig = px.scatter(mode_df, x=mode_kabco_col, y=mode_col, size="BubbleSize", color=mode_col, size_max=42, title="Travel mode severity bubble chart", hover_data={"Count": True, "BubbleSize": False})
            fig.update_traces(marker=dict(sizemin=7, opacity=0.80, line=dict(width=1, color="white")))
            figures.append(("Travel mode severity bubble chart", _polish_figure(fig), mode_df))
        road_kabco, road_col, road_kabco_col = _road_class_kabco_table(crashes, st_obj=None)
        if not road_kabco.empty:
            tree = road_kabco.copy(); tree["All crashes"] = "All crashes"
            fig = px.treemap(tree, path=["All crashes", road_col, road_kabco_col], values="Count", color=road_kabco_col, color_discrete_map=KABCO_COLOR_MAP, title=f"Road class and {road_kabco_col} treemap")
            figures.append((f"Road class and {road_kabco_col} treemap", _polish_figure(fig), tree))
        crash_kabco, crash_type_col, crash_kabco_col = _crash_type_kabco_table(crashes)
        if not crash_kabco.empty:
            pivot = crash_kabco.pivot_table(index=crash_type_col, columns=crash_kabco_col, values="Count", aggfunc="sum", fill_value=0)
            fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Crash type by {crash_kabco_col}", labels=dict(x=crash_kabco_col, y="Crash type", color="Crash count"), color_continuous_scale="YlOrRd")
            figures.append((f"Crash type by {crash_kabco_col}", _polish_figure(fig), pivot.reset_index()))
    if density is not None:
        metric = "CrashDensity" if "CrashDensity" in density.columns else _default_metric(_numeric_cols(density))
        rank_df, unit_col, value_col = _rank_units_for_chart(density, metric, 15) if metric else (pd.DataFrame(), None, None)
        if not rank_df.empty:
            plot_df = rank_df.sort_values(value_col, ascending=True)
            fig = px.bar(plot_df, y=unit_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title="Top spatial units by crash density")
            fig.update_layout(yaxis_title="Spatial unit ID", xaxis_title="Crash density")
            figures.append(("Top spatial units by crash density", _polish_figure(fig), rank_df))
    # HIN priority is intentionally exported as a report table, not as a confusing bar chart.
    return figures

# --- V35 field-mapping overrides -------------------------------------------------
# These final overrides make dashboard/report logic use the user-confirmed crash
# field mapping and canonical Dashboard* columns when available.  This reduces
# dependence on hard-coded agency-specific column names.
try:
    from modules.crash_field_mapping import normalize_kabco_value, parse_month_value
except Exception:  # pragma: no cover
    normalize_kabco_value = lambda v: v
    parse_month_value = lambda v: None


def _mapped_col(df, semantic_name):
    try:
        mapping = st.session_state.get("crash_field_mapping", {})
    except Exception:
        mapping = {}
    c = mapping.get(semantic_name, "") if isinstance(mapping, dict) else ""
    return c if df is not None and c in df.columns else None


def _crash_type_col(df):
    if df is None:
        return None
    if "DashboardCrashType" in df.columns:
        return "DashboardCrashType"
    mapped = _mapped_col(df, "crash_type")
    if mapped:
        return mapped
    return (
        _find_col(df, ["crash", "type"])
        or _find_col(df, ["collision", "type"])
        or _find_col(df, ["manner"])
        or _find_col(df, ["man", "coll"])
        or _find_col(df, ["type"])
    )


def _kabco_col(df):
    if df is None:
        return None
    if "DashboardKABCO" in df.columns:
        return "DashboardKABCO"
    mapped = _mapped_col(df, "severity")
    if mapped:
        return mapped
    for c in df.columns:
        lower = str(c).lower().replace("_", "")
        if lower == "kabco" or "kabco" in lower:
            return c
    for c in df.columns:
        lower = str(c).lower()
        if "severity" in lower or "injury" in lower:
            return c
    return None


def _time_col(df):
    if df is None:
        return None
    if "DashboardCrashDate" in df.columns:
        return "DashboardCrashDate"
    mapped = _mapped_col(df, "crash_date")
    if mapped:
        return mapped
    return (
        _find_col(df, ["date"])
        or _find_col(df, ["time"])
        or _find_col(df, ["datetime"])
    )


def _year_series_from_crashes(crashes):
    if crashes is None or getattr(crashes, "empty", True):
        return None
    if "DashboardCrashYear" in crashes.columns:
        y = pd.to_numeric(crashes["DashboardCrashYear"], errors="coerce")
        y = y.where((y >= 1900) & (y <= 2100))
        if y.notna().any():
            return y.dropna().astype(int).astype(str)
    mapped = _mapped_col(crashes, "crash_year")
    if mapped:
        y = pd.to_numeric(crashes[mapped], errors="coerce")
        y = y.where((y >= 1900) & (y <= 2100))
        if y.notna().any():
            return y.dropna().astype(int).astype(str)
    date_col = _time_col(crashes)
    if date_col:
        dt = pd.to_datetime(crashes[date_col], errors="coerce")
        if dt.notna().any():
            return dt.dt.year.dropna().astype(int).astype(str)
    return None


def _month_trend_table(crashes):
    """Return Jan-Dec crash counts with one line per year using mapped fields."""
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), "Month", "Crash count", "Year"
    work = pd.DataFrame(index=crashes.index)
    if "DashboardCrashYear" in crashes.columns:
        work["Year"] = pd.to_numeric(crashes["DashboardCrashYear"], errors="coerce")
    else:
        year_col = _mapped_col(crashes, "crash_year") or _find_col(crashes, ["year"])
        if year_col:
            work["Year"] = pd.to_numeric(crashes[year_col], errors="coerce")
    if "DashboardCrashMonth" in crashes.columns:
        work["MonthNum"] = pd.to_numeric(crashes["DashboardCrashMonth"], errors="coerce")
    else:
        month_col = _mapped_col(crashes, "crash_month") or _find_col(crashes, ["month"])
        if month_col:
            work["MonthNum"] = crashes[month_col].map(parse_month_value)
    if "Year" not in work or "MonthNum" not in work:
        date_col = _time_col(crashes)
        if date_col:
            dt = pd.to_datetime(crashes[date_col], errors="coerce")
            if "Year" not in work:
                work["Year"] = dt.dt.year
            if "MonthNum" not in work:
                work["MonthNum"] = dt.dt.month
    if "Year" not in work or "MonthNum" not in work:
        return pd.DataFrame(), "Month", "Crash count", "Year"
    work["Year"] = pd.to_numeric(work["Year"], errors="coerce")
    work["MonthNum"] = pd.to_numeric(work["MonthNum"], errors="coerce")
    work = work[(work["Year"] >= 1900) & (work["Year"] <= 2100) & (work["MonthNum"] >= 1) & (work["MonthNum"] <= 12)].copy()
    if work.empty:
        return pd.DataFrame(), "Month", "Crash count", "Year"
    work["Year"] = work["Year"].astype(int).astype(str)
    work["MonthNum"] = work["MonthNum"].astype(int)
    out = work.groupby(["Year", "MonthNum"], dropna=False).size().reset_index(name="Crash count")
    years = sorted(out["Year"].unique(), key=lambda x: int(x) if str(x).isdigit() else str(x))
    full = pd.MultiIndex.from_product([years, range(1, 13)], names=["Year", "MonthNum"]).to_frame(index=False)
    out = full.merge(out, on=["Year", "MonthNum"], how="left").fillna({"Crash count": 0})
    out["Crash count"] = out["Crash count"].astype(int)
    out["Month"] = out["MonthNum"].map(lambda m: MONTH_ORDER[int(m) - 1])
    return out[["Year", "Month", "Crash count"]], "Month", "Crash count", "Year"


def _year_kabco_table(crashes):
    """Return crash counts by year and normalized KABCO for stacked annual bars."""
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame(), "Year", "KABCO"
    years = _year_series_from_crashes(crashes)
    kabco_col = _kabco_col(crashes)
    if years is None or not kabco_col or kabco_col not in crashes.columns:
        return pd.DataFrame(), "Year", "KABCO"
    work = pd.DataFrame({"Year": years.reindex(crashes.index), "KABCO": crashes[kabco_col].map(normalize_kabco_value)})
    work = work.dropna(subset=["Year"])
    work = work[work["KABCO"].isin(["K", "A", "B", "C", "O"])]
    if work.empty:
        return pd.DataFrame(), "Year", "KABCO"
    out = work.groupby(["Year", "KABCO"], dropna=False).size().reset_index(name="Count")
    years_sorted = sorted(out["Year"].unique(), key=lambda x: int(x) if str(x).isdigit() else str(x))
    full = pd.MultiIndex.from_product([years_sorted, ["K", "A", "B", "C", "O"]], names=["Year", "KABCO"]).to_frame(index=False)
    out = full.merge(out, on=["Year", "KABCO"], how="left").fillna({"Count": 0})
    out["Count"] = out["Count"].astype(int)
    return out, "Year", "KABCO"


def _severity_count_columns(df):
    """Find person-count severity columns, preferring canonical Dashboard fields."""
    if df is None:
        return {}
    canonical = {
        "K": "DashboardFatalities",
        "A": "DashboardSeriousInjuries",
        "B": "DashboardMinorInjuries",
        "C": "DashboardPossibleInjuries",
        "O": "DashboardNoInjury",
    }
    found = {k: v for k, v in canonical.items() if v in df.columns}
    if any(v for v in found.values()):
        return found
    mapping = {
        "K": _mapped_col(df, "fatalities"),
        "A": _mapped_col(df, "serious_injuries"),
        "B": _mapped_col(df, "minor_injuries"),
        "C": _mapped_col(df, "possible_injuries"),
        "O": _mapped_col(df, "no_injury"),
    }
    found = {k: v for k, v in mapping.items() if v}
    if found:
        return found
    # Fallback to prior keyword logic if available from older definitions.
    result = {}
    names = {str(c).lower().replace("_", "").replace(" ", ""): c for c in df.columns}
    patterns = {
        "K": ["fatalities", "fatals", "fatalcount", "fatalinjuries", "killed"],
        "A": ["levelainjuries", "seriousinjuries", "ainjuries", "incapacitating"],
        "B": ["levelbinjuries", "binjuries", "nonincapacitating", "minorinjuries"],
        "C": ["levelcinjuries", "cinjuries", "possibleinjuries", "complaintofinjury"],
        "O": ["uninjured", "noinjury", "pdo", "propertydamageonly"],
    }
    for code, pats in patterns.items():
        for p in pats:
            for n, c in names.items():
                if p in n:
                    result[code] = c
                    break
            if code in result:
                break
    return result


def _summary_kpi_values(crashes):
    """Crash/person summary using mapped/canonical severity fields."""
    vals = {
        "Total crashes": int(len(crashes)) if crashes is not None else 0,
        "Fatal crashes": 0,
        "Fatalities": 0,
        "Serious injury crashes": 0,
        "Serious injuries": 0,
    }
    if crashes is None or getattr(crashes, "empty", True):
        return vals
    sev_cols = _severity_count_columns(crashes)
    if sev_cols.get("K"):
        fatal_vals = pd.to_numeric(crashes[sev_cols["K"]], errors="coerce").fillna(0)
        vals["Fatalities"] = int(fatal_vals.sum())
        vals["Fatal crashes"] = _unique_crash_count(crashes, fatal_vals > 0) if "_unique_crash_count" in globals() else int((fatal_vals > 0).sum())
    if sev_cols.get("A"):
        serious_vals = pd.to_numeric(crashes[sev_cols["A"]], errors="coerce").fillna(0)
        vals["Serious injuries"] = int(serious_vals.sum())
        vals["Serious injury crashes"] = _unique_crash_count(crashes, serious_vals > 0) if "_unique_crash_count" in globals() else int((serious_vals > 0).sum())
    kabco_col = _kabco_col(crashes)
    if kabco_col and kabco_col in crashes.columns:
        k = crashes[kabco_col].map(normalize_kabco_value)
        if vals["Fatal crashes"] == 0:
            vals["Fatal crashes"] = _unique_crash_count(crashes, k == "K") if "_unique_crash_count" in globals() else int((k == "K").sum())
        if vals["Serious injury crashes"] == 0:
            vals["Serious injury crashes"] = _unique_crash_count(crashes, k == "A") if "_unique_crash_count" in globals() else int((k == "A").sum())
        if vals["Fatalities"] == 0:
            vals["Fatalities"] = vals["Fatal crashes"]
        if vals["Serious injuries"] == 0:
            vals["Serious injuries"] = vals["Serious injury crashes"]
    return vals


# --- V36 FARS, KPI, severity-summary, and report-image fallbacks ---------------
def _is_fars_fatal_only_dataset(df):
    if df is None or getattr(df, "empty", True):
        return False
    if "DashboardFatalOnlySource" in df.columns:
        try:
            return bool(pd.Series(df["DashboardFatalOnlySource"]).fillna(False).astype(bool).any())
        except Exception:
            return True
    if "CrashSource" in df.columns:
        try:
            return df["CrashSource"].dropna().astype(str).str.upper().eq("FARS").any()
        except Exception:
            return True
    try:
        src = st.session_state.get("crash_source_label", "")
        if "FARS" in str(src).upper():
            return True
    except Exception:
        pass
    return False


def _format_kpi_value(value):
    if value is None:
        return "N/A"
    if isinstance(value, str):
        return value
    try:
        if pd.isna(value):
            return "N/A"
    except Exception:
        pass
    try:
        return f"{int(round(float(value))):,}"
    except Exception:
        return str(value)


def _summary_kpi_values(crashes):
    """Crash/person summary using mapped/canonical severity fields.

    Total crashes, fatal crashes, and serious-injury crashes count unique
    crash/case IDs when an ID field exists. Fatalities and serious injuries
    sum the person-count fields selected in Crash field mapping.

    If the dataset is explicitly marked as FARS fatal-only, serious-injury
    metrics are shown as N/A only when no serious-injury column was mapped.
    """
    vals = {
        "Total crashes": 0,
        "Fatal crashes": 0,
        "Fatalities": 0,
        "Serious injury crashes": 0,
        "Serious injuries": 0,
    }
    if crashes is None or getattr(crashes, "empty", True):
        return vals

    total = _unique_crash_count(crashes) if "_unique_crash_count" in globals() else int(len(crashes))
    vals["Total crashes"] = int(total)

    sev_cols = _severity_count_columns(crashes)
    mapped_fatal_col = _mapped_col(crashes, "fatalities")
    mapped_serious_col = _mapped_col(crashes, "serious_injuries")

    fatal_col = mapped_fatal_col or sev_cols.get("K")
    serious_col = mapped_serious_col or sev_cols.get("A")

    fatal_vals = None
    if fatal_col and fatal_col in crashes.columns:
        fatal_vals = pd.to_numeric(crashes[fatal_col], errors="coerce").fillna(0)
        vals["Fatalities"] = int(fatal_vals.sum())
        vals["Fatal crashes"] = (
            _unique_crash_count(crashes, fatal_vals > 0)
            if "_unique_crash_count" in globals()
            else int((fatal_vals > 0).sum())
        )

    serious_vals = None
    if serious_col and serious_col in crashes.columns:
        serious_vals = pd.to_numeric(crashes[serious_col], errors="coerce").fillna(0)
        vals["Serious injuries"] = int(serious_vals.sum())
        vals["Serious injury crashes"] = (
            _unique_crash_count(crashes, serious_vals > 0)
            if "_unique_crash_count" in globals()
            else int((serious_vals > 0).sum())
        )

    kabco_col = _kabco_col(crashes)
    if kabco_col and kabco_col in crashes.columns:
        k = crashes[kabco_col].map(normalize_kabco_value)
        if vals["Fatal crashes"] == 0:
            vals["Fatal crashes"] = (
                _unique_crash_count(crashes, k == "K")
                if "_unique_crash_count" in globals()
                else int((k == "K").sum())
            )
        if vals["Serious injury crashes"] == 0:
            vals["Serious injury crashes"] = (
                _unique_crash_count(crashes, k == "A")
                if "_unique_crash_count" in globals()
                else int((k == "A").sum())
            )
        if vals["Fatalities"] == 0 and not mapped_fatal_col:
            vals["Fatalities"] = vals["Fatal crashes"]
        if vals["Serious injuries"] == 0 and not mapped_serious_col:
            vals["Serious injuries"] = vals["Serious injury crashes"]

    if _is_fars_fatal_only_dataset(crashes):
        # FARS Accident rows are fatal crashes. However, if a Fatalities field
        # is mapped, use that field for both fatality persons and the fatal
        # crash flag instead of blindly making Fatal crashes = Total crashes.
        if not fatal_col:
            vals["Fatal crashes"] = int(total)
            vals["Fatalities"] = int(total)
        elif vals["Fatal crashes"] == 0 and vals["Fatalities"] == 0:
            vals["Fatal crashes"] = int(total)
            vals["Fatalities"] = int(total)

        # Serious-injury fields are not part of the FARS Accident fatal-only
        # table. Show N/A only when the user did not map/select a serious-
        # injury count column. If they selected one, show its numeric result.
        if not mapped_serious_col:
            vals["Serious injury crashes"] = "N/A"
            vals["Serious injuries"] = "N/A"

    return vals


def _render_kpi_strip(st, crashes):
    vals = _summary_kpi_values(crashes)
    st.markdown("<div class='dashboard-section-title'>Crash summary <span>selected study period and filters</span></div>", unsafe_allow_html=True)
    c1, c2, c3, c4, c5 = st.columns(5)
    icons = ["🚗", "🛑", "💔", "⚠️", "🏥"]
    for col, (label, value), icon in zip([c1, c2, c3, c4, c5], vals.items(), icons):
        with col:
            st.markdown(
                f"""
                <div style='border:1px solid #e5e7eb;border-radius:14px;padding:13px 14px;background:#ffffff;box-shadow:0 1px 4px rgba(15,23,42,.05)'>
                  <div style='font-size:1.45rem'>{icon}</div>
                  <div style='font-size:.80rem;color:#64748b'>{html.escape(label)}</div>
                  <div style='font-size:1.65rem;font-weight:700;color:#0f172a'>{html.escape(_format_kpi_value(value))}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )


def _spatial_unit_severity_table(tables, top_n=20):
    """Each row is one spatial unit; K/A/B/C/O columns contain counts.

    If true person-count injury columns are mapped, sum those counts.  If the
    dataset only has one KABCO/severity value per crash, count records by that
    KABCO value.  FARS is fatal-only, so it will show K counts only.
    """
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame()
    unit_col = _unit_col(crashes) or _normal_col(crashes, ["UnitID", "IntersectionID", "CorridorID", "SegmentID", "SegID"])
    if not unit_col:
        return pd.DataFrame()
    work = _drop_geometry(crashes).copy()
    sev_cols = _severity_count_columns(work)
    use_person_counts = False
    if sev_cols and not _is_fars_fatal_only_dataset(work):
        for col in sev_cols.values():
            if col in work.columns and pd.to_numeric(work[col], errors="coerce").fillna(0).sum() > 0:
                use_person_counts = True
                break
    if use_person_counts:
        agg = work[[unit_col] + list(sev_cols.values())].copy()
        for code, col in sev_cols.items():
            agg[code] = pd.to_numeric(agg[col], errors="coerce").fillna(0)
        rows = agg.groupby(unit_col, dropna=False)[[c for c in ["K", "A", "B", "C", "O"] if c in agg.columns]].sum().reset_index()
    else:
        kabco = _kabco_col(work)
        if not kabco or kabco not in work.columns:
            return pd.DataFrame()
        work["KABCO_Normalized"] = work[kabco].map(normalize_kabco_value)
        rows = work.groupby([unit_col, "KABCO_Normalized"], dropna=False).size().reset_index(name="Count")
        rows = rows[rows["KABCO_Normalized"].isin(["K", "A", "B", "C", "O"])]
        if rows.empty:
            return pd.DataFrame()
        rows = rows.pivot_table(index=unit_col, columns="KABCO_Normalized", values="Count", aggfunc="sum", fill_value=0).reset_index()
    for code in ["K", "A", "B", "C", "O"]:
        if code not in rows.columns:
            rows[code] = 0
    rows["Total"] = rows[["K", "A", "B", "C", "O"]].sum(axis=1)
    rows = rows.sort_values("Total", ascending=False).head(top_n).reset_index(drop=True)
    rows.insert(0, "Rank", range(1, len(rows) + 1))
    rows = rows.rename(columns={unit_col: "Spatial unit id"})
    return rows[["Rank", "Spatial unit id", "K", "A", "B", "C", "O", "Total"]]


def _matplotlib_figure_to_png_bytes(fig):
    """Best-effort static image fallback when Plotly/Kaleido is unavailable."""
    try:
        import matplotlib.pyplot as plt
        import numpy as np
        traces = list(getattr(fig, "data", []) or [])
        if not traces:
            return None
        layout = getattr(fig, "layout", None)
        title = ""
        try:
            title = layout.title.text or ""
        except Exception:
            title = ""
        mfig, ax = plt.subplots(figsize=(10.5, 5.8))
        first_type = getattr(traces[0], "type", "")
        if first_type == "heatmap":
            z = np.array(traces[0].z, dtype=float)
            im = ax.imshow(z, aspect="auto")
            x = list(traces[0].x) if traces[0].x is not None else list(range(z.shape[1]))
            y = list(traces[0].y) if traces[0].y is not None else list(range(z.shape[0]))
            ax.set_xticks(range(len(x))); ax.set_xticklabels([str(v) for v in x], rotation=35, ha="right")
            ax.set_yticks(range(len(y))); ax.set_yticklabels([str(v) for v in y])
            mfig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
        elif first_type == "pie":
            labels = [str(v) for v in traces[0].labels]
            values = [float(v) for v in traces[0].values]
            ax.pie(values, labels=labels, autopct=lambda p: f"{p:.0f}%" if p >= 4 else "")
            ax.axis("equal")
        elif first_type == "treemap":
            labels = [str(v) for v in traces[0].labels]
            values = [float(v) for v in traces[0].values]
            pairs = sorted(zip(labels, values), key=lambda x: x[1], reverse=True)[:15]
            ax.barh([p[0] for p in pairs][::-1], [p[1] for p in pairs][::-1])
            ax.set_xlabel("Count")
        elif first_type == "scatter":
            for tr in traces:
                x = list(tr.x) if tr.x is not None else []
                y = list(tr.y) if tr.y is not None else []
                mode = str(getattr(tr, "mode", "") or "")
                name = str(getattr(tr, "name", "") or "")
                if "lines" in mode:
                    ax.plot(x, y, marker="o", label=name if name else None)
                else:
                    size = getattr(getattr(tr, "marker", None), "size", None)
                    if size is None:
                        ax.scatter(x, y, label=name if name else None)
                    else:
                        sizes = np.array(size, dtype=float)
                        ax.scatter(x, y, s=np.maximum(sizes, 20), alpha=0.75, label=name if name else None)
        else:
            # Bar fallback supports stacked/grouped Plotly Express bars.
            horizontal = False
            try:
                horizontal = str(getattr(traces[0], "orientation", "") or "").lower() == "h"
            except Exception:
                horizontal = False
            if horizontal:
                labels = [str(v) for v in traces[0].y]
                left = np.zeros(len(labels))
                for tr in traces:
                    vals = np.array(tr.x, dtype=float)
                    ax.barh(labels, vals, left=left, label=str(getattr(tr, "name", "") or ""))
                    left += vals
            else:
                labels = [str(v) for v in traces[0].x]
                bottom = np.zeros(len(labels))
                for tr in traces:
                    vals = np.array(tr.y, dtype=float)
                    ax.bar(labels, vals, bottom=bottom, label=str(getattr(tr, "name", "") or ""))
                    bottom += vals
                ax.tick_params(axis="x", rotation=35)
        if title:
            ax.set_title(title)
        try:
            if any(str(getattr(t, "name", "") or "") for t in traces):
                ax.legend(loc="best", fontsize=8)
        except Exception:
            pass
        mfig.tight_layout()
        buf = io.BytesIO()
        mfig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
        plt.close(mfig)
        buf.seek(0)
        return buf.getvalue()
    except Exception:
        return None


def _figure_to_png_bytes(fig):
    try:
        return pio.to_image(_polish_figure(fig), format="png", width=1200, height=720, scale=2)
    except Exception:
        return _matplotlib_figure_to_png_bytes(fig)


# --- V37 structured report, dynamic filename, and methodology sections --------
def _clean_report_text(value, default="Unknown"):
    try:
        text = str(value).strip()
    except Exception:
        text = ""
    if not text or text.lower() in ["none", "nan", "null"]:
        return default
    # OSM display names can be very long. Keep the city/study-area readable.
    if "," in text and len(text) > 70:
        text = text.split(",")[0].strip()
    return text


def _report_study_area_name(tables=None):
    for key in ["area_name", "osm_place_query", "selected_city_name", "city_name"]:
        try:
            val = st.session_state.get(key, None)
            if val:
                return _clean_report_text(val, "Study Area")
        except Exception:
            pass
    try:
        density = (tables or {}).get("Crash density results")
        if density is not None and "City" in density.columns:
            vals = density["City"].dropna().astype(str).unique().tolist()
            vals = [v for v in vals if v and v.lower() not in ["nan", "none"]]
            if vals:
                return _clean_report_text(vals[0], "Study Area")
    except Exception:
        pass
    return "Study Area"


def _report_analysis_type(tables=None):
    try:
        val = st.session_state.get("analysis_type", None)
        if val:
            return _clean_report_text(val, "Safety Analysis")
    except Exception:
        pass
    try:
        density = (tables or {}).get("Crash density results")
        if density is not None and "UnitType" in density.columns:
            vals = density["UnitType"].dropna().astype(str).unique().tolist()
            if vals:
                return _clean_report_text(vals[0], "Safety Analysis")
    except Exception:
        pass
    return "Safety Analysis"


def _report_unit_phrase(tables=None):
    atype = _report_analysis_type(tables).lower()
    if "intersection" in atype:
        return "signalized intersection"
    if "corridor" in atype:
        return "corridor"
    if "segment" in atype or "road" in atype:
        return "road segment"
    return "spatial unit"


def _report_title(tables=None):
    city = _report_study_area_name(tables)
    unit = _report_unit_phrase(tables)
    if unit == "signalized intersection":
        return f"{city} Signalized Intersection Safety Analysis Report"
    if unit == "corridor":
        return f"{city} Corridor Safety Analysis Report"
    if unit == "road segment":
        return f"{city} Road Segment Safety and HIN Screening Report"
    return f"{city} Safety Analysis Report"


def _report_docx_filename(tables=None):
    name = _report_title(tables).lower()
    safe = "".join(ch if ch.isalnum() else "_" for ch in name)
    safe = "_".join([part for part in safe.split("_") if part])
    return f"{safe}.docx"


def _report_user_email():
    for key in ["user_email", "auth_email", "auth_user"]:
        try:
            val = st.session_state.get(key, None)
            if val:
                return _clean_report_text(val, "Not available")
        except Exception:
            pass
    return "Not available"


def _road_source_for_report():
    try:
        val = st.session_state.get("road_source_label", None)
        if val:
            return str(val)
    except Exception:
        pass
    if st.session_state.get("osm_raw_roads") is not None:
        return "OSM roads downloaded from OpenStreetMap"
    if st.session_state.get("tiger_roads_file") is not None:
        return "TIGER roads + PLACE boundary"
    if st.session_state.get("selected_roads") is not None:
        return "Uploaded/custom or previously prepared road network"
    return "Not available"


def _signal_source_for_report():
    try:
        val = st.session_state.get("signal_source_label", None)
        if val:
            return str(val)
    except Exception:
        pass
    sigs = st.session_state.get("signals_clean", None)
    if sigs is not None and not getattr(sigs, "empty", True):
        return "OSM traffic signals or uploaded signal points"
    return "Not available / not used"


def _crash_source_for_report(crashes=None):
    try:
        val = st.session_state.get("crash_source_label", None)
        if val:
            return str(val)
    except Exception:
        pass
    if crashes is not None and _is_fars_fatal_only_dataset(crashes):
        return "FARS Accident CSV fatal-crash data"
    if crashes is not None and not getattr(crashes, "empty", True):
        return "Uploaded crash file"
    return "Not available"


def _filter_summary_for_report():
    parts = []
    try:
        road_col = st.session_state.get("analysis_road_class_col", None)
        road_vals = st.session_state.get("analysis_road_class_values", None)
        if road_col and road_vals:
            parts.append(f"Road class filter: {road_col} in {list(road_vals)}")
    except Exception:
        pass
    try:
        mapping = st.session_state.get("crash_field_mapping", None)
        if mapping:
            used = {k: v for k, v in dict(mapping).items() if v}
            if used:
                parts.append("Confirmed crash field mapping was used for dashboard and report summaries.")
    except Exception:
        pass
    return parts or ["No optional filters were documented in the dashboard session."]


def _add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Description"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)
    return table


def _add_dataframe_table(doc, df, max_rows=20):
    table_df = _safe_dataframe_for_display(df.head(max_rows).copy())
    if table_df.empty:
        return False
    table = doc.add_table(rows=1, cols=len(table_df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(table_df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(table_df.columns):
            cells[i].text = str(row[col])
    return True


def _report_introduction_text(tables=None):
    city = _report_study_area_name(tables)
    unit = _report_unit_phrase(tables)
    if unit == "signalized intersection":
        return (
            f"This report summarizes signalized-intersection safety conditions for {city}. "
            "The analysis uses crash records assigned to signalized intersection spatial units and summarizes crash density, crash patterns, severity, and priority locations."
        )
    if unit == "corridor":
        return (
            f"This report summarizes corridor safety conditions for {city}. "
            "The analysis uses road corridors defined from selected roads and traffic-signal context, assigns crashes to corridor spatial units, and summarizes crash density, severity, and priority corridors."
        )
    if unit == "road segment":
        return (
            f"This report summarizes road-segment safety conditions for {city}. "
            "The analysis includes crash-density screening and, when Sliding Window/HIN results are available, HIN priority index results for identifying high-priority segments or windows."
        )
    return (
        f"This report summarizes spatial-unit safety conditions for {city}. "
        "It includes crash patterns, crash density, maps, and decision-ready priority tables from the dashboard."
    )


def _methodology_texts(tables=None):
    unit = _report_unit_phrase(tables)
    common = [
        "Load or download the road network, study-area boundary, traffic-signal data when needed, and crash records.",
        "Generate route mileposts using FromMile and ToMile so road segments and corridors have consistent linear reference fields.",
        "Optionally filter the roadway network by road class/type before later steps. If a road-class filter is enabled, later spatial-unit creation uses the filtered road network.",
        "Confirm crash field mapping so the dashboard and report use the correct crash ID, date/year, crash type, KABCO/severity, injury-count, and mode fields.",
        "Optionally filter crashes by available attributes such as year, crash type, or severity before assigning them to spatial units.",
    ]
    if unit == "signalized intersection":
        specific = [
            "Generate or upload traffic-signal points. OSM signals are de-duplicated using the selected duplicate-distance threshold and snapped/filtered to nearby selected roads using the selected road-distance threshold.",
            "Create signalized-intersection spatial units around cleaned signal points using the configured intersection buffer distance.",
            "Spatially assign filtered crashes to the signalized-intersection units.",
            "Calculate crash counts and crash density for each signalized intersection and prepare ranking tables and maps."
        ]
    elif unit == "corridor":
        specific = [
            "Generate or upload traffic-signal points and identify candidate corridor signals.",
            "Build corridors from selected roads and nearby/associated signals using the configured minimum-signal, search-buffer, and corridor-width settings.",
            "Allow corridor review and optional corridor dropping before final corridor spatial units are used.",
            "Spatially assign filtered crashes to final corridor units and calculate crash counts and crash density for each corridor."
        ]
    elif unit == "road segment":
        specific = [
            "Create road-segment spatial units from uploaded/existing road segments, equal-length segments, or selected segment-generation settings.",
            "Spatially assign filtered crashes to segment units and calculate crash count and crash density.",
            "When Sliding Window/HIN is run, move a fixed-length analysis window along each route at the selected step length.",
            "Score each window using the selected metric such as crash count or EPDO, then transfer the maximum overlapping window score to output segments/windows.",
            "Rank segments/windows using the HIN priority index and selected threshold settings."
        ]
    else:
        specific = [
            "Create the selected spatial units, assign filtered crashes to those units, calculate crash density, and prepare dashboard results."
        ]
    return common + specific


def _add_data_section(doc, tables, crashes):
    doc.add_heading("Data", level=1)
    roads = st.session_state.get("selected_roads", None)
    signals = st.session_state.get("signals_clean", None)
    density = tables.get("Crash density results")
    rows = [
        ("Road data", f"{_road_source_for_report()}; selected road features: {len(roads):,}" if roads is not None else _road_source_for_report()),
        ("Signal data", f"{_signal_source_for_report()}; signal points: {len(signals):,}" if signals is not None and not getattr(signals, "empty", True) else _signal_source_for_report()),
        ("Crash data", f"{_crash_source_for_report(crashes)}; records used: {len(crashes):,}" if crashes is not None else _crash_source_for_report(crashes)),
        ("Spatial-unit results", f"{len(density):,} units with crash-density results" if density is not None else "Not available"),
    ]
    _add_key_value_table(doc, rows)
    doc.add_paragraph("Required road fields depend on the selected method. Uploaded road networks should include valid line geometry, a route/name field, and a stable segment ID field. The app can auto-detect common fields but users should confirm route and ID fields when prompted.")
    doc.add_paragraph("Uploaded signal data, if used, should include a signal ID field and latitude/longitude fields. OSM signal generation does not require an upload, but the result should be reviewed because OSM signal coverage varies by location.")
    doc.add_paragraph("Uploaded crash data should include a crash/case ID, latitude, longitude, crash date or year, crash type or manner of collision, and severity/KABCO or injury-count fields. FARS Accident data is fatal-crash-only and uses the FARS fatality count field for total people killed.")
    for item in _filter_summary_for_report():
        doc.add_paragraph(item, style=None)


def _add_methodology_section(doc, tables):
    doc.add_heading("Methodology", level=1)
    for i, text in enumerate(_methodology_texts(tables), start=1):
        doc.add_paragraph(f"{i}. {text}")


def _add_limitations_section(doc):
    doc.add_heading("Limitations and data-quality notes", level=1)
    notes = [
        "OSM traffic-signal points are contributed data and may be incomplete, outdated, duplicated, or offset from the actual intersection location. Cleaned signal results should be reviewed before final use.",
        "OSM road classifications are based on OSM highway tags, not necessarily an official functional-class system. Local agency roadway data is preferred when official classification is required.",
        "TIGER roads provide broad coverage but may not include the same level of local roadway detail, lane information, or classification accuracy as agency-maintained centerlines.",
        "Uploaded crash/FARS files vary by agency and format. The Crash Field Mapping panel should be confirmed so the app uses the correct ID, severity, crash type, date, and injury-count fields.",
        "FARS Accident data includes fatal crashes only. It should not be interpreted as a complete all-crash dataset and does not support serious-injury crash totals from the Accident table alone.",
        "Large datasets can exceed browser, memory, upload, or Streamlit Cloud limits. Large cities or statewide/national datasets should be clipped to the study area and simplified for map display before final dashboard/report export.",
        "Crash-density and HIN results depend on the selected spatial-unit definition, filters, thresholds, crash geocoding accuracy, and roadway segmentation method. Results should be reviewed with engineering judgment."
    ]
    for note in notes:
        doc.add_paragraph(note)


def _add_bubble_size_note_if_needed(doc, figures):
    titles = [t for t, _, _ in figures]
    if any("bubble" in str(t).lower() for t in titles):
        doc.add_paragraph(
            "Travel mode severity bubble chart note: BubbleSize is calculated as 8 + 8 × sqrt(Count). "
            "This square-root scaling keeps small mode/severity categories visible while still making larger categories appear larger. "
            "The actual crash count is reported in the table and chart tooltip."
        )


def _export_dashboard_docx(tables, selected_blocks, selected_maps, extra_figures=None, maps=None, overlay_layers=None, report_timezone=None):
    if Document is None:
        return None
    doc = Document()
    title = _report_title(tables)
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {_report_time_text(report_timezone)}")
    doc.add_paragraph(f"User email: {_report_user_email()}")

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(_report_introduction_text(tables))

    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    _add_data_section(doc, tables, crashes)
    _add_methodology_section(doc, tables)
    _add_limitations_section(doc)

    doc.add_heading("Results and visualization", level=1)
    kpis = _summary_kpi_values(crashes)
    doc.add_heading("Crash summary", level=2)
    _add_key_value_table(doc, [(k, _format_kpi_value(v)) for k, v in kpis.items()])

    figures = _build_default_figures(tables) + (extra_figures or [])
    _add_bubble_size_note_if_needed(doc, figures)
    for fig_title, fig, data in figures:
        if selected_blocks and fig_title not in selected_blocks:
            continue
        doc.add_heading(str(fig_title), level=2)
        img = _figure_to_png_bytes(_polish_figure(fig))
        if img:
            doc.add_picture(io.BytesIO(img), width=Inches(6.5))
        else:
            doc.add_paragraph("Chart image could not be generated in this environment. The summary table is included below.")
        table_df = _safe_dataframe_for_display(data.copy())
        if not table_df.empty:
            doc.add_paragraph("Summary table")
            _add_dataframe_table(doc, table_df, max_rows=20)

    if selected_maps:
        doc.add_heading("Selected map layers", level=2)
        for m in selected_maps:
            doc.add_heading(str(m), level=3)
            if maps and m in maps:
                map_png = _static_map_png(maps[m], str(m), overlay_layers=overlay_layers)
                if map_png:
                    doc.add_picture(io.BytesIO(map_png), width=Inches(6.5))
                else:
                    doc.add_paragraph("Static map image could not be generated.")
            else:
                doc.add_paragraph("Map layer selected in dashboard builder.")

    doc.add_heading("Decision-ready result tables", level=2)
    for table_name, df in _report_tables(tables).items():
        doc.add_heading(str(table_name), level=3)
        if not _add_dataframe_table(doc, df, max_rows=25):
            doc.add_paragraph("No records available.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# --- Dashboard-only HIN visualization and summary additions ------------------
# These overrides add dashboard cards, plots, map choices, and report-table
# columns. They do not change any workflow calculation or stored analysis result.

def _dashboard_year_count(crashes):
    """Return number of crash years for annualized dashboard rates."""
    try:
        years = _year_series_from_crashes(crashes)
        if years is None:
            return 1.0
        vals = pd.to_numeric(years, errors="coerce").dropna()
        if vals.empty:
            return 1.0
        return float(max(int(vals.max()) - int(vals.min()) + 1, 1))
    except Exception:
        return 1.0


def _dashboard_length_col(df):
    return _normal_col(
        df,
        [
            "Length_Miles",
            "Length_Mi",
            "SegmentLength_Mile",
            "WindowLength_Miles",
            "CorridorLength_Mile",
            "CorridorLength_Miles",
            "length_mi",
            "Miles",
        ],
    )


def _dashboard_hin_metric(hin):
    if hin is None or getattr(hin, "empty", True):
        return None
    if "HIN_Priority_Index" in hin.columns:
        return "HIN_Priority_Index"
    return _default_metric(_numeric_cols(hin))


def _hin_ka_series(df):
    """Return a row-level K+A/KSI series when available in HIN-style tables."""
    if df is None or getattr(df, "empty", True):
        return pd.Series(dtype="float64")

    combined_candidates = [
        "KA_Crashes",
        "K_A_Crashes",
        "KSI_Count",
        "KSI_Crashes",
        "Fatal_Injury_Count",
        "FatalAndSeriousInjuryCrashes",
        "Fatal_Serious_Injury_Count",
        "Fatal_Serious_Crashes",
    ]
    for col in combined_candidates:
        if col in df.columns:
            return pd.to_numeric(df[col], errors="coerce").fillna(0)

    k_col = _normal_col(
        df,
        [
            "K_Crashes",
            "K_Count",
            "Fatal_Crashes",
            "FatalCrashCount",
            "Fatal_Count",
        ],
    )
    a_col = _normal_col(
        df,
        [
            "A_Crashes",
            "A_Count",
            "Serious_Injury_Crashes",
            "SeriousInjuryCrashCount",
            "Serious_Injury_Count",
        ],
    )
    if k_col or a_col:
        k = pd.to_numeric(df[k_col], errors="coerce").fillna(0) if k_col else 0
        a = pd.to_numeric(df[a_col], errors="coerce").fillna(0) if a_col else 0
        return k + a

    return pd.Series([0] * len(df), index=df.index, dtype="float64")


def _total_ka_from_crashes(crashes):
    if crashes is None or getattr(crashes, "empty", True):
        return 0
    try:
        sev_cols = _severity_count_columns(crashes)
        if sev_cols:
            k = pd.to_numeric(crashes[sev_cols.get("K")], errors="coerce").fillna(0).sum() if sev_cols.get("K") else 0
            a = pd.to_numeric(crashes[sev_cols.get("A")], errors="coerce").fillna(0).sum() if sev_cols.get("A") else 0
            return int(k + a)
        kabco = _kabco_col(crashes)
        if kabco and kabco in crashes.columns:
            vals = crashes[kabco].map(normalize_kabco_value).astype(str).str.upper()
            return int(vals.isin(["K", "A"]).sum())
    except Exception:
        pass
    return 0


def _selected_hin_subset(hin, metric, method, top_percent=10.0, top_n=20, index_threshold=50.0):
    """Select a dashboard HIN subset without altering workflow results."""
    if hin is None or getattr(hin, "empty", True) or metric not in hin.columns:
        return pd.DataFrame()
    work = hin.copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    length_col = _dashboard_length_col(work)
    if length_col:
        work[length_col] = pd.to_numeric(work[length_col], errors="coerce").fillna(0)
    else:
        length_col = "__unit_length__"
        work[length_col] = 1.0
    sorted_work = work.sort_values(metric, ascending=False).copy()
    total_mi = float(sorted_work[length_col].sum()) if sorted_work[length_col].sum() else float(len(sorted_work))

    if method == "Top percent of miles":
        limit = total_mi * float(top_percent) / 100.0
        selected = sorted_work[sorted_work[length_col].cumsum() <= limit].copy()
        if selected.empty and not sorted_work.empty:
            selected = sorted_work.head(1).copy()
        return selected
    if method == "Top number of segments/windows":
        return sorted_work.head(int(top_n)).copy()
    if method == "HIN index threshold":
        return sorted_work[sorted_work[metric] >= float(index_threshold)].copy()
    if method == "Above average HIN index":
        avg = pd.to_numeric(sorted_work[metric], errors="coerce").mean()
        return sorted_work[sorted_work[metric] >= avg].copy()
    if method == "Above median HIN index":
        med = pd.to_numeric(sorted_work[metric], errors="coerce").median()
        return sorted_work[sorted_work[metric] >= med].copy()
    return sorted_work.head(20).copy()


def _add_dashboard_rate_columns(df, crashes=None):
    """Add display-only crash/mile/year fields to a result copy."""
    if df is None or getattr(df, "empty", True):
        return df
    out = df.copy()
    length_col = _dashboard_length_col(out)
    count_col = _crash_count_col(out)
    years = _dashboard_year_count(crashes)
    if length_col and count_col:
        length = pd.to_numeric(out[length_col], errors="coerce").replace(0, pd.NA)
        count = pd.to_numeric(out[count_col], errors="coerce").fillna(0)
        out["Crash_per_Mile"] = (count / length).astype("float64").round(3)
        out["Crash_per_Mile_per_Year"] = (count / length / years).astype("float64").round(3)
    ka = _hin_ka_series(out)
    if len(ka) == len(out) and length_col:
        length = pd.to_numeric(out[length_col], errors="coerce").replace(0, pd.NA)
        out["KA_Crashes"] = pd.to_numeric(ka, errors="coerce").fillna(0).round(3)
        out["KA_per_Mile_per_Year"] = (pd.to_numeric(ka, errors="coerce").fillna(0) / length / years).astype("float64").round(3)
    return out


def _hin_distribution_figures(hin):
    figures = []
    metric = _dashboard_hin_metric(hin)
    if hin is None or getattr(hin, "empty", True) or not metric:
        return figures
    work = _drop_geometry(hin).copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce")
    work = work[work[metric].notna()].copy()
    if work.empty:
        return figures

    fig = px.histogram(
        work,
        x=metric,
        nbins=20,
        title="HIN index distribution",
        labels={metric: "HIN priority index"},
    )
    fig.update_layout(xaxis_title="HIN priority index", yaxis_title="Segment/window count")
    figures.append(("HIN index distribution", _polish_figure(fig), work[[metric]].copy()))

    route_col = _dashboard_route_col(work)
    length_col = _dashboard_length_col(work)
    if route_col:
        route_work = work.copy()
        route_work[route_col] = route_work[route_col].fillna("Unknown").astype(str)
        if length_col:
            route_work[length_col] = pd.to_numeric(route_work[length_col], errors="coerce").fillna(0)
            route_summary = route_work.groupby(route_col, dropna=False).agg(
                Mean_HIN=(metric, "mean"),
                Max_HIN=(metric, "max"),
                Segment_Count=(metric, "size"),
                Miles=(length_col, "sum"),
            ).reset_index()
        else:
            route_summary = route_work.groupby(route_col, dropna=False).agg(
                Mean_HIN=(metric, "mean"),
                Max_HIN=(metric, "max"),
                Segment_Count=(metric, "size"),
            ).reset_index()
        route_summary = route_summary.sort_values("Mean_HIN", ascending=False).head(15)
        if not route_summary.empty:
            fig = px.bar(
                route_summary.sort_values("Mean_HIN", ascending=True),
                y=route_col,
                x="Mean_HIN",
                orientation="h",
                hover_data=[c for c in ["Max_HIN", "Segment_Count", "Miles"] if c in route_summary.columns],
                title="Average HIN index by route",
            )
            fig.update_layout(yaxis_title="Route", xaxis_title="Average HIN priority index")
            figures.append(("Average HIN index by route", _polish_figure(fig), route_summary))
    return figures


def _hin_ka_bubble_figure(hin):
    metric = _dashboard_hin_metric(hin)
    if hin is None or getattr(hin, "empty", True) or not metric:
        return None, pd.DataFrame()
    work = _drop_geometry(hin).copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    count_col = _crash_count_col(work)
    if count_col:
        work["Total_Crashes"] = pd.to_numeric(work[count_col], errors="coerce").fillna(0)
    else:
        work["Total_Crashes"] = 0
    work["KA_Crashes"] = pd.to_numeric(_hin_ka_series(work), errors="coerce").fillna(0)
    if work["KA_Crashes"].sum() <= 0:
        return None, pd.DataFrame()
    route_col = _dashboard_route_col(work)
    id_col = _normal_col(work, ["RiskSegmentID", "WindowID", "SlidingWindowID", "SegmentID", "UnitID", "SourceSegmentID", "CorridorID"])
    if id_col is None:
        id_col = "DashboardUnitID"
        work[id_col] = [f"HIN_{i + 1}" for i in range(len(work))]
    color_col = None
    if route_col and work[route_col].nunique(dropna=True) <= 15:
        color_col = route_col
    hover_cols = [c for c in [id_col, route_col, _dashboard_length_col(work), "Total_Crashes", "KA_Crashes"] if c and c in work.columns]
    fig = px.scatter(
        work,
        x=metric,
        y="KA_Crashes",
        size="Total_Crashes" if work["Total_Crashes"].sum() > 0 else None,
        color=color_col,
        hover_data=hover_cols,
        size_max=42,
        title="K+A crashes vs HIN priority index",
    )
    fig.update_traces(marker=dict(sizemin=6, opacity=0.78, line=dict(width=1, color="white")))
    fig.update_layout(xaxis_title="HIN priority index", yaxis_title="K+A / KSI crashes")
    return _polish_figure(fig), work


def _render_hin_dashboard_charts(st, hin):
    if hin is None or getattr(hin, "empty", True):
        return
    st.markdown("<div class='dashboard-section-title'>HIN index diagnostics <span>distribution, route patterns, and severity relationship</span></div>", unsafe_allow_html=True)
    figures = _hin_distribution_figures(hin)
    bubble_fig, bubble_df = _hin_ka_bubble_figure(hin)
    if not figures and bubble_fig is None:
        st.info("HIN diagnostic charts need HIN priority index values and, for the bubble chart, available K+A/KSI fields.")
        return
    cols = st.columns(2)
    chart_items = figures[:]
    if bubble_fig is not None:
        chart_items.append(("K+A crashes vs HIN priority index", bubble_fig, bubble_df))
    for i, (title, fig, data) in enumerate(chart_items[:4]):
        with cols[i % 2]:
            fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35))
            st.plotly_chart(_polish_figure(fig), width="stretch", key=f"hin_diag_{_safe_name(title)}_{i}")


def _render_hin_network_summary(st, hin, crashes):
    """Dashboard-only HIN summary with custom threshold and K+A card."""
    if hin is None or getattr(hin, "empty", True):
        return
    metric = _dashboard_hin_metric(hin)
    if not metric:
        return

    st.markdown("<div class='dashboard-section-title'>High Injury Network summary <span>custom HIN threshold, miles, crashes, and K+A/KSI capture</span></div>", unsafe_allow_html=True)
    work = hin.copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    length_col = _dashboard_length_col(work)
    if length_col:
        work[length_col] = pd.to_numeric(work[length_col], errors="coerce").fillna(0)
    else:
        length_col = "__unit_length__"
        work[length_col] = 1.0

    mode_col, control_col, c_miles, c_crashes, c_ka = st.columns([1.35, 1.05, 1, 1, 1])
    method = mode_col.selectbox(
        "High-risk network threshold",
        [
            "Top percent of miles",
            "Top number of segments/windows",
            "HIN index threshold",
            "Above average HIN index",
            "Above median HIN index",
        ],
        index=0,
        key="hin_summary_threshold_mode_v38",
    )

    top_percent = 10.0
    top_n = 20
    index_threshold = 50.0
    with control_col:
        if method == "Top percent of miles":
            top_percent = st.number_input(
                "Top percent",
                min_value=1.0,
                max_value=100.0,
                value=float(st.session_state.get("hin_summary_top_percent", 10.0)),
                step=1.0,
                key="hin_summary_top_percent",
            )
        elif method == "Top number of segments/windows":
            top_n = st.number_input(
                "Top N",
                min_value=1,
                max_value=max(int(len(work)), 1),
                value=min(20, max(int(len(work)), 1)),
                step=1,
                key="hin_summary_top_n",
            )
        elif method == "HIN index threshold":
            index_threshold = st.number_input(
                "Minimum HIN index",
                min_value=0.0,
                max_value=100.0,
                value=float(st.session_state.get("hin_summary_index_threshold", 50.0)),
                step=5.0,
                key="hin_summary_index_threshold",
            )
        elif method == "Above average HIN index":
            st.metric("Average", f"{work[metric].mean():,.2f}")
        elif method == "Above median HIN index":
            st.metric("Median", f"{work[metric].median():,.2f}")

    selected = _selected_hin_subset(
        work,
        metric,
        method,
        top_percent=top_percent,
        top_n=top_n,
        index_threshold=index_threshold,
    )

    total_mi = float(work[length_col].sum()) if work[length_col].sum() else float(len(work))
    high_mi = float(selected[length_col].sum()) if not selected.empty else 0.0
    pct_mi = high_mi / total_mi * 100 if total_mi else 0.0

    count_col = _crash_count_col(work)
    total_crashes = int(pd.to_numeric(work[count_col], errors="coerce").fillna(0).sum()) if count_col else (len(crashes) if crashes is not None else 0)
    high_crashes = int(pd.to_numeric(selected[count_col], errors="coerce").fillna(0).sum()) if count_col and not selected.empty else 0
    pct_crash = high_crashes / total_crashes * 100 if total_crashes else 0.0

    total_ka_series = _hin_ka_series(work)
    selected_ka_series = _hin_ka_series(selected)
    total_ka = int(pd.to_numeric(total_ka_series, errors="coerce").fillna(0).sum()) if len(total_ka_series) else _total_ka_from_crashes(crashes)
    high_ka = int(pd.to_numeric(selected_ka_series, errors="coerce").fillna(0).sum()) if len(selected_ka_series) else 0
    pct_ka = high_ka / total_ka * 100 if total_ka else 0.0

    c_miles.metric("High-risk miles", f"{high_mi:,.2f} mi", f"{pct_mi:,.1f}% of analyzed miles")
    c_crashes.metric("Crashes on selected HIN", f"{high_crashes:,}", f"{pct_crash:,.1f}% of assigned crashes")
    c_ka.metric("K+A / KSI on selected HIN", f"{high_ka:,}", f"{pct_ka:,.1f}% of K+A/KSI")

    selected_display = _add_dashboard_rate_columns(selected, crashes=crashes)
    with st.expander("Selected HIN summary rows", expanded=False):
        preview_cols = []
        for c in [
            _normal_col(selected_display, ["RiskSegmentID", "WindowID", "SegmentID", "UnitID", "CorridorID"]),
            _dashboard_route_col(selected_display),
            _dashboard_length_col(selected_display),
            _crash_count_col(selected_display),
            "KA_Crashes",
            metric,
            "Crash_per_Mile_per_Year",
            "KA_per_Mile_per_Year",
        ]:
            if c and c in selected_display.columns and c not in preview_cols:
                preview_cols.append(c)
        if preview_cols:
            st.dataframe(_safe_dataframe_for_display(selected_display[preview_cols].head(50)), width="stretch", hide_index=True)
        else:
            st.info("No displayable HIN rows are available for the selected threshold.")

    st.caption("These dashboard controls only summarize and visualize the existing HIN results. They do not recalculate HIN scores or change workflow outputs.")
    _render_hin_dashboard_charts(st, hin)


def _hin_table_for_display(hin, metric, top_n=20):
    """Decision table for top HIN windows/segments with display-only rates."""
    if hin is None or getattr(hin, "empty", True) or metric not in hin.columns:
        return pd.DataFrame()
    crashes = None
    try:
        crashes = _available_tables(st).get("Assigned crashes", _available_tables(st).get("Uploaded crashes"))
    except Exception:
        crashes = None
    work = _add_dashboard_rate_columns(hin.copy(), crashes=crashes)
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work = work.sort_values(metric, ascending=False).head(top_n).reset_index(drop=True)

    id_col = _normal_col(work, ["RiskSegmentID", "WindowID", "SlidingWindowID", "SegmentID", "UnitID", "SourceSegmentID", "CorridorID"])
    route_col = _dashboard_route_col(work)
    length_col = _dashboard_length_col(work)
    from_col = _normal_col(work, ["FromMile", "From_Mile", "from_mile", "BeginMile", "StartMile", "WindowFromMile"])
    to_col = _normal_col(work, ["ToMile", "To_Mile", "to_mile", "EndMile", "WindowToMile"])
    count_col = _crash_count_col(work)

    out = pd.DataFrame()
    out["Rank"] = range(1, len(work) + 1)
    out["SegID"] = work[id_col].astype(str).values if id_col else [f"HIN_{i + 1}" for i in range(len(work))]
    out["Seg/window length"] = pd.to_numeric(work[length_col], errors="coerce").round(3).values if length_col else ""
    out["From mile"] = pd.to_numeric(work[from_col], errors="coerce").round(3).values if from_col else ""
    out["To mile"] = pd.to_numeric(work[to_col], errors="coerce").round(3).values if to_col else ""
    out["Route"] = work[route_col].astype(str).values if route_col else ""
    out["Crash count"] = pd.to_numeric(work[count_col], errors="coerce").fillna(0).round(0).astype(int).values if count_col else ""
    out["K+A / KSI crashes"] = pd.to_numeric(work["KA_Crashes"], errors="coerce").fillna(0).round(0).astype(int).values if "KA_Crashes" in work.columns else ""
    out["HIN index"] = pd.to_numeric(work[metric], errors="coerce").round(3).values
    if "Crash_per_Mile" in work.columns:
        out["Crashes/mile"] = pd.to_numeric(work["Crash_per_Mile"], errors="coerce").round(3).values
    if "Crash_per_Mile_per_Year" in work.columns:
        out["Crashes/mile/year"] = pd.to_numeric(work["Crash_per_Mile_per_Year"], errors="coerce").round(3).values
    if "KA_per_Mile_per_Year" in work.columns:
        out["K+A/mile/year"] = pd.to_numeric(work["KA_per_Mile_per_Year"], errors="coerce").round(3).values
    return out


def _top_density_export_table(density, top_n=20):
    """Decision-ready crash-density table with display-only rate columns."""
    if density is None or getattr(density, "empty", True):
        return pd.DataFrame()
    crashes = None
    try:
        crashes = _available_tables(st).get("Assigned crashes", _available_tables(st).get("Uploaded crashes"))
    except Exception:
        crashes = None
    df = _add_dashboard_rate_columns(_drop_geometry(density).copy(), crashes=crashes)
    unit_col = _unit_col(df) or _normal_col(df, ["UnitID", "IntersectionID", "CorridorID", "SegmentID", "Route"])
    unit_type_col = _normal_col(df, ["UnitType", "IntersectionType", "CorridorType", "SegmentType"])
    city_col = _normal_col(df, ["City", "city_name"])
    length_col = _dashboard_length_col(df)
    count_col = _crash_count_col(df)
    density_col = _normal_col(df, ["CrashDensity", "Crash_Density", "crash_density"])
    if density_col is None:
        density_col = _default_metric(_numeric_cols(df))
    if density_col:
        df[density_col] = pd.to_numeric(df[density_col], errors="coerce").fillna(0)
        df = df.sort_values(density_col, ascending=False)
    use = df.head(top_n).reset_index(drop=True)
    out = pd.DataFrame({"Rank": range(1, len(use) + 1)})
    out["Spatial unit id"] = use[unit_col].astype(str) if unit_col else use.index.astype(str)
    out["Unit type"] = use[unit_type_col].astype(str) if unit_type_col else ""
    out["City"] = use[city_col].astype(str) if city_col else ""
    out["Length_mi"] = pd.to_numeric(use[length_col], errors="coerce").round(3) if length_col else ""
    road1_col = _normal_col(use, ["RoadName1", "Road1", "Route1", "Street1", "FromRoad"])
    road2_col = _normal_col(use, ["RoadName2", "Road2", "Route2", "Street2", "ToRoad"])
    route_col = _normal_col(use, ["Route", "FULLNAME", "RoadName", "RouteName", "CorridorRoute", "RouteName_Calc"])
    if road1_col and road2_col:
        out["Road 1"] = use[road1_col].astype(str)
        out["Road 2"] = use[road2_col].astype(str)
    elif route_col:
        out["Route name"] = use[route_col].astype(str)
    out["Crash count"] = pd.to_numeric(use[count_col], errors="coerce").fillna(0).astype(int) if count_col else ""
    out["Crash density"] = pd.to_numeric(use[density_col], errors="coerce").round(3) if density_col else ""
    for src, label in [
        ("Crash_per_Mile", "Crashes/mile"),
        ("Crash_per_Mile_per_Year", "Crashes/mile/year"),
        ("KA_per_Mile_per_Year", "K+A/mile/year"),
    ]:
        if src in use.columns:
            out[label] = pd.to_numeric(use[src], errors="coerce").round(3)
    return out


def _export_tables_only(tables, top_n=20):
    """Report-ready tables with dashboard-only rate columns."""
    out = {}
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))
    if crashes is not None:
        type_col = _crash_type_col(crashes)
        if type_col:
            out["Crash type summary"] = _aggregate(crashes, type_col, None, "Count", top_n)
        years = _year_series_from_crashes(crashes)
        if years is not None:
            ydf = pd.DataFrame({"Year": years})
            ydf = ydf[ydf["Year"].ne("Unknown")]
            out["Crash year summary"] = ydf.groupby("Year", dropna=False).size().reset_index(name="Count").sort_values("Year")
        sev = _spatial_unit_severity_table(tables, top_n=top_n)
        if not sev.empty:
            out["Severity summary by spatial unit"] = sev
    if density is not None:
        top_density = _top_density_export_table(density, top_n=top_n)
        if not top_density.empty:
            out["Top crash-density spatial units"] = top_density
    if hin is not None:
        metric = _dashboard_hin_metric(hin)
        top_hin = _hin_table_for_display(hin, metric, top_n) if metric else pd.DataFrame()
        if not top_hin.empty:
            out["Top HIN/risk spatial units"] = top_hin
    return {name: _safe_dataframe_for_display(df) for name, df in out.items()}


def _available_maps(st):
    """Dashboard map layers, including display-only HIN threshold maps."""
    maps = {}
    density = st.session_state.get("spatial_units_density_map")
    if density is not None and not getattr(density, "empty", True):
        maps["Crash density map"] = _repair_gdf_crs(density, st)

    results = st.session_state.get("section7_results")
    if results is not None:
        risk_segments = results.get("risk_segments")
        if risk_segments is not None and not getattr(risk_segments, "empty", True):
            hin_map = _repair_gdf_crs(risk_segments, st)
            maps["HIN priority map"] = hin_map
            metric = _dashboard_hin_metric(hin_map)
            if metric:
                work = hin_map.copy()
                work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
                if not work.empty:
                    avg = work[metric].mean()
                    med = work[metric].median()
                    maps["HIN above average map"] = work[work[metric] >= avg].copy()
                    maps["HIN above median map"] = work[work[metric] >= med].copy()

    corridors = st.session_state.get("final_corridors", st.session_state.get("corridors"))
    if corridors is not None and not getattr(corridors, "empty", True):
        maps["Corridor map"] = _repair_gdf_crs(corridors, st)

    return {k: v for k, v in maps.items() if v is not None and not getattr(v, "empty", True)}


def _build_default_figures(tables):
    """Default dashboard/report figures with added HIN diagnostic charts."""
    figures = []
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))
    if crashes is not None:
        year_kabco, year_col, kabco_col = _year_kabco_table(crashes)
        if not year_kabco.empty:
            fig = px.bar(year_kabco, x=year_col, y="Count", color=kabco_col, color_discrete_map=KABCO_COLOR_MAP, category_orders={kabco_col: ["K", "A", "B", "C", "O"]}, title=f"Crashes by year and {kabco_col}")
            fig.update_layout(barmode="stack", xaxis_title="Year", yaxis_title="Crash count")
            figures.append((f"Crashes by year and {kabco_col}", _apply_kabco_trace_colors(_polish_figure(fig)), year_kabco))
        type_col = _crash_type_col(crashes)
        if type_col:
            type_df = _aggregate(crashes, type_col, None, "Count", 12)
            pie = px.pie(type_df, names=type_col, values="Count", hole=0.38, title=f"Crash type share by {type_col}")
            figures.append((f"Crash type share by {type_col}", _polish_figure(pie), type_df))
        monthly_df, period_col, value_col, color_col = _month_trend_table(crashes)
        if not monthly_df.empty:
            fig = px.line(monthly_df, x=period_col, y=value_col, color=color_col, markers=True, category_orders={period_col: MONTH_ORDER}, title="Monthly crash trend by year")
            figures.append(("Monthly crash trend by year", _polish_figure(fig), monthly_df))
        mode_df, mode_kabco_col, mode_col = _mode_severity_bubble_table(crashes)
        if not mode_df.empty:
            fig = px.scatter(mode_df, x=mode_kabco_col, y=mode_col, size="BubbleSize" if "BubbleSize" in mode_df.columns else "Count", color=mode_col, size_max=42, title="Travel mode severity bubble chart", hover_data={"Count": True})
            fig.update_traces(marker=dict(sizemin=7, opacity=0.80, line=dict(width=1, color="white")))
            figures.append(("Travel mode severity bubble chart", _polish_figure(fig), mode_df))
        road_kabco, road_col, road_kabco_col = _road_class_kabco_table(crashes, st_obj=None)
        if not road_kabco.empty:
            tree = road_kabco.copy(); tree["All crashes"] = "All crashes"
            fig = px.treemap(tree, path=["All crashes", road_col, road_kabco_col], values="Count", color=road_kabco_col, color_discrete_map=KABCO_COLOR_MAP, title=f"Road class and {road_kabco_col} treemap")
            figures.append((f"Road class and {road_kabco_col} treemap", _polish_figure(fig), tree))
        crash_kabco, crash_type_col, crash_kabco_col = _crash_type_kabco_table(crashes)
        if not crash_kabco.empty:
            pivot = crash_kabco.pivot_table(index=crash_type_col, columns=crash_kabco_col, values="Count", aggfunc="sum", fill_value=0)
            fig = px.imshow(pivot, text_auto=True, aspect="auto", title=f"Crash type by {crash_kabco_col}", labels=dict(x=crash_kabco_col, y="Crash type", color="Crash count"), color_continuous_scale="YlOrRd")
            figures.append((f"Crash type by {crash_kabco_col}", _polish_figure(fig), pivot.reset_index()))
    if density is not None:
        metric = "CrashDensity" if "CrashDensity" in density.columns else _default_metric(_numeric_cols(density))
        rank_df, unit_col, value_col = _rank_units_for_chart(density, metric, 15) if metric else (pd.DataFrame(), None, None)
        if not rank_df.empty:
            plot_df = rank_df.sort_values(value_col, ascending=True)
            fig = px.bar(plot_df, y=unit_col, x=value_col, orientation="h", hover_data=_context_cols_for_hover(plot_df), title="Top spatial units by crash density")
            fig.update_layout(yaxis_title="Spatial unit ID", xaxis_title="Crash density")
            figures.append(("Top spatial units by crash density", _polish_figure(fig), rank_df))
    if hin is not None:
        figures.extend(_hin_distribution_figures(hin))
        bubble_fig, bubble_df = _hin_ka_bubble_figure(hin)
        if bubble_fig is not None and not bubble_df.empty:
            figures.append(("K+A crashes vs HIN priority index", bubble_fig, bubble_df))
        metric = _dashboard_hin_metric(hin)
        hin_table = _hin_table_for_display(hin, metric, 20) if metric else pd.DataFrame()
        if not hin_table.empty:
            fig = px.bar(hin_table.sort_values("HIN index", ascending=True), y="SegID", x="HIN index", orientation="h", title="Top HIN segments/windows by HIN priority index")
            fig.update_layout(yaxis_title="SegID", xaxis_title="HIN priority index")
            figures.append(("Top HIN segments/windows table", _polish_figure(fig), hin_table))
    return figures


# --- V39 dashboard-only additions: HIN summary map support, KSI mapping, route comparison, generated-data ZIP ---
# Display/export helpers only. No workflow calculations are rerun or overwritten.


def _dashboard_numeric_series(df, col):
    if df is None or col is None or col not in df.columns:
        return pd.Series([0] * (len(df) if df is not None else 0), index=(df.index if df is not None else None), dtype="float64")
    return pd.to_numeric(df[col], errors="coerce").fillna(0)


def _dashboard_k_col(df):
    return _normal_col(
        df,
        [
            "K_Crashes", "K_Count", "Fatal_Crashes", "FatalCrashCount",
            "Fatal_Count", "Fatal", "K", "KSI_K",
        ],
    )


def _dashboard_a_col(df):
    return _normal_col(
        df,
        [
            "A_Crashes", "A_Count", "Serious_Injury_Crashes",
            "SeriousInjuryCrashCount", "Serious_Injury_Count", "Serious_Injury",
            "A", "KSI_A",
        ],
    )


def _hin_ka_series(df):
    """Return row-level KSI (K+A) values when available or mapped by the user.

    This is dashboard-only. It does not recalculate or overwrite HIN outputs.
    """
    if df is None or getattr(df, "empty", True):
        return pd.Series(dtype="float64")

    combined_key = st.session_state.get("dashboard_ksi_combined_col")
    k_key = st.session_state.get("dashboard_ksi_k_col")
    a_key = st.session_state.get("dashboard_ksi_a_col")

    if combined_key and combined_key in df.columns:
        return _dashboard_numeric_series(df, combined_key)
    if (k_key and k_key in df.columns) or (a_key and a_key in df.columns):
        return _dashboard_numeric_series(df, k_key) + _dashboard_numeric_series(df, a_key)

    combined_candidates = [
        "KSI_Crashes", "KSI_Count", "KSI", "KA_Crashes", "K_A_Crashes",
        "Fatal_Injury_Count", "FatalAndSeriousInjuryCrashes",
        "Fatal_Serious_Injury_Count", "Fatal_Serious_Crashes",
    ]
    for col in combined_candidates:
        if col in df.columns:
            return _dashboard_numeric_series(df, col)

    k_col = _dashboard_k_col(df)
    a_col = _dashboard_a_col(df)
    if k_col or a_col:
        return _dashboard_numeric_series(df, k_col) + _dashboard_numeric_series(df, a_col)

    return pd.Series([0] * len(df), index=df.index, dtype="float64")


def _render_ksi_mapping_controls(st, df):
    """Let users map K and A columns for dashboard KSI cards/charts only."""
    if df is None or getattr(df, "empty", True):
        return
    numeric_cols = [c for c in df.columns if c != "geometry" and pd.api.types.is_numeric_dtype(pd.to_numeric(df[c], errors="coerce"))]
    # Keep a less strict fallback for mixed numeric/object columns.
    if not numeric_cols:
        numeric_cols = [c for c in df.columns if c != "geometry"]
    auto_ksi = _hin_ka_series(df)
    with st.expander("KSI (K+A) crash field mapping for dashboard", expanded=(float(auto_ksi.sum()) <= 0 and len(df) > 0)):
        st.caption("Use this only when the HIN result table already contains K, A, or combined KSI columns but the dashboard could not auto-detect them. This changes dashboard summaries only.")
        options = [""] + list(numeric_cols)
        combined_default = st.session_state.get("dashboard_ksi_combined_col", "")
        k_default = st.session_state.get("dashboard_ksi_k_col", "")
        a_default = st.session_state.get("dashboard_ksi_a_col", "")
        c1, c2, c3 = st.columns(3)
        with c1:
            st.selectbox("Combined KSI (K+A) column", options, index=options.index(combined_default) if combined_default in options else 0, key="dashboard_ksi_combined_col")
        with c2:
            st.selectbox("K fatal crashes column", options, index=options.index(k_default) if k_default in options else 0, key="dashboard_ksi_k_col")
        with c3:
            st.selectbox("A serious injury crashes column", options, index=options.index(a_default) if a_default in options else 0, key="dashboard_ksi_a_col")


def _add_dashboard_rate_columns(df, crashes=None):
    """Add display-only crash/mile/year and KSI (K+A)/mile/year fields."""
    if df is None or getattr(df, "empty", True):
        return df
    out = df.copy()
    length_col = _dashboard_length_col(out)
    count_col = _crash_count_col(out)
    years = _dashboard_year_count(crashes)
    out["Dashboard_Analysis_Years"] = years
    if length_col and count_col:
        length = pd.to_numeric(out[length_col], errors="coerce").replace(0, pd.NA)
        count = pd.to_numeric(out[count_col], errors="coerce").fillna(0)
        out["Crash_per_Mile"] = (count / length).astype("float64").round(3)
        out["Crash_per_Mile_per_Year"] = (count / length / years).astype("float64").round(3)
    ksi = _hin_ka_series(out)
    if len(ksi) == len(out):
        out["KSI_Crashes_Dashboard"] = pd.to_numeric(ksi, errors="coerce").fillna(0).round(3)
        if length_col:
            length = pd.to_numeric(out[length_col], errors="coerce").replace(0, pd.NA)
            out["KSI_per_Mile_per_Year"] = (pd.to_numeric(ksi, errors="coerce").fillna(0) / length / years).astype("float64").round(3)
    return out


def _hin_distribution_figures(hin):
    """HIN diagnostic figures: line distribution plus route mean/median comparison."""
    figures = []
    metric = _dashboard_hin_metric(hin)
    if hin is None or getattr(hin, "empty", True) or not metric:
        return figures
    work = _drop_geometry(hin).copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce")
    work = work[work[metric].notna()].copy()
    if work.empty:
        return figures

    dist = work[[metric]].sort_values(metric).reset_index(drop=True)
    if len(dist) == 1:
        dist["Percentile"] = 100.0
    else:
        dist["Percentile"] = (dist.index + 1) / len(dist) * 100.0
    fig = px.line(
        dist,
        x="Percentile",
        y=metric,
        markers=False,
        title="HIN index distribution curve",
        labels={"Percentile": "Segment/window percentile", metric: "HIN priority index"},
    )
    fig.update_layout(xaxis_title="Segment/window percentile", yaxis_title="HIN priority index")
    figures.append(("HIN index distribution curve", _polish_figure(fig), dist))

    route_col = _dashboard_route_col(work)
    length_col = _dashboard_length_col(work)
    if route_col:
        route_work = work.copy()
        route_work[route_col] = route_work[route_col].fillna("Unknown").astype(str)
        agg_kwargs = dict(
            Average_HIN=(metric, "mean"),
            Median_HIN=(metric, "median"),
            Max_HIN=(metric, "max"),
            Segment_Count=(metric, "size"),
        )
        if length_col:
            route_work[length_col] = pd.to_numeric(route_work[length_col], errors="coerce").fillna(0)
            agg_kwargs["Miles"] = (length_col, "sum")
        route_summary = route_work.groupby(route_col, dropna=False).agg(**agg_kwargs).reset_index()
        route_summary = route_summary.sort_values("Average_HIN", ascending=False).head(15)
        if not route_summary.empty:
            plot_df = route_summary.copy()
            plot_df["Average_HIN_Left"] = pd.to_numeric(plot_df["Average_HIN"], errors="coerce").fillna(0)
            plot_df["Median_HIN_Right"] = pd.to_numeric(plot_df["Median_HIN"], errors="coerce").fillna(0)
            plot_df = plot_df.sort_values("Average_HIN", ascending=True)
            fig = px.bar(
                plot_df,
                y=route_col,
                x=["Average_HIN_Left", "Median_HIN_Right"],
                orientation="h",
                barmode="group",
                hover_data=[c for c in ["Average_HIN", "Median_HIN", "Max_HIN", "Segment_Count", "Miles"] if c in plot_df.columns],
                title="Average and median HIN index by route",
            )
            for trace in fig.data:
                if trace.name == "Average_HIN_Left":
                    trace.name = "Average HIN"
                    trace.text = [f"{abs(v):.1f}" for v in trace.x]
                elif trace.name == "Median_HIN_Right":
                    trace.name = "Median HIN"
                    trace.text = [f"{v:.1f}" for v in trace.x]
                trace.textposition = "inside"
            max_val = max(float(plot_df["Average_HIN"].max() or 0), float(plot_df["Median_HIN"].max() or 0), 1.0)
            tick_vals = [0, max_val / 2, max_val]
            fig.update_layout(
                yaxis_title="Route",
                xaxis_title="HIN priority index",
                legend_title="Statistic",
                bargap=0.25,
                xaxis=dict(tickmode="array", tickvals=tick_vals, ticktext=[f"{v:.0f}" for v in tick_vals], range=[0, max_val * 1.05]),
            )
            figures.append(("Average and median HIN index by route", _polish_figure(fig), route_summary))
    return figures


def _hin_ka_bubble_figure(hin):
    metric = _dashboard_hin_metric(hin)
    if hin is None or getattr(hin, "empty", True) or not metric:
        return None, pd.DataFrame()
    work = _drop_geometry(hin).copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    count_col = _crash_count_col(work)
    work["Total_Crashes"] = pd.to_numeric(work[count_col], errors="coerce").fillna(0) if count_col else 0
    work["KSI_Crashes_Dashboard"] = pd.to_numeric(_hin_ka_series(work), errors="coerce").fillna(0)
    if work["KSI_Crashes_Dashboard"].sum() <= 0:
        return None, pd.DataFrame()
    route_col = _dashboard_route_col(work)
    id_col = _normal_col(work, ["RiskSegmentID", "WindowID", "SlidingWindowID", "SegmentID", "UnitID", "SourceSegmentID", "CorridorID"])
    if id_col is None:
        id_col = "DashboardUnitID"
        work[id_col] = [f"HIN_{i + 1}" for i in range(len(work))]
    color_col = route_col if route_col and work[route_col].nunique(dropna=True) <= 15 else None
    hover_cols = [c for c in [id_col, route_col, _dashboard_length_col(work), "Total_Crashes", "KSI_Crashes_Dashboard"] if c and c in work.columns]
    fig = px.scatter(
        work,
        x=metric,
        y="KSI_Crashes_Dashboard",
        size="Total_Crashes" if work["Total_Crashes"].sum() > 0 else None,
        color=color_col,
        hover_data=hover_cols,
        size_max=42,
        title="KSI (K+A) crashes vs HIN priority index",
    )
    fig.update_traces(marker=dict(sizemin=6, opacity=0.78, line=dict(width=1, color="white")))
    fig.update_layout(xaxis_title="HIN priority index", yaxis_title="KSI (K+A) crashes")
    return _polish_figure(fig), work


def _render_hin_dashboard_charts(st, hin):
    if hin is None or getattr(hin, "empty", True):
        return
    st.markdown("<div class='dashboard-section-title'>HIN index diagnostics <span>distribution curve, route average/median comparison, and KSI relationship</span></div>", unsafe_allow_html=True)
    figures = _hin_distribution_figures(hin)
    bubble_fig, bubble_df = _hin_ka_bubble_figure(hin)
    if not figures and bubble_fig is None:
        st.info("HIN diagnostic charts need HIN priority index values and, for the bubble chart, available or mapped KSI (K+A) fields.")
        return
    chart_items = figures[:]
    if bubble_fig is not None:
        chart_items.append(("KSI (K+A) crashes vs HIN priority index", bubble_fig, bubble_df))
    cols = st.columns(2)
    for i, (title, fig, data) in enumerate(chart_items[:4]):
        with cols[i % 2]:
            fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35))
            st.plotly_chart(_polish_figure(fig), width="stretch", key=f"hin_diag_{_safe_name(title)}_{i}")


def _render_hin_network_summary(st, hin, crashes):
    """Dashboard-only HIN summary with custom threshold and KSI (K+A) card."""
    if hin is None or getattr(hin, "empty", True):
        return
    metric = _dashboard_hin_metric(hin)
    if not metric:
        return
    st.markdown("<div class='dashboard-section-title'>High Injury Network summary <span>custom HIN threshold, miles, crashes, and KSI (K+A) capture</span></div>", unsafe_allow_html=True)
    work = hin.copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    _render_ksi_mapping_controls(st, work)
    length_col = _dashboard_length_col(work)
    if length_col:
        work[length_col] = pd.to_numeric(work[length_col], errors="coerce").fillna(0)
    else:
        length_col = "__unit_length__"
        work[length_col] = 1.0

    mode_col, control_col, c_miles, c_crashes, c_ksi = st.columns([1.35, 1.05, 1, 1, 1])
    method = mode_col.selectbox(
        "High-risk network threshold",
        ["Top percent of miles", "Top number of segments/windows", "HIN index threshold", "Above average HIN index", "Above median HIN index"],
        index=0,
        key="hin_summary_threshold_mode_v39",
    )
    top_percent = 10.0
    top_n = 20
    index_threshold = 50.0
    with control_col:
        if method == "Top percent of miles":
            top_percent = st.number_input("Top percent", min_value=1.0, max_value=100.0, value=float(st.session_state.get("hin_summary_top_percent", 10.0)), step=1.0, key="hin_summary_top_percent")
        elif method == "Top number of segments/windows":
            top_n = st.number_input("Top N", min_value=1, max_value=max(int(len(work)), 1), value=min(20, max(int(len(work)), 1)), step=1, key="hin_summary_top_n")
        elif method == "HIN index threshold":
            index_threshold = st.number_input("Minimum HIN index", min_value=0.0, max_value=100.0, value=float(st.session_state.get("hin_summary_index_threshold", 50.0)), step=5.0, key="hin_summary_index_threshold")
        elif method == "Above average HIN index":
            st.metric("Average", f"{work[metric].mean():,.2f}")
        elif method == "Above median HIN index":
            st.metric("Median", f"{work[metric].median():,.2f}")

    selected = _selected_hin_subset(work, metric, method, top_percent=top_percent, top_n=top_n, index_threshold=index_threshold)
    total_mi = float(work[length_col].sum()) if work[length_col].sum() else float(len(work))
    high_mi = float(selected[length_col].sum()) if not selected.empty else 0.0
    pct_mi = high_mi / total_mi * 100 if total_mi else 0.0

    count_col = _crash_count_col(work)
    total_crashes = int(pd.to_numeric(work[count_col], errors="coerce").fillna(0).sum()) if count_col else (len(crashes) if crashes is not None else 0)
    high_crashes = int(pd.to_numeric(selected[count_col], errors="coerce").fillna(0).sum()) if count_col and not selected.empty else 0
    pct_crash = high_crashes / total_crashes * 100 if total_crashes else 0.0

    total_ksi_series = _hin_ka_series(work)
    selected_ksi_series = _hin_ka_series(selected)
    total_ksi = int(pd.to_numeric(total_ksi_series, errors="coerce").fillna(0).sum()) if len(total_ksi_series) else _total_ka_from_crashes(crashes)
    high_ksi = int(pd.to_numeric(selected_ksi_series, errors="coerce").fillna(0).sum()) if len(selected_ksi_series) else 0
    pct_ksi = high_ksi / total_ksi * 100 if total_ksi else 0.0

    c_miles.metric("High-risk miles", f"{high_mi:,.2f} mi", f"{pct_mi:,.1f}% of analyzed miles")
    c_crashes.metric("Crashes on selected HIN", f"{high_crashes:,}", f"{pct_crash:,.1f}% of assigned crashes")
    c_ksi.metric("KSI (K+A) on selected HIN", f"{high_ksi:,}", f"{pct_ksi:,.1f}% of KSI (K+A)")

    selected_display = _add_dashboard_rate_columns(selected, crashes=crashes)
    with st.expander("Selected HIN summary rows", expanded=False):
        preview_cols = []
        for c in [
            _normal_col(selected_display, ["RiskSegmentID", "WindowID", "SegmentID", "UnitID", "CorridorID"]),
            _dashboard_route_col(selected_display),
            _dashboard_length_col(selected_display), _crash_count_col(selected_display),
            "KSI_Crashes_Dashboard", metric, "Dashboard_Analysis_Years", "Crash_per_Mile_per_Year", "KSI_per_Mile_per_Year",
        ]:
            if c and c in selected_display.columns and c not in preview_cols:
                preview_cols.append(c)
        if preview_cols:
            st.dataframe(_safe_dataframe_for_display(selected_display[preview_cols].head(50)), width="stretch", hide_index=True)
        else:
            st.info("No displayable HIN rows are available for the selected threshold.")
    years = _dashboard_year_count(crashes)
    st.caption(f"Crashes/mile/year fields use the crash data year span detected in the loaded crashes: {years:g} year(s). These controls summarize existing HIN results only; they do not recalculate HIN scores.")
    _render_hin_dashboard_charts(st, hin)


def _hin_table_for_display(hin, metric, top_n=20):
    """Decision table for top HIN windows/segments with display-only KSI and rates."""
    if hin is None or getattr(hin, "empty", True) or metric not in hin.columns:
        return pd.DataFrame()
    crashes = None
    try:
        crashes = _available_tables(st).get("Assigned crashes", _available_tables(st).get("Uploaded crashes"))
    except Exception:
        crashes = None
    work = _add_dashboard_rate_columns(hin.copy(), crashes=crashes)
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work = work.sort_values(metric, ascending=False).head(top_n).reset_index(drop=True)
    id_col = _normal_col(work, ["RiskSegmentID", "WindowID", "SlidingWindowID", "SegmentID", "UnitID", "SourceSegmentID", "CorridorID"])
    route_col = _dashboard_route_col(work)
    length_col = _dashboard_length_col(work)
    from_col = _normal_col(work, ["FromMile", "From_Mile", "from_mile", "BeginMile", "StartMile", "WindowFromMile"])
    to_col = _normal_col(work, ["ToMile", "To_Mile", "to_mile", "EndMile", "WindowToMile"])
    count_col = _crash_count_col(work)
    out = pd.DataFrame()
    out["Rank"] = range(1, len(work) + 1)
    out["SegID"] = work[id_col].astype(str).values if id_col else [f"HIN_{i + 1}" for i in range(len(work))]
    out["Seg/window length"] = pd.to_numeric(work[length_col], errors="coerce").round(3).values if length_col else ""
    out["From mile"] = pd.to_numeric(work[from_col], errors="coerce").round(3).values if from_col else ""
    out["To mile"] = pd.to_numeric(work[to_col], errors="coerce").round(3).values if to_col else ""
    out["Route"] = work[route_col].astype(str).values if route_col else ""
    out["Crash count"] = pd.to_numeric(work[count_col], errors="coerce").fillna(0).round(0).astype(int).values if count_col else ""
    out["KSI (K+A) crashes"] = pd.to_numeric(work["KSI_Crashes_Dashboard"], errors="coerce").fillna(0).round(0).astype(int).values if "KSI_Crashes_Dashboard" in work.columns else ""
    out["HIN index"] = pd.to_numeric(work[metric], errors="coerce").round(3).values
    if "Crash_per_Mile" in work.columns:
        out["Crashes/mile"] = pd.to_numeric(work["Crash_per_Mile"], errors="coerce").round(3).values
    if "Crash_per_Mile_per_Year" in work.columns:
        out["Crashes/mile/year"] = pd.to_numeric(work["Crash_per_Mile_per_Year"], errors="coerce").round(3).values
    if "KSI_per_Mile_per_Year" in work.columns:
        out["KSI (K+A)/mile/year"] = pd.to_numeric(work["KSI_per_Mile_per_Year"], errors="coerce").round(3).values
    if "Dashboard_Analysis_Years" in work.columns:
        out["Crash years used"] = pd.to_numeric(work["Dashboard_Analysis_Years"], errors="coerce").round(1).values
    return out


def _top_density_export_table(density, top_n=20):
    """Decision-ready crash-density table with display-only rate columns."""
    if density is None or getattr(density, "empty", True):
        return pd.DataFrame()
    crashes = None
    try:
        crashes = _available_tables(st).get("Assigned crashes", _available_tables(st).get("Uploaded crashes"))
    except Exception:
        crashes = None
    df = _add_dashboard_rate_columns(_drop_geometry(density).copy(), crashes=crashes)
    unit_col = _unit_col(df) or _normal_col(df, ["UnitID", "IntersectionID", "CorridorID", "SegmentID", "Route"])
    unit_type_col = _normal_col(df, ["UnitType", "IntersectionType", "CorridorType", "SegmentType"])
    city_col = _normal_col(df, ["City", "city_name"])
    length_col = _dashboard_length_col(df)
    count_col = _crash_count_col(df)
    density_col = _normal_col(df, ["CrashDensity", "Crash_Density", "crash_density"])
    if density_col is None:
        density_col = _default_metric(_numeric_cols(df))
    if density_col:
        df[density_col] = pd.to_numeric(df[density_col], errors="coerce").fillna(0)
        df = df.sort_values(density_col, ascending=False)
    use = df.head(top_n).reset_index(drop=True)
    out = pd.DataFrame({"Rank": range(1, len(use) + 1)})
    out["Spatial unit id"] = use[unit_col].astype(str) if unit_col else use.index.astype(str)
    out["Unit type"] = use[unit_type_col].astype(str) if unit_type_col else ""
    out["City"] = use[city_col].astype(str) if city_col else ""
    out["Length_mi"] = pd.to_numeric(use[length_col], errors="coerce").round(3) if length_col else ""
    route_col = _normal_col(use, ["Route", "FULLNAME", "RoadName", "RouteName", "CorridorRoute", "RouteName_Calc"])
    if route_col:
        out["Route name"] = use[route_col].astype(str)
    out["Crash count"] = pd.to_numeric(use[count_col], errors="coerce").fillna(0).astype(int) if count_col else ""
    out["Crash density"] = pd.to_numeric(use[density_col], errors="coerce").round(3) if density_col else ""
    for src, label in [("Crash_per_Mile", "Crashes/mile"), ("Crash_per_Mile_per_Year", "Crashes/mile/year"), ("KSI_per_Mile_per_Year", "KSI (K+A)/mile/year")]:
        if src in use.columns:
            out[label] = pd.to_numeric(use[src], errors="coerce").round(3)
    if "Dashboard_Analysis_Years" in use.columns:
        out["Crash years used"] = pd.to_numeric(use["Dashboard_Analysis_Years"], errors="coerce").round(1)
    return out


def _download_generated_data_zip(st, tables):
    """Export only workflow-generated data in a clean CSV/GeoJSON structure.

    Folder structure inside the ZIP:
        csv/roads
        csv/signals
        csv/spatial_units_crashes
        csv/sliding_windows
        geojson/roads
        geojson/signals
        geojson/spatial_units_crashes
        geojson/sliding_windows

    The export intentionally excludes original uploaded source files/raw crash
    uploads. It includes only app-generated/processed outputs needed for review,
    GIS use, and report backup.
    """
    buffer = io.BytesIO()
    written_paths = set()
    manifest = []

    def clean_folder(value):
        value = str(value or "").strip().lower()
        value = re.sub(r"[^a-z0-9]+", "_", value).strip("_")
        return value or "data"

    def add_bytes(zf, path, data, description=""):
        if data is None:
            return False
        path = str(path).replace("\\", "/")
        if path in written_paths:
            return False
        zf.writestr(path, data)
        written_paths.add(path)
        if description:
            manifest.append(f"{path} - {description}")
        else:
            manifest.append(path)
        return True

    def _with_point_lat_lon(obj):
        """Add Latitude/Longitude fields for point GeoDataFrames before CSV export."""
        if obj is None or getattr(obj, "empty", False):
            return obj
        if not hasattr(obj, "geometry") or "geometry" not in obj.columns:
            return obj
        try:
            gdf = _repair_gdf_crs(obj, st).copy()
            gdf_ll = gdf.to_crs(4326) if getattr(gdf, "crs", None) is not None else gdf
            geom = gdf_ll.geometry
            if geom.geom_type.isin(["Point"]).all():
                if "Longitude" not in gdf.columns and "Long" not in gdf.columns and "Lon" not in gdf.columns:
                    gdf["Longitude"] = geom.x
                if "Latitude" not in gdf.columns and "Lat" not in gdf.columns:
                    gdf["Latitude"] = geom.y
            return gdf
        except Exception:
            return obj

    def simplify_roads(obj):
        """Export only analysis road segment attributes and geometry."""
        if obj is None or getattr(obj, "empty", False):
            return obj
        df = obj.copy()
        preferred = [
            st.session_state.get("segment_id_col"),
            st.session_state.get("route_col"),
            "OSMEdgeID", "LINEARID", "FACILITYID", "SegmentID", "SegID", "UnitID",
            "Route", "FULLNAME", "RouteNameOSM", "RouteName_Calc",
            "FromMile", "ToMile", "SegmentLength_Mile",
            "FunctionalClass", "RoadClass", "RoadType", "RoadStyleClass",
            "MTFCC", "RTTYP", "OSMHighway", "geometry",
        ]
        cols = []
        for c in preferred:
            if c and c in df.columns and c not in cols:
                cols.append(c)
        return df[cols].copy() if cols else df

    def simplify_signals(obj):
        """Export the same cleaned signal table shown in the workflow.

        CSV columns are limited to SignalID, City, Latitude, and Longitude.
        GeoJSON keeps the same attributes plus point geometry.
        """
        if obj is None or getattr(obj, "empty", False):
            return obj

        df = _with_point_lat_lon(obj).copy()

        if "SignalID" not in df.columns:
            df["SignalID"] = range(1, len(df) + 1)

        if "City" not in df.columns:
            df["City"] = st.session_state.get("area_name", "")

        lat_col = _normal_col(df, ["Latitude", "Lat", "LAT", "Y"])
        lon_col = _normal_col(df, ["Longitude", "Long", "Lon", "LON", "X"])

        if lat_col and lat_col != "Latitude":
            df["Latitude"] = pd.to_numeric(df[lat_col], errors="coerce")
        if lon_col and lon_col != "Longitude":
            df["Longitude"] = pd.to_numeric(df[lon_col], errors="coerce")

        preferred = [
            "SignalID",
            "City",
            "Latitude",
            "Longitude",
            "geometry",
        ]
        cols = [c for c in preferred if c in df.columns]
        return df[cols].copy() if cols else df

    def simplify_corridors(obj):
        """Export the same final corridor table shown in the workflow.

        CSV columns are limited to CorridorID, Route, and City. GeoJSON keeps
        the corridor geometry when available.
        """
        if obj is None or getattr(obj, "empty", False):
            return obj

        df = obj.copy()

        corridor_col = _normal_col(df, ["CorridorID", "CorridorId", "corridor_id"])
        route_col = _normal_col(df, ["Route", "RouteName", "FULLNAME", "CorridorRoute", "Route_Normalized"])
        city_col = _normal_col(df, ["City", "city", "area_name"])

        if corridor_col and corridor_col != "CorridorID":
            df["CorridorID"] = df[corridor_col]
        if route_col and route_col != "Route":
            df["Route"] = df[route_col]
        if city_col and city_col != "City":
            df["City"] = df[city_col]
        if "City" not in df.columns:
            df["City"] = st.session_state.get("area_name", "")

        preferred = ["CorridorID", "Route", "City", "geometry"]
        cols = [c for c in preferred if c in df.columns]
        return df[cols].copy() if cols else df

    def simplify_spatial_units_density(obj):
        """Export only necessary spatial-unit density fields plus geometry."""
        if obj is None or getattr(obj, "empty", False):
            return obj

        df = obj.copy()

        unit_col = _normal_col(df, ["UnitID", "SpatialUnitID", "IntersectionID", "CorridorID", "SegmentID"])
        unit_type_col = _normal_col(df, ["UnitType", "AnalysisType", "SpatialUnitType", "Type"])
        segment_col = _normal_col(df, ["SegmentID", "SegID", "OSMEdgeID", "LINEARID", "FACILITYID"])
        crash_col = _normal_col(df, ["CrashCount", "Crash_Count", "TotalCrashes", "Crash count"])
        length_col = _normal_col(df, ["Length_Miles", "Length_Mile", "SegmentLength_Mile", "CorridorLength_Mile", "Miles"])
        area_col = _normal_col(df, ["Area_SqMi", "AreaSqMi", "Area_Sq_Mi", "AreaSquareMiles"])
        density_col = _normal_col(df, ["CrashDensity", "Crash_Density", "Density", "Crash density"])

        rename_map = {}
        for src, dst in [
            (unit_col, "UnitID"),
            (unit_type_col, "UnitType"),
            (segment_col, "SegmentID"),
            (crash_col, "CrashCount"),
            (length_col, "Length_Miles"),
            (area_col, "Area_SqMi"),
            (density_col, "CrashDensity"),
        ]:
            if src and src in df.columns and src != dst:
                rename_map[src] = dst
        if rename_map:
            df = df.rename(columns=rename_map)

        if "UnitType" not in df.columns:
            df["UnitType"] = st.session_state.get("analysis_type", "")
        if "SegmentID" not in df.columns:
            df["SegmentID"] = ""

        preferred = [
            "UnitID",
            "UnitType",
            "SegmentID",
            "CrashCount",
            "Length_Miles",
            "Area_SqMi",
            "CrashDensity",
            "geometry",
        ]
        cols = [c for c in preferred if c in df.columns]
        return df[cols].copy() if cols else df

    def add_hin_index_to_windows(obj):
        """Ensure sliding-window exports include an explicit HIN_Index column."""
        if obj is None or getattr(obj, "empty", False):
            return obj
        df = obj.copy()
        if "HIN_Index" not in df.columns:
            if "HIN_Priority_Index" in df.columns:
                df["HIN_Index"] = pd.to_numeric(df["HIN_Priority_Index"], errors="coerce")
            else:
                try:
                    df["HIN_Index"] = _window_hin_index_series(df).values
                except Exception:
                    metric = _dashboard_hin_metric(df)
                    if metric and metric in df.columns:
                        vals = pd.to_numeric(df[metric], errors="coerce").fillna(0)
                        max_val = float(vals.max()) if len(vals) else 0.0
                        df["HIN_Index"] = vals / max_val * 100.0 if max_val > 0 else vals
        if "HIN_Priority_Index" not in df.columns and "HIN_Index" in df.columns:
            df["HIN_Priority_Index"] = df["HIN_Index"]
        return df

    def build_route_hin_corridors(results):
        """Create one HIN corridor per route for export.

        The raw risk_corridors output can split one route into many corridor IDs.
        For the clean export, each route is one corridor with a numeric CorridorID.
        """
        if not isinstance(results, dict):
            return None
        src = results.get("risk_windows")
        if src is None or getattr(src, "empty", True):
            src = results.get("risk_segments")
        if src is None or getattr(src, "empty", True):
            return None
        df = add_hin_index_to_windows(src).copy()
        route_col = _section7_route_col_from_tables(df, None) or _normal_col(
            df,
            ["Route", "FULLNAME", "RouteName", "RoadName", "RouteName_Calc", "CorridorRoute", "Name"],
        )
        if not route_col or route_col not in df.columns:
            return None
        metric = "HIN_Index" if "HIN_Index" in df.columns else _dashboard_hin_metric(df)
        if metric and metric in df.columns:
            df[metric] = pd.to_numeric(df[metric], errors="coerce").fillna(0)
        length_col = _dashboard_length_col(df)
        if length_col and length_col in df.columns:
            df[length_col] = pd.to_numeric(df[length_col], errors="coerce").fillna(0)

        rows = []
        is_geo = hasattr(df, "geometry") and "geometry" in df.columns
        for i, (route, grp) in enumerate(df.groupby(route_col, dropna=False), start=1):
            row = {
                "CorridorID": i,
                "Route": str(route),
                "Segment_Count": int(len(grp)),
            }
            if metric and metric in grp.columns:
                row["Max_HIN"] = float(pd.to_numeric(grp[metric], errors="coerce").fillna(0).max())
                row["Avg_HIN_Index"] = float(pd.to_numeric(grp[metric], errors="coerce").fillna(0).mean())
                row["Median_HIN_Index"] = float(pd.to_numeric(grp[metric], errors="coerce").fillna(0).median())
            if length_col and length_col in grp.columns:
                row["Total_Window_Miles"] = float(pd.to_numeric(grp[length_col], errors="coerce").fillna(0).sum())
            if is_geo:
                try:
                    row["geometry"] = grp.geometry.union_all()
                except Exception:
                    try:
                        row["geometry"] = grp.geometry.unary_union
                    except Exception:
                        pass
            rows.append(row)
        if not rows:
            return None
        if is_geo and any("geometry" in r for r in rows):
            try:
                return gpd.GeoDataFrame(rows, geometry="geometry", crs=getattr(df, "crs", None))
            except Exception:
                return pd.DataFrame(rows)
        return pd.DataFrame(rows)

    def add_df(zf, folder, name, obj, description="", csv_only=False, geojson_only=False):
        if obj is None or getattr(obj, "empty", False):
            return 0
        folder = clean_folder(folder)
        safe_name = _safe_name(name)
        count = 0
        if not geojson_only:
            try:
                csv = _safe_dataframe_for_display(obj).to_csv(index=False).encode("utf-8")
                if add_bytes(
                    zf,
                    f"csv/{folder}/{safe_name}.csv",
                    csv,
                    description or f"CSV export for {name}",
                ):
                    count += 1
            except Exception as exc:
                add_bytes(
                    zf,
                    f"csv/{folder}/{safe_name}_export_error.txt",
                    str(exc).encode("utf-8"),
                    f"CSV export error for {name}",
                )
                count += 1
        if not csv_only:
            try:
                if hasattr(obj, "geometry") and "geometry" in obj.columns:
                    gdf = _repair_gdf_crs(obj, st)
                    gdf = _safe_geojson_gdf(gdf)
                    if gdf is not None and not getattr(gdf, "empty", True):
                        geojson = gdf.to_json().encode("utf-8")
                        if add_bytes(
                            zf,
                            f"geojson/{folder}/{safe_name}.geojson",
                            geojson,
                            description or f"GeoJSON export for {name}",
                        ):
                            count += 1
            except Exception as exc:
                add_bytes(
                    zf,
                    f"geojson/{folder}/{safe_name}_export_error.txt",
                    str(exc).encode("utf-8"),
                    f"GeoJSON export error for {name}",
                )
                count += 1
        return count

    def add_report_table(zf, folder, name, obj):
        return add_df(
            zf,
            folder,
            f"report_table_{name}",
            obj,
            f"Report/dashboard table: {name}",
            csv_only=True,
        )

    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        # 1. Roads: analysis road segments only, with route/milepost/length fields.
        roads = st.session_state.get("selected_roads")
        if roads is None or getattr(roads, "empty", True):
            roads = st.session_state.get("base_roads")
        add_df(
            zf,
            "roads",
            "road_segments_from_to_mile",
            simplify_roads(roads),
            "Analysis road segments with FromMile, ToMile, and SegmentLength_Mile.",
        )

        # 2. Signals: cleaned signal table only, with SignalID, lat/lon, and city.
        add_df(
            zf,
            "signals",
            "cleaned_signal_table",
            simplify_signals(st.session_state.get("signals_clean")),
            "Cleaned/de-duplicated signal points with SignalID, latitude, longitude, and city.",
        )

        # 3. Corridors, spatial-unit density results, assigned crash points, and summary tables.
        corridors = st.session_state.get("final_corridors")
        if corridors is None or getattr(corridors, "empty", True):
            corridors = st.session_state.get("corridors")
        add_df(
            zf,
            "spatial_units_crashes",
            "corridor_table",
            simplify_corridors(corridors),
            "Final corridor table with CorridorID, Route, City, and corridor geometry.",
        )

        # The separate spatial_units table is intentionally not exported because
        # the density results file contains the necessary unit ID/type/count/rate
        # fields plus geometry.
        add_df(
            zf,
            "spatial_units_crashes",
            "spatial_units_density_results",
            simplify_spatial_units_density(st.session_state.get("spatial_units_density_map")),
            "Spatial unit density results with only UnitID, UnitType, SegmentID, CrashCount, Length_Miles, Area_SqMi, and CrashDensity.",
        )
        add_df(
            zf,
            "spatial_units_crashes",
            "crashes_assigned_to_units",
            _with_point_lat_lon(st.session_state.get("assigned_crashes")),
            "Crash point records assigned to generated spatial units.",
        )
        add_df(
            zf,
            "spatial_units_crashes",
            "kabco_summary",
            st.session_state.get("kabco_result"),
            "KABCO crash summary table generated by the workflow.",
            csv_only=True,
        )
        add_df(
            zf,
            "spatial_units_crashes",
            "classified_results",
            st.session_state.get("classified"),
            "Classified crash/spatial-unit result table generated by the workflow.",
        )

        # Include report-ready tables, but not raw uploaded crash data and not duplicates of core HIN/window files.
        try:
            report_tables = _report_tables(tables or {}, top_n=100000)
        except Exception:
            report_tables = {}
        for table_name, table_obj in (report_tables or {}).items():
            lower = str(table_name).lower()
            if "uploaded crash" in lower or "raw crash" in lower:
                continue
            # Avoid duplicate exports of the core sliding-window/HIN files below.
            if any(term in lower for term in ["sliding_window_table", "hin_segments", "hin_corridors"]):
                continue
            if "sliding" in lower or "hin" in lower:
                folder = "sliding_windows"
            else:
                folder = "spatial_units_crashes"
            add_report_table(zf, folder, table_name, table_obj)

        # 4. Sliding windows / HIN: window table, HIN segments, and one route-level HIN corridor file.
        results = st.session_state.get("section7_results")
        if isinstance(results, dict):
            add_df(
                zf,
                "sliding_windows",
                "sliding_window_table",
                add_hin_index_to_windows(results.get("risk_windows")),
                "Sliding-window table with route, window, crash count, score, and HIN index.",
            )
            add_df(
                zf,
                "sliding_windows",
                "hin_segments",
                add_hin_index_to_windows(results.get("risk_segments")),
                "HIN segment/window GeoJSON and table with HIN index.",
            )
            add_df(
                zf,
                "sliding_windows",
                "hin_corridors",
                build_route_hin_corridors(results),
                "Route-level HIN corridor summary: one numeric CorridorID per route.",
            )
            # Do not export sliding_window_assigned_crashes. It is an internal intermediate table.

        readme = """HIN generated data export

This ZIP contains processed outputs generated by the app. It intentionally does not include original uploaded source files or raw uploaded crash files.

Folder structure
- csv/roads and geojson/roads: analysis road segments only, including FromMile, ToMile, and SegmentLength_Mile when available.
- csv/signals and geojson/signals: cleaned/de-duplicated signal points with SignalID, Latitude, Longitude, and City.
- csv/spatial_units_crashes and geojson/spatial_units_crashes: final corridors, spatial-unit density results, crash points assigned to units, KABCO/report tables. The separate spatial_units table is not exported because the density results file contains the necessary unit attributes.
- csv/sliding_windows and geojson/sliding_windows: sliding-window/HIN tables and geometries. The assigned-crash intermediate table is intentionally not exported.

Manifest
"""
        add_bytes(
            zf,
            "README_export_manifest.txt",
            (readme + "\n".join(manifest)).encode("utf-8"),
            "Export manifest.",
        )

    if len(written_paths) <= 1:
        return None
    buffer.seek(0)
    return buffer.getvalue()

def _available_maps(st):
    """Dashboard map layers, including display-only HIN threshold maps."""
    maps = {}
    density = st.session_state.get("spatial_units_density_map")
    if density is not None and not getattr(density, "empty", True):
        maps["Crash density map"] = _repair_gdf_crs(density, st)
    results = st.session_state.get("section7_results")
    if results is not None:
        risk_segments = results.get("risk_segments")
        if risk_segments is not None and not getattr(risk_segments, "empty", True):
            hin_map = _repair_gdf_crs(risk_segments, st)
            maps["HIN priority map"] = hin_map
            metric = _dashboard_hin_metric(hin_map)
            if metric:
                work = hin_map.copy()
                work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
                avg = work[metric].mean()
                med = work[metric].median()
                maps["HIN above average map"] = work[work[metric] >= avg].copy()
                maps["HIN above median map"] = work[work[metric] >= med].copy()
    corridors = st.session_state.get("final_corridors", st.session_state.get("corridors"))
    if corridors is not None and not getattr(corridors, "empty", True):
        maps["Corridor map"] = _repair_gdf_crs(corridors, st)
    return {k: v for k, v in maps.items() if v is not None and not getattr(v, "empty", True)}


# --- V40 dashboard-only corrections: safe KSI, clearer HIN diagnostics ---
# These override earlier dashboard helper functions only. They do not change
# HIN/sliding-window calculations or stored workflow results.


def _dashboard_crash_id_col(df):
    if df is None:
        return None
    mapped_col = st.session_state.get("mapped_crash_id_col", "")
    if mapped_col and mapped_col in df.columns:
        return mapped_col
    return _normal_col(
        df,
        [
            "DashboardCrashID", "SourceCrashID", "CrashID", "CrashId",
            "crash_id", "CRASH_ID", "Crash_ID", "CaseID", "CASE_ID",
            "ST_CASE", "st_case", "OBJECTID", "ObjectID", "FID",
            "RecordID", "UnitID", "Crash_Key",
        ],
    )


def _strict_hin_ksi_series(df):
    """Return explicit row-level KSI only when the HIN table truly has it.

    Do not auto-detect generic one-letter K/A columns. In some result tables,
    those names can refer to unrelated fields or duplicated window values,
    which produced impossible KSI totals.
    """
    if df is None or getattr(df, "empty", True):
        return pd.Series(dtype="float64")

    combined_key = st.session_state.get("dashboard_ksi_combined_col")
    k_key = st.session_state.get("dashboard_ksi_k_col")
    a_key = st.session_state.get("dashboard_ksi_a_col")

    if combined_key and combined_key in df.columns:
        return _dashboard_numeric_series(df, combined_key)
    if (k_key and k_key in df.columns) or (a_key and a_key in df.columns):
        return _dashboard_numeric_series(df, k_key) + _dashboard_numeric_series(df, a_key)

    combined_candidates = [
        "KSI_Crashes", "KSI_Count", "KSI_Crash_Count", "KSI_Total",
        "KA_Crashes", "K_A_Crashes", "KA_Count", "KA_Total",
        "FatalAndSeriousInjuryCrashes", "Fatal_Serious_Injury_Count",
        "Fatal_Serious_Crashes",
    ]
    for col in combined_candidates:
        if col in df.columns:
            return _dashboard_numeric_series(df, col)

    k_candidates = [
        "K_Crashes", "K_Count", "Fatal_Crashes", "FatalCrashCount", "Fatal_Count",
    ]
    a_candidates = [
        "A_Crashes", "A_Count", "Serious_Injury_Crashes",
        "SeriousInjuryCrashCount", "Serious_Injury_Count",
    ]
    k_col = _normal_col(df, k_candidates)
    a_col = _normal_col(df, a_candidates)
    if k_col or a_col:
        return _dashboard_numeric_series(df, k_col) + _dashboard_numeric_series(df, a_col)

    return pd.Series([0] * len(df), index=df.index, dtype="float64")


def _hin_ka_series(df):
    """Compatibility wrapper using strict KSI detection."""
    return _strict_hin_ksi_series(df)


def _crash_level_ksi_series(crashes):
    """Return one KSI value per crash row using crash-level fields.

    This is used for dashboard cards and avoids summing HIN window rows.
    """
    if crashes is None or getattr(crashes, "empty", True):
        return pd.Series(dtype="float64")

    work = crashes.copy()

    combined_col = st.session_state.get("dashboard_crash_ksi_combined_col")
    k_col_user = st.session_state.get("dashboard_crash_ksi_k_col")
    a_col_user = st.session_state.get("dashboard_crash_ksi_a_col")
    kabco_user = st.session_state.get("dashboard_crash_ksi_kabco_col")

    if combined_col and combined_col in work.columns:
        return pd.to_numeric(work[combined_col], errors="coerce").fillna(0)

    if (k_col_user and k_col_user in work.columns) or (a_col_user and a_col_user in work.columns):
        k = pd.to_numeric(work[k_col_user], errors="coerce").fillna(0) if k_col_user in work.columns else 0
        a = pd.to_numeric(work[a_col_user], errors="coerce").fillna(0) if a_col_user in work.columns else 0
        return k + a

    if kabco_user and kabco_user in work.columns:
        vals = work[kabco_user].map(normalize_kabco_value).astype(str).str.upper()
        return vals.isin(["K", "A"]).astype(int)

    sev_cols = _severity_count_columns(work)
    if sev_cols:
        k = pd.to_numeric(work[sev_cols.get("K")], errors="coerce").fillna(0) if sev_cols.get("K") else 0
        a = pd.to_numeric(work[sev_cols.get("A")], errors="coerce").fillna(0) if sev_cols.get("A") else 0
        return k + a

    kabco = _kabco_col(work)
    if kabco and kabco in work.columns:
        vals = work[kabco].map(normalize_kabco_value).astype(str).str.upper()
        return vals.isin(["K", "A"]).astype(int)

    return pd.Series([0] * len(work), index=work.index, dtype="float64")


def _total_ka_from_crashes(crashes):
    vals = _crash_level_ksi_series(crashes)
    if vals is None or len(vals) == 0:
        return 0
    return int(pd.to_numeric(vals, errors="coerce").fillna(0).sum())


def _render_crash_ksi_mapping_controls(st, crashes):
    """Crash-level KSI mapping controls used by dashboard only."""
    if crashes is None or getattr(crashes, "empty", True):
        return

    cols = [c for c in crashes.columns if c != "geometry"]
    numeric_cols = [c for c in cols if pd.to_numeric(crashes[c], errors="coerce").notna().any()]
    options_all = [""] + cols
    options_num = [""] + numeric_cols
    auto_total = _total_ka_from_crashes(crashes)

    with st.expander("KSI (K+A) crash field mapping for dashboard", expanded=(auto_total <= 0)):
        st.caption(
            "Use this only if the dashboard cannot auto-detect the crash severity fields. "
            "This reads the uploaded/assigned crash table, not the HIN score table, so totals should match the Crash summary KPI."
        )
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            default = st.session_state.get("dashboard_crash_ksi_kabco_col", "")
            st.selectbox(
                "KABCO/severity column",
                options_all,
                index=options_all.index(default) if default in options_all else 0,
                key="dashboard_crash_ksi_kabco_col",
            )
        with c2:
            default = st.session_state.get("dashboard_crash_ksi_combined_col", "")
            st.selectbox(
                "Combined KSI (K+A) count column",
                options_num,
                index=options_num.index(default) if default in options_num else 0,
                key="dashboard_crash_ksi_combined_col",
            )
        with c3:
            default = st.session_state.get("dashboard_crash_ksi_k_col", "")
            st.selectbox(
                "K fatal count column",
                options_num,
                index=options_num.index(default) if default in options_num else 0,
                key="dashboard_crash_ksi_k_col",
            )
        with c4:
            default = st.session_state.get("dashboard_crash_ksi_a_col", "")
            st.selectbox(
                "A serious injury count column",
                options_num,
                index=options_num.index(default) if default in options_num else 0,
                key="dashboard_crash_ksi_a_col",
            )
        st.caption(f"Detected total KSI (K+A) crashes from crash table: {_total_ka_from_crashes(crashes):,}")


def _candidate_unit_id_cols(df):
    if df is None or getattr(df, "empty", True):
        return []
    names = [
        "RiskSegmentID", "SlidingWindowID", "WindowID", "SegmentID", "SegID",
        "SourceSegmentID", "UnitID", "SpatialUnitID", "IntersectionID",
        "CorridorID", "Corridor_Key", "AssignedUnitID", "MatchedUnitID",
        "NearestSegmentID", "RoadSegmentID",
    ]
    out = []
    for c in names:
        if c in df.columns and c not in out:
            out.append(c)
    # Include columns that look like IDs but avoid crash ID.
    for c in df.columns:
        cl = str(c).lower()
        if c not in out and "id" in cl and "crash" not in cl and c != "geometry":
            out.append(c)
    return out


def _selected_hin_crash_subset(selected_hin, crashes):
    """Try to identify crash rows linked to selected HIN rows by existing IDs.

    Returns (subset, method_note). If no trusted linkage exists, returns
    (None, explanation). No new assignment is performed here.
    """
    if selected_hin is None or getattr(selected_hin, "empty", True):
        return None, "No selected HIN rows."
    if crashes is None or getattr(crashes, "empty", True):
        return None, "No assigned/uploaded crash table available."

    selected_cols = _candidate_unit_id_cols(selected_hin)
    crash_cols = _candidate_unit_id_cols(crashes)
    for sc in selected_cols:
        svals = selected_hin[sc].dropna().astype(str).str.strip()
        svals = set(v for v in svals if v and v.lower() not in ["nan", "none"])
        if not svals:
            continue
        for cc in crash_cols:
            cvals = crashes[cc].dropna().astype(str).str.strip()
            overlap = set(cvals.unique()).intersection(svals)
            if overlap:
                subset = crashes[cvals.isin(svals)].copy()
                return subset, f"Matched selected HIN rows to crashes using {sc} ↔ {cc}."
    return None, "No existing crash-to-HIN ID linkage was found; selected-HIN crash/KSI capture is not computed from window totals."


def _safe_unique_crash_count(crashes):
    if crashes is None or getattr(crashes, "empty", True):
        return 0
    cid = _dashboard_crash_id_col(crashes)
    if cid and cid in crashes.columns:
        return int(crashes[cid].dropna().astype(str).nunique())
    return int(len(crashes))


def _add_dashboard_rate_columns(df, crashes=None):
    """Add display-only crash/mile/year and KSI (K+A)/mile/year fields safely."""
    if df is None or getattr(df, "empty", True):
        return df
    out = df.copy()
    length_col = _dashboard_length_col(out)
    count_col = _crash_count_col(out)
    years = _dashboard_year_count(crashes)
    out["Dashboard_Analysis_Years"] = years
    if length_col and count_col:
        length = pd.to_numeric(out[length_col], errors="coerce").replace(0, pd.NA)
        count = pd.to_numeric(out[count_col], errors="coerce").fillna(0)
        out["Crash_per_Mile"] = (count / length).astype("float64").round(3)
        out["Crash_per_Mile_per_Year"] = (count / length / years).astype("float64").round(3)
    ksi = _strict_hin_ksi_series(out)
    if len(ksi) == len(out) and float(pd.to_numeric(ksi, errors="coerce").fillna(0).sum()) > 0:
        out["KSI_Crashes_Dashboard"] = pd.to_numeric(ksi, errors="coerce").fillna(0).round(3)
        if length_col:
            length = pd.to_numeric(out[length_col], errors="coerce").replace(0, pd.NA)
            out["KSI_per_Mile_per_Year"] = (pd.to_numeric(ksi, errors="coerce").fillna(0) / length / years).astype("float64").round(3)
    return out


def _hin_distribution_figures(hin):
    """Clear HIN diagnostics: score-bin distribution plus route mean/median."""
    figures = []
    metric = _dashboard_hin_metric(hin)
    if hin is None or getattr(hin, "empty", True) or not metric:
        return figures
    work = _drop_geometry(hin).copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce")
    work = work[work[metric].notna()].copy()
    if work.empty:
        return figures

    # Binned line chart is easier to read than sorted percentile when most
    # segments have HIN = 0 and a few segments jump to 100.
    max_score = max(float(work[metric].max()), 1.0)
    bin_width = 5 if max_score <= 100 else max(max_score / 20.0, 1.0)
    bins = list(pd.interval_range(start=0, end=max_score + bin_width, freq=bin_width, closed="left"))
    if not bins:
        bins = pd.interval_range(start=0, end=100, freq=5, closed="left")
    tmp = work.copy()
    tmp["HIN score range"] = pd.cut(tmp[metric], bins=bins, include_lowest=True)
    dist = tmp.groupby("HIN score range", observed=False).size().reset_index(name="Segment/window count")
    dist["HIN score range"] = dist["HIN score range"].astype(str)
    dist["Bin midpoint"] = [i * bin_width + bin_width / 2 for i in range(len(dist))]
    fig = px.line(
        dist,
        x="Bin midpoint",
        y="Segment/window count",
        markers=True,
        title="HIN index distribution by score range",
        labels={"Bin midpoint": "HIN priority index", "Segment/window count": "Segment/window count"},
    )
    fig.update_layout(xaxis_title="HIN priority index", yaxis_title="Segment/window count")
    figures.append(("HIN index distribution by score range", _polish_figure(fig), dist))

    route_col = _dashboard_route_col(work)
    length_col = _dashboard_length_col(work)
    if route_col:
        route_work = work.copy()
        route_work[route_col] = route_work[route_col].fillna("Unknown").astype(str)
        agg_kwargs = dict(
            Average_HIN=(metric, "mean"),
            Median_HIN=(metric, "median"),
            Max_HIN=(metric, "max"),
            Segment_Count=(metric, "size"),
        )
        if length_col:
            route_work[length_col] = pd.to_numeric(route_work[length_col], errors="coerce").fillna(0)
            agg_kwargs["Miles"] = (length_col, "sum")
        route_summary = route_work.groupby(route_col, dropna=False).agg(**agg_kwargs).reset_index()
        route_summary = route_summary.sort_values("Average_HIN", ascending=False).head(15)
        if not route_summary.empty:
            plot_df = route_summary.copy()
            plot_df["Average_HIN_Left"] = pd.to_numeric(plot_df["Average_HIN"], errors="coerce").fillna(0)
            plot_df["Median_HIN_Right"] = pd.to_numeric(plot_df["Median_HIN"], errors="coerce").fillna(0)
            plot_df = plot_df.sort_values("Average_HIN", ascending=True)
            fig = px.bar(
                plot_df,
                y=route_col,
                x=["Average_HIN_Left", "Median_HIN_Right"],
                orientation="h",
                barmode="group",
                hover_data=[c for c in ["Average_HIN", "Median_HIN", "Max_HIN", "Segment_Count", "Miles"] if c in plot_df.columns],
                title="Average and median HIN index by route",
            )
            for trace in fig.data:
                if trace.name == "Average_HIN_Left":
                    trace.name = "Average HIN"
                    trace.text = [f"{abs(v):.1f}" for v in trace.x]
                elif trace.name == "Median_HIN_Right":
                    trace.name = "Median HIN"
                    trace.text = [f"{v:.1f}" for v in trace.x]
                trace.textposition = "inside"
            max_val = max(float(plot_df["Average_HIN"].max() or 0), float(plot_df["Median_HIN"].max() or 0), 1.0)
            tick_vals = [0, max_val / 2, max_val]
            fig.update_layout(
                yaxis_title="Route",
                xaxis_title="HIN priority index",
                legend_title="Statistic",
                bargap=0.25,
                xaxis=dict(tickmode="array", tickvals=tick_vals, ticktext=[f"{v:.0f}" for v in tick_vals], range=[0, max_val * 1.05]),
            )
            figures.append(("Average and median HIN index by route", _polish_figure(fig), route_summary))
    return figures


def _hin_ka_bubble_figure(hin):
    """Replace confusing point bubble chart with score-band summary.

    If trusted row-level KSI is unavailable, show crash counts by HIN band and
    skip KSI rather than plotting impossible KSI totals.
    """
    metric = _dashboard_hin_metric(hin)
    if hin is None or getattr(hin, "empty", True) or not metric:
        return None, pd.DataFrame()
    work = _drop_geometry(hin).copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    count_col = _crash_count_col(work)
    if not count_col:
        return None, pd.DataFrame()
    work["Total crashes"] = pd.to_numeric(work[count_col], errors="coerce").fillna(0)
    ksi = _strict_hin_ksi_series(work)
    has_ksi = len(ksi) == len(work) and float(pd.to_numeric(ksi, errors="coerce").fillna(0).sum()) > 0
    if has_ksi:
        work["KSI (K+A) crashes"] = pd.to_numeric(ksi, errors="coerce").fillna(0)

    max_score = max(float(work[metric].max()), 1.0)
    bin_width = 10 if max_score <= 100 else max(max_score / 10.0, 1.0)
    bins = pd.interval_range(start=0, end=max_score + bin_width, freq=bin_width, closed="left")
    work["HIN score range"] = pd.cut(work[metric], bins=bins, include_lowest=True).astype(str)
    agg = work.groupby("HIN score range", observed=False).agg(**{"Total crashes": ("Total crashes", "sum")}).reset_index()
    if has_ksi:
        ksi_agg = work.groupby("HIN score range", observed=False)["KSI (K+A) crashes"].sum().reset_index()
        agg = agg.merge(ksi_agg, on="HIN score range", how="left")
        plot_cols = ["Total crashes", "KSI (K+A) crashes"]
        title = "Crashes and KSI (K+A) by HIN score range"
    else:
        plot_cols = ["Total crashes"]
        title = "Crashes by HIN score range"

    fig = px.bar(
        agg,
        x="HIN score range",
        y=plot_cols,
        barmode="group",
        title=title,
    )
    fig.update_layout(xaxis_title="HIN priority index range", yaxis_title="Crash count", legend_title="Measure")
    return _polish_figure(fig), agg


def _render_hin_dashboard_charts(st, hin):
    if hin is None or getattr(hin, "empty", True):
        return
    st.markdown(
        "<div class='dashboard-section-title'>HIN index diagnostics <span>distribution by score range, route average/median comparison, and crash/KSI relationship</span></div>",
        unsafe_allow_html=True,
    )
    st.caption(
        "These charts use all existing HIN result rows and are not affected by the map display choice such as above-average or above-median."
    )
    figures = _hin_distribution_figures(hin)
    relation_fig, relation_df = _hin_ka_bubble_figure(hin)
    chart_items = figures[:]
    if relation_fig is not None:
        chart_items.append(("Crashes/KSI by HIN score range", relation_fig, relation_df))
    if not chart_items:
        st.info("HIN diagnostic charts need HIN priority index values.")
        return
    cols = st.columns(2)
    for i, (title, fig, data) in enumerate(chart_items[:4]):
        with cols[i % 2]:
            fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35))
            st.plotly_chart(_polish_figure(fig), width="stretch", key=f"hin_diag_v40_{_safe_name(title)}_{i}")


def _render_hin_network_summary(st, hin, crashes):
    """Dashboard-only HIN summary with safer crash/KSI capture."""
    if hin is None or getattr(hin, "empty", True):
        return
    metric = _dashboard_hin_metric(hin)
    if not metric:
        return
    st.markdown("<div class='dashboard-section-title'>High Injury Network summary <span>custom HIN threshold, miles, crashes, and KSI (K+A) capture</span></div>", unsafe_allow_html=True)
    work = hin.copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)

    _render_crash_ksi_mapping_controls(st, crashes)

    length_col = _dashboard_length_col(work)
    if length_col:
        work[length_col] = pd.to_numeric(work[length_col], errors="coerce").fillna(0)
    else:
        length_col = "__unit_length__"
        work[length_col] = 1.0

    mode_col, control_col, c_miles, c_crashes, c_ksi = st.columns([1.35, 1.05, 1, 1, 1])
    method = mode_col.selectbox(
        "High-risk network threshold",
        ["Top percent of miles", "Top number of segments/windows", "HIN index threshold", "Above average HIN index", "Above median HIN index"],
        index=0,
        key="hin_summary_threshold_mode_v40",
    )
    top_percent = 10.0
    top_n = 20
    index_threshold = 50.0
    with control_col:
        if method == "Top percent of miles":
            top_percent = st.number_input("Top percent", min_value=1.0, max_value=100.0, value=float(st.session_state.get("hin_summary_top_percent", 10.0)), step=1.0, key="hin_summary_top_percent")
        elif method == "Top number of segments/windows":
            top_n = st.number_input("Top N", min_value=1, max_value=max(int(len(work)), 1), value=min(20, max(int(len(work)), 1)), step=1, key="hin_summary_top_n")
        elif method == "HIN index threshold":
            index_threshold = st.number_input("Minimum HIN index", min_value=0.0, max_value=100.0, value=float(st.session_state.get("hin_summary_index_threshold", 50.0)), step=5.0, key="hin_summary_index_threshold")
        elif method == "Above average HIN index":
            st.metric("Average", f"{work[metric].mean():,.2f}")
        elif method == "Above median HIN index":
            st.metric("Median", f"{work[metric].median():,.2f}")

    selected = _selected_hin_subset(work, metric, method, top_percent=top_percent, top_n=top_n, index_threshold=index_threshold)
    total_mi = float(work[length_col].sum()) if work[length_col].sum() else float(len(work))
    high_mi = float(selected[length_col].sum()) if not selected.empty else 0.0
    pct_mi = high_mi / total_mi * 100 if total_mi else 0.0

    crash_subset, capture_note = _selected_hin_crash_subset(selected, crashes)
    total_crashes = _safe_unique_crash_count(crashes) if crashes is not None else 0
    if crash_subset is not None:
        high_crashes = _safe_unique_crash_count(crash_subset)
        pct_crash = high_crashes / total_crashes * 100 if total_crashes else 0.0
        crash_value = f"{high_crashes:,}"
        crash_delta = f"{pct_crash:,.1f}% of assigned crashes"
        high_ksi = _total_ka_from_crashes(crash_subset)
        total_ksi = _total_ka_from_crashes(crashes)
        pct_ksi = high_ksi / total_ksi * 100 if total_ksi else 0.0
        ksi_value = f"{high_ksi:,}"
        ksi_delta = f"{pct_ksi:,.1f}% of KSI (K+A)"
    else:
        count_col = _crash_count_col(selected)
        if count_col and not selected.empty and total_crashes:
            # Display row-based count as an estimate, capped only for the percent.
            row_crashes = float(pd.to_numeric(selected[count_col], errors="coerce").fillna(0).sum())
            pct_crash = min(row_crashes, total_crashes) / total_crashes * 100
            crash_value = f"{row_crashes:,.0f}"
            crash_delta = f"{pct_crash:,.1f}% of loaded crashes; row-based"
        else:
            crash_value = "N/A"
            crash_delta = "No crash linkage"
        ksi_value = "N/A"
        ksi_delta = "Needs crash-to-HIN linkage"

    c_miles.metric("High-risk miles", f"{high_mi:,.2f} mi", f"{pct_mi:,.1f}% of analyzed miles")
    c_crashes.metric("Crashes on selected HIN", crash_value, crash_delta)
    c_ksi.metric("KSI (K+A) on selected HIN", ksi_value, ksi_delta)

    selected_display = _add_dashboard_rate_columns(selected, crashes=crashes)
    with st.expander("Selected HIN summary rows", expanded=False):
        preview_cols = []
        for c in [
            _normal_col(selected_display, ["RiskSegmentID", "WindowID", "SegmentID", "UnitID", "CorridorID"]),
            _dashboard_route_col(selected_display),
            _dashboard_length_col(selected_display), _crash_count_col(selected_display),
            "KSI_Crashes_Dashboard", metric, "Dashboard_Analysis_Years", "Crash_per_Mile_per_Year", "KSI_per_Mile_per_Year",
        ]:
            if c and c in selected_display.columns and c not in preview_cols:
                preview_cols.append(c)
        if preview_cols:
            st.dataframe(_safe_dataframe_for_display(selected_display[preview_cols].head(50)), width="stretch", hide_index=True)
        else:
            st.info("No displayable HIN rows are available for the selected threshold.")
    years = _dashboard_year_count(crashes)
    st.caption(
        f"Crashes/mile/year uses the detected crash data year span: {years:g} year(s). {capture_note} "
        "Dashboard controls summarize existing HIN results only; they do not recalculate HIN scores."
    )
    _render_hin_dashboard_charts(st, hin)


def _hin_table_for_display(hin, metric, top_n=20):
    """Decision table for top HIN windows/segments with safe display rates."""
    if hin is None or getattr(hin, "empty", True) or metric not in hin.columns:
        return pd.DataFrame()
    crashes = None
    try:
        crashes = _available_tables(st).get("Assigned crashes", _available_tables(st).get("Uploaded crashes"))
    except Exception:
        crashes = None
    work = _add_dashboard_rate_columns(hin.copy(), crashes=crashes)
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work = work.sort_values(metric, ascending=False).head(top_n).reset_index(drop=True)
    id_col = _normal_col(work, ["RiskSegmentID", "WindowID", "SlidingWindowID", "SegmentID", "UnitID", "SourceSegmentID", "CorridorID"])
    route_col = _dashboard_route_col(work)
    length_col = _dashboard_length_col(work)
    from_col = _normal_col(work, ["FromMile", "From_Mile", "from_mile", "BeginMile", "StartMile", "WindowFromMile"])
    to_col = _normal_col(work, ["ToMile", "To_Mile", "to_mile", "EndMile", "WindowToMile"])
    count_col = _crash_count_col(work)
    out = pd.DataFrame()
    out["Rank"] = range(1, len(work) + 1)
    out["SegID"] = work[id_col].astype(str).values if id_col else [f"HIN_{i + 1}" for i in range(len(work))]
    out["Seg/window length"] = pd.to_numeric(work[length_col], errors="coerce").round(3).values if length_col else ""
    out["From mile"] = pd.to_numeric(work[from_col], errors="coerce").round(3).values if from_col else ""
    out["To mile"] = pd.to_numeric(work[to_col], errors="coerce").round(3).values if to_col else ""
    out["Route"] = work[route_col].astype(str).values if route_col else ""
    out["Crash count"] = pd.to_numeric(work[count_col], errors="coerce").fillna(0).round(0).astype(int).values if count_col else ""
    if "KSI_Crashes_Dashboard" in work.columns:
        out["KSI (K+A) crashes"] = pd.to_numeric(work["KSI_Crashes_Dashboard"], errors="coerce").fillna(0).round(0).astype(int).values
    out["HIN index"] = pd.to_numeric(work[metric], errors="coerce").round(3).values
    if "Crash_per_Mile" in work.columns:
        out["Crashes/mile"] = pd.to_numeric(work["Crash_per_Mile"], errors="coerce").round(3).values
    if "Crash_per_Mile_per_Year" in work.columns:
        out["Crashes/mile/year"] = pd.to_numeric(work["Crash_per_Mile_per_Year"], errors="coerce").round(3).values
    if "KSI_per_Mile_per_Year" in work.columns:
        out["KSI (K+A)/mile/year"] = pd.to_numeric(work["KSI_per_Mile_per_Year"], errors="coerce").round(3).values
    if "Dashboard_Analysis_Years" in work.columns:
        out["Crash years used"] = pd.to_numeric(work["Dashboard_Analysis_Years"], errors="coerce").round(1).values
    return out


# --- V41 dashboard-only corrections: exact HIN crash/KSI counts by segment geometry ---
# These helpers override earlier dashboard-only functions. They do not change any
# workflow calculation result; they only calculate dashboard/report display fields
# from existing HIN geometries and existing crash records.


def _dashboard_hin_source_gdf(hin=None):
    """Return the HIN GeoDataFrame from session when the dashboard table lost geometry."""
    try:
        if hin is not None and "geometry" in getattr(hin, "columns", []) and not getattr(hin, "empty", True):
            return hin.copy()
    except Exception:
        pass

    try:
        results = st.session_state.get("section7_results") or {}
        for key in ["risk_segments", "risk_windows", "risk_corridors"]:
            gdf = results.get(key)
            if gdf is not None and not getattr(gdf, "empty", True) and "geometry" in gdf.columns:
                if hin is None or len(gdf) == len(hin) or key == "risk_segments":
                    return gdf.copy()
    except Exception:
        pass

    return hin.copy() if hin is not None else pd.DataFrame()


def _dashboard_crashes_source_gdf(crashes=None):
    """Return crash GeoDataFrame for HIN dashboard counting.

    Prefer the sliding-window assigned crashes because they are the exact crash
    records used by the HIN workflow. Fall back to other existing crash tables.
    """
    candidates = []
    try:
        results = st.session_state.get("section7_results") or {}
        candidates.append(results.get("assigned_crashes"))
    except Exception:
        pass
    try:
        candidates.append(st.session_state.get("section7_crashes_for_map"))
    except Exception:
        pass
    try:
        candidates.append(st.session_state.get("assigned_crashes"))
    except Exception:
        pass
    try:
        candidates.append(st.session_state.get("crashes"))
    except Exception:
        pass
    candidates.append(crashes)

    for cand in candidates:
        try:
            if cand is not None and not getattr(cand, "empty", True) and "geometry" in cand.columns:
                return cand.copy()
        except Exception:
            pass
    return crashes.copy() if crashes is not None else pd.DataFrame()


def _strict_hin_ksi_series(df):
    """Do not infer KSI from HIN/window score rows.

    KSI (K+A) must come from crash-level severity records, not from HIN row
    fields, because HIN rows can contain window scores or duplicated summaries.
    """
    if df is None or getattr(df, "empty", True):
        return pd.Series(dtype="float64")
    return pd.Series([0] * len(df), index=df.index, dtype="float64")


def _hin_ka_series(df):
    return _strict_hin_ksi_series(df)


def _render_crash_ksi_mapping_controls(st, crashes):
    """KSI mapping removed.

    The dashboard now uses crash-level KABCO/severity fields already available
    from the workflow. This avoids asking users to map non-existent combined
    KSI columns and prevents accidental use of injury-count fields as crash
    counts.
    """
    return


def _crash_level_ksi_series(crashes):
    """Return one 0/1 KSI flag per crash row.

    KSI means a crash is fatal or serious-injury (K or A).  The preferred
    source is the mapped/canonical crash-level severity field.  If the dataset
    does not have a single KABCO field, use the mapped Fatalities and Serious
    Injuries person-count fields as crash-level flags: any value greater than
    zero means that crash is K or A.  The returned value is still 0/1 per crash
    row, not a sum of injured people.
    """
    if crashes is None or getattr(crashes, "empty", True):
        return pd.Series(dtype="float64")

    work = crashes.copy()

    kabco = _kabco_col(work)
    if kabco and kabco in work.columns:
        vals = work[kabco].map(normalize_kabco_value).astype(str).str.upper()
        if vals.isin(["K", "A"]).any():
            return vals.isin(["K", "A"]).astype(int)

    fatal_col = _mapped_dashboard_col(
        work,
        "fatalities",
        ["DashboardFatalities", "Fatalities", "FATALITIES", "Fatals", "FATALS", "Fatal", "K"],
    )
    serious_col = _mapped_dashboard_col(
        work,
        "serious_injuries",
        ["DashboardSeriousInjuries", "SeriousInjuries", "Serious_Injuries", "Level_A_Injuries", "A_Injuries", "A"],
    )

    flags = pd.Series([0] * len(work), index=work.index, dtype="int64")
    if fatal_col and fatal_col in work.columns:
        fatal_vals = pd.to_numeric(work[fatal_col], errors="coerce").fillna(0)
        flags = flags | (fatal_vals > 0).astype(int)
    if serious_col and serious_col in work.columns:
        serious_vals = pd.to_numeric(work[serious_col], errors="coerce").fillna(0)
        flags = flags | (serious_vals > 0).astype(int)
    if flags.sum() > 0:
        return flags.astype(int)

    k_col = _normal_col(work, ["K_Crash", "K_Crashes", "Fatal_Crash", "Fatal_Crashes", "Crash_K", "Is_K"])
    a_col = _normal_col(work, ["A_Crash", "A_Crashes", "Serious_Injury_Crash", "Serious_Injury_Crashes", "Crash_A", "Is_A"])
    if k_col or a_col:
        k = pd.to_numeric(work[k_col], errors="coerce").fillna(0) if k_col else 0
        a = pd.to_numeric(work[a_col], errors="coerce").fillna(0) if a_col else 0
        out = pd.to_numeric(k + a, errors="coerce").fillna(0)
        return (out > 0).astype(int)

    return pd.Series([0] * len(work), index=work.index, dtype="float64")


def _total_ka_from_crashes(crashes):
    vals = _crash_level_ksi_series(crashes)
    if vals is None or len(vals) == 0:
        return 0
    return int(pd.to_numeric(vals, errors="coerce").fillna(0).sum())


def _dashboard_hin_row_id_col(df):
    return "__Dashboard_HIN_RowID__"


def _prepare_hin_gdf_for_counting(hin):
    gdf = _dashboard_hin_source_gdf(hin)
    if gdf is None or getattr(gdf, "empty", True):
        return gdf
    gdf = gdf.copy().reset_index(drop=True)
    gdf[_dashboard_hin_row_id_col(gdf)] = range(len(gdf))
    return gdf


def _dashboard_hin_crash_join(hin, crashes=None):
    """Assign each crash to one nearest HIN row for dashboard-only counts."""
    if gpd is None:
        return pd.DataFrame()

    hin_gdf = _prepare_hin_gdf_for_counting(hin)
    crash_gdf = _dashboard_crashes_source_gdf(crashes)

    if hin_gdf is None or getattr(hin_gdf, "empty", True) or "geometry" not in hin_gdf.columns:
        return pd.DataFrame()
    if crash_gdf is None or getattr(crash_gdf, "empty", True) or "geometry" not in crash_gdf.columns:
        return pd.DataFrame()

    try:
        if not isinstance(hin_gdf, gpd.GeoDataFrame):
            hin_gdf = gpd.GeoDataFrame(hin_gdf, geometry="geometry")
        if not isinstance(crash_gdf, gpd.GeoDataFrame):
            crash_gdf = gpd.GeoDataFrame(crash_gdf, geometry="geometry")
        if hin_gdf.crs is None:
            hin_gdf = hin_gdf.set_crs(4326, allow_override=True)
        if crash_gdf.crs is None:
            crash_gdf = crash_gdf.set_crs(4326, allow_override=True)
        hin_m = hin_gdf.to_crs(epsg=3857)
        crash_m = crash_gdf.to_crs(epsg=3857)
    except Exception:
        return pd.DataFrame()

    row_id = _dashboard_hin_row_id_col(hin_m)
    snap_ft = float(st.session_state.get("section7_crash_snap_dist_ft", 150.0))
    snap_m = snap_ft * 0.3048

    keep_cols = [row_id, "geometry"]
    try:
        joined = gpd.sjoin_nearest(
            crash_m,
            hin_m[keep_cols],
            how="inner",
            max_distance=snap_m,
            distance_col="Dashboard_DistToHIN_M",
        )
    except Exception:
        return pd.DataFrame()

    if joined is None or joined.empty:
        return pd.DataFrame()

    # Preserve one row per crash. sjoin_nearest can return ties; keep nearest/first.
    cid = _dashboard_crash_id_col(joined) or "__Dashboard_Crash_RowID__"
    if cid == "__Dashboard_Crash_RowID__":
        joined[cid] = range(len(joined))
    if "Dashboard_DistToHIN_M" in joined.columns:
        joined = joined.sort_values("Dashboard_DistToHIN_M")
    joined = joined.drop_duplicates(subset=[cid], keep="first").copy()
    joined["Dashboard_KSI_Flag"] = pd.to_numeric(_crash_level_ksi_series(joined), errors="coerce").fillna(0).astype(int).values
    return joined


def _hin_row_crash_stats(hin, crashes=None):
    """Return exact dashboard crash/KSI counts for each HIN row."""
    hin_gdf = _prepare_hin_gdf_for_counting(hin)
    if hin_gdf is None or getattr(hin_gdf, "empty", True):
        return pd.DataFrame()
    row_id = _dashboard_hin_row_id_col(hin_gdf)
    base = pd.DataFrame({row_id: hin_gdf[row_id].values})

    joined = _dashboard_hin_crash_join(hin_gdf, crashes=crashes)
    if joined is None or joined.empty or row_id not in joined.columns:
        base["Dashboard_Crash_Count"] = 0
        base["KSI_Crashes_Dashboard"] = 0
        return base

    cid = _dashboard_crash_id_col(joined) or "__Dashboard_Crash_RowID__"
    if cid not in joined.columns:
        joined[cid] = range(len(joined))

    grouped = joined.groupby(row_id, dropna=False).agg(
        Dashboard_Crash_Count=(cid, "nunique"),
        KSI_Crashes_Dashboard=("Dashboard_KSI_Flag", "sum"),
    ).reset_index()

    out = base.merge(grouped, on=row_id, how="left")
    out["Dashboard_Crash_Count"] = pd.to_numeric(out["Dashboard_Crash_Count"], errors="coerce").fillna(0).astype(int)
    out["KSI_Crashes_Dashboard"] = pd.to_numeric(out["KSI_Crashes_Dashboard"], errors="coerce").fillna(0).astype(int)
    return out


def _add_dashboard_rate_columns(df, crashes=None):
    """Add display-only exact crash/mile/year and KSI/mile/year fields.

    For HIN rows with geometry, crash count and KSI are counted by assigning
    existing crash records to the nearest HIN segment/window. This fixes the
    previous dashboard issue where HIN window scores were mislabeled as segment
    crash totals.
    """
    if df is None or getattr(df, "empty", True):
        return df

    work_src = _dashboard_hin_source_gdf(df)
    out = work_src.copy().reset_index(drop=True)
    row_id = _dashboard_hin_row_id_col(out)
    out[row_id] = range(len(out))

    crash_source = _dashboard_crashes_source_gdf(crashes)
    stats = _hin_row_crash_stats(out, crashes=crash_source)
    if stats is not None and not stats.empty:
        out = out.drop(columns=["Dashboard_Crash_Count", "KSI_Crashes_Dashboard"], errors="ignore")
        out = out.merge(stats, on=row_id, how="left")
    else:
        out["Dashboard_Crash_Count"] = 0
        out["KSI_Crashes_Dashboard"] = 0

    out["Dashboard_Crash_Count"] = pd.to_numeric(out["Dashboard_Crash_Count"], errors="coerce").fillna(0).astype(int)
    out["KSI_Crashes_Dashboard"] = pd.to_numeric(out["KSI_Crashes_Dashboard"], errors="coerce").fillna(0).astype(int)

    years = _dashboard_year_count(crash_source)
    out["Dashboard_Analysis_Years"] = years

    length_col = _dashboard_length_col(out)
    if length_col:
        length = pd.to_numeric(out[length_col], errors="coerce").replace(0, pd.NA)
        out["Crash_per_Mile"] = (out["Dashboard_Crash_Count"] / length).astype("float64").replace([float("inf"), -float("inf")], 0).fillna(0).round(3)
        out["Crash_per_Mile_per_Year"] = (out["Dashboard_Crash_Count"] / length / years).astype("float64").replace([float("inf"), -float("inf")], 0).fillna(0).round(3)
        out["KSI_per_Mile_per_Year"] = (out["KSI_Crashes_Dashboard"] / length / years).astype("float64").replace([float("inf"), -float("inf")], 0).fillna(0).round(3)

    return out


def _selected_hin_crash_subset(selected_hin, crashes):
    """Return unique crash records assigned to selected HIN geometry."""
    joined = _dashboard_hin_crash_join(selected_hin, crashes=crashes)
    if joined is None or joined.empty:
        return pd.DataFrame(), "Dashboard counted crashes by spatially assigning existing crash records to selected HIN rows; no crashes matched the selected HIN within the snap distance."
    cid = _dashboard_crash_id_col(joined) or "__Dashboard_Crash_RowID__"
    if cid not in joined.columns:
        joined[cid] = range(len(joined))
    subset = joined.drop_duplicates(subset=[cid], keep="first").copy()
    return subset, "Dashboard counted unique crashes by assigning existing crash records to the selected HIN segment/window geometries."


def _hin_ka_bubble_figure(hin):
    """Simple scatter: KSI crash count versus HIN priority index."""
    if hin is None or getattr(hin, "empty", True):
        return None, pd.DataFrame()
    metric = _dashboard_hin_metric(hin)
    if not metric:
        return None, pd.DataFrame()

    crashes = _dashboard_crashes_source_gdf(None)
    work = _add_dashboard_rate_columns(hin, crashes=crashes)
    if work is None or getattr(work, "empty", True) or metric not in work.columns:
        return None, pd.DataFrame()

    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work["KSI (K+A) crashes"] = pd.to_numeric(work.get("KSI_Crashes_Dashboard", 0), errors="coerce").fillna(0).astype(int)
    work["Crash count"] = pd.to_numeric(work.get("Dashboard_Crash_Count", 0), errors="coerce").fillna(0).astype(int)

    route_col = _dashboard_route_col(work)
    id_col = _normal_col(work, ["RiskSegmentID", "WindowID", "SlidingWindowID", "SegmentID", "UnitID", "SourceSegmentID", "CorridorID"])
    length_col = _dashboard_length_col(work)
    hover_cols = [c for c in [id_col, route_col, length_col, "Crash count", "KSI (K+A) crashes"] if c and c in work.columns]

    fig = px.scatter(
        work,
        x=metric,
        y="KSI (K+A) crashes",
        hover_data=hover_cols,
        title="KSI (K+A) crashes vs HIN priority index",
    )
    fig.update_traces(marker=dict(size=8, opacity=0.65))
    fig.update_layout(xaxis_title="HIN priority index", yaxis_title="KSI (K+A) crash count")
    return _polish_figure(fig), work[[c for c in [metric, "KSI (K+A) crashes", "Crash count", route_col, id_col] if c and c in work.columns]].copy()


def _render_hin_dashboard_charts(st, hin):
    if hin is None or getattr(hin, "empty", True):
        return
    st.markdown(
        "<div class='dashboard-section-title'>HIN index diagnostics <span>distribution, route average/median comparison, and KSI scatter</span></div>",
        unsafe_allow_html=True,
    )
    st.caption(
        "These charts use existing HIN result rows. They are not affected by the Step 4 map display choice such as above-average or above-median."
    )
    figures = _hin_distribution_figures(_dashboard_hin_source_gdf(hin))
    scatter_fig, scatter_df = _hin_ka_bubble_figure(hin)
    chart_items = figures[:]
    if scatter_fig is not None:
        chart_items.append(("KSI scatter", scatter_fig, scatter_df))
    if not chart_items:
        st.info("HIN diagnostic charts need HIN priority index values.")
        return
    cols = st.columns(2)
    for i, (title, fig, data) in enumerate(chart_items[:4]):
        with cols[i % 2]:
            fig.update_layout(height=_chart_height(), margin=dict(l=20, r=20, t=45, b=35))
            st.plotly_chart(_polish_figure(fig), width="stretch", key=f"hin_diag_v41_{_safe_name(title)}_{i}")


def _render_hin_network_summary(st, hin, crashes):
    """Dashboard HIN summary using exact crash/KSI counts from selected geometries."""
    if hin is None or getattr(hin, "empty", True):
        return
    metric = _dashboard_hin_metric(hin)
    if not metric:
        return

    st.markdown("<div class='dashboard-section-title'>High Injury Network summary <span>custom HIN threshold, miles, crashes, and KSI (K+A) capture</span></div>", unsafe_allow_html=True)
    work = _dashboard_hin_source_gdf(hin).copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)

    length_col = _dashboard_length_col(work)
    if length_col:
        work[length_col] = pd.to_numeric(work[length_col], errors="coerce").fillna(0)
    else:
        length_col = "__unit_length__"
        work[length_col] = 1.0

    mode_col, control_col, c_miles, c_crashes, c_ksi = st.columns([1.35, 1.05, 1, 1, 1])
    method = mode_col.selectbox(
        "High-risk network threshold",
        ["Top percent of miles", "Top number of segments/windows", "HIN index threshold", "Above average HIN index", "Above median HIN index"],
        index=0,
        key="hin_summary_threshold_mode_v41",
    )
    top_percent = 10.0
    top_n = 20
    index_threshold = 50.0
    with control_col:
        if method == "Top percent of miles":
            top_percent = st.number_input("Top percent", min_value=1.0, max_value=100.0, value=float(st.session_state.get("hin_summary_top_percent", 10.0)), step=1.0, key="hin_summary_top_percent")
        elif method == "Top number of segments/windows":
            top_n = st.number_input("Top N", min_value=1, max_value=max(int(len(work)), 1), value=min(20, max(int(len(work)), 1)), step=1, key="hin_summary_top_n_v41")
        elif method == "HIN index threshold":
            index_threshold = st.number_input("Minimum HIN index", min_value=0.0, max_value=100.0, value=float(st.session_state.get("hin_summary_index_threshold", 50.0)), step=5.0, key="hin_summary_index_threshold")
        elif method == "Above average HIN index":
            st.metric("Average", f"{work[metric].mean():,.2f}")
        elif method == "Above median HIN index":
            st.metric("Median", f"{work[metric].median():,.2f}")

    selected = _selected_hin_subset(work, metric, method, top_percent=top_percent, top_n=top_n, index_threshold=index_threshold)
    total_mi = float(work[length_col].sum()) if work[length_col].sum() else float(len(work))
    high_mi = float(selected[length_col].sum()) if not selected.empty else 0.0
    pct_mi = high_mi / total_mi * 100 if total_mi else 0.0

    crash_source = _dashboard_crashes_source_gdf(crashes)
    total_crashes = _safe_unique_crash_count(crash_source)

    selected_is_full_network = False
    try:
        selected_is_full_network = len(selected) >= len(work) and len(work) > 0
    except Exception:
        selected_is_full_network = False
    if selected_is_full_network:
        crash_subset = crash_source.copy() if crash_source is not None else pd.DataFrame()
        capture_note = "The selected HIN includes 100% of analyzed HIN rows, so the dashboard uses all loaded HIN crash records for the selected-network crash and KSI capture cards."
    else:
        crash_subset, capture_note = _selected_hin_crash_subset(selected, crash_source)

    high_crashes = _safe_unique_crash_count(crash_subset)
    pct_crash = high_crashes / total_crashes * 100 if total_crashes else 0.0

    total_ksi = _total_ka_from_crashes(crash_source)
    high_ksi = _total_ka_from_crashes(crash_subset)
    pct_ksi = high_ksi / total_ksi * 100 if total_ksi else 0.0

    c_miles.metric("High-risk miles", f"{high_mi:,.2f} mi", f"{pct_mi:,.1f}% of analyzed miles")
    c_crashes.metric("Crashes on selected HIN", f"{high_crashes:,}", f"{pct_crash:,.1f}% of loaded HIN crashes")
    c_ksi.metric("KSI (K+A) on selected HIN", f"{high_ksi:,}", f"{pct_ksi:,.1f}% of KSI (K+A)")

    selected_display = _add_dashboard_rate_columns(selected, crashes=crash_source)
    with st.expander("Selected HIN summary rows", expanded=False):
        preview_cols = []
        for c in [
            _normal_col(selected_display, ["RiskSegmentID", "WindowID", "SegmentID", "UnitID", "CorridorID"]),
            _dashboard_route_col(selected_display),
            _dashboard_length_col(selected_display), "Dashboard_Crash_Count", "KSI_Crashes_Dashboard",
            metric, "Dashboard_Analysis_Years", "Crash_per_Mile_per_Year", "KSI_per_Mile_per_Year",
        ]:
            if c and c in selected_display.columns and c not in preview_cols:
                preview_cols.append(c)
        if preview_cols:
            st.dataframe(_safe_dataframe_for_display(selected_display[preview_cols].head(50)), width="stretch", hide_index=True)
        else:
            st.info("No displayable HIN rows are available for the selected threshold.")

    years = _dashboard_year_count(crash_source)
    st.caption(
        f"Crashes/mile/year uses the detected crash data year span: {years:g} year(s). {capture_note} "
        "Dashboard controls summarize existing HIN results only; they do not recalculate HIN scores."
    )
    _render_hin_dashboard_charts(st, work)


def _hin_table_for_display(hin, metric, top_n=20):
    """Decision table for top HIN rows with exact dashboard crash/KSI counts."""
    if hin is None or getattr(hin, "empty", True) or metric not in hin.columns:
        return pd.DataFrame()
    crash_source = _dashboard_crashes_source_gdf(None)
    work = _add_dashboard_rate_columns(hin.copy(), crashes=crash_source)
    if metric not in work.columns:
        return pd.DataFrame()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work = work.sort_values(metric, ascending=False).head(top_n).reset_index(drop=True)

    id_col = _normal_col(work, ["RiskSegmentID", "WindowID", "SlidingWindowID", "SegmentID", "UnitID", "SourceSegmentID", "CorridorID"])
    route_col = _dashboard_route_col(work)
    length_col = _dashboard_length_col(work)
    from_col = _normal_col(work, ["FromMile", "From_Mile", "from_mile", "BeginMile", "StartMile", "WindowFromMile", "Win_From_Mi"])
    to_col = _normal_col(work, ["ToMile", "To_Mile", "to_mile", "EndMile", "WindowToMile", "Win_To_Mi"])

    out = pd.DataFrame()
    out["Rank"] = range(1, len(work) + 1)
    out["SegID"] = work[id_col].astype(str).values if id_col else [f"HIN_{i + 1}" for i in range(len(work))]
    out["Seg/window length"] = pd.to_numeric(work[length_col], errors="coerce").round(3).values if length_col else ""
    out["From mile"] = pd.to_numeric(work[from_col], errors="coerce").round(3).values if from_col else ""
    out["To mile"] = pd.to_numeric(work[to_col], errors="coerce").round(3).values if to_col else ""
    out["Route"] = work[route_col].astype(str).values if route_col else ""
    out["Crash count"] = pd.to_numeric(work["Dashboard_Crash_Count"], errors="coerce").fillna(0).round(0).astype(int).values if "Dashboard_Crash_Count" in work.columns else ""
    out["KSI (K+A) crashes"] = pd.to_numeric(work["KSI_Crashes_Dashboard"], errors="coerce").fillna(0).round(0).astype(int).values if "KSI_Crashes_Dashboard" in work.columns else ""
    out["HIN index"] = pd.to_numeric(work[metric], errors="coerce").round(3).values
    if "Crash_per_Mile" in work.columns:
        out["Crashes/mile"] = pd.to_numeric(work["Crash_per_Mile"], errors="coerce").round(3).values
    if "Crash_per_Mile_per_Year" in work.columns:
        out["Crashes/mile/year"] = pd.to_numeric(work["Crash_per_Mile_per_Year"], errors="coerce").round(3).values
    if "KSI_per_Mile_per_Year" in work.columns:
        out["KSI (K+A)/mile/year"] = pd.to_numeric(work["KSI_per_Mile_per_Year"], errors="coerce").round(3).values
    if "Dashboard_Analysis_Years" in work.columns:
        out["Crash years used"] = pd.to_numeric(work["Dashboard_Analysis_Years"], errors="coerce").round(1).values
    return out


# --- V42 dashboard-only corrections: use actual sliding-window rows for HIN dashboard ---
# The HIN dashboard/ranking should summarize the sliding windows created by
# modules/sliding_window.py.  It should not re-snap crashes to the display
# lines, because window/segment line geometries can produce zero matches when
# points are not exactly on the line.  Instead, use the module outputs:
#   section7_results["risk_windows"]      -> each window, with window score/count
#   section7_results["assigned_crashes"]  -> crash records snapped to route with Route_Pos_M
# This does not change the sliding-window calculation; it only displays its
# existing outputs with additional KSI/rate fields.


def _section7_results_for_dashboard():
    try:
        return st.session_state.get("section7_results") or {}
    except Exception:
        return {}


def _section7_route_col_from_tables(windows=None, crashes=None):
    route_session = None
    try:
        route_session = st.session_state.get("section7_route_col_s7") or st.session_state.get("route_col")
    except Exception:
        route_session = None
    candidates = [route_session, "Route", "FULLNAME", "RouteName", "RouteName_Calc", "RoadName", "Name"]
    for table in [windows, crashes]:
        if table is None or getattr(table, "empty", True):
            continue
        for col in candidates:
            if col and col in table.columns:
                return col
    return None


def _dashboard_hin_source_gdf(hin=None):
    """Prefer the real sliding-window rows for dashboard HIN summaries.

    The user-facing HIN ranking represents windows, so use risk_windows when it
    exists.  Fall back to the supplied table or risk_segments only when windows
    are unavailable.
    """
    results = _section7_results_for_dashboard()
    for key in ["risk_windows", "risk_segments", "risk_corridors"]:
        try:
            gdf = results.get(key)
            if gdf is not None and not getattr(gdf, "empty", True):
                return gdf.copy()
        except Exception:
            pass
    try:
        if hin is not None and not getattr(hin, "empty", True):
            return hin.copy()
    except Exception:
        pass
    return pd.DataFrame()


def _dashboard_length_col(df):
    return _normal_col(
        df,
        [
            "Window_Length_Mi",
            "WindowLength_Miles",
            "Length_Miles",
            "Length_Mi",
            "Seg_Length_Mi",
            "SegmentLength_Mile",
            "CorridorLength_Mile",
            "CorridorLength_Miles",
            "length_mi",
            "Miles",
        ],
    )


def _dashboard_hin_metric(hin):
    if hin is None or getattr(hin, "empty", True):
        return None
    for col in ["HIN_Priority_Index", "Risk_Score", "Window_HIN_Index"]:
        if col in hin.columns:
            return col
    if "Window_Score" in hin.columns:
        return "Window_Score"
    return _default_metric(_numeric_cols(hin))


def _window_hin_index_series(windows):
    if windows is None or getattr(windows, "empty", True):
        return pd.Series(dtype="float64")
    if "HIN_Priority_Index" in windows.columns:
        return pd.to_numeric(windows["HIN_Priority_Index"], errors="coerce").fillna(0)
    score_col = "Window_Score" if "Window_Score" in windows.columns else _dashboard_hin_metric(windows)
    if not score_col or score_col not in windows.columns:
        return pd.Series([0] * len(windows), index=windows.index, dtype="float64")
    score = pd.to_numeric(windows[score_col], errors="coerce").fillna(0)
    max_score = float(score.max()) if len(score) else 0.0
    if max_score <= 0:
        return pd.Series([0] * len(windows), index=windows.index, dtype="float64")
    return score / max_score * 100.0


def _prepare_window_dashboard_table(hin=None, crashes=None):
    """Return HIN windows with crash count, KSI, HIN index, and rates.

    Crash count is the window Crash_Count generated by modules/sliding_window.py.
    KSI is counted from section7 assigned crash records whose Route_Pos_M falls
    inside the window on the same route.  This matches the sliding-window process
    and avoids a second geometry snap in the dashboard.
    """
    windows = _dashboard_hin_source_gdf(hin)
    if windows is None or getattr(windows, "empty", True):
        return pd.DataFrame()

    work = windows.copy().reset_index(drop=True)
    metric_existing = _dashboard_hin_metric(work)
    work["Dashboard_HIN_Index"] = _window_hin_index_series(work).values
    if "HIN_Priority_Index" not in work.columns:
        work["HIN_Priority_Index"] = work["Dashboard_HIN_Index"]

    if "Dashboard_Crash_Count" not in work.columns:
        if "Crash_Count" in work.columns:
            work["Dashboard_Crash_Count"] = pd.to_numeric(work["Crash_Count"], errors="coerce").fillna(0).astype(int)
        elif "CrashCount" in work.columns:
            work["Dashboard_Crash_Count"] = pd.to_numeric(work["CrashCount"], errors="coerce").fillna(0).astype(int)
        else:
            work["Dashboard_Crash_Count"] = 0

    work["KSI_Crashes_Dashboard"] = 0

    results = _section7_results_for_dashboard()
    assigned = crashes
    if assigned is None or getattr(assigned, "empty", True):
        assigned = results.get("assigned_crashes")
    if assigned is None or getattr(assigned, "empty", True):
        assigned = _dashboard_crashes_source_gdf(crashes)

    route_col = _section7_route_col_from_tables(work, assigned)

    if (
        assigned is not None
        and not getattr(assigned, "empty", True)
        and route_col
        and route_col in work.columns
        and route_col in assigned.columns
        and "Route_Pos_M" in assigned.columns
        and "Win_Start_M" in work.columns
        and "Win_End_M" in work.columns
    ):
        ac = assigned.copy()
        ac["Dashboard_KSI_Flag"] = pd.to_numeric(_crash_level_ksi_series(ac), errors="coerce").fillna(0).astype(int).values
        ac["Route_Pos_M"] = pd.to_numeric(ac["Route_Pos_M"], errors="coerce")
        work["Win_Start_M"] = pd.to_numeric(work["Win_Start_M"], errors="coerce")
        work["Win_End_M"] = pd.to_numeric(work["Win_End_M"], errors="coerce")

        ksi_counts = []
        # Keep module crash counts for Crash_Count.  KSI is dashboard-only.
        for _, row in work.iterrows():
            route_value = row.get(route_col)
            start_m = row.get("Win_Start_M")
            end_m = row.get("Win_End_M")
            if pd.isna(start_m) or pd.isna(end_m):
                ksi_counts.append(0)
                continue
            mask = (
                ac[route_col].astype(str).eq(str(route_value))
                & (ac["Route_Pos_M"] >= float(start_m))
                & (ac["Route_Pos_M"] < float(end_m))
            )
            ksi_counts.append(int(ac.loc[mask, "Dashboard_KSI_Flag"].sum()))
        work["KSI_Crashes_Dashboard"] = ksi_counts

    years = _dashboard_year_count(assigned)
    work["Dashboard_Analysis_Years"] = years
    length_col = _dashboard_length_col(work)
    if length_col:
        length = pd.to_numeric(work[length_col], errors="coerce").replace(0, pd.NA)
        work["Crash_per_Mile"] = (pd.to_numeric(work["Dashboard_Crash_Count"], errors="coerce").fillna(0) / length).astype("float64").replace([float("inf"), -float("inf")], 0).fillna(0).round(3)
        work["Crash_per_Mile_per_Year"] = (pd.to_numeric(work["Dashboard_Crash_Count"], errors="coerce").fillna(0) / length / years).astype("float64").replace([float("inf"), -float("inf")], 0).fillna(0).round(3)
        work["KSI_per_Mile_per_Year"] = (pd.to_numeric(work["KSI_Crashes_Dashboard"], errors="coerce").fillna(0) / length / years).astype("float64").replace([float("inf"), -float("inf")], 0).fillna(0).round(3)

    return work


def _add_dashboard_rate_columns(df, crashes=None):
    """Dashboard display fields only; use sliding-window outputs when present."""
    prepared = _prepare_window_dashboard_table(df, crashes=crashes)
    if prepared is not None and not getattr(prepared, "empty", True):
        return prepared
    return df


def _selected_hin_crash_subset(selected_hin, crashes):
    """Return unique crash records falling inside selected sliding windows."""
    if selected_hin is None or getattr(selected_hin, "empty", True):
        return pd.DataFrame(), "No selected HIN windows."

    results = _section7_results_for_dashboard()
    assigned = crashes
    if assigned is None or getattr(assigned, "empty", True):
        assigned = results.get("assigned_crashes")
    if assigned is None or getattr(assigned, "empty", True):
        return pd.DataFrame(), "No sliding-window assigned crash table was available."

    route_col = _section7_route_col_from_tables(selected_hin, assigned)
    if not route_col or route_col not in selected_hin.columns or route_col not in assigned.columns:
        return pd.DataFrame(), "Selected-HIN crash capture needs a route column in both selected windows and assigned crashes."
    if "Route_Pos_M" not in assigned.columns or "Win_Start_M" not in selected_hin.columns or "Win_End_M" not in selected_hin.columns:
        return pd.DataFrame(), "Selected-HIN crash capture needs Route_Pos_M plus window start/end fields."

    ac = assigned.copy()
    ac["Route_Pos_M"] = pd.to_numeric(ac["Route_Pos_M"], errors="coerce")
    pieces = []
    for _, row in selected_hin.iterrows():
        start_m = pd.to_numeric(row.get("Win_Start_M"), errors="coerce")
        end_m = pd.to_numeric(row.get("Win_End_M"), errors="coerce")
        if pd.isna(start_m) or pd.isna(end_m):
            continue
        mask = (
            ac[route_col].astype(str).eq(str(row.get(route_col)))
            & (ac["Route_Pos_M"] >= float(start_m))
            & (ac["Route_Pos_M"] < float(end_m))
        )
        if mask.any():
            pieces.append(ac.loc[mask].copy())

    if not pieces:
        return pd.DataFrame(), "No crash records fell inside the selected HIN windows."

    subset = pd.concat(pieces, ignore_index=True)
    cid = _dashboard_crash_id_col(subset) or "CrashID_S7"
    if cid not in subset.columns:
        subset["CrashID_S7"] = range(1, len(subset) + 1)
        cid = "CrashID_S7"
    subset = subset.drop_duplicates(subset=[cid], keep="first").copy()
    return subset, "Dashboard counted unique crash records whose route position falls inside the selected sliding windows."


def _safe_unique_crash_count(crashes):
    if crashes is None or getattr(crashes, "empty", True):
        return 0
    cid = _dashboard_crash_id_col(crashes) or "CrashID_S7"
    if cid in crashes.columns:
        return int(crashes[cid].dropna().astype(str).nunique())
    return int(len(crashes))


def _hin_distribution_figures(hin):
    """HIN score distribution by score range.

    Y-axis is the number of windows/segments in each HIN index range. This is
    the clearer replacement for the earlier percentile curve.
    """
    figures = []
    work = _prepare_window_dashboard_table(hin)
    if work is None or getattr(work, "empty", True):
        return figures
    metric = "HIN_Priority_Index" if "HIN_Priority_Index" in work.columns else _dashboard_hin_metric(work)
    if not metric or metric not in work.columns:
        return figures
    values = pd.to_numeric(work[metric], errors="coerce").fillna(0).clip(lower=0, upper=100)
    bins = list(range(0, 105, 5))
    labels = [f"{i}-{i + 5}" for i in bins[:-1]]
    dist = pd.DataFrame({"HIN score range": pd.cut(values, bins=bins, labels=labels, include_lowest=True, right=True)})
    dist = dist.groupby("HIN score range", observed=False).size().reset_index(name="Window/segment count")
    dist["HIN score midpoint"] = [i + 2.5 for i in bins[:-1]]
    fig = px.line(
        dist,
        x="HIN score midpoint",
        y="Window/segment count",
        markers=True,
        title="HIN index distribution by score range",
    )
    fig.update_layout(
        xaxis_title="HIN priority index range",
        yaxis_title="Number of windows/segments",
    )
    figures.append(("HIN index distribution by score range", _polish_figure(fig), dist))

    route_col = _section7_route_col_from_tables(work, None)
    if route_col and route_col in work.columns:
        route_stats = work.copy()
        route_stats[metric] = pd.to_numeric(route_stats[metric], errors="coerce").fillna(0)
        route_stats = route_stats.groupby(route_col, dropna=False)[metric].agg(["mean", "median"]).reset_index()
        route_stats = route_stats.sort_values("mean", ascending=False).head(15)
        route_stats["Average HIN"] = route_stats["mean"]
        route_stats["Median HIN"] = route_stats["median"]
        plot_df = route_stats[[route_col, "Average HIN", "Median HIN"]].copy()
        long_df = plot_df.melt(id_vars=[route_col], value_vars=["Average HIN", "Median HIN"], var_name="Metric", value_name="Value")
        fig2 = px.bar(long_df, x="Value", y=route_col, color="Metric", orientation="h", title="Average and median HIN by route")
        fig2.update_layout(xaxis_title="HIN priority index", yaxis_title="Route", barmode="group")
        figures.append(("Average and median HIN by route", _polish_figure(fig2), route_stats))
    return figures


def _hin_ka_bubble_figure(hin):
    """Simple scatter: each window/segment is one point."""
    work = _prepare_window_dashboard_table(hin)
    if work is None or getattr(work, "empty", True):
        return None, pd.DataFrame()
    metric = "HIN_Priority_Index" if "HIN_Priority_Index" in work.columns else _dashboard_hin_metric(work)
    if not metric or metric not in work.columns:
        return None, pd.DataFrame()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work["KSI (K+A) crashes"] = pd.to_numeric(work.get("KSI_Crashes_Dashboard", 0), errors="coerce").fillna(0).astype(int)
    work["Crash count"] = pd.to_numeric(work.get("Dashboard_Crash_Count", 0), errors="coerce").fillna(0).astype(int)
    route_col = _section7_route_col_from_tables(work, None)
    id_col = _normal_col(work, ["SegID", "WindowID", "SlidingWindowID", "SegmentID", "UnitID", "SourceSegmentID", "CorridorID"])
    length_col = _dashboard_length_col(work)
    hover_cols = [c for c in [id_col, route_col, length_col, "Crash count", "KSI (K+A) crashes", "Win_From_Mi", "Win_To_Mi"] if c and c in work.columns]
    fig = px.scatter(
        work,
        x=metric,
        y="KSI (K+A) crashes",
        hover_data=hover_cols,
        title="KSI (K+A) crashes vs HIN priority index",
    )
    fig.update_traces(marker=dict(size=7, opacity=0.65))
    fig.update_layout(xaxis_title="HIN priority index", yaxis_title="KSI (K+A) crash count")
    return _polish_figure(fig), work[[c for c in [metric, "KSI (K+A) crashes", "Crash count", route_col, id_col] if c and c in work.columns]].copy()


def _hin_table_for_display(hin, metric, top_n=20):
    """Decision table for top sliding windows with module crash counts and dashboard KSI."""
    work = _prepare_window_dashboard_table(hin)
    if work is None or getattr(work, "empty", True):
        return pd.DataFrame()
    metric = "HIN_Priority_Index" if "HIN_Priority_Index" in work.columns else metric
    if not metric or metric not in work.columns:
        return pd.DataFrame()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
    work = work.sort_values(metric, ascending=False).head(top_n).reset_index(drop=True)

    id_col = _normal_col(work, ["SegID", "WindowID", "SlidingWindowID", "SegmentID", "UnitID", "SourceSegmentID", "CorridorID"])
    route_col = _section7_route_col_from_tables(work, None)
    length_col = _dashboard_length_col(work)
    from_col = _normal_col(work, ["Win_From_Mi", "FromMile", "From_Mile", "from_mile", "BeginMile", "StartMile", "WindowFromMile"])
    to_col = _normal_col(work, ["Win_To_Mi", "ToMile", "To_Mile", "to_mile", "EndMile", "WindowToMile"])

    out = pd.DataFrame()
    out["Rank"] = range(1, len(work) + 1)
    out["SegID"] = work[id_col].astype(str).values if id_col else [f"HIN_{i + 1}" for i in range(len(work))]
    out["Seg/window length"] = pd.to_numeric(work[length_col], errors="coerce").round(3).values if length_col else ""
    out["From mile"] = pd.to_numeric(work[from_col], errors="coerce").round(3).values if from_col else ""
    out["To mile"] = pd.to_numeric(work[to_col], errors="coerce").round(3).values if to_col else ""
    out["Route"] = work[route_col].astype(str).values if route_col else ""
    out["Crash count"] = pd.to_numeric(work["Dashboard_Crash_Count"], errors="coerce").fillna(0).round(0).astype(int).values
    out["KSI (K+A) crashes"] = pd.to_numeric(work["KSI_Crashes_Dashboard"], errors="coerce").fillna(0).round(0).astype(int).values
    out["HIN index"] = pd.to_numeric(work[metric], errors="coerce").round(3).values
    out["Crashes/mile/year"] = pd.to_numeric(work.get("Crash_per_Mile_per_Year", 0), errors="coerce").fillna(0).round(3).values
    out["KSI (K+A)/mile/year"] = pd.to_numeric(work.get("KSI_per_Mile_per_Year", 0), errors="coerce").fillna(0).round(3).values
    out["Crash years used"] = pd.to_numeric(work.get("Dashboard_Analysis_Years", 1), errors="coerce").fillna(1).round(1).values
    return out


# --- V43 dashboard-only fix: safe sliding-window HIN summary ---
# This override fixes the dashboard crash where the summary used a metric name
# from one table but then summarized another table.  It also keeps selected-HIN
# crash/KSI capture tied to the sliding-window assigned-crash table.


def _section7_assigned_crashes_for_dashboard(crashes=None):
    results = _section7_results_for_dashboard()
    assigned = results.get("assigned_crashes")
    if assigned is not None and not getattr(assigned, "empty", True):
        return assigned
    if crashes is not None and not getattr(crashes, "empty", True):
        return crashes
    return _dashboard_crashes_source_gdf(crashes)


def _render_hin_network_summary(st, hin, crashes):
    """High Injury Network summary using prepared sliding-window rows.

    The important detail is that the metric and selected rows must come from
    the same prepared table.  This avoids KeyError: HIN_Priority_Index when the
    incoming HIN table uses a different column name than risk_windows.
    """
    work = _prepare_window_dashboard_table(hin, crashes=crashes)
    if work is None or getattr(work, "empty", True):
        return

    metric = "HIN_Priority_Index" if "HIN_Priority_Index" in work.columns else _dashboard_hin_metric(work)
    if not metric or metric not in work.columns:
        st.info("HIN summary needs a HIN priority index or window score field.")
        return

    st.markdown(
        "<div class='dashboard-section-title'>High Injury Network summary "
        "<span>custom HIN threshold, miles, crashes, and KSI (K+A) capture</span></div>",
        unsafe_allow_html=True,
    )

    work = work.copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)

    length_col = _dashboard_length_col(work)
    if length_col:
        work[length_col] = pd.to_numeric(work[length_col], errors="coerce").fillna(0)
    else:
        length_col = "__unit_length__"
        work[length_col] = 1.0

    mode_col, control_col, c_miles, c_crashes, c_ksi = st.columns([1.35, 1.05, 1, 1, 1])

    method = mode_col.selectbox(
        "High-risk network threshold",
        [
            "Top percent of miles",
            "Top number of segments/windows",
            "HIN index threshold",
            "Above average HIN index",
            "Above median HIN index",
        ],
        index=0,
        key="hin_summary_threshold_mode_v43",
    )

    top_percent = 10.0
    top_n = 20
    index_threshold = 50.0

    with control_col:
        if method == "Top percent of miles":
            top_percent = st.number_input(
                "Top percent",
                min_value=1.0,
                max_value=100.0,
                value=float(st.session_state.get("hin_summary_top_percent", 10.0)),
                step=1.0,
                key="hin_summary_top_percent_v43",
            )
        elif method == "Top number of segments/windows":
            top_n = st.number_input(
                "Top N",
                min_value=1,
                max_value=max(int(len(work)), 1),
                value=min(20, max(int(len(work)), 1)),
                step=1,
                key="hin_summary_top_n_v43",
            )
        elif method == "HIN index threshold":
            index_threshold = st.number_input(
                "Minimum HIN index",
                min_value=0.0,
                max_value=100.0,
                value=float(st.session_state.get("hin_summary_index_threshold", 50.0)),
                step=5.0,
                key="hin_summary_index_threshold_v43",
            )
        elif method == "Above average HIN index":
            st.metric("Average", f"{work[metric].mean():,.2f}")
        elif method == "Above median HIN index":
            st.metric("Median", f"{work[metric].median():,.2f}")

    selected = _selected_hin_subset(
        work,
        metric,
        method,
        top_percent=top_percent,
        top_n=top_n,
        index_threshold=index_threshold,
    )

    total_mi = float(pd.to_numeric(work[length_col], errors="coerce").fillna(0).sum())
    high_mi = float(pd.to_numeric(selected[length_col], errors="coerce").fillna(0).sum()) if not selected.empty else 0.0
    pct_mi = high_mi / total_mi * 100 if total_mi else 0.0

    assigned = _section7_assigned_crashes_for_dashboard(crashes)
    crash_subset, capture_note = _selected_hin_crash_subset(selected, assigned)

    total_crashes = _safe_unique_crash_count(assigned)
    high_crashes = _safe_unique_crash_count(crash_subset)
    pct_crash = high_crashes / total_crashes * 100 if total_crashes else 0.0

    total_ksi = _total_ka_from_crashes(assigned)
    high_ksi = _total_ka_from_crashes(crash_subset)
    pct_ksi = high_ksi / total_ksi * 100 if total_ksi else 0.0

    c_miles.metric("High-risk miles", f"{high_mi:,.2f} mi", f"{pct_mi:,.1f}% of analyzed miles")
    c_crashes.metric("Crashes on selected HIN", f"{high_crashes:,}", f"{pct_crash:,.1f}% of assigned sliding-window crashes")
    c_ksi.metric("KSI (K+A) on selected HIN", f"{high_ksi:,}", f"{pct_ksi:,.1f}% of KSI (K+A)")

    with st.expander("Selected HIN summary rows", expanded=False):
        selected_display = _add_dashboard_rate_columns(selected, crashes=assigned)
        preview_cols = []
        for c in [
            _normal_col(selected_display, ["RiskSegmentID", "WindowID", "SlidingWindowID", "SegmentID", "UnitID", "CorridorID"]),
            _dashboard_route_col(selected_display),
            _dashboard_length_col(selected_display),
            _normal_col(selected_display, ["Win_Start_M", "FromMile", "From_Mile", "WindowFromMile", "Win_From_Mi"]),
            _normal_col(selected_display, ["Win_End_M", "ToMile", "To_Mile", "WindowToMile", "Win_To_Mi"]),
            "Dashboard_Crash_Count",
            "KSI_Crashes_Dashboard",
            metric,
            "Crash_per_Mile_per_Year",
            "KSI_per_Mile_per_Year",
            "Dashboard_Analysis_Years",
        ]:
            if c and c in selected_display.columns and c not in preview_cols:
                preview_cols.append(c)
        if preview_cols:
            st.dataframe(_safe_dataframe_for_display(selected_display[preview_cols].head(50)), width="stretch", hide_index=True)
        else:
            st.info("No displayable HIN rows are available for the selected threshold.")

    years = _dashboard_year_count(assigned)
    st.caption(
        f"Crashes/mile/year uses the detected crash data year span: {years:g} year(s). "
        f"{capture_note} Dashboard controls summarize existing HIN results only; they do not recalculate HIN scores."
    )

    _render_hin_dashboard_charts(st, work)

# --- V10 report recommendation and KSI summary overrides -------------------
# Report/dashboard display-export only. These helpers do not change workflow
# spatial assignment, crash-density, corridor, or sliding-window calculations.

def _ksi_summary_by_spatial_unit_table(tables, top_n=20):
    """Report table focused on KSI (K+A) by spatial unit.

    The old severity summary could include B/C/O.  For the Word report, this
    table is intentionally limited to K, A, and KSI (K+A) because it is used as
    an injury-severity priority summary.
    """
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    if crashes is None or getattr(crashes, "empty", True):
        return pd.DataFrame()

    unit_col = _unit_col(crashes) or _normal_col(
        crashes,
        ["UnitID", "IntersectionID", "CorridorID", "SegmentID", "SegID"],
    )
    if not unit_col:
        return pd.DataFrame()

    work = _drop_geometry(crashes).copy()
    sev_cols = _severity_count_columns(work)
    use_person_counts = False
    if sev_cols and not _is_fars_fatal_only_dataset(work):
        for col in [sev_cols.get("K"), sev_cols.get("A")]:
            if col and col in work.columns and pd.to_numeric(work[col], errors="coerce").fillna(0).sum() > 0:
                use_person_counts = True
                break

    if use_person_counts:
        rows = pd.DataFrame({unit_col: work[unit_col].astype(str)})
        rows["K"] = pd.to_numeric(work[sev_cols.get("K")], errors="coerce").fillna(0) if sev_cols.get("K") else 0
        rows["A"] = pd.to_numeric(work[sev_cols.get("A")], errors="coerce").fillna(0) if sev_cols.get("A") else 0
        rows = rows.groupby(unit_col, dropna=False)[["K", "A"]].sum().reset_index()
    else:
        kabco = _kabco_col(work)
        if not kabco or kabco not in work.columns:
            return pd.DataFrame()
        work["KABCO_Normalized"] = work[kabco].map(normalize_kabco_value).astype(str).str.upper()
        work = work[work["KABCO_Normalized"].isin(["K", "A"])].copy()
        if work.empty:
            return pd.DataFrame()
        rows = (
            work.groupby([unit_col, "KABCO_Normalized"], dropna=False)
            .size()
            .reset_index(name="Count")
            .pivot_table(index=unit_col, columns="KABCO_Normalized", values="Count", aggfunc="sum", fill_value=0)
            .reset_index()
        )
        for col in ["K", "A"]:
            if col not in rows.columns:
                rows[col] = 0

    rows["K"] = pd.to_numeric(rows["K"], errors="coerce").fillna(0).astype(int)
    rows["A"] = pd.to_numeric(rows["A"], errors="coerce").fillna(0).astype(int)
    rows["KSI (K+A)"] = rows["K"] + rows["A"]

    density = tables.get("Crash density results")
    if density is not None and not getattr(density, "empty", True):
        d = _drop_geometry(density).copy()
        d_unit = _unit_col(d) or _normal_col(d, ["UnitID", "IntersectionID", "CorridorID", "SegmentID", "SegID"])
        keep = [c for c in [d_unit, "UnitType", "City", "CrashCount", "CrashDensity"] if c and c in d.columns]
        if d_unit and keep:
            rows = rows.merge(d[keep].drop_duplicates(subset=[d_unit]), left_on=unit_col, right_on=d_unit, how="left")
            if d_unit != unit_col and d_unit in rows.columns:
                rows = rows.drop(columns=[d_unit])

    rows = rows.sort_values(["KSI (K+A)", "K", "A"], ascending=False).head(top_n).reset_index(drop=True)
    rows.insert(0, "Rank", range(1, len(rows) + 1))

    out_cols = ["Rank", unit_col]
    for c in ["UnitType", "City", "CrashCount", "CrashDensity", "K", "A", "KSI (K+A)"]:
        if c in rows.columns and c not in out_cols:
            out_cols.append(c)
    out = rows[out_cols].copy()
    if unit_col != "UnitID":
        out = out.rename(columns={unit_col: "UnitID"})
    return out


def _export_tables_only(tables, top_n=20):
    """Report-ready tables with KSI-only severity summary."""
    out = {}
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density = tables.get("Crash density results")
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))

    if crashes is not None:
        type_col = _crash_type_col(crashes)
        if type_col:
            out["Crash type summary"] = _aggregate(crashes, type_col, None, "Count", top_n)
        years = _year_series_from_crashes(crashes)
        if years is not None:
            ydf = pd.DataFrame({"Year": years})
            ydf = ydf[ydf["Year"].ne("Unknown")]
            if not ydf.empty:
                out["Crash year summary"] = ydf.groupby("Year", dropna=False).size().reset_index(name="Count").sort_values("Year")
        ksi = _ksi_summary_by_spatial_unit_table(tables, top_n=top_n)
        if not ksi.empty:
            out["KSI summary by spatial unit"] = ksi

    if density is not None:
        top_density = _top_density_export_table(density, top_n=top_n)
        if not top_density.empty:
            out["Top crash-density spatial units"] = top_density

    if hin is not None:
        metric = _dashboard_hin_metric(hin)
        top_hin = _hin_table_for_display(hin, metric, top_n) if metric else pd.DataFrame()
        if not top_hin.empty:
            out["Top HIN/risk spatial units"] = top_hin

    return {name: _safe_dataframe_for_display(df) for name, df in out.items()}


def _report_tables(tables, top_n=20):
    """Report-ready tables for dashboard preview/export."""
    return _export_tables_only(tables, top_n=top_n)


def _clean_route_name_for_report(value):
    text = str(value or "").strip()
    if not text or text.lower() in ["nan", "none", "unknown"]:
        return ""
    return text


def _join_report_names(names, max_items=5):
    vals = []
    seen = set()
    for n in names:
        t = _clean_route_name_for_report(n)
        if t and t.lower() not in seen:
            vals.append(t)
            seen.add(t.lower())
        if len(vals) >= max_items:
            break
    if not vals:
        return ""
    if len(vals) == 1:
        return vals[0]
    if len(vals) == 2:
        return f"{vals[0]} and {vals[1]}"
    return ", ".join(vals[:-1]) + f", and {vals[-1]}"


def _top_hin_routes_for_recommendations(tables, top_n=5):
    hin = tables.get("HIN risk segments", tables.get("HIN corridors"))
    if hin is None or getattr(hin, "empty", True):
        return pd.DataFrame()
    df = _drop_geometry(hin).copy()
    metric = _dashboard_hin_metric(df)
    route_col = _normal_col(df, ["Route", "FULLNAME", "RoadName", "RouteName", "RouteName_Calc", "CorridorRoute"])
    length_col = _dashboard_length_col(df)
    if not metric or not route_col:
        return pd.DataFrame()
    df[metric] = pd.to_numeric(df[metric], errors="coerce").fillna(0)
    df[route_col] = df[route_col].map(_clean_route_name_for_report)
    df = df[df[route_col].ne("")].copy()
    if df.empty:
        return pd.DataFrame()
    agg_dict = {
        "Max_HIN": (metric, "max"),
        "Avg_HIN": (metric, "mean"),
        "Segment_Window_Count": (metric, "size"),
    }
    if length_col and length_col in df.columns:
        df[length_col] = pd.to_numeric(df[length_col], errors="coerce").fillna(0)
        agg_dict["Miles"] = (length_col, "sum")
    out = df.groupby(route_col, dropna=False).agg(**agg_dict).reset_index().rename(columns={route_col: "Route"})
    out = out.sort_values(["Max_HIN", "Avg_HIN", "Segment_Window_Count"], ascending=False).head(top_n).reset_index(drop=True)
    return out




def _intersection_route_lookup_from_session():
    """Build SignalID/UnitID -> readable Route 1 / Route 2 labels.

    Intersection buffers preserve SignalID where available.  The signal-to-road
    table stores one row per nearby route for that signal.  This helper uses
    those existing workflow tables only for report/dashboard labels; it does
    not change spatial-unit creation or crash assignment.
    """
    try:
        signals = st.session_state.get("signals_with_corridor", None)
    except Exception:
        signals = None

    if signals is None or getattr(signals, "empty", True):
        return {}

    sig = _drop_geometry(signals).copy()
    sig_id_col = _normal_col(sig, ["SignalID", "signal_id", "intersection_id", "IntersectionID"])
    route_col = _normal_col(sig, ["Route", "RouteName", "FULLNAME", "RoadName", "Route_Normalized"])

    if not sig_id_col or not route_col or sig_id_col not in sig.columns or route_col not in sig.columns:
        return {}

    sig[sig_id_col] = sig[sig_id_col].astype(str).str.strip()
    sig[route_col] = sig[route_col].map(_clean_route_name_for_report)
    sig = sig[(sig[sig_id_col] != "") & (sig[route_col] != "")].copy()

    if sig.empty:
        return {}

    lookup = {}
    for sid, group in sig.groupby(sig_id_col, dropna=False):
        routes = []
        seen = set()
        for value in group[route_col].tolist():
            route = _clean_route_name_for_report(value)
            key = route.lower()
            if route and key not in seen:
                routes.append(route)
                seen.add(key)
        if not routes:
            continue
        label = " / ".join(routes[:2]) if len(routes) >= 2 else routes[0]
        sid_text = str(sid).strip()
        lookup[sid_text] = label
        # Common UnitID pattern for intersection buffers.
        if sid_text.replace(".", "", 1).isdigit():
            try:
                sid_int = str(int(float(sid_text)))
                lookup[sid_int] = label
                lookup[f"INT_{sid_int}"] = label
                lookup[f"Intersection {sid_int}"] = label
            except Exception:
                pass
    return lookup


def _intersection_label_from_report_row(row):
    """Return a route-pair label for an intersection row if it can be recovered."""
    lookup = _intersection_route_lookup_from_session()
    if not lookup:
        return ""

    cols = list(getattr(row, "index", []))
    candidate_cols = [
        "SignalID",
        "signal_id",
        "intersection_id",
        "IntersectionID",
        "UnitID",
        "Spatial unit id",
    ]

    for cand in candidate_cols:
        for col in cols:
            key = str(col).lower().replace("_", "").replace(" ", "")
            ckey = str(cand).lower().replace("_", "").replace(" ", "")
            if key == ckey or ckey in key:
                raw = str(row.get(col, "")).strip()
                if not raw:
                    continue
                candidates = [raw]
                # INT_92 -> 92
                m = re.search(r"(\d+)", raw)
                if m:
                    candidates.extend([m.group(1), f"INT_{m.group(1)}"])
                # 92.0 -> 92
                try:
                    as_int = str(int(float(raw)))
                    candidates.extend([as_int, f"INT_{as_int}"])
                except Exception:
                    pass
                for value in candidates:
                    if value in lookup:
                        return lookup[value]
    return ""

def _location_label_for_report_row(row, unit_type=""):
    """Build a readable location label for report recommendations.

    Intersections should show the two crossing road names when available.
    Corridors/segments should show the route name when available.  This helper
    is report text only and does not change any analysis calculations.
    """
    cols = list(getattr(row, "index", []))

    def val_from(candidates):
        for cand in candidates:
            for col in cols:
                key = str(col).lower().replace("_", "").replace(" ", "")
                ckey = str(cand).lower().replace("_", "").replace(" ", "")
                if key == ckey or ckey in key:
                    value = _clean_route_name_for_report(row.get(col, ""))
                    if value:
                        return value
        return ""

    road1 = val_from(["RoadName1", "Road1", "Route1", "Street1", "FromRoad", "PrimaryRoad", "Approach1"])
    road2 = val_from(["RoadName2", "Road2", "Route2", "Street2", "ToRoad", "CrossStreet", "Approach2"])
    if road1 and road2:
        return f"{road1} / {road2}"

    is_intersection = "intersection" in str(unit_type or "").lower()
    unit_text = val_from(["UnitID", "IntersectionID", "Spatial unit id"])
    if unit_text.upper().startswith("INT_"):
        is_intersection = True
    try:
        if str(st.session_state.get("analysis_type", "")).lower().startswith("intersection"):
            is_intersection = True
    except Exception:
        pass

    if is_intersection:
        enriched = _intersection_label_from_report_row(row)
        if enriched:
            return enriched

    route = val_from(["Route", "FULLNAME", "RoadName", "RouteName", "CorridorRoute", "RouteName_Calc", "FacilityName"])
    if route:
        return route

    unit = val_from(["UnitID", "IntersectionID", "CorridorID", "SegmentID", "SegID", "Spatial unit id"])
    return unit or "selected location"


def _top_density_locations_for_recommendations(tables, top_n=6):
    """Top crash-density locations with analysis-aware labels.

    For intersection reports, Location is built from Route 1 / Route 2 when
    those columns exist.  For corridor reports, Location is the route/corridor
    name.  For segment reports, Location falls back to route or unit ID.
    """
    density = tables.get("Crash density results")
    if density is None or getattr(density, "empty", True):
        return pd.DataFrame()

    df = _drop_geometry(density).copy()
    density_col = _normal_col(df, ["CrashDensity", "Crash_Density", "crash_density"])
    count_col = _crash_count_col(df)
    unit_col = _unit_col(df) or _normal_col(df, ["UnitID", "IntersectionID", "CorridorID", "SegmentID", "SegID"])
    unit_type_col = _normal_col(df, ["UnitType", "IntersectionType", "CorridorType", "SegmentType"])

    if not density_col:
        return pd.DataFrame()

    df[density_col] = pd.to_numeric(df[density_col], errors="coerce").fillna(0)
    df = df.sort_values(density_col, ascending=False).head(top_n).copy()
    if df.empty:
        return pd.DataFrame()

    locations = []
    for _, row in df.iterrows():
        unit_type = str(row.get(unit_type_col, "")) if unit_type_col else ""
        locations.append(_location_label_for_report_row(row, unit_type=unit_type))

    out = pd.DataFrame({
        "Location": locations,
        "CrashDensity": pd.to_numeric(df[density_col], errors="coerce").fillna(0).values,
    })
    if count_col and count_col in df.columns:
        out["CrashCount"] = pd.to_numeric(df[count_col], errors="coerce").fillna(0).astype(int).values
    if unit_col and unit_col in df.columns:
        out["UnitID"] = df[unit_col].astype(str).values
    return out.reset_index(drop=True)


def _top_crash_patterns_for_recommendations(crashes, top_n=3):
    if crashes is None or getattr(crashes, "empty", True):
        return []
    type_col = _crash_type_col(crashes)
    if not type_col or type_col not in crashes.columns:
        return []
    s = crashes[type_col].dropna().astype(str).str.strip()
    s = s[~s.str.lower().isin(["", "nan", "none", "unknown"])]
    if s.empty:
        return []
    return list(s.value_counts().head(top_n).index)


def _recommendation_priority_context(tables):
    """Return analysis-specific priority language for the report.

    HIN is used only for segment/HIN reports.  Intersection and corridor
    reports use the crash-density / spatial-unit results so the recommendation
    text matches the workflow that generated the report.
    """
    analysis_type = _report_analysis_type(tables).lower()
    density_locs = _top_density_locations_for_recommendations(tables, top_n=6)

    if "intersection" in analysis_type:
        names = _join_report_names(density_locs["Location"].tolist() if not density_locs.empty else [], max_items=6)
        first = _clean_route_name_for_report(density_locs["Location"].iloc[0]) if not density_locs.empty else "the highest-ranked intersection"
        return {
            "mode": "intersection",
            "heading1": "1. Prioritize high-risk intersections for detailed safety review",
            "priority_type": "Intersection",
            "priority_issue": "High crash-density / KSI screening priority",
            "first_location": first,
            "location_names": names,
            "paragraph1": (
                f"Begin detailed safety review at {names}, because these signalized intersections appear in the highest crash-density or injury-priority rankings. "
                "Review each intersection by Route 1 / Route 2, crash diagram, crash severity, signal phasing, turn-lane operations, queueing, pedestrian/bicycle conditions, lighting, and sight distance before selecting countermeasures."
                if names else
                "Use the highest-ranked signalized intersections from the result table as the starting point for detailed safety diagnosis. Review each location by Route 1 / Route 2 before selecting countermeasures."
            ),
            "paragraph2": "Use the intersection ranking table to identify near-term spot-treatment candidates. These are location-specific priorities rather than corridor-level HIN priorities.",
            "follow_up": "Intersection crash diagram and operations review",
            "countermeasure": "Signal operations, turn-lane review, pedestrian/bicycle treatments, visibility, and quick-build spot safety treatments",
        }

    if "corridor" in analysis_type:
        names = _join_report_names(density_locs["Location"].tolist() if not density_locs.empty else [], max_items=6)
        first = _clean_route_name_for_report(density_locs["Location"].iloc[0]) if not density_locs.empty else "the highest-ranked corridor"
        return {
            "mode": "corridor",
            "heading1": "1. Prioritize high-risk corridors for detailed safety review",
            "priority_type": "Corridor",
            "priority_issue": "High corridor crash-density / injury screening priority",
            "first_location": first,
            "location_names": names,
            "paragraph1": (
                f"Begin corridor-level safety diagnosis on {names}, because these route corridors appear in the highest crash-density or injury-priority rankings. "
                "For each route, review crash diagrams, signal spacing, access spacing, turn-lane operations, queueing, speed environment, lighting, and pedestrian/bicycle context before selecting countermeasures."
                if names else
                "Use the highest-ranked route corridors from the result table as the starting point for detailed safety diagnosis."
            ),
            "paragraph2": "Use the corridor ranking table to identify the first route corridors for follow-up review, and separately note any short high-density locations that may be suitable for spot treatments.",
            "follow_up": "Corridor safety diagnosis",
            "countermeasure": "Signal operations, access management, speed management, intersection review, and systemic corridor treatments",
        }

    # Segment / HIN report
    hin_routes = _top_hin_routes_for_recommendations(tables, top_n=5)
    route_names = _join_report_names(hin_routes["Route"].tolist() if not hin_routes.empty else [], max_items=5)
    first_route = _clean_route_name_for_report(hin_routes["Route"].iloc[0]) if not hin_routes.empty else "the highest-ranked HIN segment or route"
    return {
        "mode": "segment",
        "heading1": "1. Prioritize HIN segments and routes for detailed safety review",
        "priority_type": "HIN segment / route",
        "priority_issue": "High HIN screening score",
        "first_location": first_route,
        "location_names": route_names,
        "paragraph1": (
            f"Begin detailed safety diagnosis on {route_names}, because these routes appear in the highest HIN/risk rankings. "
            f"Treat {first_route} as an initial priority because it has the highest HIN screening score in the current results."
            if route_names else
            "Use the highest-ranked HIN segments, windows, or routes from the result table as the starting point for detailed safety diagnosis."
        ),
        "paragraph2": "Review high crash-density locations separately from the HIN route priorities so the implementation plan can address both systemic injury risk and localized crash clusters.",
        "follow_up": "Segment / corridor safety diagnosis",
        "countermeasure": "Speed management, access management, signal operations, intersection review, and systemic safety treatments",
    }


def _add_recommendations_section(doc, tables):
    """Add data-driven, workflow-aware screening recommendations."""
    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    density_locs = _top_density_locations_for_recommendations(tables, top_n=6)
    patterns = _top_crash_patterns_for_recommendations(crashes, top_n=3)
    context = _recommendation_priority_context(tables)

    doc.add_heading("Key recommendations and suggested next steps", level=1)
    doc.add_paragraph(
        "The following recommendations are generated from the current screening results. "
        "They are intended to identify where detailed safety diagnosis should begin; they do not replace field review, crash-diagram review, or engineering judgment."
    )

    doc.add_heading(context["heading1"], level=2)
    doc.add_paragraph(context["paragraph1"])

    doc.add_heading("2. Use the ranking results to identify near-term follow-up locations", level=2)
    density_names = _join_report_names(density_locs["Location"].tolist() if not density_locs.empty else [], max_items=6)
    if context["mode"] == "intersection":
        if density_names:
            doc.add_paragraph(
                f"Review high crash-density intersections such as {density_names} as near-term candidates for site-level safety review. "
                "These locations should be confirmed using crash diagrams, field observations, and signal operations review."
            )
        else:
            doc.add_paragraph("Use the intersection crash-density ranking table to identify near-term site-level safety review candidates.")
    elif context["mode"] == "corridor":
        if density_names:
            doc.add_paragraph(
                f"Review high-priority corridors such as {density_names} for corridor-level diagnosis. "
                "Within each route, identify whether the crash pattern is corridor-wide or concentrated at specific intersections or short segments."
            )
        else:
            doc.add_paragraph("Use the corridor crash-density ranking table to identify route corridors for near-term follow-up review.")
    else:
        if density_names:
            doc.add_paragraph(
                f"Review high crash-density locations such as {density_names} as potential near-term spot-treatment candidates. "
                "These locations should be considered separately from the HIN segment or route priorities so the implementation plan can address both systemic risk and localized crash clusters."
            )
        else:
            doc.add_paragraph("Use the crash-density ranking table to identify near-term spot-treatment candidates alongside the HIN screening results.")

    doc.add_heading("3. Target dominant crash patterns after location-level validation", level=2)
    pattern_text = _join_report_names(patterns, max_items=3)
    if pattern_text:
        doc.add_paragraph(
            f"The leading crash patterns in the current dataset include {pattern_text}. "
            "Use these patterns as a starting point for diagnosis, then confirm the pattern at each priority location before selecting treatments."
        )
    else:
        doc.add_paragraph("Review crash type, severity, and location patterns at each priority location before selecting treatments.")

    doc.add_heading("4. Validate data and assumptions before programming projects", level=2)
    doc.add_paragraph(
        "Confirm crash geocoding, crash-year coverage, severity/KSI mapping, roadway segmentation or unit definitions, selected road-class filters, and threshold settings before using the results for funding or project programming. "
        "Locations already improved or already programmed should be flagged during the follow-up review."
    )

    doc.add_heading("5. Convert the screening into an action matrix", level=2)
    doc.add_paragraph(
        "Prepare a short implementation matrix listing each priority location, the main safety issue, likely contributing factors, recommended follow-up review, potential countermeasure category, timeframe, and responsible lead."
    )
    first_location = context.get("first_location") or "Highest-ranked location"
    second_location = _clean_route_name_for_report(density_locs["Location"].iloc[0]) if not density_locs.empty else "Highest crash-density location"
    matrix_rows = [
        [first_location, context["priority_type"], context["priority_issue"], context["follow_up"], context["countermeasure"], "Near-term study", "Agency to assign"],
    ]
    if second_location and second_location.lower() != str(first_location).lower():
        matrix_rows.append([second_location, "Spot location", "Localized crash concentration", "Site-level crash review", "Quick-build or spot safety treatment", "Short term", "Agency to assign"])
    action_df = pd.DataFrame(matrix_rows, columns=["Priority location", "Priority type", "Main safety issue", "Follow-up review", "Potential countermeasure category", "Timeframe", "Lead"])
    _add_dataframe_table(doc, action_df, max_rows=10)

def _export_dashboard_docx(tables, selected_blocks, selected_maps, extra_figures=None, maps=None, overlay_layers=None, report_timezone=None):
    """Word report export with data-driven recommendations and KSI-only table names."""
    if Document is None:
        return None
    doc = Document()
    title = _report_title(tables)
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {_report_time_text(report_timezone)}")
    doc.add_paragraph(f"User email: {_report_user_email()}")

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(_report_introduction_text(tables))

    crashes = tables.get("Assigned crashes", tables.get("Uploaded crashes"))
    _add_data_section(doc, tables, crashes)
    _add_methodology_section(doc, tables)
    _add_limitations_section(doc)

    doc.add_heading("Results and visualization", level=1)
    kpis = _summary_kpi_values(crashes)
    doc.add_heading("Crash summary", level=2)
    _add_key_value_table(doc, [(k, _format_kpi_value(v)) for k, v in kpis.items()])

    figures = _build_default_figures(tables) + (extra_figures or [])
    _add_bubble_size_note_if_needed(doc, figures)
    for fig_title, fig, data in figures:
        if selected_blocks and fig_title not in selected_blocks:
            continue
        doc.add_heading(str(fig_title), level=2)
        img = _figure_to_png_bytes(_polish_figure(fig))
        if img:
            doc.add_picture(io.BytesIO(img), width=Inches(6.5))
        else:
            doc.add_paragraph("Chart image could not be generated in this environment. The summary table is included below.")
        table_df = _safe_dataframe_for_display(data.copy())
        if not table_df.empty:
            doc.add_paragraph("Summary table")
            _add_dataframe_table(doc, table_df, max_rows=20)

    if selected_maps:
        doc.add_heading("Selected map layers", level=2)
        for m in selected_maps:
            doc.add_heading(str(m), level=3)
            if maps and m in maps:
                map_png = _static_map_png(maps[m], str(m), overlay_layers=overlay_layers)
                if map_png:
                    doc.add_picture(io.BytesIO(map_png), width=Inches(6.5))
                else:
                    doc.add_paragraph("Static map image could not be generated.")
            else:
                doc.add_paragraph("Map layer selected in dashboard builder.")

    _add_recommendations_section(doc, tables)

    doc.add_heading("Decision-ready result tables", level=2)
    for table_name, df in _report_tables(tables).items():
        doc.add_heading(str(table_name), level=3)
        if not _add_dataframe_table(doc, df, max_rows=25):
            doc.add_paragraph("No records available.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- V16 workflow-specific dashboard/report map fixes -------------------------
# These overrides only change dashboard/report map selection and static map
# metric selection. They do not change crash assignment, crash density, corridor,
# signal, or HIN calculations.

def _metric_for_map(gdf, map_name):
    cols = list(getattr(gdf, "columns", []))
    name = str(map_name or "").lower()

    if "hin" in name:
        for c in ["HIN_Priority_Index", "RiskScore", "CrashDensity", "CrashCount", "Crash_Count"]:
            if c in cols:
                return c

    if "count" in name:
        for c in ["CrashCount", "Crash_Count", "TotalCrashes", "Total_Crashes", "CrashDensity", "CrashDensity_per_mile"]:
            if c in cols:
                return c

    if "density" in name:
        for c in ["CrashDensity", "CrashDensity_per_mile", "CrashCount", "Crash_Count"]:
            if c in cols:
                return c

    for c in ["CrashDensity", "HIN_Priority_Index", "CrashCount", "Crash_Count"]:
        if c in cols:
            return c

    try:
        nums = _numeric_cols(_drop_geometry(gdf))
        return nums[0] if nums else None
    except Exception:
        return None


def _available_maps(st):
    """Return only maps that match the current workflow.

    This prevents stale intersection maps or context-only corridor layers from
    appearing in a corridor report after the user previously ran a different
    workflow in the same session.
    """
    maps = {}
    density = st.session_state.get("spatial_units_density_map")
    analysis_type = str(
        st.session_state.get(
            "analysis_type",
            st.session_state.get("spatial_unit_selector", st.session_state.get("spatial_unit", ""))
        )
    ).lower()

    if density is not None and not getattr(density, "empty", True):
        density_map = _repair_gdf_crs(density, st)

        if "intersection" in analysis_type:
            # Intersection reports should show one map: crash count by intersection.
            maps["Crash count map"] = density_map

        elif "corridor" in analysis_type:
            # Corridor reports should show both crash count and crash density,
            # using the corridor spatial-unit result layer, not the separate
            # corridor-context layer.
            maps["Crash count map"] = density_map
            maps["Crash density map"] = density_map

        else:
            # Segment and generic workflows keep crash density as the primary
            # spatial-unit map. HIN maps are added below when available.
            maps["Crash density map"] = density_map

    results = st.session_state.get("section7_results")
    if results is not None and "intersection" not in analysis_type and "corridor" not in analysis_type:
        risk_segments = results.get("risk_segments") if isinstance(results, dict) else None
        if risk_segments is not None and not getattr(risk_segments, "empty", True):
            hin_map = _repair_gdf_crs(risk_segments, st)
            maps["HIN priority map"] = hin_map
            metric = _dashboard_hin_metric(hin_map)
            if metric:
                work = hin_map.copy()
                work[metric] = pd.to_numeric(work[metric], errors="coerce").fillna(0)
                avg = work[metric].mean()
                med = work[metric].median()
                maps["HIN above average map"] = work[work[metric] >= avg].copy()
                maps["HIN above median map"] = work[work[metric] >= med].copy()

    return {
        k: v
        for k, v in maps.items()
        if v is not None and not getattr(v, "empty", True)
    }


def _dashboard_default_map_selection(maps):
    names = list(maps.keys())
    preferred = [
        name for name in ["Crash count map", "Crash density map", "HIN priority map"]
        if name in maps
    ]
    return preferred[:2] if preferred else names[: min(2, len(names))]


def _sanitize_dashboard_map_state(st, maps):
    """Remove map names that no longer exist for the active workflow."""
    key = "dash_builder_map_layers"
    valid_names = list(maps.keys())
    previous = st.session_state.get(key, None)
    if previous is None:
        return
    if not isinstance(previous, (list, tuple, set)):
        previous = [previous]
    cleaned = [name for name in previous if name in valid_names]
    if not cleaned:
        cleaned = _dashboard_default_map_selection(maps)
    if list(previous) != cleaned:
        st.session_state[key] = cleaned


def _render_dashboard_builder(st, tables):
    maps = _available_maps(st)
    _sanitize_dashboard_map_state(st, maps)

    custom_figures = st.session_state.get("dashboard_custom_figures", [])
    default_figures = _build_default_figures(tables) + custom_figures
    figure_titles = [title for title, _, _ in default_figures]

    st.markdown("<div class='dashboard-section-title'>Dashboard builder <span>choose charts, tables, and map layers for one review page</span></div>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([0.27, 0.27, 0.46], gap="large")
    with c1:
        st.markdown("**Charts and figures**")
        selected_blocks = st.multiselect("Select dashboard charts", figure_titles, default=figure_titles[: min(5, len(figure_titles))], key="dash_builder_chart_blocks")
        include_tables = st.checkbox("Include ranking/data table previews", value=True, key="dash_builder_include_tables")
        if custom_figures and st.button("Clear added custom charts", key="clear_custom_dashboard_charts"):
            st.session_state["dashboard_custom_figures"] = []
            st.rerun()
    with c2:
        st.markdown("**Map layers**")
        if maps:
            selected_maps = st.multiselect(
                "Select dashboard maps",
                list(maps.keys()),
                default=_dashboard_default_map_selection(maps),
                key="dash_builder_map_layers",
                help="Dashboard maps are read-only and limited to the active workflow, so old intersection/corridor maps do not carry into the report.",
            )
        else:
            selected_maps = []
            st.caption("No dashboard maps are available for the current workflow yet.")

        overlay_sources = _workflow_overlay_sources(st)
        selected_overlays = st.multiselect(
            "Optional workflow layers on dashboard maps",
            list(overlay_sources.keys()),
            default=[name for name in ["Roads"] if name in overlay_sources],
            key="dash_builder_overlay_layers",
            help="Only selected context layers are included in the dashboard and report maps. Signals are not included unless you select Signals.",
        )
    with c3:
        st.markdown("**Exports**")
        st.caption("Export the dashboard as a static PNG summary or a Word report with charts, map summaries, and decision-ready tables.")
        report_timezone = st.selectbox("Report time zone", ["America/Denver", "Local/server time", "UTC", "America/Chicago", "America/Los_Angeles", "America/New_York"], index=0, key="dashboard_report_timezone", help="Streamlit Cloud often runs in UTC. Choose the local project timezone so the report timestamp matches your expected local time.")
        report_tz_value = None if report_timezone == "Local/server time" else report_timezone
        d1, d2, d3 = st.columns(3)
        with d1:
            png_bytes = _export_summary_image(tables, "png", extra_figures=custom_figures)
            if png_bytes:
                st.download_button("Download PNG", data=png_bytes, file_name="hin_dashboard_summary.png", mime="image/png", key="dash_export_png")
            else:
                st.caption("PNG needs kaleido")
        with d2:
            selected_maps = [m for m in selected_maps if m in maps]
            docx_bytes = _export_dashboard_docx(tables, selected_blocks, selected_maps, extra_figures=custom_figures, maps=maps, overlay_layers={name: overlay_sources[name] for name in selected_overlays if name in overlay_sources}, report_timezone=report_tz_value)
            if docx_bytes is not None:
                st.download_button("Download Word report", data=docx_bytes, file_name=_report_docx_filename(tables), mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dash_export_docx")
            else:
                st.info("Install python-docx for Word export.")
        with d3:
            data_zip_bytes = _download_generated_data_zip(st, tables)
            if data_zip_bytes:
                st.download_button(
                    "Download all generated data ZIP",
                    data=data_zip_bytes,
                    file_name="hin_generated_data_export.zip",
                    mime="application/zip",
                    key="dash_export_generated_data_zip",
                    help="Exports workflow-generated tables and GIS layers as CSV and GeoJSON files. This does not change analysis results.",
                )
            else:
                st.caption("Run workflow steps before exporting data.")

    st.markdown("<div class='dashboard-section-title'>Generated dashboard</div>", unsafe_allow_html=True)
    chart_titles = set(selected_blocks)
    for i in range(0, len(default_figures), 2):
        cols = st.columns(2)
        for j, col in enumerate(cols):
            idx = i + j
            if idx >= len(default_figures):
                continue
            title, fig, data = default_figures[idx]
            if title not in chart_titles:
                continue
            with col:
                fig.update_layout(height=330, margin=dict(l=20, r=20, t=45, b=35))
                st.plotly_chart(_polish_figure(fig), width="stretch", key=f"dash_generated_fig_{idx}")

    selected_maps = [m for m in selected_maps if m in maps]
    if selected_maps:
        map_cols = st.columns(min(2, len(selected_maps)))
        for i, map_name in enumerate(selected_maps[:2]):
            with map_cols[i % len(map_cols)]:
                st.markdown(f"**{map_name}**")
                _render_dashboard_map(
                    st,
                    map_name,
                    maps[map_name],
                    key=f"dash_map_{_safe_name(map_name)}_{i}",
                    height=420,
                    overlay_layers={name: overlay_sources[name] for name in selected_overlays if name in overlay_sources},
                )

    if include_tables:
        st.markdown("**Dashboard table preview**")
        compact_tables = _report_tables(tables)
        if compact_tables:
            table_name = st.selectbox("Preview table", list(compact_tables.keys()), key="dash_builder_preview_table")
            st.dataframe(_safe_dataframe_for_display(compact_tables[table_name]).head(25), width="stretch", hide_index=True)
        else:
            st.info("No report-ready result tables are available yet.")
